"""
Produce flop_trade_model_v21.docx from v8.docx.

v21: Based on v20 with the following changes:
  - Reliability-adjusted results lead Section 6 (preferred specification up front)
  - Observed-tariff ranking (Iran cheapest) demoted to illustrative comparison
  - Governance section trimmed (no redundant reliability repeat)
  - λ* subscript fix (_msup instead of _msubsup with empty sub)
  - Eq (5) inlined, equations renumbered (6→5, 7→6)
  - USD/kWh → $/kWh consistency
  - "Hardware generation" renamed to "GPU vintage"
  - Appendix E: construction cost regression table
  - Results preview paragraph added to Introduction
  - Manual edits integrated: "AI compute", "compute facilities", active voice,
    staffing paragraph removed, citations shortened to et al., "Our model",
    Section 3.1 opening trimmed, Costinot reference removed from Section 6

v20: Capacity-Constrained Ricardian Model restructuring.
  - Merged Sections 3+4+5 into unified Model (Section 3)
    3.1 Production Technology, 3.2 Trade Costs, 3.3 Demand,
    3.4 Sourcing and Market Equilibrium
  - New Section 4: Equilibrium Properties (Propositions 1-5, purely theoretical)
  - Renumbered: 5=Data, 6=Calibration, 7=Conclusion
  - Capacity constraints K_bar_j: training supply stack, market-clearing p_T,
    Ricardian rents, shadow values mu_j (inline), HHI_T < 1
  - New Appendix B: model derivation (B.1-B.6, 5 display equations)
  - Training/inference demand split: alpha parameter
  - 6 main-text display equations (p_T inlined), 5 propositions
"""

import csv
import pathlib
import sys
import io
import copy
from datetime import datetime
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

DOCS = pathlib.Path(r"F:\onedrive\__documents\papers\FLOPsExport\Documents")
DATA = pathlib.Path(r"F:\onedrive\__documents\papers\FLOPsExport\Data")

TAU = 0.0008
LAMBDA = 0.10
PHI = 1.08
DELTA_PUE = 0.015
THETA_REF = 15.0
GAMMA = 0.700
GPU_PRICE = 25_000
GPU_LIFE = 3
GPU_UTIL = 0.70
DC_LIFE = 15
H_YR = 365.25 * 24
RHO = GPU_PRICE / (GPU_LIFE * H_YR * GPU_UTIL)
GPU_TDP_KW = 0.700  # GPU thermal design power in kW (same as GAMMA)
ETA = 0.15          # Amortized networking cost per GPU-hour ($/hr)

# Training share of compute demand (industry estimate)
ALPHA = 0.50

# Total global compute demand in GPU-hours per period
# Calibrated to ~10 million data center GPUs at ~70% utilization (~8,766 hrs/yr)
Q_TOTAL = 60_000_000_000   # 60 billion GPU-hours

# Grid capacity CSV unit correction (kWh→GWh used /1e9 instead of /1e6)
K_BAR_SCALE = 1000

# Cost-reflective electricity prices for cost-recovery adjustment ($/kWh)
# Replacement = estimated LRMC of dominant generation at opportunity-cost fuel price
# Sources: IMF (2025) explicit fossil-fuel subsidy estimates; Lazard (2025) LCOE v17
SUBSIDY_ADJ = {
    'IRN': 0.085,  # Gas CCGT at export-parity fuel cost; IMF 2025, Lazard CCGT
    'TKM': 0.070,  # Gas CCGT at export-parity fuel cost; IMF 2025
    'DZA': 0.065,  # Gas CCGT at near-export parity; IMF 2025
    'EGY': 0.080,  # Gas/oil, reform target; IMF 2025, Egypt subsidy reform
    'UZB': 0.090,  # Gas, WB reform target; World Bank 2024
    'QAT': 0.100,  # Gas CCGT opportunity cost of LNG export; MEI GCC study
    'SAU': 0.100,  # Gas/oil opportunity cost; MEI GCC study
    'ARE': 0.095,  # Gas CCGT; MEI GCC study
    'RUS': 0.065,  # Gas/coal/nuclear mix; IMF 2025
    'KAZ': 0.085,  # Coal at cost recovery; IMF 2025
    'NGA': 0.080,  # Gas/hydro mix; IMF 2025
    'ZAF': 0.095,  # Coal, Eskom cost recovery; NERSA 2025 tariff path
    'ETH': 0.050,  # Hydro cost-recovery target; IMF 2025
}


def recompute_costs(cal, gpu_price=None, gpu_util=None,
                    p_E_delta=0.0, pue_cap=None, subsidy_adj=None):
    """Re-derive c_j from CSV primitives with parameter overrides."""
    gp = gpu_price or GPU_PRICE
    gu = gpu_util or GPU_UTIL
    rho = gp / (GPU_LIFE * H_YR * gu)
    costs = {}
    for row in cal:
        iso = row["iso3"]
        p_E = float(row["p_E_usd_kwh"])
        if subsidy_adj and iso in subsidy_adj:
            p_E = subsidy_adj[iso]
        p_E += p_E_delta
        theta = float(row["theta_summer_C"])
        pue = PHI + DELTA_PUE * max(0, theta - THETA_REF)
        if pue_cap is not None:
            pue = min(pue, pue_cap)
        c_elec = pue * GAMMA * p_E
        c_constr = float(row["c_j_construction"])
        costs[iso] = c_elec + rho + ETA + c_constr
    return costs


def run_sensitivity(cal, omega, dc_k, k_bar, sanctioned):
    """Run sensitivity analysis across parameter scenarios.

    Returns list of scenario result dicts with rankings, equilibrium
    outcomes, and Spearman rank correlations vs. the baseline.
    """
    scenarios = [
        ("Baseline calibration",                         {}),
        ("Electricity price +$0.01/kWh (\u224810% above mean)", {"p_E_delta": +0.01}),
        ("Electricity price \u2212$0.01/kWh (\u224810% below mean)", {"p_E_delta": -0.01}),
        ("GPU hardware cost +20% ($30,000/unit)",        {"gpu_price": 30_000}),
        ("GPU hardware cost \u221220% ($20,000/unit)",   {"gpu_price": 20_000}),
        ("Cooling efficiency cap (PUE \u2264 1.20)",     {"pue_cap": 1.20}),
    ]

    def _solve_mini(supply_stack_s, costs_s):
        """Standalone capacity-constrained training equilibrium solver."""
        p_T = supply_stack_s[0][1]
        for _ in range(30):
            Q_TX = 0
            for iso in dc_k:
                if iso in costs_s:
                    if costs_s[iso] > p_T:
                        Q_TX += ALPHA * omega.get(iso, 0) * Q_TOTAL
            cum_cap = 0
            found = False
            p_T_new = p_T
            for iso_j, c_j, k_j in supply_stack_s:
                if iso_j in sanctioned:
                    continue
                cum_cap += k_j * ALPHA
                if cum_cap >= Q_TX and Q_TX > 0:
                    p_T_new = c_j
                    found = True
                    break
            if found and abs(p_T_new - p_T) < 0.0001:
                p_T = p_T_new
                break
            if found:
                p_T = p_T_new
        # Count exporters and HHI
        shares = {}
        remaining = Q_TX
        for iso_j, c_j, k_j in supply_stack_s:
            if iso_j in sanctioned:
                continue
            if c_j > p_T:
                break
            ca = min(k_j * ALPHA, remaining)
            if ca > 0:
                shares[iso_j] = ca
                remaining -= ca
            if remaining <= 0:
                break
        total_exp = sum(shares.values())
        hhi = sum((s / total_exp) ** 2 for s in shares.values()) if total_exp > 0 else 1.0
        return p_T, len(shares), hhi

    def _spearman(rank_a, rank_b, isos):
        """Spearman rank correlation between two ranking dicts."""
        n = len(isos)
        if n < 2:
            return 1.0
        d_sq = sum((rank_a[iso] - rank_b[iso]) ** 2 for iso in isos)
        return 1 - 6 * d_sq / (n * (n ** 2 - 1))

    # Run all scenarios
    results = []
    baseline_rank = None
    baseline_top5 = None

    for label, kwargs in scenarios:
        costs_s = recompute_costs(cal, subsidy_adj=SUBSIDY_ADJ, **kwargs)
        ranked = sorted(costs_s.items(), key=lambda x: x[1])
        rank_map = {iso: r for r, (iso, _) in enumerate(ranked, 1)}
        top5 = [iso for iso, _ in ranked[:5]]

        # Build supply stack
        stack = sorted(
            [(iso, costs_s[iso], k_bar.get(iso, 1e12))
             for iso in costs_s if iso in k_bar],
            key=lambda x: x[1]
        )
        p_T, n_exp, hhi = _solve_mini(stack, costs_s)

        if baseline_rank is None:
            baseline_rank = rank_map
            baseline_top5 = top5

        common = set(rank_map) & set(baseline_rank)
        rho_s = _spearman(rank_map, baseline_rank, common)
        top5_match = (top5 == baseline_top5)

        result = {
            "label": label, "p_T": p_T, "n_exporters": n_exp,
            "hhi_T": hhi, "rank_corr": rho_s, "top5": top5,
            "top5_unchanged": top5_match, "kwargs": kwargs,
        }
        results.append(result)
        print(f"  Sensitivity [{label}]: p_T=${p_T:.3f}, n_exp={n_exp}, "
              f"HHI={hhi:.4f}, ρ={rho_s:.4f}, top5={'same' if top5_match else 'CHANGED'}")

    return results


def _ordinal(n):
    """Return ordinal string for integer n (e.g. 1 -> '1st', 23 -> '23rd')."""
    s = ('th', 'st', 'nd', 'rd') + ('th',) * 6
    if 11 <= (n % 100) <= 13:
        return f'{n}th'
    return f'{n}{s[n % 10]}'


def _num_word(n):
    """Return English word for small integers, digit string otherwise."""
    words = {0: 'zero', 1: 'one', 2: 'two', 3: 'three', 4: 'four',
             5: 'five', 6: 'six', 7: 'seven', 8: 'eight', 9: 'nine',
             10: 'ten', 11: 'eleven', 12: 'twelve', 13: 'thirteen',
             14: 'fourteen', 15: 'fifteen', 16: 'sixteen', 17: 'seventeen',
             18: 'eighteen', 19: 'nineteen', 20: 'twenty'}
    return words.get(n, str(n))


# Fonts
CAMBRIA_MATH = 'Cambria Math'
TIMES_NEW_ROMAN = 'Times New Roman'

# Colors
HEADING_BLUE = RGBColor(0x2F, 0x54, 0x96)
LINK_COLOR = '1F3864'

# XML namespaces
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'
SPACE_PRESERVE = 'preserve'

# Table formatting
TABLE_WIDTH_PCT = '5000'

# ═══════════════════════════════════════════════════════════════════════
# OMML HELPERS
# ═══════════════════════════════════════════════════════════════════════


def _mr(text, italic=True):
    r = OxmlElement('m:r')
    rPr = OxmlElement('m:rPr')
    sty = OxmlElement('m:sty')
    sty.set(qn('m:val'), 'i' if italic else 'p')
    rPr.append(sty)
    r.append(rPr)
    wrPr = OxmlElement('w:rPr')
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:ascii'), CAMBRIA_MATH)
    rF.set(qn('w:hAnsi'), CAMBRIA_MATH)
    wrPr.append(rF)
    r.append(wrPr)
    t = OxmlElement('m:t')
    t.set(XML_SPACE, SPACE_PRESERVE)
    t.text = text
    r.append(t)
    return r


def _v(text):
    return _mr(text, True)


def _t(text):
    return _mr(text, False)


def _msub(base, sub, base_italic=True, sub_italic=True):
    el = OxmlElement('m:sSub')
    el.append(OxmlElement('m:sSubPr'))
    e = OxmlElement('m:e')
    e.append(_mr(base, base_italic))
    el.append(e)
    s = OxmlElement('m:sub')
    s.append(_mr(sub, sub_italic))
    el.append(s)
    return el


def _msup(base, sup, base_italic=True, sup_italic=True):
    el = OxmlElement('m:sSup')
    el.append(OxmlElement('m:sSupPr'))
    e = OxmlElement('m:e')
    e.append(_mr(base, base_italic))
    el.append(e)
    s = OxmlElement('m:sup')
    s.append(_mr(sup, sup_italic))
    el.append(s)
    return el


def _mbar(base, base_italic=True):
    """Overbar accent using OMML <m:bar> element (renders better than combining macron)."""
    el = OxmlElement('m:bar')
    barPr = OxmlElement('m:barPr')
    pos = OxmlElement('m:pos')
    pos.set(qn('m:val'), 'top')
    barPr.append(pos)
    el.append(barPr)
    e = OxmlElement('m:e')
    e.append(_mr(base, base_italic))
    el.append(e)
    return el


def _mbar_sub(base, sub, base_italic=True, sub_italic=True):
    """Barred base with subscript: properly nested as sSub(bar(base), sub)."""
    el = OxmlElement('m:sSub')
    el.append(OxmlElement('m:sSubPr'))
    e = OxmlElement('m:e')
    e.append(_mbar(base, base_italic))
    el.append(e)
    s = OxmlElement('m:sub')
    s.append(_mr(sub, sub_italic))
    el.append(s)
    return el


def _msubsup(base, sub, sup):
    """Subscript-superscript combo."""
    el = OxmlElement('m:sSubSup')
    el.append(OxmlElement('m:sSubSupPr'))
    e = OxmlElement('m:e')
    e.append(_mr(base))
    el.append(e)
    s = OxmlElement('m:sub')
    s.append(_mr(sub))
    el.append(s)
    u = OxmlElement('m:sup')
    u.append(_mr(sup))
    el.append(u)
    return el


def _nary(char, sub_parts, sup_parts, e_parts):
    """N-ary operator (summation, product) with sub/sup limits."""
    el = OxmlElement('m:nary')
    pr = OxmlElement('m:naryPr')
    ch = OxmlElement('m:chr')
    ch.set(qn('m:val'), char)
    pr.append(ch)
    if not sup_parts:
        supHide = OxmlElement('m:supHide')
        supHide.set(qn('m:val'), '1')
        pr.append(supHide)
    el.append(pr)
    sub = OxmlElement('m:sub')
    for p in sub_parts:
        sub.append(p)
    el.append(sub)
    sup = OxmlElement('m:sup')
    for p in sup_parts:
        sup.append(p)
    el.append(sup)
    e = OxmlElement('m:e')
    for p in e_parts:
        e.append(p)
    el.append(e)
    return el


def _limlow(e_parts, lim_parts):
    """Lower limit (argmin, min, lim) with limit underneath."""
    el = OxmlElement('m:limLow')
    el.append(OxmlElement('m:limLowPr'))
    e = OxmlElement('m:e')
    for p in e_parts:
        e.append(p)
    el.append(e)
    lim = OxmlElement('m:lim')
    for p in lim_parts:
        lim.append(p)
    el.append(lim)
    return el


def omath(p, parts):
    om = OxmlElement('m:oMath')
    for part in parts:
        om.append(part)
    p._element.append(om)


def omath_display(doc, body, cursor, parts, eq_num=None):
    """Display equation in a borderless 2-column table: centered equation + right-aligned number."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove all borders
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'none')
        e.set(qn('w:sz'), '0')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), 'auto')
        borders.append(e)
    tblPr.append(borders)
    # Full-width table
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), TABLE_WIDTH_PCT)
    tblW.set(qn('w:type'), 'pct')
    old_w = tblPr.find(qn('w:tblW'))
    if old_w is not None:
        tblPr.remove(old_w)
    tblPr.append(tblW)
    # Column widths: equation 85%, number 15%
    for j, w in enumerate([8100, 1400]):
        tc = tbl.cell(0, j)._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(w))
        tcW.set(qn('w:type'), 'dxa')
        old = tcPr.find(qn('w:tcW'))
        if old is not None:
            tcPr.remove(old)
        tcPr.append(tcW)
    # Equation in first cell, centered
    p0 = tbl.cell(0, 0).paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr0 = p0._element.get_or_add_pPr()
    sp0 = OxmlElement('w:spacing')
    sp0.set(qn('w:before'), '60')
    sp0.set(qn('w:after'), '60')
    pPr0.append(sp0)
    om = OxmlElement('m:oMath')
    for part in parts:
        om.append(part)
    p0._element.append(om)
    # Number in second cell, right-aligned
    p1 = tbl.cell(0, 1).paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr1 = p1._element.get_or_add_pPr()
    sp1 = OxmlElement('w:spacing')
    sp1.set(qn('w:before'), '60')
    sp1.set(qn('w:after'), '60')
    pPr1.append(sp1)
    if eq_num:
        # Add bookmark target so in-text "equation (N)" mentions can link here
        bm_name = f'Eq{eq_num}'
        _eq_clean = eq_num.replace('.', '').replace('B', '90').replace('a', '01').replace('b', '02')
        bm_id_val = 800 + int(_eq_clean)
        p1._element.append(make_bookmark(bm_id_val, bm_name))
        p1.add_run(f'({eq_num})')
        p1._element.append(make_bookmark_end(bm_id_val))
    # Vertical center cell content
    for j in range(2):
        tc = tbl.cell(0, j)._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
    tbl_el = tbl._tbl
    body.remove(tbl_el)
    cursor.addnext(tbl_el)
    return None, tbl_el

# ═══════════════════════════════════════════════════════════════════════
# PARAGRAPH HELPERS
# ═══════════════════════════════════════════════════════════════════════


def mkp(doc, body, cursor, space_before=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(space_before if space_before is not None else 0)
    p.paragraph_format.space_after = Pt(8)
    el = p._element
    body.remove(el)
    cursor.addnext(el)
    return p, el


def mkh(doc, body, cursor, text, level=1):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    el = p._element
    body.remove(el)
    cursor.addnext(el)
    return el


def add_italic(p, text):
    """Add an italic run to paragraph p."""
    r = p.add_run(text)
    r.italic = True
    return r


def add_page_break(doc, body, after_el):
    """Insert a page break paragraph after after_el. Returns the new element."""
    pb_p = doc.add_paragraph()
    pb_p.paragraph_format.space_before = Pt(0)
    pb_p.paragraph_format.space_after = Pt(0)
    pb_run = pb_p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    pb_run._element.append(br)
    pb_el = pb_p._element
    body.remove(pb_el)
    after_el.addnext(pb_el)
    return pb_el


def make_bookmark(bm_id, name):
    """Create a w:bookmarkStart element."""
    bs = OxmlElement('w:bookmarkStart')
    bs.set(qn('w:id'), str(bm_id))
    bs.set(qn('w:name'), name)
    return bs


def make_bookmark_end(bm_id):
    """Create a w:bookmarkEnd element."""
    be = OxmlElement('w:bookmarkEnd')
    be.set(qn('w:id'), str(bm_id))
    return be


def make_hyperlink(anchor, text, rPr_orig=None, color=LINK_COLOR):
    """Create a w:hyperlink element with blue underlined text."""
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), anchor)
    hl.set(qn('w:history'), '1')
    r = OxmlElement('w:r')
    if rPr_orig is not None:
        new_rPr = copy.deepcopy(rPr_orig)
    else:
        new_rPr = OxmlElement('w:rPr')
    clr = OxmlElement('w:color')
    clr.set(qn('w:val'), color)
    uu = OxmlElement('w:u')
    uu.set(qn('w:val'), 'single')
    new_rPr.append(clr)
    new_rPr.append(uu)
    r.append(new_rPr)
    t = OxmlElement('w:t')
    t.set(XML_SPACE, SPACE_PRESERVE)
    t.text = text
    r.append(t)
    hl.append(r)
    return hl


# ═══════════════════════════════════════════════════════════════════════
# FOOTNOTE HELPER
# ═══════════════════════════════════════════════════════════════════════
_fn_xml = [None]   # cached parsed footnotes XML
_fn_part = [None]  # cached footnotes Part

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def init_footnotes(doc):
    """Parse the footnotes part from the document package and remove old content footnotes."""
    for rel in doc.part.rels.values():
        if 'footnotes' in rel.reltype:
            _fn_part[0] = rel.target_part
            _fn_xml[0] = etree.fromstring(_fn_part[0].blob)
            # Remove existing content footnotes (keep IDs 0 and -1 = Word separators)
            for fn in list(_fn_xml[0]):
                fid = fn.get(f'{{{W_NS}}}id', '')
                if fid not in ('0', '-1', ''):
                    _fn_xml[0].remove(fn)
            return
    print("  Warning: no footnotes part found in template")


def make_footnote(p, fn_text, fn_id):
    """Add a footnote at the end of paragraph p."""
    if _fn_xml[0] is None:
        return
    fn_el = etree.SubElement(_fn_xml[0], f'{{{W_NS}}}footnote')
    fn_el.set(f'{{{W_NS}}}id', str(fn_id))
    fn_p = etree.SubElement(fn_el, f'{{{W_NS}}}p')
    fn_pPr = etree.SubElement(fn_p, f'{{{W_NS}}}pPr')
    fn_pStyle = etree.SubElement(fn_pPr, f'{{{W_NS}}}pStyle')
    fn_pStyle.set(f'{{{W_NS}}}val', 'FootnoteText')
    # Auto-number mark
    fn_r1 = etree.SubElement(fn_p, f'{{{W_NS}}}r')
    fn_rPr1 = etree.SubElement(fn_r1, f'{{{W_NS}}}rPr')
    fn_rStyle1 = etree.SubElement(fn_rPr1, f'{{{W_NS}}}rStyle')
    fn_rStyle1.set(f'{{{W_NS}}}val', 'FootnoteReference')
    etree.SubElement(fn_r1, f'{{{W_NS}}}footnoteRef')
    # Footnote text
    fn_r2 = etree.SubElement(fn_p, f'{{{W_NS}}}r')
    fn_t2 = etree.SubElement(fn_r2, f'{{{W_NS}}}t')
    fn_t2.set(XML_SPACE, SPACE_PRESERVE)
    fn_t2.text = ' ' + fn_text
    # Footnote reference in main text
    fn_ref_r = OxmlElement('w:r')
    fn_ref_rPr = OxmlElement('w:rPr')
    fn_ref_rStyle = OxmlElement('w:rStyle')
    fn_ref_rStyle.set(qn('w:val'), 'FootnoteReference')
    fn_ref_rPr.append(fn_ref_rStyle)
    fn_ref_r.append(fn_ref_rPr)
    fn_ref_el = OxmlElement('w:footnoteReference')
    fn_ref_el.set(qn('w:id'), str(fn_id))
    fn_ref_r.append(fn_ref_el)
    p._element.append(fn_ref_r)


def flush_footnotes():
    """Write cached footnotes XML back to the part."""
    if _fn_part[0] is not None and _fn_xml[0] is not None:
        _fn_part[0]._blob = etree.tostring(
            _fn_xml[0], xml_declaration=True, encoding='UTF-8', standalone=True)

# ═══════════════════════════════════════════════════════════════════════
# CITATION CROSS-REFERENCE SYSTEM
# ═══════════════════════════════════════════════════════════════════════


# Map: citation text as it appears in-text -> bookmark key
# (cite_author, year, key, ref_anchor)
# cite_author: author string as it appears in inline citations
# year: publication year
# key: unique bookmark key
# ref_anchor: string to match in the reference list for back-linking
CITATIONS = [
    ('Epoch AI', '2024', 'EpochAI2024', 'Epoch AI. (2024)'),
    ('Deloitte', '2025', 'Deloitte2025', 'Deloitte. (2025)'),
    ('Deloitte', '2020', 'Deloitte2020', 'Deloitte and Google. (2020)'),
    ('IEA', '2025', 'IEA2025', 'IEA. (2025)'),
    ('Ohlin', '1933', 'Ohlin1933', 'Ohlin, B.'),
    ('Biglaiser, Cr\u00E9mer, and Mantovani', '2024', 'Biglaiser2024', 'Biglaiser, G.'),
    ('Stojkoski et al.', '2024', 'Stojkoski2024', 'Stojkoski, V.'),
    ('World Bank', '2025', 'WorldBank2025', 'World Bank. (2025)'),
    ('Hausmann, Hwang, and Rodrik', '2007', 'Hausmann2007', 'Hausmann, R.'),
    ('Uptime Institute', '2024', 'UptimeInstitute2024',
     'Uptime Institute. (2024)'),
    ('Firebird', '2026', 'Firebird2026', 'Firebird. (2026)'),
    ('Flucker, Tozer, and Whitehead', '2013', 'Flucker2013', 'Flucker, S.'),
    ('Oltmanns, Krcmarik, and Gatti', '2021', 'Oltmanns2021', 'Oltmanns, J.'),
    ('Liu et al.', '2023', 'Liu2023', 'Liu, Z.'),
    ('Goldfarb and Trefler', '2018', 'Goldfarb2018', 'Goldfarb, A.'),
    ('Korinek and Stiglitz', '2021', 'Korinek2021', 'Korinek, A.'),
    ('UNCTAD', '2025', 'UNCTAD2025', 'UNCTAD. (2025)'),
    ('Grossman and Rossi-Hansberg', '2008', 'Grossman2008', 'Grossman, G.'),
    ('Hummels and Schaur', '2013', 'Hummels2013', 'Hummels, D.'),
    ('Brainard', '1997', 'Brainard1997', 'Brainard, S.'),
    ('Helpman, Melitz, and Yeaple', '2004', 'HMY2004', 'Helpman, E.'),
    ('Lim\u00E3o and Venables', '2001', 'Limao2001', 'Lim\u00E3o, N.'),
    ('Eurostat', '2025', 'Eurostat2025', 'Eurostat. (2025)'),
    ('EIA', '2025', 'EIA2025', 'EIA. (2025)'),
    ('Hersbach et al.', '2020', 'Hersbach2020', 'Hersbach, H.'),
    ('Turner & Townsend', '2025', 'TurnerTownsend2025',
     'Turner & Townsend. (2025)'),
    ('WonderNetwork', '2024', 'WonderNetwork2024', 'WonderNetwork. (2024)'),
    ('NVIDIA', '2024', 'NVIDIA2024', 'NVIDIA. (2024)'),
    ('Krugman', '1991', 'Krugman1991', 'Krugman, P.'),
    ('World Bank', '2024', 'WorldBank2024', 'World Bank. (2024)'),
    ('GlobalPetrolPrices', '2025', 'GlobalPetrolPrices2025',
     'GlobalPetrolPrices. (2025)'),
    ('Cloudscene', '2025', 'Cloudscene2025', 'Cloudscene. (2025)'),
    ('Sastry, Heim, et al.', '2024', 'Sastry2024', 'Sastry, G.'),
    ('Lehdonvirta, Wu, and Hawkins', '2024', 'Lehdonvirta2024',
     'Lehdonvirta, V.'),
    ('Pilz, Mahmood, and Heim', '2025', 'Pilz2025', 'Pilz, K.'),
    ('Turner Lee and West', '2025', 'TurnerLee2025', 'Turner Lee, N.'),
    ('IMF', '2025', 'IMF2025', 'IMF. (2025)'),
    ('Lazard', '2025', 'Lazard2025', 'Lazard. (2025)'),
    ('U.S. DOJ and FTC', '2010', 'DOJFTC2010',
     'U.S. Department of Justice'),
    ('Eaton and Kortum', '2002', 'EatonKortum2002', 'Eaton, J.'),
    ('Dornbusch, Fischer, and Samuelson', '1977', 'DFS1977', 'Dornbusch, R.'),
    ('Arkolakis, Costinot, and Rodr\u00EDguez-Clare', '2012', 'ACR2012',
     'Arkolakis, C.'),
    ('van der Ploeg', '2011', 'vanderPloeg2011', 'van der Ploeg, F.'),
    ('Barroso, H\u00F6lzle, and Ranganathan', '2018', 'Barroso2018', 'Barroso, L.'),
]

# Auto-generate CITE_MAP: both "Author (Year)" and "Author Year" forms
CITE_MAP = {}
for _auth, _yr, _key, _ in CITATIONS:
    CITE_MAP[f'{_auth} ({_yr})'] = _key   # narrative: Author (Year)
    CITE_MAP[f'{_auth} {_yr}'] = _key      # parenthetical: Author Year
CITE_MAP['World Bank'] = 'WorldBank2024'   # bare mention without year
CITE_MAP['Cloudscene'] = 'Cloudscene2025'  # bare mention without year

# Auto-generate REF_KEY_MAP for back-linking from reference list
REF_KEY_MAP = {_key: _anchor for _, _, _key, _anchor in CITATIONS}


def link_citations_pass(body, cite_map, bm_id):
    """Single pass: find citation text in runs and replace with bookmark+hyperlink. Returns count."""
    count = 0
    # Sort by length descending so longer citations match first
    sorted_cites = sorted(cite_map.items(), key=lambda x: -len(x[0]))
    for p_el in list(body.findall(qn('w:p'))):
        for child in list(p_el):
            if child.tag != qn('w:r'):
                continue
            t_el = child.find(qn('w:t'))
            if t_el is None or not t_el.text:
                continue
            text = t_el.text
            for cite_text, key in sorted_cites:
                if cite_text not in text:
                    continue
                idx = text.index(cite_text)
                before = text[:idx]
                after = text[idx + len(cite_text):]
                rPr_orig = child.find(qn('w:rPr'))
                # Modify current run to "before" text only
                t_el.text = before
                t_el.set(XML_SPACE, SPACE_PRESERVE)
                ins = child
                # bookmarkStart
                bm_start = make_bookmark(bm_id[0], f'{key}txt')
                ins.addnext(bm_start)
                ins = bm_start
                # hyperlink (blue underline)
                hyperlink = make_hyperlink(key, cite_text, rPr_orig)
                ins.addnext(hyperlink)
                ins = hyperlink
                # bookmarkEnd
                bm_end = make_bookmark_end(bm_id[0])
                ins.addnext(bm_end)
                ins = bm_end
                bm_id[0] += 1
                # after text
                if after:
                    ra = OxmlElement('w:r')
                    if rPr_orig is not None:
                        ra.append(copy.deepcopy(rPr_orig))
                    ta = OxmlElement('w:t')
                    ta.set(XML_SPACE, SPACE_PRESERVE)
                    ta.text = after
                    ra.append(ta)
                    ins.addnext(ra)
                count += 1
                break  # one per run per pass
    return count


# ═══════════════════════════════════════════════════════════════════════
# ITALIC JOURNAL / BOOK TITLES IN REFERENCES
# ═══════════════════════════════════════════════════════════════════════
ITALIC_IN_REFS = {
    'Brainard': 'American Economic Review',
    'Cloudscene': 'Global Data Center Directory',
    'Deloitte': 'Deloitte Insights',
    'EIA': 'Electric Power Monthly',
    'Eurostat': 'Electricity Prices for Non-Household Consumers (nrg_pc_205)',
    'Flucker': 'Building Services Engineering Research and Technology',
    'GlobalPetrolPrices': 'Electricity Prices Around the World',
    'Goldfarb': 'The Economics of Artificial Intelligence',
    'Grossman': 'American Economic Review',
    'Hausmann': 'Journal of Economic Growth',
    'Helpman': 'American Economic Review',
    'Hersbach': 'Quarterly Journal of the Royal Meteorological Society',
    'Hummels': 'American Economic Review',
    'Korinek': 'NBER Working Paper',
    'Krugman': 'Journal of Political Economy',
    'Lim\u00E3o': 'World Bank Economic Review',
    'Liu': 'Proceedings of ACM e-Energy',
    'NVIDIA': 'NVIDIA H100 Tensor Core GPU Datasheet',
    'Oltmanns': 'Journal of Property Investment & Finance',
    'Turner': 'Data Centre Construction Cost Index 2025',
    'UNCTAD': 'Technology and Innovation Report 2025',
    'WonderNetwork': 'Global Ping Statistics',
    'World Bank': 'World Development Indicators',
    'Lehdonvirta': 'Proceedings of the AAAI/ACM Conference on AI, Ethics, and Society',
    'Pilz': 'AI\u2019s Power Requirements Under Exponential Growth',
    'IMF': 'IMF Working Paper',
    'Lazard': 'Lazard\u2019s Levelized Cost of Energy Analysis, Version 17.0',
    'U.S. Department': 'Horizontal Merger Guidelines',
    'Eaton': 'Econometrica',
    'Dornbusch': 'American Economic Review',
    'Arkolakis': 'American Economic Review',
    'van der Ploeg': 'Journal of Economic Literature',
    'Barroso': 'The Datacenter as a Computer',
    'Ohlin': 'Interregional and International Trade',
    'Biglaiser': 'Toulouse School of Economics Working Paper',
    'Stojkoski': 'WTO Staff Working Paper',
    'World Bank. (2025)': 'Digital Progress and Trends Report 2025',
}


def find_italic_portion(full_rt):
    """Find the journal/book title that should be italicized in a reference."""
    for author_start, italic_text in ITALIC_IN_REFS.items():
        if full_rt.startswith(author_start) and italic_text in full_rt:
            return italic_text
    return None


def _write_ref_segments(p, text, italic_portion):
    """Write reference text with italic journal/book title."""
    if italic_portion and italic_portion in text:
        idx = text.index(italic_portion)
        if idx > 0:
            p.add_run(text[:idx])
        r = p.add_run(italic_portion)
        r.italic = True
        after = text[idx + len(italic_portion):]
        if after:
            p.add_run(after)
    else:
        p.add_run(text)


def add_table(doc, body, after_el, headers, rows, col_widths=None, title=None):
    if title:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after = Pt(3)
        tp.paragraph_format.first_line_indent = Inches(0)
        run = tp.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        tbl_el = tp._element
        body.remove(tbl_el)
        after_el.addnext(tbl_el)
        after_el = tbl_el
    nr = len(rows) + 1
    nc = len(headers)
    table = doc.add_table(rows=nr, cols=nc)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Remove all table-level borders
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    # AutoFit to window: 100% page width
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), TABLE_WIDTH_PCT)
    tblW.set(qn('w:type'), 'pct')
    # Academic-style horizontal rules on header row (top + bottom)

    def _cell_border(tc, sides, style='single'):
        tcPr = tc.get_or_add_tcPr()
        tcB = OxmlElement('w:tcBorders')
        for s in sides:
            b = OxmlElement(f'w:{s}')
            b.set(qn('w:val'), style)
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            tcB.append(b)
        tcPr.append(tcB)
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ""
        pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pp.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        _cell_border(c._tc, ['top', 'bottom'])
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = table.cell(i + 1, j)
            c.text = ""
            pp = c.paragraphs[0]
            if j >= 2:
                pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = pp.add_run(str(val))
            run.font.size = Pt(8)
            # Double bottom border on last data row
            if i == len(rows) - 1:
                _cell_border(c._tc, ['bottom'], style='double')
    if col_widths:
        for j, w in enumerate(col_widths):
            for i in range(nr):
                c = table.cell(i, j)
                tcPr = c._tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(w))
                tcW.set(qn('w:type'), 'dxa')
                old = tcPr.find(qn('w:tcW'))
                if old is not None:
                    tcPr.remove(old)
                tcPr.append(tcW)
    for row in table.rows:
        for cell in row.cells:
            for pp in cell.paragraphs:
                pPr = pp._element.get_or_add_pPr()
                sp = OxmlElement('w:spacing')
                sp.set(qn('w:before'), '10')
                sp.set(qn('w:after'), '10')
                pPr.append(sp)
    tbl_el = table._tbl
    body.remove(tbl_el)
    after_el.addnext(tbl_el)
    return tbl_el


def write_title_and_abstract(doc, body, all_el, hmap):
    print("Rewriting title and abstract...")
    # Replace title (first element — no previous, so clear and rewrite in place)
    title_el = all_el[0]
    for child in list(title_el):
        if child.tag != qn('w:pPr'):
            title_el.remove(child)
    title_p = doc.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.first_line_indent = Inches(0)
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(0)
    r_title = title_p.add_run('Selling FLOPs:\nCompute Exports as a New Industry for Developing Countries')
    r_title.bold = False
    r_title.font.size = Pt(16)
    r_title.font.name = TIMES_NEW_ROMAN

    # Add author name
    author_p, author_el = mkp(doc, body, title_el, space_before=12)
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.paragraph_format.space_after = Pt(12)
    r_author = author_p.add_run('Michael Lokshin')
    r_author.italic = True
    make_footnote(author_p,
                  'This paper\u2019s findings, interpretations, and conclusions are entirely those of the '
                  'author and do not necessarily represent the views of the author\u2019s employer, the '
                  'World Bank, its Executive Directors, or the countries they represent. '
                  'Michael Lokshin: mlokshin@worldbank.org', 1)

    # Version stamp
    ver_p, ver_el = mkp(doc, body, author_el, space_before=2)
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_p.paragraph_format.space_after = Pt(12)
    r_ver = ver_p.add_run(f'v21  \u2014  {datetime.now().strftime("%B %d, %Y  %H:%M")}')
    r_ver.font.size = Pt(9)
    r_ver.font.color.rgb = RGBColor(128, 128, 128)
    r_ver.font.name = TIMES_NEW_ROMAN

    # Replace Abstract heading + text with single paragraph
    # Remove old Abstract heading
    abs_heading = hmap['abs']
    abs_text = all_el[2]
    body.remove(abs_heading)
    body.remove(abs_text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    r_abs_label = p.add_run('Abstract')
    r_abs_label.bold = True
    p.add_run(
        ': This paper develops a trade model in which AI compute is produced and traded '
        'internationally. Latency-insensitive AI training can be offshored '
        'to the lowest-cost producers, while latency-sensitive inference favors proximity '
        'to users; a sovereignty premium captures governments\u2019 preference for domestic data '
        'processing. Calibration across 86 countries shows that energy-abundant economies '
        'have a comparative advantage in training compute, while regional inference hubs '
        'form around major demand centers. Because hardware costs are globally uniform, '
        'cross-country cost differences are small, making institutional quality, reliability, '
        'and policy constraints decisive for location of compute facilities. For energy-rich developing '
        'countries with limited export diversification, compute exports offer a pathway to '
        'convert natural resources into high-value digital services and integrate into the '
        'global economy.'
    )
    el = p._element
    body.remove(el)
    ver_el.addnext(el)
    abs_text_el = el

    # JEL classification and keywords after abstract
    p_jel, jel_el = mkp(doc, body, abs_text_el, space_before=12)
    p_jel.paragraph_format.left_indent = Inches(0.5)
    p_jel.paragraph_format.right_indent = Inches(0.5)
    p_jel.paragraph_format.line_spacing = 1.0
    r_jel_label = p_jel.add_run('JEL Classification: ')
    r_jel_label.bold = True
    p_jel.add_run('F14, F18, L86, O14, O33, Q40')

    p_kw, kw_el = mkp(doc, body, jel_el, space_before=2)
    p_kw.paragraph_format.left_indent = Inches(0.5)
    p_kw.paragraph_format.right_indent = Inches(0.5)
    p_kw.paragraph_format.line_spacing = 1.0
    r_kw_label = p_kw.add_run('Keywords: ')
    r_kw_label.bold = True
    p_kw.add_run(
        'compute trade, FLOPs, artificial intelligence, data centers, '
        'comparative advantage, electricity costs, developing countries'
    )

    return title_el, author_el, ver_el, abs_text_el, kw_el


def write_introduction(doc, body, hmap):
    print("Inserting Section 1: Introduction...")
    cur = mkh(doc, body, hmap['1'].getprevious(), '1. Introduction', level=1)

    # Para 1: AI compute demand + electricity footprint (consolidated)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The expansion of artificial intelligence drives the demand for computational '
        'resources. The compute used to train the largest AI '
        'models has been doubling every six months since 2010 (Epoch AI 2024). '
        'Data centers accounted for approximately 1.5% of global electricity demand in '
        '2024\u2014more than the electricity consumption of France\u2014a share projected '
        'to more than double by 2030 '
        '(IEA 2025). '
        'AI-oriented facilities are qualitatively different from traditional cloud or enterprise '
        'data centers. They deploy thousands of GPUs at power densities of 40\u2013100 kW per rack '
        '(versus 5\u201310 kW in conventional facilities), and can consume over 500,000 gallons of cooling '
        'water per day (Turner Lee and West 2025).'
    )
    # footnote 2 removed (unclear)

    # Para 3: FLOP exporting as value chain upgrading
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'This surge in demand for compute creates a new type of export opportunity. '
        'This paper refers to the production of compute services in one country '
        'for consumption in another as '
    )
    add_italic(p, 'FLOP exporting')
    p.add_run(
        '. FLOP exporting is a form of value chain upgrading\u2014moving from low-value to '
        'higher-value activities within an industry. Rather than '
        'exporting raw energy resources as primary commodities, '
        'countries can convert electricity into a higher value-added digital service. '
        'For energy-rich developing countries, FLOP exporting offers a route up the '
        'value chain without the heavy industrialization traditionally required for '
        'such upgrading (Hausmann, Hwang, and Rodrik 2007).'
    )

    # Para 4: ECA opportunity
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The opportunity is particularly relevant for the Europe and Central Asia (ECA) region. '
        'Several ECA countries, including Turkmenistan, Kyrgyzstan, and the countries of the '
        'South Caucasus, have among the world\u2019s lowest electricity prices but limited '
        'integration into the global digital economy. Building data centers in these locations '
        'and selling compute services to high-cost markets could generate export revenue, attract '
        'foreign investment, and accelerate the development of digital infrastructure. '
        'Export-oriented capacity can also later serve the domestic market as local AI '
        'demand grows.'
    )

    # Para 6: Real data center plans + profit estimate
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Recent megaprojects across Africa, the Middle East, and Central Asia suggest that FLOP exporting is already technically and commercially '
        'feasible. Armenia is deploying 50,000 GPUs in a $4 billion '
        'AI megaproject (Firebird 2026), while Kenya, Saudi Arabia, Morocco, Malaysia, and '
        'Indonesia have each attracted billion-dollar data center commitments.'
    )
    make_footnote(p, 'Microsoft and G42 announced a $1 billion geothermal-powered data center '
                  'in Kenya (2024); AWS committed $5.3 billion to a cloud region in Saudi Arabia (2024); '
                  'Morocco allocated $1.1 billion under its Digital Morocco 2030 strategy; '
                  'Microsoft ($2.2 billion) and Google ($2 billion) announced data centers in Malaysia (2024); '
                  'Microsoft committed $1.7 billion to cloud and AI infrastructure in Indonesia (2024).', 4)
    p.add_run(
        ' Cloud computing exports already exceed $9 billion annually, with the United States '
        'accounting for 87% of the global total (World Bank 2025). '
        'A 40 MW data center in Kyrgyzstan could '
        'generate annual revenue of $630\u2013950 million at wholesale contract rates, '
        'equivalent to over 15% of Kyrgyzstan\u2019s $3.8 billion in goods exports (World Bank 2024).'
    )
    make_footnote(p, 'At $0.038/kWh electricity, a 40 MW facility houses approximately '
                  '53,000 GPUs with production costs of $453 million per year. A Kyrgyz operator '
                  'would most likely sell at wholesale or long-term contract rates of roughly '
                  '$0.80\u20131.20/GPU-hour, yielding gross revenue of $630\u2013950 million. '
                  'Hyperscaler retail rates ($2.00\u20132.50/GPU-hour) represent an upper bound that '
                  'is unlikely for a new market entrant. Even at the wholesale lower bound, '
                  'this exceeds 15% of Kyrgyzstan\u2019s $3.8 billion in goods exports (2024).', 5)

    # Para 9: First paper + contributions
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Recent work examines compute governance and the geography of AI infrastructure '
        '(Sastry et al. 2024, Lehdonvirta et al. 2024, '
        'Pilz et al. 2025), but no formal trade model of compute exists. '
        'This paper offers the first such model, treating FLOPs as commodities produced and exported '
        'according to Ricardian comparative advantage. '
        'The paper makes three contributions. It decomposes the cost of a FLOP into electricity, hardware, and construction '
        'components, and introduces an iceberg trade cost for inference that captures latency '
        'degradation, as well as a sovereignty premium for domestic production preference. '
        'Second, it calibrates the model for 86 countries using data on electricity prices, '
        'climate, data center construction costs, and inter-country network latency, '
        'correcting for energy subsidies that distort headline cost rankings. '
        'Third, it characterizes the resulting trade regimes (which countries export, which '
        'import, and which adopt hybrid strategies) and shows how the sovereignty premium '
        'determines the boundary between domestic and foreign sourcing.'
    )

    # Para 10: Two main results preview
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The analysis in this paper demonstrates two main results that support the '
        'paper\u2019s central argument. First, hardware amortization accounts for over '
        '80 percent of the cost of a GPU-hour and is identical across countries, '
        'compressing the total cost spread to roughly 20 percent across 86 countries. '
        'Electricity and construction, the only location-specific inputs, operate on a '
        'narrow residual. Second, when empirically grounded reliability penalties\u2014'
        'capturing grid outages, governance quality, and sanctions exposure\u2014are applied, '
        'the cost ranking changes substantially. Several energy-abundant but '
        'institutionally fragile economies fall out of the top tier, replaced by countries '
        'with more reliable grids and stronger institutions at moderately higher energy '
        'costs. Together, these results carry a clear policy message: cheap electricity is '
        'necessary but not sufficient for compute exporting; durable comparative advantage '
        'requires credible reliability and institutional quality.'
    )

    # Para 11: Roadmap
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The remainder of the paper is organized as follows. Section 2 reviews the related '
        'literature. Section 3 develops the model, covering production technology, trade costs, demand, '
        'and the capacity-constrained market equilibrium. Section 4 derives the equilibrium '
        'properties, including propositions on country taxonomy, concentration, sovereignty '
        'thresholds, and the nesting of training within inference exporters. Section 5 '
        'describes the data. Section 6 calibrates the model and discusses the results. '
        'Section 7 concludes.'
    )


def write_literature(doc, body, hmap):
    print("Inserting Section 2: Related Literature...")
    cur = mkh(doc, body, hmap['1'].getprevious(), '2. Related Literature', level=1)

    # Para 1: AI comparative advantage + value chain upgrading (merged)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Goldfarb and Trefler (2018) argue that AI shifts comparative advantage toward '
        'countries with data, human capital, and institutional capacity. Our model introduces '
        'a complementary mechanism where comparative advantage in compute '
    )
    add_italic(p, 'production')
    p.add_run(
        ' depends on electricity costs and climate, so resource-rich countries could become '
        'compute exporters without domestic AI research industries. Korinek and Stiglitz (2021) '
        'suggest that developing countries could be left behind in the AI '
        'revolution. FLOP exporting offers a pathway for energy-rich developing '
        'countries to participate in that revolution. The concept of FLOP exporting as value chain upgrading '
        'connects to Hausmann, Hwang, and Rodrik (2007), who show that what a country exports '
        'matters for growth. Lim\u00E3o and Venables (2001) demonstrate that infrastructure quality '
        'determines trade costs. In our model, network infrastructure plays the analogous role '
        'for digital trade. Krugman (1991) shows that increasing returns and transport costs '
        'interact to produce geographic concentration, and analogous centripetal forces '
        '(network effects, colocation with internet exchanges, customer proximity) '
        'favor incumbent data center hubs. The model abstracts from these agglomeration '
        'economies to isolate cost-based comparative advantage.'
    )

    # Para 2: Data center location literature
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Several studies examine the determinants of data center location. '
        'Flucker, Tozer, and Whitehead (2013) show that climate affects data center '
        'cooling costs. Oltmanns, Krcmarik, and Gatti (2021) model data center location as a '
        'function of electricity prices, climate, connectivity, and political stability. '
        'Liu et al. (2023) study data center placement under renewable energy constraints. '
        'These studies focus on where firms should build data centers. '
        'In international trade theory, Brainard (1997) formalizes the proximity-concentration '
        'trade-off between serving a market locally and concentrating production abroad, '
        'and Helpman, Melitz, and Yeaple (2004) extend this to heterogeneous firms choosing between '
        'exporting and FDI.'
    )

    # Para 3: Compute governance literature
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Several papers address compute governance directly. Sastry, Heim, et al. (2024) '
        'argue that compute is well-suited for regulation because governments can track '
        'the number of chips in circulation, restrict who can buy them, and measure how much computation they '
        'perform. Lehdonvirta, Wu, and Hawkins (2024) '
        'map the global geography of cloud GPU infrastructure, distinguishing a '
        '\u201CCompute North\u201D with training-capable hardware from a '
        '\u201CCompute South\u201D limited to inference-grade chips. '
        'Pilz, Mahmood, and Heim (2025) project that AI data center power demand '
        'could reach 327 GW by 2030 and that domestic power shortages may push '
        'compute infrastructure abroad. The World Bank (2025) documents the resulting '
        'global compute divide: high-income countries hold 77% of colocation data center '
        'capacity and account for 87% of cloud computing exports (Stojkoski et al. 2024), '
        'but the report offers descriptive evidence without a formal framework linking '
        'production costs to trade patterns. On the industrial organization side, '
        'Biglaiser, Cr\u00E9mer, and Mantovani (2024) survey the economics of cloud '
        'markets\u2014switching costs, egress fees, and platform competition among '
        'hyperscalers\u2014but the supply-side question of where compute is produced and '
        'whether developing countries can become competitive exporters has not been addressed.'
    )


def write_production_technology(doc, body, hmap):
    print("Rewriting Section 3.1 (Production Technology, merged with Section 3 opening)...")

    # Clear everything between Section 3 heading and Section 3.2 heading,
    # preserving the Section 3.1 heading element
    all_now = list(body)
    s1i = all_now.index(hmap['1'])
    s12i = all_now.index(hmap['1.2'])
    s11_el = hmap['1.1']
    for el in all_now[s1i + 1:s12i]:
        if el is not s11_el:
            body.remove(el)
    cur = hmap['1']  # start after Section 3 heading

    # Para 1: linking paragraph from lit review to model (before 3.1 subtitle)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The existing literature documents where compute infrastructure is located and who '
        'controls it, but no formal framework links production costs to trade patterns. '
        'This section models compute as a tradable good with '
        'country-specific production costs, a delivery cost that depends on whether the '
        'workload is training (latency-insensitive) or inference (latency-sensitive), and '
        'a sovereignty premium reflecting governments\u2019 preference for domestic production.'
    )

    cur = hmap['1.1']  # continue after 3.1 subtitle

    # Para 2: formal setup
    p, cur = mkp(doc, body, cur)
    p.add_run('Consider ')
    omath(p, [_v('N')])
    p.add_run(
        ' countries, each capable of producing compute services. The unit cost of '
        'producing one GPU-hour of compute in country '
    )
    omath(p, [_v('j')])
    p.add_run(
        ' depends on three inputs\u2014electricity, hardware, and data center construction.'
    )
    make_footnote(p,
                  'A floating-point operation (FLOP) is a single arithmetic computation. '
                  'Computing power is measured in petaFLOP/s (10\u00b9\u2075 FLOPs per second). '
                  'A current-generation NVIDIA H100 GPU delivers approximately '
                  '1 petaFLOP/s at 16-bit precision. In retail markets, inference is '
                  'increasingly priced per token (dollars per million tokens) and training '
                  'per GPU-hour or per job. Because tokens per GPU-hour are determined by '
                  'model architecture and serving software\u2014not by country of production\u2014'
                  'the choice of unit does not affect cross-country cost comparisons.', 6)
    # Bump all subsequent footnotes by 1
    # (old fn 6 becomes 7, old fn 7 becomes 8, etc.)

    # PUE inlined (no display equation)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'A data center consumes more electricity than its GPUs alone draw, because '
        'cooling, power distribution, and lighting add overhead. '
        'This overhead is measured by the '
    )
    add_italic(p, 'power usage effectiveness')
    p.add_run(' (PUE), the ratio of total facility power to IT equipment power '
              '(Flucker, Tozer, and Whitehead 2013). '
              'PUE is modeled as a linear function of peak summer temperature: ')
    omath(p, [_t('PUE('), _msub('\u03B8', 'j'), _t(') = '),
              _v('\u03C6'), _t(' + '),
              _v('\u03B4'), _t(' \u00b7 max(0, '),
              _msub('\u03B8', 'j'), _t(' \u2212 '),
              _mbar('\u03B8'), _t(')')])
    p.add_run(
        ', where '
    )
    omath(p, [_v('\u03C6')])
    p.add_run(' is the baseline PUE in cold climates and ')
    omath(p, [_v('\u03B4')])
    p.add_run(
        ' is the sensitivity per \u00b0C above the reference temperature '
    )
    omath(p, [_mbar('\u03B8')])
    p.add_run(
        '. PUE ranges from 1.08 in cold climates to over 1.4 in hot ones.'
    )
    make_footnote(p, 'The linear PUE model is a simplification. Modern liquid and immersion '
                  'cooling technologies can achieve PUE \u2248 1.2 even in hot climates, flattening the '
                  'temperature\u2013PUE relationship. The robustness check in Section 6 confirms that '
                  'the results are insensitive to this specification. '
                  'Google\u2019s 2024 sustainability report indicates a fleet-wide trailing '
                  'twelve-month PUE of 1.10.', 7)

    # Equation lead-in
    p, cur = mkp(doc, body, cur)
    p.add_run('The total cost per GPU-hour in country ')
    omath(p, [_v('j')])
    p.add_run(' is:')
    p.paragraph_format.space_after = Pt(2)

    # Equation (2): cost function (with networking η)
    _, cur = omath_display(doc, body, cur, [
        _msub('c', 'j'), _t(' = '),
        _t('PUE('), _msub('\u03B8', 'j'), _t(') \u00b7 '),
        _v('\u03B3'), _t(' \u00b7 '),
        _msub('p', 'E,j'), _t(' + '),
        _v('\u03C1'), _t(' + '),
        _v('\u03B7'), _t(' + '),
        _msub('p', 'L,j'), _t(' / ('),
        _v('D'), _t(' \u00b7 '), _v('H'), _t('),'),
    ], eq_num='1')

    # Equation explanation (streamlined — no "first term/second term" redundancy)
    p, cur = mkp(doc, body, cur)
    p.add_run('where ')
    omath(p, [_v('\u03B3')])
    p.add_run(
        ' is GPU power draw (kW), '
    )
    omath(p, [_msub('p', 'E,j')])
    p.add_run(' is the electricity price ($/kWh), ')
    omath(p, [_v('\u03C1'), _t(' = '), _msub('P', 'GPU'),
              _t(' / ('), _v('L'), _t(' \u00b7 '), _v('H'),
              _t(' \u00b7 '), _v('\u03B2'), _t(')')])
    p.add_run(
        ' is amortized hardware cost per GPU-hour '
        '('
    )
    omath(p, [_msub('P', 'GPU')])
    p.add_run(' = purchase price, ')
    omath(p, [_v('L')])
    p.add_run(' = lifetime in years, ')
    omath(p, [_v('H')])
    p.add_run(' = 8,766 hours per year, ')
    omath(p, [_v('\u03B2')])
    p.add_run(' = utilization rate),')
    make_footnote(p, 'For the NVIDIA H100: $25,000 / (3 years \u00d7 8,766 hours/year \u00d7 70% '
                  'utilization) \u2248 $1.36/hr. Street prices have fallen to $18,000\u2013$22,000 '
                  'as of late 2025. Each GPU draws approximately 700 watts.', 9)
    p.add_run(' ')
    omath(p, [_v('\u03B7')])
    p.add_run(
        ' is amortized networking cost (high-speed interconnect such as InfiniBand), '
        'and the last term amortizes construction costs '
    )
    omath(p, [_msub('p', 'L,j')])
    p.add_run(
        ' ($/W of IT capacity) over the facility lifetime '
    )
    omath(p, [_v('D')])
    p.add_run('. Both ')
    omath(p, [_v('\u03C1')])
    p.add_run(' and ')
    omath(p, [_v('\u03B7')])
    p.add_run(
        ' are determined in global hardware markets and are common across countries.'
    )
    make_footnote(p, 'China is developing an alternative domestic chip stack based on '
                  'Huawei\u2019s Ascend series (910B/910C) and other domestic accelerators. If these '
                  'achieve comparable FLOPs per watt at lower prices, China\u2019s effective \u03C1 could '
                  'diverge from the NVIDIA-based benchmark used here, potentially improving its '
                  'cost position despite export controls.', 10)
    p.add_run(
        ' Cross-country variation in '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' is therefore driven by electricity prices, climate (through PUE), '
        'and construction costs.'
    )

    # Endowment paragraph
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Countries export goods intensive in their '
        'abundant factors (Ohlin 1933). For compute production, the relevant endowment is not electricity '
        'per se but the natural resources that generate it\u2014hydropower reservoirs '
        '(Kyrgyzstan, Ethiopia, Georgia), oil and gas (Iran, Turkmenistan, Qatar), solar '
        'irradiance (North Africa, the Gulf), and geothermal energy (Kenya, Iceland). '
        'The electricity price '
    )
    omath(p, [_msub('p', 'E,j')])
    p.add_run(
        ' in equation (1) is therefore a reduced-form expression for country '
    )
    omath(p, [_v('j')])
    p.add_run('\u2019s energy resource endowment.')


def write_trade_costs(doc, body, hmap):
    print("Rewriting Section 3.2 (Trade Costs)...")

    all_now = list(body)
    s12i = all_now.index(hmap['1.2'])
    s2i = all_now.index(hmap['2'])
    for el in all_now[s12i + 1:s2i]:
        body.remove(el)
    cur = hmap['1.2']

    # Redefine two service types
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Following the trade-in-tasks framework of Grossman and Rossi-Hansberg (2008) '
        'and the iceberg cost structure of Eaton and Kortum (2002), '
        'countries produce and trade two types of compute services that differ in their '
        'offshoring costs. '
    )
    add_italic(p, 'Training services')
    p.add_run(
        ' encompass batch workloads such as model training, fine-tuning, and large-scale data '
        'processing. Training a state-of-the-art AI model typically takes weeks to months on '
        'thousands of GPUs. The client ships its data to a data center, the computation '
        'executes locally, and the output is returned to the client. Since neither input nor output is '
        'time-sensitive, network latency plays no role. '
    )
    add_italic(p, 'Inference services')
    p.add_run(
        ' encompass real-time workloads such as chatbot responses, autonomous decisions, and interactive '
        'agents. Each query must travel to the server and back within milliseconds, so the '
        'service degrades with delivery delay.'
    )

    # Latency definition
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Latency')
    p.add_run(', denoted ')
    omath(p, [_msub('l', 'jk')])
    p.add_run(
        ', is the round-trip time for a data packet to travel from seller '
    )
    omath(p, [_v('j')])
    p.add_run(' to buyer ')
    omath(p, [_v('k')])
    p.add_run(
        ' and back, measured in milliseconds (ms). '
        'Within a country, latency is typically 5\u201310 ms, while across continents it can exceed '
        '150 ms. For training, the workload ships to the seller, so effective latency is zero.'
    )

    # Sovereignty premium definition
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Governments and firms may prefer to process data domestically for reasons of national '
        'security, regulatory compliance, or political preference. This is captured by a '
    )
    add_italic(p, 'sovereignty premium')
    p.add_run(' ')
    omath(p, [_v('\u03BB'), _t(' \u2265 0')])
    p.add_run(
        ', which acts as a proportional markup on the cost of foreign-sourced compute. '
        'When a country sources compute from a foreign seller, the effective cost is '
        'inflated by the factor '
    )
    omath(p, [_t('(1 + '), _v('\u03BB'), _t(')')])
    p.add_run(
        '. The sovereignty premium is zero for domestic production.'
    )

    # Equation (3): delivered cost with ξ_j
    p, cur = mkp(doc, body, cur)
    p.add_run('The delivered cost of service ')
    omath(p, [_v('s'), _t(' \u2208 {'), _v('T'), _t(', '), _v('I'), _t('}')])
    p.add_run(' from seller ')
    omath(p, [_v('j')])
    p.add_run(' to buyer ')
    omath(p, [_v('k')])
    p.add_run(' is:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msub('P', 's'), _t('('), _v('j'), _t(', '), _v('k'),
        _t(') = (1 + '), _msub('\u03BB', 'jk'),
        _t(') \u00b7 (1 + '), _msub('\u03C4', 's'), _t(' \u00b7 '),
        _msub('l', 'jk'), _t(') \u00b7 '),
        _msub('c', 'j'), _t(' / '),
        _msub('\u03BE', 'j'), _t(','),
    ], eq_num='2')

    p, cur = mkp(doc, body, cur)
    p.add_run('where ')
    omath(p, [_msub('\u03BB', 'jk'), _t(' = '), _v('\u03BB')])
    p.add_run(' if ')
    omath(p, [_v('j'), _t(' \u2260 '), _v('k')])
    p.add_run(' and ')
    omath(p, [_msub('\u03BB', 'jk'), _t(' = 0')])
    p.add_run(' if ')
    omath(p, [_v('j'), _t(' = '), _v('k')])
    p.add_run('. ')
    omath(p, [_msub('\u03BE', 'j'), _t(' \u2208 (0, 1]')])
    p.add_run(
        ' is a reliability index that captures institutional factors, '
        'such as grid outages, regulatory '
        'unpredictability, and sanctions risk, that reduce effective output. '
        'It approaches one for countries with stable institutions and reliable power, '
        'and falls well below one where outages or sanctions reduce effective delivery. '
        'The parameter '
    )
    omath(p, [_v('\u03C4')])
    p.add_run(
        ' measures the rate of quality degradation per millisecond of round-trip latency, with '
    )
    omath(p, [_msub('\u03C4', 'T'), _t(' = 0')])
    p.add_run(' and ')
    omath(p, [_msub('\u03C4', 'I'), _t(' = '), _v('\u03C4'), _t(' > 0')])
    p.add_run(
        '. This iceberg formulation parallels Hummels and Schaur (2013), who estimate that each '
        'day of shipping time is equivalent to a tariff; here, milliseconds replace days. '
        'For training ('
    )
    omath(p, [_msub('\u03C4', 'T'), _t(' = 0')])
    p.add_run(
        '), the delivered cost reduces to the production cost plus the sovereignty '
        'markup; distance plays no role. '
        'Inference also faces a hard latency ceiling: beyond a threshold '
    )
    omath(p, [_mbar('l')])
    p.add_run(
        ' (typically 200\u2013300 ms for interactive applications), the service becomes '
        'unusable regardless of price, modeled as '
    )
    omath(p, [_msub('P', 'I'), _t('('), _v('j'), _t(', '), _v('k'),
              _t(') = \u221E')])
    p.add_run(' if ')
    omath(p, [_msub('l', 'jk'), _t(' > '), _mbar('l')])
    p.add_run('.')


def renumber_sections(hmap):
    print("Renumbering sections...")
    # v20 structure: 1=Intro, 2=Lit, 3=Model(3.1,3.2), 4=Equil Props, 5=Data, 6=Calib, 7=Concl
    # v8 headings: 1→3 (Model), 1.1→3.1, 1.2→3.2, 2→4 (Equil Props), 4→6 (Calib), 5→7 (Concl)
    # Section 3 (Make-or-Buy) content will be absorbed; heading removed by write functions
    renumber = [
        ('1.2', '1.2', '3.2'), ('1.1', '1.1', '3.1'), ('1', '1.', '3.'),
        ('2', '2.', '4.'),
        ('4', '4.', '6.'), ('5', '5.', '7.'),
    ]
    for key, old, new in renumber:
        if key in hmap:
            el = hmap[key]
            for t in el.findall(f'.//{qn("w:t")}'):
                if t.text and old in t.text:
                    t.text = t.text.replace(old, new, 1)
                    break
    # Rename Section 3.1 heading
    if '1.1' in hmap:
        for t in hmap['1.1'].findall(f'.//{qn("w:t")}'):
            if t.text and 'Production Technology' in t.text:
                t.text = t.text.replace('Production Technology',
                                        'Production Technology and Cost Structure')
                break


def write_demand(doc, body, hmap, demand_data):
    """Section 3.3 Demand — moved from old Section 5 (Make-or-Buy)."""
    print("Inserting Section 3.3 (Demand)...")

    # Insert 3.3 heading after the end of Section 3.2 content (before old Section 2 heading)
    s2 = hmap['2']  # was v8 "2. Comparative Advantage", now will be "4. Equilibrium Properties"
    cur = mkh(doc, body, s2.getprevious(), '3.3 Global Compute Demand', level=2)

    # Demand specification: Equation (4)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The model is closed by specifying demand for compute services. Let '
    )
    omath(p, [_msub('q', 'k')])
    p.add_run(
        ' denote the volume of compute purchased by buyer '
    )
    omath(p, [_v('k')])
    p.add_run(
        '. The paper measures compute demand '
        'using installed data center capacity in megawatts (MW), compiled from industry sources as follows:'
    )
    p.paragraph_format.space_after = Pt(2)

    # Equation (4): q_k = ω_k · Q
    _, cur = omath_display(doc, body, cur, [
        _msub('q', 'k'), _t(' = '),
        _msub('\u03C9', 'k'), _t(' \u00b7 '), _v('Q'),
        _t(',     '),
        _msub('\u03C9', 'k'), _t(' = '),
        _msub('M', 'k'), _t(' / '),
        _nary('\u2211', [_v("k\u2032")], [],
              [_msub('M', "k\u2032")]), _t(','),
    ], eq_num='3')

    p, cur = mkp(doc, body, cur)
    p.add_run('where ')
    omath(p, [_v('Q')])
    p.add_run(
        ' is total global compute spending and '
    )
    omath(p, [_msub('\u03C9', 'k')])
    p.add_run(
        ' is country '
    )
    omath(p, [_v('k')])
    p.add_run(
        '\u2019s share of global demand, measured by its share of installed data center '
        'capacity (MW).'
    )
    make_footnote(p,
                  'Installed capacity is preferable to GDP as a demand proxy because '
                  'compute consumption is driven by data center infrastructure, not aggregate '
                  'income. Ireland and the Netherlands, for example, host far more capacity '
                  'per capita than their GDP shares would predict, while large economies like '
                  'India and Brazil account for modest shares of global data center power.',
                  20)
    p.add_run(
        ' Since all results below depend only on demand '
    )
    add_italic(p, 'shares')
    p.add_run(
        ', not on the absolute level '
    )
    omath(p, [_v('Q')])
    p.add_run(
        ', the calibration does not require an estimate of total global compute spending.'
    )

    # Training/inference split
    p, cur = mkp(doc, body, cur)
    p.add_run('Demand splits between training and inference. Training demand is ')
    omath(p, [_msub('q', 'Tk'), _t(' = '), _v('\u03B1'),
              _t(' \u00b7 '), _msub('q', 'k')])
    p.add_run(' and inference demand is ')
    omath(p, [_msub('q', 'Ik'), _t(' = (1 \u2212 '), _v('\u03B1'),
              _t(') \u00b7 '), _msub('q', 'k')])
    p.add_run(', where ')
    omath(p, [_v('\u03B1'), _t(' \u2208 (0, 1)')])
    p.add_run(
        ' is the exogenous training share. '
        'The parameter '
    )
    omath(p, [_v('\u03B1')])
    p.add_run(
        ' should be interpreted as the share of compute that is fully '
        'latency-insensitive and freely offshorable; the effective offshorable share may '
        'be smaller as intermediate workloads (agentic inference, fine-tuning) grow.'
    )
    make_footnote(p,
                  'Emerging workload categories, notably agentic inference (long-running, multi-step '
                  'reasoning tasks) and fine-tuning (rapid iterative retraining on proprietary data), '
                  'occupy a middle ground, tolerating moderate latency but requiring sustained GPU '
                  'allocation and proximity to data. Using installed capacity to proxy demand is a '
                  'static assumption; endogenizing demand, for instance proportional to GDP or digital '
                  'adoption, is a natural extension.', 21)


def write_sourcing_and_equilibrium(doc, body, hmap, demand_data):
    """Section 3.4 Sourcing and Market Equilibrium — simplified (3 display equations)."""
    print("Inserting Section 3.4 (Sourcing and Market Equilibrium)...")

    s2 = hmap['2']
    cur = mkh(doc, body, s2.getprevious(), '3.4 Sourcing and Market Equilibrium', level=2)

    # Sourcing rule: Equation (4)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'For each service type '
    )
    omath(p, [_v('s'), _t(' \u2208 {'), _v('T'), _t(', '), _v('I'), _t('}')])
    p.add_run(', each buyer ')
    omath(p, [_v('k')])
    p.add_run(' chooses the source that minimizes the delivered cost:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msubsup('j', 's', '*'), _t('('), _v('k'),
        _t(') = '),
        _limlow([_t('arg min')], [_v('j')]),
        _t(' '), _msub('P', 's'), _t('('), _v('j'), _t(', '), _v('k'), _t(').'),
    ], eq_num='4')

    # Capacity ceiling — defined before training market (which references it)
    p, cur = mkp(doc, body, cur)
    p.add_run('Each country ')
    omath(p, [_v('j')])
    p.add_run(' is characterized by a capacity ceiling ')
    omath(p, [_mbar_sub('K', 'j')])
    p.add_run(
        ', measured in GPU-hours per period, representing the maximum volume of compute '
        'the country can supply. This ceiling reflects the joint constraint of grid '
        'electricity availability, institutional capacity for data center permitting and '
        'construction, and access to GPU financing.'
    )

    # Training market
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Training market. ')
    p.add_run('Since ')
    omath(p, [_msub('\u03C4', 'T'), _t(' = 0')])
    p.add_run(
        ', training is a homogeneous good with no distance-related quality degradation. '
        'Country '
    )
    omath(p, [_v('k')])
    p.add_run(' imports training whenever the world price, even after adding the sovereignty '
              'premium, is cheaper than producing domestically: ')
    omath(p, [_t('(1 + '), _v('\u03BB'), _t(') \u00b7 '),
              _msub('p', 'T'), _t(' < '), _msub('c', 'k')])
    p.add_run(
        ', where '
    )
    omath(p, [_msub('p', 'T')])
    p.add_run(
        ' is the competitive world training price. '
        'Following Dornbusch, Fischer, and Samuelson (1977), '
        'countries are ranked by production cost: '
    )
    omath(p, [_msub('c', '(1)'), _t(' \u2264 '), _msub('c', '(2)'),
              _t(' \u2264 \u2026 \u2264 '), _msub('c', '(N)')])
    p.add_run(
        '. In the capacity-constrained equilibrium, exporters fill demand in cost order, with '
        'country (1) supplying up to its capacity, then country (2), and so on. '
        'The marginal training exporter '
    )
    omath(p, [_msub('m', 'T')])
    p.add_run(
        ' is the index at which cumulative capacity first meets export demand. '
        'The equilibrium training price equals the marginal exporter\u2019s cost: '
    )
    # p_T = c_{(m_T)} — inline (was display Eq 5)
    # Build subscript manually: c with subscript containing "(m_T)"
    c_sub = OxmlElement('m:sSub')
    c_sub.append(OxmlElement('m:sSubPr'))
    c_e = OxmlElement('m:e')
    c_e.append(_mr('c', True))
    c_sub.append(c_e)
    c_sub_content = OxmlElement('m:sub')
    c_sub_content.append(_mr('(', False))
    mt_sub = OxmlElement('m:sSub')
    mt_sub.append(OxmlElement('m:sSubPr'))
    mt_e = OxmlElement('m:e')
    mt_e.append(_mr('m', True))
    mt_sub.append(mt_e)
    mt_s = OxmlElement('m:sub')
    mt_s.append(_mr('T', True))
    mt_sub.append(mt_s)
    c_sub_content.append(mt_sub)
    c_sub_content.append(_mr(')', False))
    c_sub.append(c_sub_content)
    omath(p, [_msub('p', 'T'), _t(' = '), c_sub])
    p.add_run('.')

    # Rents and shadow value (K̄_j already defined before training market)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Without capacity constraints, '
    )
    omath(p, [_msub('m', 'T'), _t(' = 1')])
    p.add_run(
        ' and the cheapest country serves all demand at its own cost, earning zero rent. '
        'With binding capacity constraints, '
    )
    omath(p, [_msub('m', 'T'), _t(' > 1')])
    p.add_run(
        ', the price rises to the cost of the marginal entrant, and all infra-marginal '
        'exporters earn strictly positive Ricardian rents: '
    )
    omath(p, [_msub('\u03C0', 'Tj'), _t(' = ('), _msub('p', 'T'),
              _t(' \u2212 '), _msub('c', 'j'), _t(') \u00b7 '),
              _msub('K', 'Tj')])
    p.add_run(
        '. For a capacity-constrained exporter, the shadow value '
    )
    omath(p, [_msub('\u03BC', 'j')])
    p.add_run(
        ' of the grid constraint equals the margin on the least profitable active use, '
        'measuring the incremental gain from relaxing the capacity ceiling by one GPU-hour.'
    )

    # Inference: Equation (6)  [was display Eq 7, renumbered after inlining p_T]
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Inference market. ')
    p.add_run('Since ')
    omath(p, [_msub('\u03C4', 'I'), _t(' = '), _v('\u03C4'), _t(' > 0')])
    p.add_run(
        ', inference suffers distance-dependent quality degradation. '
        'The inference market for buyer '
    )
    omath(p, [_v('k')])
    p.add_run(
        ' is localized, as only countries with latency '
    )
    omath(p, [_msub('l', 'jk'), _t(' \u2264 '), _mbar('l')])
    p.add_run(
        ' can participate, and each faces a different delivered cost. '
        'The delivered inference price for buyer '
    )
    omath(p, [_v('k')])
    p.add_run(' is:')
    p.paragraph_format.space_after = Pt(2)

    # Build l_{m_I(k), k} and c_{m_I(k)} with (k) INSIDE the subscript
    # l subscripted with "m_I(k), k"
    l_sub = OxmlElement('m:sSub')
    l_sub.append(OxmlElement('m:sSubPr'))
    l_e = OxmlElement('m:e'); l_e.append(_mr('l', True)); l_sub.append(l_e)
    l_s = OxmlElement('m:sub')
    l_s.append(_msub('m', 'I'))
    l_s.append(_mr('(', False)); l_s.append(_mr('k', True)); l_s.append(_mr('),\u2009', False))
    l_s.append(_mr('k', True))
    l_sub.append(l_s)

    # c subscripted with "m_I(k)"
    c_sub2 = OxmlElement('m:sSub')
    c_sub2.append(OxmlElement('m:sSubPr'))
    c_e2 = OxmlElement('m:e'); c_e2.append(_mr('c', True)); c_sub2.append(c_e2)
    c_s2 = OxmlElement('m:sub')
    c_s2.append(_msub('m', 'I'))
    c_s2.append(_mr('(', False)); c_s2.append(_mr('k', True)); c_s2.append(_mr(')', False))
    c_sub2.append(c_s2)

    _, cur = omath_display(doc, body, cur, [
        _msubsup('p', 'I', 'f'), _t('('), _v('k'), _t(') = (1 + '),
        _v('\u03C4'), _t(' \u00b7 '), l_sub,
        _t(') \u00b7 '), c_sub2, _t(','),
    ], eq_num='5')

    p, cur = mkp(doc, body, cur)
    p.add_run('where ')
    omath(p, [_msub('m', 'I'), _t('('), _v('k'), _t(')')])
    p.add_run(
        ' is the marginal inference supplier to '
    )
    omath(p, [_v('k')])
    p.add_run(
        ', determined by the capacity-constrained supply stack for '
    )
    omath(p, [_v('k')])
    p.add_run(
        '\u2019s inference market. '
        'Each GPU-hour of capacity is allocated to its highest-margin use, whether '
        'training exports, inference exports to various destinations, or domestic supply. '
        'The aggregate rent function (derived in Appendix B) is concave and piecewise linear '
        'in total capacity deployed.'
    )


def write_equilibrium_properties(doc, body, hmap, demand_data):
    print("Rewriting Section 4 (Equilibrium Properties)...")

    # Clear content between section 4 heading (was v8 "2") and section 5 heading (was v8 "3")
    # Also remove the old Make-or-Buy heading and its content
    all_now = list(body)
    s4 = hmap['2']
    s4_old_next = hmap['4']  # old calibration heading (now renumbered to 6)
    s4i = all_now.index(s4)
    s4_next_i = all_now.index(s4_old_next)
    # Remove everything from after Section 4 heading to before old Calibration heading
    # This removes both old Section 4 content AND old Section 5 (Make-or-Buy) heading+content
    for el in all_now[s4i + 1:s4_next_i]:
        body.remove(el)
    cur = s4

    # Also rename the heading text from "Comparative Advantage" to "Equilibrium Properties"
    for t in s4.findall(f'.//{qn("w:t")}'):
        if t.text and 'Comp' in t.text:
            t.text = t.text.replace('Comparative Advantage', 'Equilibrium Properties').replace(' Results', '')
            break

    # Introduction
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'This section derives the formal properties of the capacity-constrained '
        'equilibrium defined in Section 3. Full derivations appear in Appendix B.'
    )

    # Proposition 1: Country taxonomy
    p, cur = mkp(doc, body, cur, space_before=6)
    r = p.add_run('Proposition 1 (Country Taxonomy). ')
    r.bold = True
    p.add_run(
        'In equilibrium, each country falls into exactly one of three regimes: '
        '(i) exporter ('
    )
    omath(p, [_msub('c', 'j'), _t(' < '), _msub('p', 'T')])
    p.add_run(
        '), when the country supplies compute to the world market up to its capacity, '
        '(ii) domestic producer ('
    )
    omath(p, [_msub('p', 'T'), _t(' \u2264 '), _msub('c', 'k'),
              _t(' \u2264 (1 + '), _v('\u03BB'), _t(') \u00b7 '),
              _msub('p', 'T')])
    p.add_run(
        '), when the sovereignty premium makes domestic production cheaper than importing, '
        'and (iii) importer ('
    )
    omath(p, [_msub('c', 'k'), _t(' > (1 + '), _v('\u03BB'), _t(') \u00b7 '),
              _msub('p', 'T')])
    p.add_run(
        '), when even with the sovereignty markup, foreign compute is cheaper. '
        'The three cases partition the cost space and are mutually exclusive.'
    )

    # Proposition 2: Concentration
    p, cur = mkp(doc, body, cur, space_before=6)
    r = p.add_run('Proposition 2 (Capacity Constraints Reduce Concentration). ')
    r.bold = True
    p.add_run(
        'Define the Herfindahl\u2013Hirschman Index (HHI), '
        'a standard measure of market concentration equal to the sum of squared market shares'
    )
    make_footnote(p,
                  'The HHI ranges from '
                  '1/N (equal division among N producers) to 1 (a single producer captures '
                  'the entire market). Values above 0.25 indicate high concentration. '
                  'The index is used by the U.S. Department of Justice and Federal Trade '
                  'Commission to screen mergers and assess market power.', 11)
    p.add_run(', for training market concentration as ')
    omath(p, [_msub('HHI', 'T'), _t(' = '),
              _nary('\u2211', [_v('j')], [],
                    [_msup('(', '2', False, False)]),
              _msub('K', 'Tj'), _t('/'),
              _msubsup('Q', 'T', 'X'),
              _msup(')', '2', False, False)])
    p.add_run(
        '. Without capacity constraints, the cheapest producer captures all training demand '
        'and '
    )
    omath(p, [_msub('HHI', 'T'), _t(' = 1')])
    p.add_run(
        '. With binding capacity constraints on the cheapest producers, '
    )
    omath(p, [_msub('HHI', 'T'), _t(' < 1')])
    p.add_run(
        ', and '
    )
    omath(p, [_msub('HHI', 'T')])
    p.add_run(
        ' is strictly decreasing in the number of capacity-constrained infra-marginal '
        'exporters (U.S. DOJ and FTC 2010). Intuitively, when cheap producers hit capacity '
        'limits, residual demand spills over to costlier suppliers, spreading market shares '
        'more evenly. The proof follows from the strict Cauchy-Schwarz '
        'inequality when at least two producers hold positive market shares.'
    )

    # Proposition 3: Sovereignty threshold
    p, cur = mkp(doc, body, cur, space_before=6)
    r = p.add_run('Proposition 3 (Sovereignty Switching Threshold). ')
    r.bold = True
    p.add_run(
        'A country will bear the additional cost of domestic AI training only if its sovereignty '
        'premium is large enough to justify the price premium over cheaper foreign producers. '
        'Formally, country '
    )
    omath(p, [_v('k')])
    p.add_run(' produces training domestically if and only if ')
    omath(p, [_v('\u03BB'), _t(' \u2265 '),
              _msubsup('\u03BB', 'k', '*'), _t(' = '),
              _msub('c', 'k'), _t('/'), _msub('p', 'T'), _t(' \u2212 1')])
    p.add_run(
        '. The threshold is increasing in '
    )
    omath(p, [_msub('c', 'k')])
    p.add_run(' and decreasing in ')
    omath(p, [_msub('p', 'T')])
    p.add_run(
        '. Under capacity constraints, '
    )
    omath(p, [_msub('p', 'T'), _t(' > '), _msub('c', '(1)')])
    p.add_run(
        ', so the threshold is lower than in the unconstrained model. '
        'Capacity constraints reduce the sovereignty premium required for domestic production '
        'because a higher world price makes imports more expensive.'
    )

    # Proposition 4: Shadow value
    p, cur = mkp(doc, body, cur, space_before=6)
    r = p.add_run('Proposition 4 (Shadow Value and Grid Expansion). ')
    r.bold = True
    p.add_run(
        'For a capacity-constrained exporter, the shadow value '
    )
    omath(p, [_msub('\u03BC', 'j')])
    p.add_run(
        ' of the grid constraint equals the margin on the least profitable active use '
        '(Section 3.4). Grid expansion is warranted if and only if the amortized '
        'per-GPU-hour cost of grid infrastructure '
    )
    omath(p, [_msub('g', 'j')])
    p.add_run(' satisfies ')
    omath(p, [_msub('g', 'j'), _t(' < '), _msub('\u03BC', 'j')])
    p.add_run(
        '. This follows directly from the concavity of the rent function.'
    )

    # Proposition 5: Nesting
    p, cur = mkp(doc, body, cur, space_before=6)
    r = p.add_run('Proposition 5 (Training Exporters Nest Within Inference Exporters). ')
    r.bold = True
    p.add_run(
        'If a country is cheap enough to export training (which can be done from anywhere), '
        'it is also cheap enough to export inference to nearby demand centers. '
        'The set of training exporters is therefore a subset of inference exporters for demand centers '
        'within the latency threshold '
    )
    omath(p, [_mbar('l')])
    p.add_run(
        '. A training exporter has the globally lowest '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run(
        '. For inference to proximate demand centers, this cost advantage dominates '
        'the latency markup, so the same country wins the inference competition. '
        'Since training has no distance penalty while inference does, every country '
        'that exports training is also competitive in inference within its geographic '
        'neighborhood, but not vice versa.'
    )

    # Welfare cost
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Welfare cost of sovereignty. ')
    p.add_run(
        'The sovereignty premium imposes a welfare cost with two components, '
        'an import markup (importers pay '
    )
    omath(p, [_v('\u03BB'), _t(' \u00b7 '), _msub('p', 'T')])
    p.add_run(
        ' per unit above the competitive price) and an allocative inefficiency '
        '(countries with '
    )
    omath(p, [_msub('p', 'T'), _t(' < '), _msub('c', 'k'),
              _t(' \u2264 (1 + '), _v('\u03BB'), _t(') \u00b7 '), _msub('p', 'T')])
    p.add_run(
        ' produce domestically at above-world-price costs). '
        'Under capacity constraints, both components are smaller than in the unconstrained '
        'model because the higher world price narrows the gap between domestic and import costs.'
    )


def write_data_section(doc, body, hmap, demand_data):
    print("Inserting Section 5: Data...")

    sec6_heading = hmap['4']  # was v8 "4", now renumbered to "6"
    cur = mkh(doc, body, sec6_heading.getprevious(), '5. Data', level=1)

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Calibrating the production-cost and trade-cost parameters in equations (2) and (3) requires data on '
        'electricity prices, temperatures, construction costs, and bilateral latencies. '
    )
    add_italic(p, 'Electricity prices. ')
    p.add_run(
        'For European countries, the paper uses Eurostat industrial electricity prices in the '
        '20,000\u201369,999 MWh consumption band, which corresponds to large '
        'industrial consumers (Eurostat 2025). For non-European countries, the paper uses national '
        'regulator tariff sheets and secondary sources, including U.S. Energy Information Administration '
        '(EIA 2025) for the United States, KEPCO for South Korea, national utility tariffs for '
        'Central Asian countries (Barki Tojik, AERA, Ministry of Energy of Uzbekistan), and '
        'GlobalPetrolPrices (2025) for remaining countries. All prices are converted to $/kWh'
        'at 2024 average exchange rates.'
    )

    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Temperature and construction. ')
    p.add_run(
        'Peak summer temperature for each country is computed from ERA5 reanalysis data '
        '(Hersbach et al. 2020) as the average monthly maximum in the three warmest months, '
        'aggregated across populated grid cells. '
    )
    p.add_run(
        'Data center construction costs per watt of IT capacity are from the Turner & Townsend '
        'Data Centre Construction Cost Index 2025 (Turner & Townsend 2025), which covers 37 '
        'countries. For the remaining countries, costs are predicted using a log-linear '
        'regression of ln($/W) on GDP per capita, population, urbanization, seismic risk, '
        'and regional dummies (Appendix\u2009E). Since construction is only 3\u20136% of total per-GPU-hour '
        'costs, imputation error has limited impact on cost rankings.'
    )
    p.add_run(
        ' Costs are amortized over 15 years. The per-GPU construction cost in '
        'equation (1) equals 700 W (the GPU thermal design power) times the build cost per watt.'
    )

    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Latency. ')
    p.add_run(
        'Inter-country round-trip latency is measured using WonderNetwork\u2019s global ping dataset '
        '(WonderNetwork 2024). For each country pair, the median round-trip time (RTT) '
        'in milliseconds is used. '
        'Domestic latency defaults to 5 ms where no intra-country measurement is available. '
        'These measurements reflect today\u2019s network infrastructure. New undersea cables, '
        'terrestrial fiber, and CDN expansions could cut bilateral latencies enough to '
        'redraw inference trade patterns, opening distant low-cost producers to markets '
        'they cannot currently reach.'
    )

    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Hardware. ')
    p.add_run(
        'The calibration uses the NVIDIA H100 SXM GPU as the reference hardware platform, with list price '
        '$25,000, thermal design power 700W, economic lifetime 3 years, utilization rate 70% '
        '(Barroso, H\u00F6lzle, and Ranganathan 2018; NVIDIA 2024). '
        'Google\u2019s fleet-wide GPU utilization, '
        'after years of optimization with custom schedulers and workload packing, runs in the '
        '60\u201375% range (Barroso, H\u00F6lzle, and Ranganathan 2018). A new entrant '
        'in Central Asia would likely achieve 40\u201360% utilization in early years, which would '
        'roughly double the effective hardware cost per useful GPU-hour. '
        'This yields an amortized hardware cost '
    )
    omath(p, [_v('\u03C1'), _t(f' = ${RHO:.3f}')])
    p.add_run(
        '/hr. '
        'Networking costs are calibrated at '
    )
    omath(p, [_v('\u03B7'), _t(f' = ${ETA:.2f}')])
    p.add_run(
        '/hr, based on the amortized cost of InfiniBand interconnect fabric per GPU '
        'over the same three-year horizon (Barroso, H\u00F6lzle, and Ranganathan 2018). '
        'Like hardware, networking equipment is procured at uniform global prices, '
        'so this term does not affect cross-country cost rankings. '
        'GPU prices are assumed uniform across countries. In practice, export controls, '
        'logistics costs, insurance, and local distribution markups can raise effective GPU '
        'prices by 5\u201315% in developing countries. A 10% GPU price premium would add '
        'roughly $0.10/hr to unit costs, substantially eroding the thin cost advantages '
        'documented in Table A2. This assumption thus works in favor of developing-country '
        'exporters and should be kept in mind when interpreting the calibration results.'
    )

    # GPU vintage paragraph
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'GPU vintage. ')
    p.add_run(
        'The calibration uses the NVIDIA H100, but successor GPUs (B200, shipping 2025) '
        'draw approximately 1,000W versus 700W and deliver roughly four times the training '
        'throughput. Higher power draw widens the absolute electricity cost gap across '
        'countries, modestly strengthening developing-country comparative advantage. '
        'However, late entrants selling older-generation GPU-hours must discount against '
        'competitors with newer hardware, potentially eroding their cost edge. The cost '
        'structure is stable across generations\u2014hardware amortization remains 80\u201385 '
        'percent of total cost\u2014so the qualitative findings, including the narrow '
        'cross-country spread and the dominance of institutional factors, are robust to '
        'hardware choice.'
    )

    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Other parameters. ')
    p.add_run('The latency degradation parameter is set at ')
    omath(p, [_v('\u03C4'), _t(f' = {TAU}'), _t(' per ms')])
    p.add_run(
        f', implying that 100 ms of latency inflates inference cost by {TAU * 100:.0%}. '
        'This value is calibrated to match the observed willingness of cloud providers to '
        'invest in regional points of presence. At '
    )
    omath(p, [_v('\u03C4'), _t(f' = {TAU}'), _t(' per ms')])
    p.add_run(
        ', a latency difference of 100 ms (roughly the intercontinental round-trip between '
        'Europe and East Asia) imposes an 8% cost penalty, consistent with industry evidence '
        'that web-service revenue declines by approximately 1% per 100 ms of additional '
        'latency.'
    )
    make_footnote(p, 'Deloitte (2020) finds that a 0.1-second improvement in mobile page load '
                  'time increases retail conversion rates by 8.4% across 30 million user sessions.',
                  14)
    p.add_run(
        ' The sovereignty premium is '
    )
    omath(p, [_v('\u03BB'), _t(f' = {LAMBDA:.0%}')])
    p.add_run('. The training share of compute demand is ')
    omath(p, [_v('\u03B1'), _t(' = 0.50')])
    p.add_run(
        ', within the industry range of 0.4\u20130.6 (Deloitte 2025).'
    )
    make_footnote(p, 'The 10% sovereignty premium is conservative. Survey evidence on data '
                  'localization suggests enterprises pay 15\u201330% more for guaranteed domestic '
                  'data residency (UNCTAD 2025).', 15)

    # Reliability index
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Reliability. ')
    p.add_run('The reliability index ')
    omath(p, [_msub('\u03BE', 'j')])
    p.add_run(
        ' from equation (2) is constructed as the product of three normalized scores '
        'for governance quality (World Bank Worldwide Governance Indicators, Rule of Law '
        'percentile, rescaled to [0, 1]), grid reliability (inverse of hours without electricity '
        'per firm per year, from the World Bank Enterprise Surveys), and '
        'sanctions exposure (0 for unrestricted, 0.5 for comprehensive sanctions). '
        'For OECD countries with negligible outages and no sanctions, '
    )
    omath(p, [_msub('\u03BE', 'j'), _t(' \u2248 1')])
    p.add_run(
        '. For countries under heavy sanctions (Iran, Russia, Belarus), '
    )
    omath(p, [_msub('\u03BE', 'j'), _t(' = 0.50')])
    p.add_run(
        '. Developing countries with weak grids and governance fall in between. '
        'Note that even countries with '
    )
    omath(p, [_msub('\u03BE', 'j'), _t(' \u2248 1')])
    p.add_run(
        ' can shift in rank after reliability adjustment: penalizing low-\u03BE '
        'competitors pushes them down, mechanically raising higher-\u03BE countries. '
        'Moreover, \u03BE ranges from 0.95 to 1.00 even within the OECD, and the tight '
        'cross-country cost spread (\u224820%) means small reliability differences '
        'reshuffle adjacent ranks.'
    )

    # Demand data paragraph
    p, cur = mkp(doc, body, cur)
    top5 = demand_data["top5_labels"]
    add_italic(p, 'Demand. ')
    p.add_run(
        'Compute demand '
    )
    omath(p, [_msub('q', 'k')])
    p.add_run(
        ' is proxied by installed data center capacity in MW, '
        'as specified in equation (3). '
        'For the top 15 markets, capacity estimates come from industry reports '
        '(Synergy Research, Cushman & Wakefield, CBRE, Mordor Intelligence). '
        'For smaller markets, capacity is estimated from facility counts and regional averages. '
        'The five largest demand centers ('
        f'{top5[0][1]} ({top5[0][2] * 100:.0f}%), '
        f'{top5[1][1]} ({top5[1][2] * 100:.0f}%), '
        f'{top5[2][1]} ({top5[2][2] * 100:.1f}%), '
        f'{top5[3][1]} ({top5[3][2] * 100:.0f}%), and '
        f'{top5[4][1]} ({top5[4][2] * 100:.0f}%)) '
        f'account for {demand_data["top5_share"] * 100:.0f}% of global demand. '
        'MW capacity captures the scale of compute infrastructure more accurately than facility counts. '
        'Chinese data centers, though fewer in number (449 in Cloudscene), are substantially larger on average, '
        'and IEA (2025) data confirm that China accounts for roughly 25% of global data center electricity '
        'consumption compared with 44% for the United States. '
    )
    p._element.append(make_bookmark(103, 'Table1txt'))
    p._element.append(make_hyperlink('Table1', 'Table 1'))
    p._element.append(make_bookmark_end(103))
    p.add_run(' reports all model parameters.')


def write_calibration(doc, body, hmap, cal, reg, n_eca, n_total, all_reg, all_sov, demand_data):
    print("Replacing Section 6 (Calibration)...")

    sec7 = hmap['4']
    sec8 = hmap['5']
    all_now = list(body)
    s7i = all_now.index(sec7)
    s8i = all_now.index(sec8)
    for el in all_now[s7i + 1:s8i]:
        body.remove(el)
    cur = sec7

    # Introductory paragraph with explanation of costs
    p, cur = mkp(doc, body, cur)
    p.add_run('The model is calibrated for ')
    omath(p, [_v('N'), _t(f' = {n_total}')])
    p.add_run(f' countries ({n_eca} in ECA, {n_total - n_eca} non-ECA comparators). ')
    p.add_run('The unit cost ')
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' represents the total hourly cost of operating one GPU in country '
    )
    omath(p, [_v('j')])
    p.add_run(
        ', measured in dollars per GPU-hour ($/hr). '
        'The PUE parameters are calibrated as follows. The baseline '
    )
    omath(p, [_v('\u03C6'), _t(' = 1.08')])
    p.add_run(
        ' matches Google\u2019s reported fleet-wide PUE for facilities with free-air cooling '
        'in cold climates (Uptime Institute 2024). The sensitivity coefficient '
    )
    omath(p, [_v('\u03B4'), _t(' = 0.015')])
    p.add_run(
        ' per \u00b0C is estimated from cross-sectional variation in PUE across data center '
        'locations with different cooling loads (Liu et al. 2023). The threshold '
    )
    omath(p, [_mbar('\u03B8'), _t(' = 15\u00b0C')])
    p.add_run(
        ' is the approximate outdoor temperature above which mechanical cooling is needed, '
        'below which free-air cooling suffices. '
        'Together, these yield PUE values from 1.08 (Iceland, Scandinavia) to 1.41 (UAE), '
        'consistent with the industry range of 1.1\u20131.6 reported by the Uptime Institute.'
    )
    make_footnote(p,
                  'Capping PUE at 1.20 (simulating universal liquid cooling) '
                  'yields a Kendall rank correlation of 0.96 with the baseline rankings. The top five '
                  'countries are unchanged and the maximum rank shift is six positions. Gulf states and '
                  'North Africa gain the most (UAE moves from 26th to 20th, Qatar from 15th to 11th), '
                  'but the effect is small because electricity prices, not cooling, dominate '
                  'cross-country cost variation.', 13)

    # Main result — reliability-adjusted ranking (preferred specification)
    _xi_adj = demand_data.get("xi_adjusted", {})
    _xi_top5 = _xi_adj["top5"] if _xi_adj else []
    _xi_n_changed = _xi_adj.get("n_changed_top10", 0)

    p, cur = mkp(doc, body, cur, space_before=6)
    p._element.append(make_bookmark(100, 'TableA2txt'))
    p._element.append(make_hyperlink('TableA2', 'Table A2'))
    p._element.append(make_bookmark_end(100))
    p.add_run(' in the Appendix reports the full results for all ')
    omath(p, [_v('N'), _t(f' = {n_total}')])
    p.add_run(
        ' countries. The paper\u2019s preferred specification applies '
        'cost-recovery electricity prices (replacing subsidized tariffs with long-run '
        'marginal cost) and adjusts for institutional reliability using the index '
    )
    omath(p, [_msub('\u03BE', 'j')])
    p.add_run(
        ' defined in equation (2). Under this specification, the five cheapest '
        'producers are '
    )
    if _xi_top5:
        p.add_run(
            f'{_xi_top5[0][0]} (${_xi_top5[0][1]:.2f}/hr), '
            f'{_xi_top5[1][0]} (${_xi_top5[1][1]:.2f}/hr), '
            f'{_xi_top5[2][0]} (${_xi_top5[2][1]:.2f}/hr), '
            f'{_xi_top5[3][0]} (${_xi_top5[3][1]:.2f}/hr), '
            f'and {_xi_top5[4][0]} (${_xi_top5[4][1]:.2f}/hr). '
            f'{_xi_n_changed} of the ten cheapest countries under the pure engineering '
            'cost ranking fall out of the top ten after reliability adjustment, '
            'replaced by countries with stronger institutions and more reliable grids. '
        )
    p.add_run(
        'Because hardware and networking costs account for roughly 94% of engineering '
        'costs and are identical everywhere, the cross-country cost spread is narrow '
        '(about 20%). Even modest institutional penalties are large relative to this '
        'thin margin, so governance quality can easily dominate the cost ranking. '
        'Engineering cost advantage is therefore necessary but not sufficient for '
        'FLOP exporting. '
    )
    p._element.append(make_bookmark(121, 'Figure1txt'))
    p._element.append(make_hyperlink('Figure1', 'Figure 1'))
    p._element.append(make_bookmark_end(121))
    p.add_run(' illustrates the resulting rank reshuffling.')

    # Illustrative comparison — observed electricity tariffs
    cheapest = cal[0]
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'For comparison, under observed electricity tariffs and without reliability '
        f'adjustment, the cheapest producer is {cheapest["country"]} '
        f'(${float(cheapest["c_j_total"]):.2f}/hr), followed by '
        f'{cal[1]["country"]} (${float(cal[1]["c_j_total"]):.2f}/hr) '
        f'and {cal[2]["country"]} (${float(cal[2]["c_j_total"]):.2f}/hr). '
        'But this ranking is misleading. '
        'Iran\u2019s headline cost rests on electricity priced at '
        f'${float(cheapest["p_E_usd_kwh"]):.3f}/kWh, a figure sustained by one of the '
        'world\u2019s largest fossil fuel subsidies. Turkmenistan, Algeria, Qatar, and '
        'several other low-cost producers face similar distortions. '
        'Once subsidies are removed and institutional reliability is factored in, '
        'the ranking changes substantially.'
    )

    # Cost-recovery adjustment — framed as core analytical step
    adj_top5 = demand_data["adj_top5"]
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'To distinguish genuine comparative advantage from fiscal artifact, '
        'the calibration replaces subsidized tariffs with cost-recovery prices, defined as the '
        'long-run marginal cost (LRMC) of the dominant generation technology at '
        'opportunity-cost fuel prices (IMF 2025, Lazard 2025). This adjustment is applied to '
    )
    p.add_run(f'{demand_data["n_adjusted"]} countries')
    p.add_run(
        ' whose retail electricity prices fall below estimated LRMC. '
        'Hydropower producers (Kyrgyzstan, Canada, Norway) are not adjusted because their '
        'low prices reflect genuine resource advantages rather than fiscal transfers. '
        'The resulting cost-recovery ranking serves as the baseline for the preferred specification. '
        'The five cheapest producers become '
        f'{adj_top5[0][1]} (${adj_top5[0][2]:.2f}/hr), '
        f'{adj_top5[1][1]} (${adj_top5[1][2]:.2f}/hr), '
        f'{adj_top5[2][1]} (${adj_top5[2][2]:.2f}/hr), '
        f'{adj_top5[3][1]} (${adj_top5[3][2]:.2f}/hr), '
        f'and {adj_top5[4][1]} (${adj_top5[4][2]:.2f}/hr). '
        'A sovereignty premium '
    )
    omath(p, [_v('\u03BB'), _t(f' = {LAMBDA:.0%}')])
    p.add_run(
        ' shifts the majority of countries to '
    )
    add_italic(p, 'domestic')
    p.add_run(
        ' production. The sovereignty premium is particularly powerful for inference, '
        'since the latency markup within Europe is moderate (10\u201340 ms, adding 1\u20133%), '
        'even a small domestic preference can tip the decision away from importing.'
    )

    # Lambda calibration
    p, cur = mkp(doc, body, cur)
    p.add_run('In practice, ')
    omath(p, [_v('\u03BB')])
    p.add_run(
        ' is bilateral and heterogeneous. Between allies with mutual data adequacy '
        'agreements (e.g., EU member states), the effective sovereignty premium may be near zero. '
        'Between geopolitical adversaries, it is effectively infinite\u2014the United States '
        'would not source training from Iran regardless of cost, and current sanctions make '
        'such transactions illegal. The uniform '
    )
    omath(p, [_v('\u03BB'), _t(' = 10%')])
    p.add_run(
        ' should therefore be understood as an average over non-adversarial country pairs. '
        'In a model with bilateral '
    )
    omath(p, [_msub('\u03BB', 'jk')])
    p.add_run(
        ', sanctioned countries would be excluded from serving most demand centers. '
        'Under cost-recovery prices, Iran is already outside the top ten, so sanctions '
        'reinforce rather than alter the cost-recovery ranking.'
    )

    # Sovereignty switching thresholds (numerical examples)
    p, cur = mkp(doc, body, cur)
    ls = demand_data["lambda_star"]
    p.add_run(
        'The country-specific switching threshold '
    )
    omath(p, [_msubsup('\u03BB', 'k', '*')])
    p.add_run(
        ' from Section 4 varies widely across the calibration. Near-frontier '
        'countries, whose production costs are close to the cheapest global supplier, '
        'switch to domestic production with minimal sovereignty premia. '
        'Kyrgyzstan requires only '
    )
    omath(p, [_msup('\u03BB', '*'),
              _t(f' = {ls["KGZ"] * 100:.1f}%')])
    p.add_run(
        f', China {ls["CHN"] * 100:.1f}%, and the United States {ls["USA"] * 100:.1f}%. '
        'High-cost countries require much larger premia. Germany needs '
    )
    omath(p, [_msup('\u03BB', '*'),
              _t(f' = {ls["DEU"] * 100:.1f}%')])
    p.add_run(
        f' and Japan {ls["JPN"] * 100:.1f}%. '
        'Countries with high '
    )
    omath(p, [_msup('\u03BB', '*')])
    p.add_run(
        ' face large cost penalties from '
        'sovereignty-driven domestic sourcing.'
    )

    # Demand-weighted trade flows under capacity constraints
    n_exp = demand_data.get("n_train_exporters", 1)
    cap_hhi = demand_data.get("cap_hhi_t", 1.0)
    p_T_val = demand_data.get("p_T", 1.10)
    mu_vals = demand_data.get("mu_j", {})
    n_exp_sov = demand_data.get("n_train_exporters_sov", 1)
    cap_hhi_sov = demand_data.get("cap_hhi_t_sov", 1.0)
    p_T_sov = demand_data.get("p_T_sov", p_T_val)
    ir = demand_data["inf_revenue"]

    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Trade flows under capacity constraints. ')
    p.add_run(
        'Weighting the sourcing patterns by demand shares from equation (3) and applying '
        'capacity constraints from Section 3.4, the equilibrium training price is '
    )
    omath(p, [_msub('p', 'T'), _t(f' = ${p_T_val:.2f}')])
    p.add_run(
        '/hr, set by the marginal exporter\u2019s cost. '
        f'Training demand is served by {n_exp} exporter{"s" if n_exp > 1 else ""} '
        f'(HHI = {cap_hhi:.2f}), confirming Proposition 2. '
    )
    if mu_vals:
        top_mu = sorted(mu_vals.items(), key=lambda x: -x[1])[:3]
        mu_labels = []
        for iso, mu in top_mu:
            co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
            mu_labels.append(f'{co} (${mu:.3f}/hr)')
        p.add_run(
            'The largest shadow values of grid capacity are '
            f'{", ".join(mu_labels)}, '
            'consistent with Proposition 4. '
        )
    # Top inference exporters
    top_inf = sorted(ir.items(), key=lambda x: -x[1])
    top5_inf = top_inf[:5]
    inf_labels = []
    for iso, share in top5_inf:
        co = next(r["country"] for r in cal if r["iso3"] == iso)
        inf_labels.append(f'{co} ({share * 100:.0f}%)')
    p.add_run(
        'Inference is more dispersed, with the top five suppliers being '
        f'{", ".join(inf_labels)}, collectively accounting for '
        f'{sum(s for _, s in top5_inf) * 100:.0f}% of global inference demand '
        f'(HHI = {demand_data["hhi_i"]:.2f}). '
        'Under the 10% sovereignty premium, most training demand is served '
        'domestically, and the smaller export market is served by '
        f'{n_exp_sov} exporter{"s" if n_exp_sov > 1 else ""} (HHI = {cap_hhi_sov:.2f}) at '
        f'${p_T_sov:.2f}/hr.'
    )

    # Revenue for developing countries (dynamic, using cost-recovery inference)
    p, cur = mkp(doc, body, cur)
    kgz_clients = demand_data["kgz_inf_clients"]
    kgz_total = sum(w for _, _, w in kgz_clients)
    if kgz_total > 0:
        kgz_client_names = [
            co for _, co, _ in sorted(kgz_clients, key=lambda x: -x[2])
            if co != "Kyrgyzstan"]
        names = kgz_client_names[:3]
        if len(names) <= 2:
            kgz_list = " and ".join(names)
        else:
            kgz_list = f'{", ".join(names[:-1])}, and {names[-1]}'
        p.add_run(
            f'Among developing countries, Kyrgyzstan captures {kgz_total:.0f}% of global '
            f'inference demand by serving {kgz_list}, a large share '
            'for a country with a GDP of under $15 billion. '
        )
    # Find the largest non-self developing-country inference exporter besides KGZ
    _dev = {'DZA', 'KGZ', 'ETH', 'EGY', 'KOS', 'XKX', 'TKM', 'UZB', 'TJK',
            'ALB', 'MKD', 'GEO', 'ARM', 'MDA', 'UKR', 'BIH', 'SRB'}
    for _iso, _share in sorted(ir.items(), key=lambda x: -x[1]):
        if _iso in _dev and _iso != 'KGZ' and _share > 0.01:
            _co = next((r["country"] for r in cal if r["iso3"] == _iso), _iso)
            # Count how many countries this hub serves
            _n_served = sum(
                1 for i in demand_data.get("adj_reg", {})
                if demand_data["adj_reg"][i]["best_inf_source"] == _iso
                and i != _iso)
            if _n_served > 0:
                p.add_run(
                    f'{_co} serves as an inference hub for {_n_served} '
                    f'{"country" if _n_served == 1 else "countries"}, '
                    f'accounting for {_share * 100:.0f}% of global inference demand. '
                )
            break
    p.add_run(
        'These results are illustrative, not forecasts. They show the cost structure that '
        'would make FLOP exporting viable, not that specific countries will necessarily capture '
        'these market shares. Cheap-energy developing countries '
        'can, in principle, earn export revenue from much larger economies.'
    )

    # Counterfactual
    p, cur = mkp(doc, body, cur)
    es10 = demand_data["export_share_10"]
    es20 = demand_data["export_share_20"]
    extra = demand_data["extra_dom"]
    if es10 < 0.005:
        # Cost-recovery baseline: costs are so close that even λ=10% makes
        # nearly all countries domestic; the 10→20% comparison is uninformative.
        p.add_run(
            'Under cost-recovery pricing, the narrow cost spread means that '
            'even a 10% sovereignty premium is sufficient to make domestic '
            'training viable for nearly all countries, leaving the share of '
            'global training demand available to foreign exporters negligible. '
            f'Raising the premium to 20% shifts {extra} additional '
            f'{"country" if extra == 1 else "countries"} to domestic '
            'production, but the marginal effect is small. '
            'Inference exports are more resilient to sovereignty premia because '
            'the latency advantage of proximity partially insulates regional hubs.'
        )
    else:
        p.add_run(
            f'Doubling the sovereignty premium to 20% shifts {extra} '
            f'additional {"country" if extra == 1 else "countries"} '
            'to domestic training production, reducing '
            'the share of global training demand available to foreign producers '
            f'from {es10 * 100:.0f}% to '
            f'{es20 * 100:.0f}%. '
            'Inference exports are more resilient to sovereignty premia because '
            'the latency advantage of proximity partially insulates regional hubs.'
        )

    # Major consumer markets (using cost-recovery adj_reg)
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Major demand centers. ')
    ar = demand_data.get("adj_reg", {})
    _iso_name = {r["iso3"]: r["country"] for r in cal}
    usa_inf = ar.get('USA', {}).get('best_inf_source', 'CAN')
    usa_inf_cost = ar.get('USA', {}).get('best_inf_cost', '1.190')
    deu_inf = ar.get('DEU', {}).get('best_inf_source', 'KOS')
    deu_inf_cost = ar.get('DEU', {}).get('best_inf_cost', '1.180')
    gbr_inf = ar.get('GBR', {}).get('best_inf_source', 'GBR')
    gbr_inf_cost = ar.get('GBR', {}).get('best_inf_cost', '1.176')
    fra_inf = ar.get('FRA', {}).get('best_inf_source', 'FRA')
    fra_inf_cost = ar.get('FRA', {}).get('best_inf_cost', '1.174')
    chn_inf = ar.get('CHN', {}).get('best_inf_source', 'KGZ')
    chn_inf_cost = ar.get('CHN', {}).get('best_inf_cost', '1.161')
    p.add_run(
        'The model\u2019s predictions vary across major AI demand centers because '
        'each faces a different latency geography. '
        'For the United States, the cost-recovery optimum sources training from the cheapest '
        'available producer and inference from '
        f'{_iso_name.get(usa_inf, usa_inf)} (${float(usa_inf_cost):.2f}/hr). '
        'For Germany, inference is sourced from '
        f'{_iso_name.get(deu_inf, deu_inf)} '
        f'(${float(deu_inf_cost):.2f}/hr), '
        f'for the United Kingdom from {_iso_name.get(gbr_inf, gbr_inf)} '
        f'(${float(gbr_inf_cost):.2f}/hr), '
        f'and for France from {_iso_name.get(fra_inf, fra_inf)} '
        f'(${float(fra_inf_cost):.2f}/hr). '
        f'For China, the cheapest inference source is {_iso_name.get(chn_inf, chn_inf)} '
        f'(${float(chn_inf_cost):.2f}/hr), a bordering country with hydropower-based electricity. '
        'These patterns illustrate the model\u2019s core prediction that inference organizes around '
        'latency-bounded regional hubs, and each major market has a distinct optimal supplier '
        'determined by geography. '
        f'With a sovereignty premium of {LAMBDA:.0%}, most large economies shift to full '
        'domestic production.'
    )

    # Governance and Political Economy discussion (condensed to 3 paragraphs)
    # Para 1: Caveats + Institutional factors
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Governance and political economy. ')
    p.add_run(
        'These cost rankings assume globally uniform hardware prices and abstract from '
        'institutional heterogeneity. In practice, GPU export controls raise effective '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' for Iran, Russia, and Belarus, and grid reliability varies widely. '
        'Data center investments are large, long-lived, and immobile, so the viability of a '
        'country as a compute exporter depends on institutional factors not captured by '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' alone. Several of the cheapest producers in the calibration (Iran, Turkmenistan, '
        'Uzbekistan) rank poorly on property rights and rule of law indices, and subsidized '
        'electricity prices may be politically fragile. Effective entry barriers are therefore '
        'higher than production costs alone suggest. The World Bank (2025) frames this as '
        'the central policy choice: whether to build domestic compute capacity or secure '
        'affordable access to international cloud services, a trade-off the present model '
        'formalizes through the sovereignty premium and capacity constraints.'
    )

    # Para 2: Regulatory, geopolitical barriers, and infrastructure
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The EU\u2019s GDPR and AI Act segment the compute market along regulatory lines, '
        'reinforcing the sovereignty premium '
    )
    omath(p, [_v('\u03BB')])
    p.add_run(
        ' as a structural feature. U.S. export controls on advanced GPUs (October 2022 and '
        'October 2023 rules, expanded in 2025) create a hard binary constraint for certain '
        'jurisdictions\u2014H100-class GPUs cannot legally be shipped there at all, making the '
        'relevant question availability rather than price. For countries not under outright '
        'bans but subject to per-chip caps, the effective hardware cost '
    )
    omath(p, [_v('\u03C1')])
    p.add_run(
        ' rises through grey-market procurement, potentially offsetting any electricity cost '
        'advantage and discouraging long-term investment. '
        'As shown above, once reliability and governance are factored into costs, '
        'viable compute exporters are a strict subset of low-cost producers\u2014those '
        'that combine cheap energy with adequate institutional quality, such as the '
        'Nordic countries, Canada, and parts of the Gulf and Central Asia. '
        'Water is another constraint. Evaporative cooling consumes large volumes, and several '
        'of the cheapest producers (Iran, Turkmenistan, Egypt, Saudi Arabia) are water-scarce. '
        'Liquid cooling reduces water needs but does not eliminate them.'
    )

    # Para 3: Sovereignty skepticism
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The sovereignty premium deserves scrutiny. Some domestic processing preference '
        'is justified for genuinely confidential data (e.g., military intelligence, health records, '
        'and national statistical systems). But much of the current policy push, particularly in '
        'the EU, extends the sovereignty logic far beyond these cases to cover routine commercial '
        'computation that carries no security risk. The welfare cost is not trivial. As shown '
        'above, a 10% premium already shifts most countries toward domestic production, '
        'forgoing the cost savings from specialization. Developing countries in Central Asia and '
        'Africa are likely to follow the EU template, imposing data localization requirements '
        'that their small markets cannot efficiently serve. A policy tension arises: the same countries '
        'whose cost advantages position them as natural FLOP exporters may simultaneously erect '
        'sovereignty barriers against importing compute from their neighbors, '
        'reducing the welfare gains from regional specialization that the model predicts.'
    )

    # Subsidy sustainability and cost-recovery detail (trimmed)
    max_gap_country = demand_data["max_gap_country"]
    max_fiscal_m = demand_data["max_fiscal_transfer"] / 1e6
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Subsidy adjustment. ')
    p.add_run(
        'The cost-recovery prices are derived from country-specific LRMC estimates. '
        'For gas exporters (Iran, Turkmenistan, Algeria, Qatar), the calibration uses combined-cycle gas '
        'generation at export-parity fuel prices ($0.065\u2013$0.100/kWh). For the Gulf states, '
        'it uses the opportunity cost of domestic gas combustion relative to LNG exports. '
        'For coal-dependent producers (Kazakhstan, South Africa), the calibration uses the Eskom-style '
        'cost-recovery tariff. For Ethiopia, it uses the IMF\u2019s hydro cost-recovery target '
        '($0.050/kWh).'
    )
    make_footnote(p,
                  'The IMF estimates global fossil fuel subsidies at $6.7 trillion in 2024. '
                  'Explicit subsidies (below-cost pricing) account for 8%; the remainder reflects '
                  'unpriced environmental costs. This paper uses only the explicit component.',
                  16)
    p.add_run(
        ' The subsidy gap ranges from '
        f'${demand_data["min_gap_mwh"] / 1000:.3f} to ${demand_data["max_gap_mwh_val"] / 1000:.3f}/kWh. '
        f'for {max_gap_country}, a 100\u2009MW data center would receive roughly '
        f'${max_fiscal_m:.0f}\u2009million per year in implicit fiscal transfer. '
        'At export scale, this fiscal arithmetic becomes unsustainable. '
        'Even cost-recovery prices may understate the true resource cost. Regulated tariffs '
        'in many developing countries cover operating expenses but not the full capital cost '
        'of generation, transmission, and distribution infrastructure. State-owned '
        'enterprises (SOEs) accumulate quasi-fiscal deficits that are eventually borne by '
        'taxpayers or future consumers. Allowing large-scale FLOP exports at these prices '
        'would accelerate infrastructure depreciation while the SOE cannot finance '
        'replacement, raising a fundamental question\u2014is it politically sustainable to export '
        'compute while the domestic energy sector cannot maintain its capital stock? '
        'Governments ultimately face '
        'a choice between raising data center tariffs (eroding the cost advantage), maintaining '
        'subsidies at growing fiscal cost, or capping capacity. '
        f'Iran drops from first to {_ordinal(demand_data["adj_rank_map"]["IRN"])}. '
        f'{_num_word(demand_data["regime_changes"]).capitalize()} '
        f'{"country changes" if demand_data["regime_changes"] == 1 else "countries change"} '
        'trade regime.'
    )

    # Endogenous electricity prices caveat
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Endogenous electricity prices. ')
    p.add_run(
        'The model treats electricity prices as exogenous parameters. For large economies '
        'this is innocuous, but for the small, cheap-energy countries that rank highest in '
        'Table A2, a hyperscale data center can be large relative to the host grid. '
        'Kyrgyzstan\u2019s installed generation capacity is approximately 3,500\u2009MW, and '
        'a single 100\u2009MW facility would consume roughly 5% of national electricity output, '
        'and at the multi-facility scale implied by the model\u2019s export predictions, '
        'data centers would become the dominant industrial load. '
        'At that scale, the assumption of price-taking behaviour breaks down. Increased '
        'demand would bid up wholesale electricity prices, competing with residential heating '
        'in winter (when Kyrgyz hydropower output drops), and likely triggering regulatory '
        'intervention. The cheap electricity that attracts investment would be partially '
        'eroded by the investment itself.'
    )
    make_footnote(p,
                  'Kyrgyzstan already experiences seasonal power shortages when reservoir '
                  'levels fall. Adding several hundred MW of year-round base load would '
                  'exacerbate this constraint.',
                  17)
    p.add_run(
        ' The capacity ceiling '
    )
    omath(p, [_mbar_sub('K', 'j')])
    p.add_run(
        ' partially addresses this concern by capping each country\u2019s compute output, '
        'but within the feasible range, the model\u2019s fixed-price assumption means that '
        'the cost advantages in Table A2 are upper bounds. A general equilibrium extension '
        'with upward-sloping electricity supply curves would compress these advantages '
        'further and narrow the set of viable exporters.'
    )

    # Sensitivity analysis paragraph
    sens = demand_data.get("sensitivity", [])
    if sens:
        baseline_pT = sens[0]["p_T"]
        non_baseline = sens[1:]
        min_rho = min(s["rank_corr"] for s in non_baseline) if non_baseline else 1.0
        n_top5_same = sum(1 for s in non_baseline if s["top5_unchanged"])
        n_scenarios = len(non_baseline)
        max_pT_shift = max(abs(s["p_T"] - baseline_pT) for s in non_baseline) if non_baseline else 0
        p, cur = mkp(doc, body, cur, space_before=6)
        add_italic(p, 'Sensitivity analysis. ')
        p.add_run(
            'The cost rankings in Table A2 are robust to substantial parameter variation. '
            f'Across {_num_word(n_scenarios)} scenarios\u2014electricity prices '
            '\u00b1$0.01/kWh, GPU price \u00b120%, and PUE capped '
            f'at 1.20\u2014the Spearman rank correlation with the baseline never falls below '
            f'{min_rho:.3f}, '
            f'the top five cheapest countries are unchanged in '
            f'{_num_word(n_top5_same)} of {_num_word(n_scenarios)} scenarios, '
            f'and the training price shifts by at most ${max_pT_shift:.3f}/hr. '
            'Table A3 in Appendix C reports the full results.'
        )

    # Model extensions (condensed to 1 paragraph)
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Model extensions. ')
    p.add_run(
        'The model can be extended in several directions. It can accommodate endogenous '
        'capacity investment, allowing countries to optimally choose their capacity ceiling '
        'rather than taking grid limits as given. It can incorporate stochastic disruptions '
        'such as grid outages or political instability, giving buyers a reason to diversify '
        'workloads across providers. Demand can be segmented by latency tolerance to capture '
        'heterogeneous service requirements. Carbon pricing can introduce a \u201Cgreen premium\u201D '
        'that favors hydropower-rich countries. The framework can also accommodate strategic '
        'interaction among oligopolistic providers, and governance can enter as a multiplicative '
        'cost shifter on '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run('.')

    # Agglomeration and market structure
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Agglomeration and market structure. ')
    p.add_run(
        'The competitive framework abstracts from the industrial organization of the '
        'cloud compute market, which is dominated by a small number of hyperscalers '
        '(AWS, Azure, Google Cloud) with significant scale economies, proprietary networks, '
        'and market power. In practice, whether a country becomes a compute exporter depends '
        'not only on unit costs but on whether a hyperscaler or colocation provider chooses '
        'to invest there, a decision shaped by agglomeration economies, institutional quality, '
        'and network connectivity (Krugman 1991). The concentration of data centers in '
        'locations such as Northern Virginia reflects precisely these centripetal forces.'
    )

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The model\u2019s contribution is to identify which countries satisfy the necessary '
        'cost condition for competitive supply. Even under increasing returns, hyperscalers '
        'expanding internationally will favor locations where electricity, cooling, and '
        'construction costs are lowest, conditional on meeting minimum infrastructure '
        'thresholds. The cost ranking in Table A2 thus identifies the feasibility frontier. '
        'Countries that are cost-competitive have a prerequisite for attracting investment, '
        'though cost competitiveness alone is not sufficient. The thin margins documented '
        'above (a 20% spread between cheapest and most expensive) reinforce this point. Since '
        'unit cost advantages are modest, institutional and agglomeration factors are often '
        'decisive in determining which cost-competitive countries attract '
        'investment.'
    )


def write_conclusion(doc, body, hmap, demand_data):
    print("Rewriting Section 7 (Conclusion)...")
    sec8 = hmap['5']
    refs = hmap['refs']
    all_now = list(body)
    c8i = all_now.index(sec8)
    ri = all_now.index(refs)
    for el in all_now[c8i + 1:ri]:
        body.remove(el)

    p, cur_concl = mkp(doc, body, sec8)
    p.add_run(
        'This paper develops a capacity-constrained model of trade in computing '
        'services (FLOPs) in which countries produce and export computing capacity based on '
        'their energy-resource endowments, as reflected in electricity prices, climate, '
        'and construction costs. The model distinguishes '
        'two service types, latency-insensitive training and latency-sensitive inference, and '
        'introduces a sovereignty premium to capture governments\u2019 preference for domestic '
        'data processing. Capacity ceilings transform the classical cost-based assignment into '
        'a framework with market-clearing prices and scarcity rents. '
        'The paper calibrates the model for 86 countries using data on electricity '
        'prices, temperatures, construction costs, bilateral latencies, and grid capacity.'
    )

    p, cur_concl = mkp(doc, body, cur_concl)
    p.add_run(
        'Across 86 countries, low-energy-cost countries export training compute, '
        'while inference is served by regional hubs close enough to users to meet latency requirements. '
        'The sovereignty premium rationalizes widespread domestic investment, shifting '
        'the majority of countries from import to domestic production, '
        f'at a demand-weighted welfare cost of {demand_data["welfare_pct"]:.1f}% of '
        'average compute spending, comparable in magnitude to the 1\u201310% welfare losses '
        'from trade barriers estimated for goods trade '
        '(Arkolakis, Costinot, and Rodr\u00EDguez-Clare 2012). '
        'The model generates a country taxonomy (full importers, training exporters, '
        'inference hubs, and hybrid regimes) that maps onto observed investment patterns. '
        'This geographic structure is consistent with Lehdonvirta, Wu, and Hawkins (2024), '
        'who independently find that training-capable GPU infrastructure is concentrated in '
        'roughly 30 countries while the rest are limited to inference-grade hardware. '
        'A central finding is that electricity costs, while the main source of cross-country '
        'cost variation, are not the dominant determinant of actual data center location. '
        'Because hardware amortization is uniform and accounts for 94% of per-GPU-hour costs, '
        'the cost spread across countries is narrow\u2014about 20%. Institutional quality, '
        'sovereignty preferences, and access to GPU hardware frequently outweigh these thin '
        'cost margins.'
    )

    p, cur_concl = mkp(doc, body, cur_concl)
    p.add_run(
        'For developing countries, the results point to a new avenue for economic participation '
        'in the global economy. Countries like Kyrgyzstan, Uzbekistan, and Egypt, which rank '
        'among the cheapest FLOP producers in the calibration, could use their energy '
        'resource endowments\u2014hydropower, natural gas, and solar irradiance\u2014to convert '
        'cheap electricity into a high-value digital export without building a domestic AI '
        'research ecosystem. '
        'FLOP exporting is the digital equivalent of resource-based industrialization, but '
        'with the advantage that the underlying resource (electricity) need not deplete a '
        'finite reserve and the product '
        '(compute) serves the fastest-growing sector of the world economy. '
        'That said, the resource curse literature (van der Ploeg 2011) cautions that '
        'concentrated export revenues can produce Dutch disease, institutional degradation, '
        'and volatility. Whether FLOP exporting shares these risks depends on whether '
        'the revenues are broadly distributed or concentrated among a narrow set of actors, and on whether '
        'governments invest the proceeds in human capital and institutional development.'
    )

    p, _ = mkp(doc, body, cur_concl)
    p.add_run(
        'The policy implications are asymmetric across training and inference. '
        'Training workloads tolerate high latency, so restricting training imports '
        'raises costs without offsetting proximity gains. '
        'Inference, by contrast, is latency-sensitive, giving domestic production a genuine '
        'quality-of-service advantage\u2014though this rationale weakens for countries close '
        'to low-cost neighbors. For developing countries seeking to enter the compute export '
        'market, the binding constraints are not technological but institutional. '
        'Reliable power grids, political stability, data governance frameworks, '
        'and international connectivity determine whether cost advantages translate into actual exports.'
    )


def write_appendix(doc, body, last_ref_el, eca_cal, non_eca_cal, reg, demand_data):
    print("Inserting Appendix (Table A2)...")

    # ─── Portrait section break (ends portrait section, next page stays portrait) ───
    sep = doc.add_paragraph()
    sep.paragraph_format.space_before = Pt(0)
    sep.paragraph_format.space_after = Pt(0)
    sep_el = sep._element
    body.remove(sep_el)
    last_ref_el.addnext(sep_el)
    sect_portrait = OxmlElement('w:sectPr')
    pg_sz_p = OxmlElement('w:pgSz')
    pg_sz_p.set(qn('w:w'), '12240')   # 8.5 inches
    pg_sz_p.set(qn('w:h'), '15840')   # 11 inches
    sect_portrait.append(pg_sz_p)
    pg_mar_p = OxmlElement('w:pgMar')
    pg_mar_p.set(qn('w:top'), '1440')
    pg_mar_p.set(qn('w:right'), '1440')
    pg_mar_p.set(qn('w:bottom'), '1440')
    pg_mar_p.set(qn('w:left'), '1440')
    pg_mar_p.set(qn('w:header'), '720')
    pg_mar_p.set(qn('w:footer'), '720')
    sect_portrait.append(pg_mar_p)
    sep_pPr = sep_el.find(f'{{{W_NS}}}pPr')
    if sep_pPr is None:
        sep_pPr = etree.SubElement(sep_el, f'{{{W_NS}}}pPr')
    sep_pPr.append(sect_portrait)

    # Appendix heading (portrait page)
    cur_app = mkh(doc, body, sep_el, 'Appendix', level=1)

    # ─── Portrait section break (ends portrait for landscape Table A2) ───
    hr_a1 = doc.add_paragraph()
    hr_a1.paragraph_format.space_before = Pt(0)
    hr_a1.paragraph_format.space_after = Pt(0)
    hr_a1_el = hr_a1._element
    body.remove(hr_a1_el)
    cur_app.addnext(hr_a1_el)

    sect_a1_end = OxmlElement('w:sectPr')
    pg_sz_a1 = OxmlElement('w:pgSz')
    pg_sz_a1.set(qn('w:w'), '12240')
    pg_sz_a1.set(qn('w:h'), '15840')
    sect_a1_end.append(pg_sz_a1)
    pg_mar_a1 = OxmlElement('w:pgMar')
    pg_mar_a1.set(qn('w:top'), '1440')
    pg_mar_a1.set(qn('w:right'), '1440')
    pg_mar_a1.set(qn('w:bottom'), '1440')
    pg_mar_a1.set(qn('w:left'), '1440')
    pg_mar_a1.set(qn('w:header'), '720')
    pg_mar_a1.set(qn('w:footer'), '720')
    sect_a1_end.append(pg_mar_a1)
    hr_a1_pPr = hr_a1_el.find(f'{{{W_NS}}}pPr')
    if hr_a1_pPr is None:
        hr_a1_pPr = etree.SubElement(hr_a1_el, f'{{{W_NS}}}pPr')
    hr_a1_pPr.append(sect_a1_end)

    # ═══════════════════════════════════════════════════════════════════════
    # TABLE A2: COUNTRY-SPECIFIC CALIBRATION PARAMETERS (landscape)
    # ═══════════════════════════════════════════════════════════════════════
    print("Inserting Table A2 (Country parameters, landscape)...")

    # Table A2 title with bookmark + back-link (follows directly after A1 notes)
    tp2 = doc.add_paragraph()
    tp2.paragraph_format.space_before = Pt(6)
    tp2.paragraph_format.space_after = Pt(3)
    tp2.paragraph_format.first_line_indent = Inches(0)
    tp2._element.append(make_bookmark(104, 'TableA2'))
    hl_t = OxmlElement('w:hyperlink')
    hl_t.set(qn('w:anchor'), 'TableA2txt')
    hl_t.set(qn('w:history'), '1')
    r_t = OxmlElement('w:r')
    rPr_t = OxmlElement('w:rPr')
    b_t = OxmlElement('w:b')
    rPr_t.append(b_t)
    sz_t = OxmlElement('w:sz')
    sz_t.set(qn('w:val'), '20')
    rPr_t.append(sz_t)
    clr_t = OxmlElement('w:color')
    clr_t.set(qn('w:val'), LINK_COLOR)
    uu_t = OxmlElement('w:u')
    uu_t.set(qn('w:val'), 'single')
    rPr_t.append(clr_t)
    rPr_t.append(uu_t)
    r_t.append(rPr_t)
    t_t = OxmlElement('w:t')
    t_t.text = 'Table A2'
    r_t.append(t_t)
    hl_t.append(r_t)
    tp2._element.append(hl_t)
    tp2._element.append(make_bookmark_end(104))
    run_tt2 = tp2.add_run('. Country-specific calibration parameters')
    run_tt2.bold = True
    run_tt2.font.size = Pt(10)
    tp2_el = tp2._element
    body.remove(tp2_el)
    hr_a1_el.addnext(tp2_el)

    # Gather all country data
    omega = demand_data["omega"]
    dc_k = demand_data.get("dc_k", {})
    xi = demand_data.get("xi", {})
    adj_rank_map = demand_data.get("adj_rank_map", {})
    # Sort by cost-recovery adjusted rank
    all_cal = sorted(eca_cal + non_eca_cal,
                     key=lambda r: adj_rank_map.get(r["iso3"], 999))

    a2_headers = ["Rank", "Country", "p\u1d31\n($/kWh)", "\u03B8\u2c7c\n(\u00b0C)",
                  "PUE", "Constr.\n($/W)", "k\u0304\u2c7c\n(MW)",
                  "\u03C9\u2c7c\n(%)", "\u03BE\u2c7c",
                  "c\u2c7c\n($/hr)", "Cost-Rec.\np\u1d31 ($/kWh)", "Regime"]

    # Build row data; track which rows need bold in cost-recovery column
    a2_rows = []
    bold_cr_rows = []  # row indices (0-based) where cost-rec price is substituted
    for idx, r_row in enumerate(all_cal):
        iso = r_row["iso3"]
        co = r_row["country"]
        if len(co) > 20:
            co = co[:19] + "."
        adj_rank = adj_rank_map.get(iso, 999)
        regime = reg.get(iso, {}).get("regime", "full import")
        rs = {"full import": "import", "import training + build inference": "hybrid",
              "full domestic": "domestic"}.get(regime, regime)
        # Cost-recovery price: substituted value for 13 countries, otherwise same as p_E
        p_E_raw = float(r_row["p_E_usd_kwh"])
        cr = SUBSIDY_ADJ.get(iso)
        cr_price = cr if cr is not None else p_E_raw
        cr_str = f'${cr_price:.3f}'
        if cr is not None:
            bold_cr_rows.append(idx)
        cap = dc_k.get(iso, 5.0)
        cap_str = f'{cap:,.0f}' if cap >= 10 else f'{cap:.0f}'
        share = omega.get(iso, 0)
        xi_j = xi.get(iso, 1.0)
        a2_rows.append((
            str(adj_rank), co,
            f'${p_E_raw:.3f}',
            f'{float(r_row["theta_summer_C"]):.1f}',
            f'{float(r_row["pue"]):.2f}',
            f'${float(r_row["p_L_usd_per_W"]):.2f}',
            cap_str,
            f'{share * 100:.1f}',
            f'{xi_j:.2f}',
            f'${float(r_row["c_j_total"]):.2f}',
            cr_str, rs,
        ))

    # Column widths for landscape (9 inches usable)
    a2_cw = [380, 1400, 600, 500, 450, 550, 600, 500, 400, 600, 650, 600]
    last_a2_tbl = add_table(doc, body, tp2_el, a2_headers, a2_rows, a2_cw)

    # Post-process: bold the cost-recovery price cells for subsidized countries
    # Table rows: row 0 = header, data rows start at 1
    all_trs = last_a2_tbl.findall(f'{{{W_NS}}}tr')
    for row_idx in bold_cr_rows:
        tr = all_trs[row_idx + 1]  # skip header row
        # Column 10 = cost-recovery price
        tcs = tr.findall(f'{{{W_NS}}}tc')
        if len(tcs) > 10:
            tc = tcs[10]
            for r_el in tc.findall(f'.//{{{W_NS}}}r'):
                rPr = r_el.find(f'{{{W_NS}}}rPr')
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r_el.insert(0, rPr)
                b_el = OxmlElement('w:b')
                rPr.append(b_el)

    # Table A2 notes
    note_a2 = doc.add_paragraph()
    note_a2.paragraph_format.space_before = Pt(4)
    note_a2.paragraph_format.space_after = Pt(0)
    note_a2.paragraph_format.first_line_indent = Inches(0)
    note_a2.paragraph_format.line_spacing = 1.0
    rn = note_a2.add_run('Notes: ')
    rn.bold = True
    rn.font.size = Pt(7.5)
    rn = note_a2.add_run(
        'Countries sorted by cost-recovery adjusted rank (ascending). '
        'p\u1d31 = national electricity price for industrial/data center consumers ($/kWh). '
        '\u03B8\u2c7c = peak summer temperature (\u00b0C). '
        'PUE = Power Usage Effectiveness. '
        'Constr. = predicted data center construction cost ($/W of IT load). '
        'k\u0304\u2c7c = installed data center power capacity (MW). '
        '\u03C9\u2c7c = country share of global compute demand from equation (3). '
        '\u03BE\u2c7c = reliability index combining governance quality, grid reliability, '
        'and sanctions exposure. '
        'c\u2c7c = total hourly cost of operating one H100 GPU (electricity + '
        'hardware at $1.36/hr + networking at $0.15/hr + amortized construction). '
        'Cost-Rec. p\u1d31 = cost-recovery electricity price. '
        'For 13 countries with subsidized tariffs, this is the estimated long-run marginal cost '
        'of electricity generation (shown in bold). '
        'For all other countries, the cost-recovery price equals the observed tariff. '
        'Regime = optimal sourcing strategy from equation (4) without sovereignty premium.'
    )
    rn.font.size = Pt(7.5)
    note_a2_el = note_a2._element
    body.remove(note_a2_el)
    last_a2_tbl.addnext(note_a2_el)

    # Empty paragraph after Table A2 notes (hard return)
    hr_a2 = doc.add_paragraph()
    hr_a2.paragraph_format.space_before = Pt(0)
    hr_a2.paragraph_format.space_after = Pt(0)
    hr_a2_el = hr_a2._element
    body.remove(hr_a2_el)
    note_a2_el.addnext(hr_a2_el)

    # Landscape section break (on empty paragraph, ends landscape for portrait Appendix B)
    sect_pr = OxmlElement('w:sectPr')
    pg_sz = OxmlElement('w:pgSz')
    pg_sz.set(qn('w:w'), '15840')
    pg_sz.set(qn('w:h'), '12240')
    pg_sz.set(qn('w:orient'), 'landscape')
    sect_pr.append(pg_sz)
    pg_mar = OxmlElement('w:pgMar')
    pg_mar.set(qn('w:top'), '1440')
    pg_mar.set(qn('w:right'), '1440')
    pg_mar.set(qn('w:bottom'), '1440')
    pg_mar.set(qn('w:left'), '1440')
    pg_mar.set(qn('w:header'), '720')
    pg_mar.set(qn('w:footer'), '720')
    sect_pr.append(pg_mar)
    hr_a2_pPr = hr_a2_el.find(f'{{{W_NS}}}pPr')
    if hr_a2_pPr is None:
        hr_a2_pPr = etree.SubElement(hr_a2_el, f'{{{W_NS}}}pPr')
    hr_a2_pPr.append(sect_pr)

    return hr_a2_el


def write_model_appendix(doc, body, last_note):
    """Appendix B: Full model derivation from flops_capacity_model.md."""
    print("Inserting Appendix B (Model Derivation)...")

    pb = add_page_break(doc, body, last_note)
    cur = mkh(doc, body, pb, 'Appendix B: Model Derivation', level=1)

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'This appendix provides the full derivation of the capacity-constrained Ricardian '
        'model summarized in Sections 3\u20134.'
    )

    # B.1 Primitives
    cur = mkh(doc, body, cur, 'B.1 Primitives', level=2)
    p, cur = mkp(doc, body, cur)
    p.add_run('Each country ')
    omath(p, [_v('j')])
    p.add_run(' is endowed with a capacity ceiling ')
    omath(p, [_mbar_sub('K', 'j')])
    p.add_run(
        ' (GPU-hours per period), representing the maximum volume of compute it can supply. '
        'Country '
    )
    omath(p, [_v('j')])
    p.add_run(' faces unit production cost ')
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' from equation (1). On the demand side, total compute demand from country '
    )
    omath(p, [_v('k')])
    p.add_run(' is ')
    omath(p, [_msub('q', 'k')])
    p.add_run(' from equation (3). Training demand is ')
    omath(p, [_msub('q', 'Tk'), _t(' = '), _v('\u03B1'), _t(' \u00b7 '), _msub('q', 'k')])
    p.add_run(' and inference demand is ')
    omath(p, [_msub('q', 'Ik'), _t(' = (1 \u2212 '), _v('\u03B1'), _t(') \u00b7 '), _msub('q', 'k')])
    p.add_run('. Countries are ordered by cost: ')
    omath(p, [_msub('c', '(1)'), _t(' \u2264 '), _msub('c', '(2)'),
              _t(' \u2264 \u2026 \u2264 '), _msub('c', '(N)')])
    p.add_run('.')

    # B.2 Training Market
    cur = mkh(doc, body, cur, 'B.2 The Training Market', level=2)
    p, cur = mkp(doc, body, cur)
    p.add_run('Country ')
    omath(p, [_v('k')])
    p.add_run(' imports training if and only if ')
    omath(p, [_t('(1 + '), _v('\u03BB'), _t(') \u00b7 '), _msub('p', 'T'),
              _t(' < '), _msub('c', 'k')])
    p.add_run('. The set of training importers is ')
    omath(p, [_msub('M', 'T'), _t(' = { '), _v('k'), _t(' : '),
              _msub('c', 'k'), _t(' > (1 + '), _v('\u03BB'), _t(') \u00b7 '),
              _msub('p', 'T'), _t(' }')])
    p.add_run(' and total training export demand is ')
    omath(p, [_msubsup('Q', 'T', 'X'), _t(' = '),
              _nary('\u2211', [_v('k'), _t(' \u2208 '), _msub('M', 'T')], [],
                    [_msub('q', 'Tk')])])
    p.add_run(
        '. The marginal training exporter '
    )
    omath(p, [_msub('m', 'T')])
    p.add_run(' is defined by:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msub('m', 'T'), _t(' = min { '), _v('m'),
        _t(' : '),
        _nary('\u2211', [_v('i'), _t(' = 1')], [_v('m')],
              [_msub('K', 'T,(i)')]),
        _t(' \u2265 '), _msubsup('Q', 'T', 'X'), _t(' }.'),
    ], eq_num='B.1')

    p, cur = mkp(doc, body, cur)
    p.add_run('The equilibrium training price is ')
    omath(p, [_msub('p', 'T'), _t(' = '), _msub('c', '('),
              _msub('m', 'T'), _t(')')])
    p.add_run('. Training rent for country ')
    omath(p, [_v('j')])
    p.add_run(' with ')
    omath(p, [_msub('c', 'j'), _t(' < '), _msub('p', 'T')])
    p.add_run(' is ')
    omath(p, [_msub('\u03C0', 'Tj'), _t(' = ('), _msub('p', 'T'),
              _t(' \u2212 '), _msub('c', 'j'), _t(') \u00b7 '),
              _msub('K', 'Tj')])
    p.add_run('.')

    # B.3 Inference Market
    cur = mkh(doc, body, cur, 'B.3 The Inference Market', level=2)
    p, cur = mkp(doc, body, cur)
    p.add_run('The feasible supplier set for demand center ')
    omath(p, [_v('k')])
    p.add_run(' is ')
    omath(p, [_v('S'), _t('('), _v('k'), _t(') = { '), _v('j'), _t(' : '),
              _msub('l', 'jk'), _t(' \u2264 '), _mbar('l'), _t(' }')])
    p.add_run('. The marginal cost of delivering one effective unit of inference from ')
    omath(p, [_v('j')])
    p.add_run(' to ')
    omath(p, [_v('k')])
    p.add_run(' is:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msub('MC', 'I'), _t('('), _v('j'), _t(', '), _v('k'),
        _t(') = (1 + '), _v('\u03C4'), _t(' \u00b7 '),
        _msub('l', 'jk'), _t(') \u00b7 '), _msub('c', 'j'), _t('.'),
    ], eq_num='B.2')

    p, cur = mkp(doc, body, cur)
    p.add_run('The inference rent per GPU-hour allocated to serving ')
    omath(p, [_v('k')])
    p.add_run(' is ')
    omath(p, [_msub('r', 'I'), _t('('), _v('j'), _t(', '), _v('k'),
              _t(') = '), _msubsup('p', 'I', 'f'), _t('('), _v('k'),
              _t(') / (1 + '), _v('\u03C4'), _t(' \u00b7 '),
              _msub('l', 'jk'), _t(') \u2212 '), _msub('c', 'j')])
    p.add_run('.')

    # B.4 Capacity Allocation
    cur = mkh(doc, body, cur, 'B.4 Capacity Allocation', level=2)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Each GPU-hour is allocated to its highest-margin use. The margins per GPU-hour are: '
        'training exports '
    )
    omath(p, [_msub('r', 'T'), _t('('), _v('j'), _t(') = '),
              _msub('p', 'T'), _t(' \u2212 '), _msub('c', 'j')])
    p.add_run('; inference exports to ')
    omath(p, [_v('k')])
    p.add_run(': ')
    omath(p, [_msub('r', 'I'), _t('('), _v('j'), _t(', '), _v('k'),
              _t(') = '), _msubsup('p', 'I', 'f'), _t('('), _v('k'),
              _t(') / (1 + '), _v('\u03C4'), _t(' \u00b7 '),
              _msub('l', 'jk'), _t(') \u2212 '), _msub('c', 'j')])
    p.add_run(
        '. Total rent from operating '
    )
    omath(p, [_msub('K', 'j')])
    p.add_run(' GPU-hours is:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msub('\u03A0', 'j'), _t('('), _msub('K', 'j'),
        _t(') = '),
        _nary('\u2211', [_v('n'), _t(' = 1')], [_msub('K', 'j')],
              [_msubsup('r', 'j', '(n)')]), _t(','),
    ], eq_num='B.3')

    p, cur = mkp(doc, body, cur)
    p.add_run('which is concave and piecewise linear in ')
    omath(p, [_msub('K', 'j')])
    p.add_run('.')

    # B.5 Equilibrium and Existence (was B.6)
    cur = mkh(doc, body, cur, 'B.5 Equilibrium Definition and Existence', level=2)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'A competitive equilibrium consists of a training price '
    )
    omath(p, [_msub('p', 'T')])
    p.add_run(', inference prices ')
    omath(p, [_t('{'), _msubsup('p', 'I', 'f'), _t('('), _v('k'), _t(')}')])
    p.add_run(', and capacity allocations ')
    omath(p, [_t('{'), _msub('K', 'j'), _t('}')])
    p.add_run(
        ' such that: (i) each GPU-hour is allocated to its highest-margin use; '
        '(ii) training and inference markets clear; '
        '(iii) all allocations are feasible ('
    )
    omath(p, [_msub('K', 'j'), _t(' \u2264 '),
              _mbar_sub('K', 'j')])
    p.add_run(
        '). '
        'Existence follows from a fixed-point argument: the training supply curve is a '
        'step function with steps at '
    )
    omath(p, [_msub('c', '(i)')])
    p.add_run(
        ' and widths '
    )
    omath(p, [_mbar_sub('K', '(i)')])
    p.add_run('; intersection with the demand curve pins down ')
    omath(p, [_msub('p', 'T')])
    p.add_run('.')

    # B.6 Welfare (was B.7)
    cur = mkh(doc, body, cur, 'B.6 Welfare Cost of Sovereignty', level=2)
    p, cur = mkp(doc, body, cur)
    p.add_run('The welfare cost has two components. Import markup:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msub('DWL', 'import'), _t(' = '),
        _nary('\u2211', [_v('k'), _t(' \u2208 '), _msub('M', 'T')], [],
              [_msub('q', 'Tk'), _t(' \u00b7 '), _v('\u03BB'),
               _t(' \u00b7 '), _msub('p', 'T')]), _t('.'),
    ], eq_num='B.4')

    p, cur = mkp(doc, body, cur)
    p.add_run('Allocative inefficiency:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msub('DWL', 'alloc'), _t(' = '),
        _nary('\u2211', [_v('k'), _t(' : '), _msub('p', 'T'), _t(' < '),
                         _msub('c', 'k'), _t(' \u2264 (1+\u03BB)'),
                         _msub('p', 'T')], [],
              [_msub('q', 'Tk'), _t(' \u00b7 ('),
               _msub('c', 'k'), _t(' \u2212 '), _msub('p', 'T'), _t(').')]),
    ], eq_num='B.5')

    p, cur = mkp(doc, body, cur)
    p.add_run('Total: ')
    omath(p, [_t('DWL('), _v('\u03BB'), _t(') = '),
              _msub('DWL', 'import'), _t(' + '),
              _msub('DWL', 'alloc')])
    p.add_run(
        '. Under capacity constraints, both components are smaller because the higher '
    )
    omath(p, [_msub('p', 'T')])
    p.add_run(' narrows the gap between domestic and import costs.')

    return cur


def write_sensitivity_appendix(doc, body, last_el, demand_data):
    """Appendix C: Sensitivity Analysis with Table A3."""
    print("Inserting Appendix C (Sensitivity Analysis)...")

    pb = add_page_break(doc, body, last_el)
    cur = mkh(doc, body, pb, 'Appendix C: Sensitivity Analysis', level=1)

    sens = demand_data.get("sensitivity", [])
    if not sens:
        return cur

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Table A3 reports equilibrium outcomes under five parameter perturbations. '
        'Rankings are robust because hardware amortization accounts for approximately '
        '94 percent of total cost and is identical across countries. '
        'Only the electricity and construction components vary cross-country, and '
        'their combined share is too small for plausible perturbations to overturn the ordering.'
    )

    # --- Build parameter-change description for each scenario ---
    def _param_change(kw):
        parts = []
        if "p_E_delta" in kw:
            d = kw["p_E_delta"]
            sign = '+' if d > 0 else '\u2212'
            parts.append(f'All electricity prices {sign}${abs(d):.2f}/kWh')
        if "gpu_price" in kw:
            gp = kw["gpu_price"]
            pct = (gp - GPU_PRICE) / GPU_PRICE * 100
            parts.append(f'GPU unit price ${gp:,} ({pct:+.0f}% from ${GPU_PRICE:,})')
        if "pue_cap" in kw:
            parts.append(f'PUE capped at {kw["pue_cap"]:.2f} for all countries')
        if "gpu_util" in kw:
            parts.append(f'GPU utilization set to {kw["gpu_util"]:.0%}')
        return '; '.join(parts) if parts else '\u2014'

    headers = ['Scenario', 'Parameter change', 'p\u209c ($/hr)', 'Exporters',
               'HHI', 'Spearman \u03c1', 'Top 5']
    rows = []
    for s in sens:
        rows.append([
            s["label"],
            _param_change(s["kwargs"]) if s["kwargs"] else '\u2014',
            f'${s["p_T"]:.3f}',
            str(s["n_exporters"]),
            f'{s["hhi_T"]:.4f}',
            f'{s["rank_corr"]:.4f}',
            'Same' if s["top5_unchanged"] else 'Changed',
        ])

    col_widths = [2200, 2600, 900, 800, 800, 900, 600]
    tbl_el = add_table(
        doc, body, cur, headers, rows, col_widths=col_widths,
        title='Table A3. Sensitivity of equilibrium outcomes to parameter variation',
    )

    # Notes paragraph
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(2)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.first_line_indent = Inches(0)
    rn = note.add_run(
        'Notes: Each row re-solves the capacity-constrained equilibrium under the stated '
        'parameter change. Spearman \u03c1 is the rank correlation of country-level training costs '
        'with the baseline ordering. Top 5 indicates whether the five cheapest countries '
        'remain the same set in the same order. HHI is the Herfindahl\u2013Hirschman Index '
        'of export concentration.'
    )
    rn.font.size = Pt(7.5)
    note_el = note._element
    body.remove(note_el)
    tbl_el.addnext(note_el)

    return note_el


def write_kyrgyzstan_appendix(doc, body, last_el):
    """Appendix D: Data Center Investment Model — Kyrgyzstan."""
    print("Inserting Appendix D (Kyrgyzstan DCF)...")

    pb = add_page_break(doc, body, last_el)
    cur = mkh(doc, body, pb, 'Appendix D: Data Center Investment Model \u2014 Kyrgyzstan', level=1)

    # ── Parameters ────────────────────────────────────────────────────────
    IT_MW = 40
    PUE_KGZ = 1.08
    TOTAL_MW = IT_MW * PUE_KGZ
    LIFE = 15
    GP = 25_000
    G_LIFE = 3
    G_UTIL = 0.70
    G_TDP_W = 700
    GPUS_MW = 1_000 / G_TDP_W * 1_000
    N_GPU = int(IT_MW * GPUS_MW)
    H = 365.25 * 24
    NET_COST = 2_000
    P_ELEC = 0.038
    P_CONSTR_W = 7.83
    CONSTR = IT_MW * 1e6 * P_CONSTR_W
    STAFF = 50 * 12_000
    MAINT_PCT = 0.02
    INS_PCT = 0.005
    BW_COST = 2_400_000
    REV_HR = 2.00
    RAMP = {0: 0.0, 1: 0.40, 2: 0.60, 3: 0.70}
    TAX_R = 0.10
    GPU_DECLINE = 0.10
    ELEC_ESC = 0.02

    # WACC
    RF = 0.05; CRP = 0.04; ERP = 0.06
    COE = RF + CRP + ERP  # 15%
    COD = 0.10; DSHARE = 0.40; ESHARE = 0.60
    WACC = ESHARE * COE + DSHARE * COD * (1 - TAX_R)

    # GPU refresh schedule
    gpu_refresh = [1, 4, 7, 10, 13]
    gpu_prices = [(yr, GP * (1 - GPU_DECLINE) ** i) for i, yr in enumerate(gpu_refresh)]
    net_refresh = [1, 6, 11]

    # ── Compute year-by-year ──────────────────────────────────────────────
    years = list(range(0, LIFE + 1))
    results = []
    cum = 0
    payback = None
    for yr in years:
        cx_c = CONSTR if yr == 0 else 0
        cx_g = 0
        for gy, gp in gpu_prices:
            if yr == gy:
                cx_g = N_GPU * gp
        cx_n = N_GPU * NET_COST if yr in net_refresh else 0
        cx = cx_c + cx_g + cx_n

        if yr >= 1:
            util = RAMP.get(yr, G_UTIL)
            ep = P_ELEC * (1 + ELEC_ESC) ** (yr - 1)
            ox_e = TOTAL_MW * 1_000 * H * ep
            ox_s = STAFF * 1.03 ** (yr - 1)
            ox_m = CONSTR * MAINT_PCT
            gpu_val = 0
            for gy, gp in reversed(gpu_prices):
                if gy <= yr:
                    gpu_val = N_GPU * gp * max(0, 1 - (yr - gy) / G_LIFE)
                    break
            ox_i = (CONSTR + gpu_val) * INS_PCT
            ox_bw = BW_COST
            ox = ox_e + ox_s + ox_m + ox_i + ox_bw
            rev = N_GPU * H * util * REV_HR
            depr_c = CONSTR / LIFE
            depr_g = 0
            for gy, gp in gpu_prices:
                if gy <= yr < gy + G_LIFE:
                    depr_g = N_GPU * gp / G_LIFE
                    break
            depr = depr_c + depr_g
        else:
            util = 0; ox = 0; rev = 0; depr = 0

        ebitda = rev - ox
        ebt = ebitda - depr
        tax = max(0, ebt * TAX_R)
        ni = ebt - tax
        fcf = ni + depr - cx
        cum += fcf
        if payback is None and cum > 0 and yr >= 1:
            payback = yr
        results.append(dict(year=yr, capex=cx, revenue=rev, opex=ox,
                            ebitda=ebitda, tax=tax, ni=ni, fcf=fcf, cum=cum))

    fcf_s = [r['fcf'] for r in results]
    npv = sum(f / (1 + WACC) ** y for f, y in zip(fcf_s, years))
    lo, hi = -0.50, 2.0
    for _ in range(200):
        mid = (lo + hi) / 2
        if sum(f / (1 + mid) ** y for f, y in zip(fcf_s, years)) > 0:
            lo = mid
        else:
            hi = mid
    irr = mid

    tot_rev = sum(r['revenue'] for r in results)
    tot_cx = sum(r['capex'] for r in results)
    tot_ox = sum(r['opex'] for r in results)
    tot_elec = sum(TOTAL_MW * 1_000 * H * P_ELEC * (1 + ELEC_ESC) ** (y - 1)
                   for y in range(1, LIFE + 1))
    tot_gpu_cx = sum(N_GPU * gp for _, gp in gpu_prices)

    # ── Intro paragraph ───────────────────────────────────────────────────
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'This appendix presents a 15-year discounted cash flow (DCF) analysis for a '
        'hypothetical 40\u2009MW data center in Kyrgyzstan, the lowest-cost seller in the '
        'cost-recovery-adjusted calibration. All parameters are drawn from the calibration '
        'or from industry benchmarks.'
    )

    # ── Table A4: Facility specification ──────────────────────────────────
    specs_rows = [
        ['IT capacity', f'{IT_MW} MW'],
        ['Total power (with cooling)', f'{TOTAL_MW:.1f} MW (PUE = {PUE_KGZ:.2f})'],
        ['GPU count', f'{N_GPU:,} (H100-class, {G_TDP_W}W each)'],
        ['GPU cost / lifetime', f'${GP:,} / {G_LIFE} yr (\u221210% per generation)'],
        ['Construction cost', f'${CONSTR/1e6:.0f}M (${P_CONSTR_W:.2f}/W)'],
        ['Electricity price', f'${P_ELEC:.3f}/kWh (+2%/yr real)'],
        ['Revenue price', f'${REV_HR:.2f}/GPU-hr (wholesale)'],
        ['Utilization', f'{G_UTIL:.0%} steady-state (40% yr 1, 60% yr 2)'],
        ['WACC', f'{WACC:.1%}'],
    ]
    tbl_a4 = add_table(doc, body, cur, ['Parameter', 'Value'],
                       specs_rows, col_widths=[3500, 5300],
                       title='Table A4. Facility specification')

    # WACC note
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    rn = p.add_run(
        f'Notes: WACC = {ESHARE:.0%} \u00d7 {COE:.0%} (cost of equity) '
        f'+ {DSHARE:.0%} \u00d7 {COD:.0%} \u00d7 (1 \u2212 {TAX_R:.0%}) (after-tax debt) '
        f'= {WACC:.1%}. Cost of equity includes a {CRP:.0%} country risk premium and '
        f'{ERP:.0%} emerging-market equity premium over the {RF:.0%} risk-free rate.'
    )
    rn.font.size = Pt(7.5)
    wacc_el = p._element
    body.remove(wacc_el)
    tbl_a4.addnext(wacc_el)
    cur = wacc_el

    # ── Table A5: Year-by-year cash flow ──────────────────────────────────
    cur = add_page_break(doc, body, cur)
    cf_headers = ['Year', 'CAPEX', 'Revenue', 'OPEX', 'EBITDA', 'FCF', 'Cum.\u2009CF']
    cf_rows = []
    for r in results:
        cf_rows.append([
            str(r['year']),
            f'{r["capex"]/1e6:.1f}',
            f'{r["revenue"]/1e6:.1f}',
            f'{r["opex"]/1e6:.1f}',
            f'{r["ebitda"]/1e6:.1f}',
            f'{r["fcf"]/1e6:.1f}',
            f'{r["cum"]/1e6:.1f}',
        ])
    # Totals row
    cf_rows.append([
        'Total',
        f'{tot_cx/1e6:.1f}',
        f'{tot_rev/1e6:.1f}',
        f'{tot_ox/1e6:.1f}',
        f'{sum(r["ebitda"] for r in results)/1e6:.1f}',
        f'{sum(r["fcf"] for r in results)/1e6:.1f}',
        '',
    ])
    tbl_a5 = add_table(doc, body, cur, cf_headers, cf_rows,
                       col_widths=[700, 1400, 1400, 1400, 1400, 1400, 1100],
                       title='Table A5. Year-by-year cash flow ($ millions)')

    # ── Key metrics paragraph ─────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    p.add_run(
        f'The project yields an NPV of ${npv/1e6:,.0f}M at a {WACC:.1%} WACC, '
        f'an IRR of {irr:.1%}, and a simple payback in year\u2009{payback}. '
        f'GPU hardware accounts for ${tot_gpu_cx/1e6:.0f}M of the '
        f'${tot_cx/1e6:.0f}M total CAPEX ({tot_gpu_cx/tot_cx:.0%}), '
        f'and electricity represents {tot_elec/tot_ox:.0%} of operating costs.'
    )
    met_el = p._element
    body.remove(met_el)
    tbl_a5.addnext(met_el)
    cur = met_el

    # ── Table A6: Sensitivity analysis ────────────────────────────────────
    cur = add_page_break(doc, body, cur)
    def _run_scen(label, wacc_adj=0, price_adj=0, elec_adj=0, gpu_adj=0, util_adj=0):
        w = WACC + wacc_adj
        cfs = []
        for yr in years:
            cx = CONSTR if yr == 0 else 0
            for gy, gp in gpu_prices:
                if yr == gy:
                    cx += N_GPU * gp * (1 + gpu_adj)
            if yr in net_refresh:
                cx += N_GPU * NET_COST
            if yr >= 1:
                ep = (P_ELEC + elec_adj) * (1 + ELEC_ESC) ** (yr - 1)
                ox = (TOTAL_MW * 1000 * H * ep + STAFF * 1.03 ** (yr - 1)
                      + CONSTR * MAINT_PCT + CONSTR * INS_PCT + BW_COST)
            else:
                ox = 0
            if yr >= 1:
                u = min(max(RAMP.get(yr, G_UTIL) + util_adj, 0), 0.95)
                rv = N_GPU * H * u * (REV_HR + price_adj)
            else:
                rv = 0
            ebitda = rv - ox
            dp = CONSTR / LIFE if yr >= 1 else 0
            ebt = ebitda - dp
            tx = max(0, ebt * TAX_R)
            cfs.append(ebt - tx + dp - cx)
        npv_s = sum(cf / (1 + w) ** y for cf, y in zip(cfs, years))
        l, h = -0.50, 2.0
        for _ in range(200):
            m = (l + h) / 2
            if sum(cf / (1 + m) ** y for cf, y in zip(cfs, years)) > 0:
                l = m
            else:
                h = m
        return [label, f'${npv_s/1e6:,.0f}', f'{m:.1%}']

    sens_scenarios = [
        _run_scen('Base case'),
        _run_scen('GPU price \u221220%', gpu_adj=-0.20),
        _run_scen('GPU price +20%', gpu_adj=+0.20),
        _run_scen('Electricity +50%', elec_adj=+0.019),
        _run_scen('Electricity \u221225%', elec_adj=-0.0095),
        _run_scen('Revenue +5%', price_adj=+0.08),
        _run_scen('Revenue \u22125%', price_adj=-0.08),
        _run_scen('Utilization 80%', util_adj=+0.10),
        _run_scen('Utilization 60%', util_adj=-0.10),
        _run_scen('WACC 10%', wacc_adj=-0.026),
        _run_scen('WACC 16%', wacc_adj=+0.034),
    ]
    tbl_a6 = add_table(doc, body, cur, ['Scenario', 'NPV ($M)', 'IRR'],
                       sens_scenarios, col_widths=[3800, 2400, 2600],
                       title='Table A6. Sensitivity of investment returns to parameter variation')

    # ── Risks paragraph ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run('Risks. ')
    r.bold = True
    p.add_run(
        'Kyrgyzstan depends on the Toktogul reservoir for over 80% of electricity; '
        'seasonal drawdowns and drought years create acute power shortages. '
        'GPU procurement faces US export-control uncertainty. '
        'The reliability index assigns Kyrgyzstan a governance score of 0.50, '
        'reflecting underdeveloped contract enforcement and regulatory frameworks. '
        'Despite these risks, the engineering economics are clear: '
        'electricity at $0.038/kWh and a PUE of 1.08 yield production costs well below '
        'the global median, and the positive NPV survives all single-parameter '
        'perturbations in Table\u2009A6.'
    )
    risk_el = p._element
    body.remove(risk_el)
    tbl_a6.addnext(risk_el)

    return risk_el


def write_construction_regression_appendix(doc, body, last_el):
    """Appendix E: Construction Cost Regression."""
    import math as _math
    print("Inserting Appendix E (Construction Regression)...")

    pb = add_page_break(doc, body, last_el)
    cur = mkh(doc, body, pb, 'Appendix E: Construction Cost Regression', level=1)

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Data center construction costs per watt of IT capacity are observed for 37 countries '
        'from the Turner & Townsend Data Centre Construction Cost Index 2025 (52 markets). '
        'For the remaining countries, construction costs are predicted using the log-linear '
        'regression reported in Table\u2009A7. The dependent variable is ln($/W). '
        'Since construction accounts for only 3\u20136% of total per-GPU-hour costs, '
        'imputation error has limited impact on cost rankings.'
    )

    # ── Run regression inline ──────────────────────────────────────────────
    import numpy as _np
    _DATA = pathlib.Path(r"F:\onedrive\__documents\papers\FLOPsExport\Data")

    MARKET_TO_ISO3 = {
        "Tokyo": "JPN", "Singapore": "SGP", "Zurich": "CHE", "Osaka": "JPN",
        "Silicon Valley": "USA", "New Jersey": "USA", "Oslo": "NOR",
        "Auckland": "NZL", "Stockholm": "SWE", "Helsinki": "FIN",
        "Copenhagen": "DNK", "London": "GBR", "Vienna": "AUT",
        "Cardiff": "GBR", "Frankfurt": "DEU", "Berlin": "DEU",
        "Kuala Lumpur": "MYS", "Kingdom of Saudi Arabia": "SAU",
        "Chicago": "USA", "Jakarta": "IDN", "North Virginia": "USA",
        "Portland": "USA", "Paris": "FRA", "Amsterdam": "NLD",
        "S\u00e3o Paulo": "BRA", "Sydney": "AUS", "Lagos": "NGA",
        "Melbourne": "AUS", "Quer\u00e9taro": "MEX", "Cape Town": "ZAF",
        "Lisbon": "PRT", "Seoul": "KOR", "Johannesburg": "ZAF",
        "Bordeaux": "FRA", "Dublin": "IRL", "Madrid": "ESP",
        "Atlanta": "USA", "Montevideo": "URY", "Phoenix": "USA",
        "Columbus": "USA", "Milan": "ITA", "Nairobi": "KEN",
        "Dallas": "USA", "Charlotte": "USA", "Toronto": "CAN",
        "UAE": "ARE", "Warsaw": "POL", "Santiago": "CHL",
        "Athens": "GRC", "Bogot\u00e1": "COL", "Mumbai": "IND",
        "Shanghai": "CHN",
    }

    dcci = {}
    with open(_DATA / "dcci_2025_construction_costs.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            iso3 = MARKET_TO_ISO3[row["market"]]
            cost = float(row["usd_per_watt"])
            if iso3 in dcci:
                dcci[iso3].append(cost)
            else:
                dcci[iso3] = [cost]
    for iso3 in dcci:
        dcci[iso3] = _np.mean(dcci[iso3])

    gdp_d = {}
    with open(_DATA / "wb_gdp_per_capita_ppp_2023.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            gdp_d[row["iso3"]] = float(row["gdp_pcap_ppp_2023"])
    reg_d = {}
    with open(_DATA / "wb_country_regions.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            reg_d[row["iso3"]] = row["region"]
    urban_d = {}
    with open(_DATA / "wb_urban_share_2023.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            urban_d[row["iso3"]] = float(row["urban_share_pct"]) / 100.0
    seismic_d = {}
    with open(_DATA / "seismic_zones.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            seismic_d[row["iso3"]] = int(row["seismic_high"])

    REF_REGION = "Europe & Central Asia"
    DUMMY_REGIONS = sorted(r for r in set(reg_d.values()) if r != REF_REGION)

    matched = []
    for iso3, avg_cost in dcci.items():
        if iso3 in gdp_d and iso3 in reg_d:
            matched.append({
                "iso3": iso3, "cost": avg_cost,
                "gdp_pcap": gdp_d[iso3], "region": reg_d[iso3],
                "urban_share": urban_d.get(iso3, 0.5),
                "seismic": seismic_d.get(iso3, 0),
            })

    n = len(matched)
    k = 5 + len(DUMMY_REGIONS)
    y = _np.array([_math.log(m["cost"]) for m in matched])
    X = _np.zeros((n, k))
    col_names = ["Intercept", "ln(GDP per capita)", "ln(Population)",
                 "Urban population share",
                 "Seismic zone indicator"] + [r.split(",")[0].strip() for r in DUMMY_REGIONS]
    pop_d = {}
    with open(_DATA / "wb_population_2023.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            pop_d[row["iso3"]] = int(row["population_2023"])
    for i, m in enumerate(matched):
        X[i, 0] = 1.0
        X[i, 1] = _math.log(m["gdp_pcap"])
        X[i, 2] = _math.log(pop_d.get(m["iso3"], 1_000_000))
        X[i, 3] = m["urban_share"]
        X[i, 4] = m["seismic"]
        for j2, reg in enumerate(DUMMY_REGIONS):
            X[i, 5 + j2] = 1.0 if m["region"] == reg else 0.0

    beta = _np.linalg.lstsq(X, y, rcond=None)[0]
    y_hat = X @ beta
    resid = y - y_hat
    ss_res = _np.sum(resid ** 2)
    ss_tot = _np.sum((y - _np.mean(y)) ** 2)
    r2 = 1 - ss_res / ss_tot
    adj_r2 = 1 - (1 - r2) * (n - 1) / (n - k)
    rmse = _math.sqrt(ss_res / (n - k))
    var_beta = ss_res / (n - k) * _np.diag(_np.linalg.inv(X.T @ X))
    se = _np.sqrt(_np.maximum(var_beta, 0))

    # Build table rows
    reg_rows = []
    for j2 in range(k):
        sig = ''
        t = beta[j2] / se[j2] if se[j2] > 0 else 0
        if abs(t) > 2.576:
            sig = '***'
        elif abs(t) > 1.96:
            sig = '**'
        elif abs(t) > 1.645:
            sig = '*'
        reg_rows.append([
            col_names[j2],
            f'{beta[j2]:.3f}',
            f'({se[j2]:.3f})',
            sig,
        ])

    tbl = add_table(doc, body, cur,
                    ['Variable', 'Coefficient', 'Std. Error', ''],
                    reg_rows, col_widths=[3600, 1800, 1800, 600],
                    title='Table A7. Construction cost regression: ln($/W)')

    # Notes paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    rn = p.add_run(
        f'Notes: OLS regression on {n} countries from the Turner & Townsend DCCI 2025. '
        f'Dependent variable: ln(construction cost in $/W). '
        f'R\u00b2 = {r2:.2f}, adjusted R\u00b2 = {adj_r2:.2f}, RMSE = {rmse:.3f}. '
        f'Reference region: Europe & Central Asia. '
        f'*** p < 0.01, ** p < 0.05, * p < 0.10.'
    )
    rn.font.size = Pt(7.5)
    note_el = p._element
    body.remove(note_el)
    tbl.addnext(note_el)

    return note_el


def write_figure4b(doc, body, last_ref, demand_data):
    """Generate and embed Figure 1 (reliability rank scatter) after references."""
    print("Embedding Figure 1 (Reliability Rank Scatter)...")
    xi_adj = demand_data.get("xi_adjusted", {})
    baseline_rank = xi_adj.get("baseline_rank_map", {})
    xi_rank = xi_adj.get("xi_rank_map", {})
    iso_country = demand_data.get("iso_country", {})
    xi = demand_data.get("xi", {})
    if not baseline_rank or not xi_rank:
        return last_ref

    # Countries with active DC construction announcements (from Section 1)
    DC_ACTIVE = {
        'ARM', 'KEN', 'SAU', 'MAR', 'MYS', 'IDN', 'ARE',  # original
        'IND', 'BRA', 'MEX', 'CHN', 'THA', 'VNM', 'TUR',  # tier 1
        'PHL', 'ZAF', 'EGY', 'KAZ', 'UZB', 'NGA',         # tier 1-2
    }
    # Countries whose ξ value should be shown in the label
    XI_SHOW = {'IRN', 'PAK', 'CHN', 'FIN', 'KEN',
               'RUS', 'TJK', 'UKR', 'BIH', 'ARM', 'BGR'}

    common = [iso for iso in baseline_rank if iso in xi_rank]

    # Separate into active-construction and regular
    reg_isos = [iso for iso in common if iso not in DC_ACTIVE]
    act_isos = [iso for iso in common if iso in DC_ACTIVE]

    fig, ax = plt.subplots(figsize=(5.5, 5.5))
    # Regular countries: dots
    ax.scatter([baseline_rank[iso] + 1 for iso in reg_isos],
               [xi_rank[iso] + 1 for iso in reg_isos],
               s=20, c='#b2182b', alpha=0.7, marker='o',
               edgecolors='white', linewidth=0.3, label='Other countries', zorder=2)
    # Active construction: stars
    if act_isos:
        ax.scatter([baseline_rank[iso] + 1 for iso in act_isos],
                   [xi_rank[iso] + 1 for iso in act_isos],
                   s=60, c='#1a3a5c', alpha=0.85, marker='*',
                   edgecolors='white', linewidth=0.3,
                   label='Active DC construction', zorder=3)
    maxr = max(max(baseline_rank[iso] + 1 for iso in common),
               max(xi_rank[iso] + 1 for iso in common))
    ax.plot([1, maxr], [1, maxr], '--', color='gray', alpha=0.5, linewidth=0.8)

    # Build labels
    def _label(iso):
        name = iso_country.get(iso, iso)
        if iso in XI_SHOW:
            return f'{name} (\u03BE={xi.get(iso, 1.0):.2f})'
        return name

    try:
        from adjustText import adjust_text
        texts = []
        for iso in common:
            shift = abs(baseline_rank[iso] - xi_rank[iso])
            if (shift > 15 or baseline_rank[iso] < 5 or xi_rank[iso] < 5
                    or iso in XI_SHOW or iso in DC_ACTIVE):
                texts.append(ax.text(baseline_rank[iso] + 1, xi_rank[iso] + 1,
                                     _label(iso), fontsize=5.5, alpha=0.85))
        adjust_text(texts, ax=ax,
                    arrowprops=dict(arrowstyle='-', color='gray', alpha=0.4, lw=0.4),
                    force_text=(0.4, 0.4), expand=(1.2, 1.4))
    except ImportError:
        for iso in common:
            shift = abs(baseline_rank[iso] - xi_rank[iso])
            if (shift > 15 or baseline_rank[iso] < 5 or xi_rank[iso] < 5
                    or iso in XI_SHOW or iso in DC_ACTIVE):
                ax.annotate(_label(iso),
                            (baseline_rank[iso] + 1, xi_rank[iso] + 1),
                            fontsize=5.5, alpha=0.85)

    ax.set_xlabel('Baseline cost rank', fontsize=9)
    ax.set_ylabel('Reliability-adjusted rank', fontsize=9)
    # No legend – star/dot distinction explained in figure notes
    ax.grid(alpha=0.2)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=300, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)

    # Page break before figure
    pb_el = add_page_break(doc, body, last_ref)

    # Figure title with bookmark (outside the image)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.paragraph_format.first_line_indent = Inches(0)
    title_p._element.append(make_bookmark(120, 'Figure1'))
    hl_f1 = OxmlElement('w:hyperlink')
    hl_f1.set(qn('w:anchor'), 'Figure1txt')
    hl_f1.set(qn('w:history'), '1')
    r_f1 = OxmlElement('w:r')
    rPr_f1 = OxmlElement('w:rPr')
    b_f1 = OxmlElement('w:b')
    rPr_f1.append(b_f1)
    sz_f1 = OxmlElement('w:sz')
    sz_f1.set(qn('w:val'), '20')
    rPr_f1.append(sz_f1)
    clr_f1 = OxmlElement('w:color')
    clr_f1.set(qn('w:val'), LINK_COLOR)
    uu_f1 = OxmlElement('w:u')
    uu_f1.set(qn('w:val'), 'single')
    rPr_f1.append(clr_f1)
    rPr_f1.append(uu_f1)
    r_f1.append(rPr_f1)
    t_f1 = OxmlElement('w:t')
    t_f1.text = 'Figure 1'
    r_f1.append(t_f1)
    hl_f1.append(r_f1)
    title_p._element.append(hl_f1)
    title_p._element.append(make_bookmark_end(120))
    run_ft = title_p.add_run('. Rank change with reliability adjustment')
    run_ft.bold = True
    run_ft.font.size = Pt(10)
    title_el = title_p._element
    body.remove(title_el)
    pb_el.addnext(title_el)

    # Embed image
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(4)
    pic_p.paragraph_format.space_after = Pt(4)
    run = pic_p.add_run()
    run.add_picture(buf, width=Inches(4.5))
    pic_el = pic_p._element
    body.remove(pic_el)
    title_el.addnext(pic_el)

    # Notes (with 0.5" left and right indent)
    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(4)
    note_p.paragraph_format.space_after = Pt(6)
    note_p.paragraph_format.first_line_indent = Inches(0)
    note_p.paragraph_format.left_indent = Inches(0.5)
    note_p.paragraph_format.right_indent = Inches(0.5)
    rn1 = note_p.add_run('Notes: ')
    rn1.bold = True
    rn1.font.size = Pt(7.5)
    rn2 = note_p.add_run(
        'Each point is one country. The dashed line marks unchanged rank. '
        'Countries above the line improve their position after reliability '
        'adjustment; countries below it fall. Stars (\u2605) indicate countries '
        'with active data center construction announcements. '
        'Values in parentheses show the reliability index \u03BE. '
        'Even countries with \u03BE \u2248 1 shift off the diagonal because '
        'penalizing low-\u03BE competitors pushes them down, mechanically '
        'raising higher-\u03BE countries.'
    )
    rn2.font.size = Pt(7.5)
    note_el = note_p._element
    body.remove(note_el)
    pic_el.addnext(note_el)

    return note_el


def write_table1(doc, body, after_el, demand_data):
    """Table 1: Model parameters (formerly Table A1), placed in main body."""
    print("Inserting Table 1 (Model parameters)...")

    # Page break before Table 1
    pb_el = add_page_break(doc, body, after_el)

    # Table 1 title with bookmark
    tp1 = doc.add_paragraph()
    tp1.paragraph_format.space_before = Pt(6)
    tp1.paragraph_format.space_after = Pt(3)
    tp1.paragraph_format.first_line_indent = Inches(0)
    tp1._element.append(make_bookmark(110, 'Table1'))
    hl_a1 = OxmlElement('w:hyperlink')
    hl_a1.set(qn('w:anchor'), 'Table1txt')
    hl_a1.set(qn('w:history'), '1')
    r_a1 = OxmlElement('w:r')
    rPr_a1 = OxmlElement('w:rPr')
    b_a1 = OxmlElement('w:b')
    rPr_a1.append(b_a1)
    sz_a1 = OxmlElement('w:sz')
    sz_a1.set(qn('w:val'), '20')
    rPr_a1.append(sz_a1)
    clr_a1 = OxmlElement('w:color')
    clr_a1.set(qn('w:val'), LINK_COLOR)
    uu_a1 = OxmlElement('w:u')
    uu_a1.set(qn('w:val'), 'single')
    rPr_a1.append(clr_a1)
    rPr_a1.append(uu_a1)
    r_a1.append(rPr_a1)
    t_a1 = OxmlElement('w:t')
    t_a1.text = 'Table 1'
    r_a1.append(t_a1)
    hl_a1.append(r_a1)
    tp1._element.append(hl_a1)
    tp1._element.append(make_bookmark_end(110))
    run_tt1 = tp1.add_run('. Model parameters')
    run_tt1.bold = True
    run_tt1.font.size = Pt(10)
    tp1_el = tp1._element
    body.remove(tp1_el)
    pb_el.addnext(tp1_el)

    # Load parameters from CSV
    param_rows = []
    with open(DATA / "model_parameters.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            param_rows.append(row)

    _sym_map = {
        'gamma': '\u03B3', 'P_GPU': 'P\u1d33\u1d18\u1d1c', 'L': 'L',
        'beta': '\u03B2', 'H': 'H', 'rho': '\u03C1', 'eta': '\u03B7',
        'phi': '\u03C6', 'delta': '\u03B4', 'theta_bar': '\u03B8\u0304',
        'D': 'D', 'tau': '\u03C4', 'lambda': '\u03BB', 'alpha': '\u03B1',
        'Q': 'Q', 'xi_j': '\u03BE\u2C7C',
    }
    _source_to_bm = {
        'NVIDIA (2024)': 'NVIDIA2024',
        'Barroso et al. (2018)': 'Barroso2018',
        'Liu et al. (2023)': 'Liu2023',
        'Flucker et al. (2013)': 'Flucker2013',
        'Turner and Townsend (2025)': 'TurnerTownsend2025',
        'UNCTAD (2025)': 'UNCTAD2025',
        'Deloitte (2025)': 'Deloitte2025',
        'Epoch AI (2024)': 'EpochAI2024',
        'Google (2024)': None,
        'WGI and Enterprise Surveys': 'WorldBank2024',
    }

    n_params = len(param_rows)
    param_tbl = doc.add_table(rows=n_params + 1, cols=5)
    param_tbl.style = 'Table Grid'
    param_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tblPr = param_tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        param_tbl._tbl.insert(0, tblPr)
    old_bdr = tblPr.find(qn('w:tblBorders'))
    if old_bdr is not None:
        tblPr.remove(old_bdr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), TABLE_WIDTH_PCT)
    tblW.set(qn('w:type'), 'pct')

    _pcw = [Inches(2.3), Inches(0.6), Inches(0.6), Inches(1.0), Inches(2.0)]
    _pcw_labels = ['Parameter', 'Symbol', 'Eq.', 'Value', 'Source']

    def _cell_border(tc, sides, style='single'):
        tcPr = tc.get_or_add_tcPr()
        tcB = OxmlElement('w:tcBorders')
        for s in sides:
            b = OxmlElement(f'w:{s}')
            b.set(qn('w:val'), style)
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            tcB.append(b)
        tcPr.append(tcB)

    for j, lbl in enumerate(_pcw_labels):
        cell = param_tbl.rows[0].cells[j]
        cell.text = ''
        p_h = cell.paragraphs[0]
        p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rh = p_h.add_run(lbl)
        rh.bold = True
        rh.font.size = Pt(8.5)
        cell.width = _pcw[j]
        _cell_border(cell._tc, ['top', 'bottom'])

    for i, pr in enumerate(param_rows):
        sym_display = _sym_map.get(pr['symbol'], pr['symbol'])
        val_str = pr['value']
        if pr['unit']:
            val_str = f"{val_str} {pr['unit']}"
        if pr['symbol'] == 'P_GPU':
            val_str = f"${int(float(pr['value'])):,}"
        elif pr['symbol'] == 'rho':
            val_str = f"${RHO:.2f}/hr"
        elif pr['symbol'] == 'eta':
            val_str = f"${ETA:.2f}/hr"
        elif pr['symbol'] == 'Q':
            val_str = "6\u00d710\u00b9\u2070 GPU-hr/yr"
        eq_str = pr.get('equation', '')
        src_text = pr['source']
        src_bm = _source_to_bm.get(src_text)
        row_data = [pr['description'], sym_display, eq_str, val_str]
        for j, txt in enumerate(row_data):
            cell = param_tbl.rows[i + 1].cells[j]
            cell.text = ''
            p_c = cell.paragraphs[0]
            if j == 0:
                p_c.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            rc = p_c.add_run(txt)
            rc.font.size = Pt(8)
            cell.width = _pcw[j]
        src_cell = param_tbl.rows[i + 1].cells[4]
        src_cell.text = ''
        src_p = src_cell.paragraphs[0]
        src_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        src_cell.width = _pcw[4]
        if src_bm and src_text:
            rPr_src = OxmlElement('w:rPr')
            sz_src = OxmlElement('w:sz')
            sz_src.set(qn('w:val'), '16')
            rPr_src.append(sz_src)
            hl_src = make_hyperlink(src_bm, src_text, rPr_orig=rPr_src)
            src_p._element.append(hl_src)
        elif src_text:
            rc_src = src_p.add_run(src_text)
            rc_src.font.size = Pt(8)
        if i == n_params - 1:
            for j in range(5):
                _cell_border(param_tbl.rows[i + 1].cells[j]._tc, ['bottom'], style='double')

    for row in param_tbl.rows:
        for cell in row.cells:
            for pp in cell.paragraphs:
                pPr = pp._element.get_or_add_pPr()
                sp = OxmlElement('w:spacing')
                sp.set(qn('w:before'), '10')
                sp.set(qn('w:after'), '10')
                pPr.append(sp)

    param_tbl_el = param_tbl._tbl
    body.remove(param_tbl_el)
    tp1_el.addnext(param_tbl_el)

    # Table 1 notes
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(6)
    note.paragraph_format.first_line_indent = Inches(0)
    note.paragraph_format.line_spacing = 1.0
    rn1 = note.add_run('Notes: ')
    rn1.bold = True
    rn1.font.size = Pt(7.5)
    rn1 = note.add_run(
        'Hardware cost \u03C1 = P\u1d33\u1d18\u1d1c / (L \u00b7 H \u00b7 \u03B2). '
        'PUE(\u03B8) = \u03C6 + \u03B4 \u00b7 max(0, \u03B8 \u2212 \u03B8\u0304). '
        'RTT = round-trip time, the network delay for a data packet to travel from '
        'client to server and back, measured in milliseconds. '
        'The reliability index \u03BE\u2C7C combines governance quality, grid reliability, '
        'and sanctions exposure (equation 2). '
        'The baseline calibration sets \u03BE\u2C7C = 1 for all countries.'
    )
    rn1.font.size = Pt(7.5)
    note_el = note._element
    body.remove(note_el)
    param_tbl_el.addnext(note_el)

    return note_el


def write_references(doc, body, refs):
    print("Updating references...")

    # Page break before References heading
    add_page_break(doc, body, refs.getprevious())

    all_now = list(body)
    ri = all_now.index(refs)
    ref_els = []
    ref_txts = []
    for i in range(ri + 1, len(all_now)):
        el = all_now[i]
        if el.tag == qn('w:p'):
            t = "".join(r.text or "" for r in el.findall(f'.//{qn("w:t")}'))
            if t.strip():
                ref_txts.append(t.strip())
                ref_els.append(el)
        elif el.tag == qn('w:sectPr'):
            break

    new_refs = [
        'Barroso, L., U. H\u00F6lzle, and P. Ranganathan. (2018). '
        'The Datacenter as a Computer: Designing Warehouse-Scale Machines, '
        '3rd ed. San Rafael, CA: Morgan & Claypool.',

        'Brainard, S. (1997). \u201CAn Empirical Assessment of the Proximity-Concentration '
        'Trade-off.\u201D American Economic Review, 87(4): 520\u2013544.',

        'Cloudscene. (2025). Global Data Center Directory. cloudscene.com.',

        'Deloitte. (2025). \u201CTechnology, Media, and Telecommunications Predictions 2026.\u201D '
        'Deloitte Insights.',

        'Deloitte and Google. (2020). \u201CMilliseconds Make Millions.\u201D '
        'Deloitte Digital and Google.',

        'EIA. (2025). Electric Power Monthly. U.S. Energy Information Administration.',

        'Epoch AI. (2024). \u201CThe Training Compute of Notable AI Models.\u201D epochai.org.',

        'Eurostat. (2025). Electricity Prices for Non-Household Consumers '
        '(nrg_pc_205). Luxembourg: Eurostat.',

        'Firebird. (2026). \u201CPhase 2 of Armenia AI Megaproject, Scaling to $4 Billion '
        'and 50,000 GPUs.\u201D Press release, January 2026.',

        'Flucker, S., R. Tozer, and R. Whitehead. (2013). \u201CData Centre Energy Efficiency '
        'Analysis.\u201D Building Services Engineering Research and Technology, 34(1): 103\u2013117.',

        'GlobalPetrolPrices. (2025). Electricity Prices Around the World. '
        'globalpetrolprices.com.',

        'Goldfarb, A., and D. Trefler. (2018). \u201CAI and International Trade.\u201D '
        'In The Economics of Artificial Intelligence. Chicago: Univ. of Chicago Press, '
        'pp. 463\u2013492.',

        'Grossman, G., and E. Rossi-Hansberg. (2008). \u201CTrading Tasks: A Simple Theory '
        'of Offshoring.\u201D American Economic Review, 98(5): 1978\u20131997.',

        'Hausmann, R., J. Hwang, and D. Rodrik. (2007). \u201CWhat You Export Matters.\u201D '
        'Journal of Economic Growth, 12(1): 1\u201325.',

        'Helpman, E., M. Melitz, and S. Yeaple. (2004). \u201CExport Versus FDI with '
        'Heterogeneous Firms.\u201D American Economic Review, 94(1): 300\u2013316.',

        'Hersbach, H., et al. (2020). \u201CThe ERA5 Global Reanalysis.\u201D '
        'Quarterly Journal of the Royal Meteorological Society, 146(730): 1999\u20132049.',

        'Hummels, D., and G. Schaur. (2013). \u201CTime as a Trade Barrier.\u201D '
        'American Economic Review, 103(7): 2935\u20132959.',

        'IEA. (2025). \u201CEnergy Demand from AI.\u201D Published online at iea.org.',

        'IMF. (2025). \u201CFossil Fuel Subsidies Data: 2025 Update.\u201D '
        'IMF Working Paper WP/25/270.',

        'Korinek, A., and J. Stiglitz. (2021). \u201CAI, Globalization, and Strategies for '
        'Economic Development.\u201D NBER Working Paper No. 28453.',

        'Lazard. (2025). Lazard\u2019s Levelized Cost of Energy Analysis, Version 17.0. '
        'lazard.com.',

        'Krugman, P. (1991). \u201CIncreasing Returns and Economic Geography.\u201D '
        'Journal of Political Economy, 99(3): 483\u2013499.',

        'Lim\u00E3o, N., and A. Venables. (2001). \u201CInfrastructure, Geographical '
        'Disadvantage, Transport Costs, and Trade.\u201D '
        'World Bank Economic Review, 15(3): 451\u2013479.',

        'Liu, Z., A. Wierman, Y. Chen, B. Raber, and J. Moriarty. (2023). '
        '\u201CSustainability of Data Center Digital Twins.\u201D '
        'Proceedings of ACM e-Energy, pp. 178\u2013189.',

        'NVIDIA. (2024). NVIDIA H100 Tensor Core GPU Datasheet. nvidia.com.',

        'Oltmanns, J., D. Krcmarik, and R. Gatti. (2021). \u201CData Centre Site Selection.\u201D '
        'Journal of Property Investment & Finance, 39(1): 55\u201372.',

        'Turner & Townsend. (2025). Data Centre Construction Cost Index 2025. '
        'turnerandtownsend.com.',

        'Turner Lee, N., and D. West. (2025). \u201CThe Future of Data Centers.\u201D '
        'Brookings Institution, November 2025.',

        'UNCTAD. (2025). Technology and Innovation Report 2025. Geneva: United Nations.',

        'U.S. Department of Justice and Federal Trade Commission. (2010). '
        'Horizontal Merger Guidelines. Washington, DC.',

        'Uptime Institute. (2024). \u201CData Center Staffing: Trends and Best Practices.\u201D uptimeinstitute.com.',

        'WonderNetwork. (2024). Global Ping Statistics. wondernetwork.com.',

        'World Bank. (2024). World Development Indicators. Washington, DC.',

        'Lehdonvirta, V., B. Wu, and Z. Hawkins. (2024). \u201CCompute North vs. Compute South: '
        'The Uneven Possibilities of Compute-Based AI Governance Around the Globe.\u201D '
        'Proceedings of the AAAI/ACM Conference on AI, Ethics, and Society, 7(1): 828\u2013838.',

        'Pilz, K., Y. Mahmood, and L. Heim. (2025). AI\u2019s Power Requirements Under '
        'Exponential Growth. Santa Monica, CA: RAND Corporation, RR-A3572-1.',

        'Sastry, G., L. Heim, et al. (2024). \u201CComputing Power and the Governance of '
        'Artificial Intelligence.\u201D arXiv:2402.08797.',

        'Eaton, J., and S. Kortum. (2002). \u201CTechnology, Geography, and Trade.\u201D '
        'Econometrica, 70(5): 1741\u20131779.',

        'Dornbusch, R., S. Fischer, and P. Samuelson. (1977). \u201CComparative Advantage, '
        'Trade, and Payments in a Ricardian Model with a Continuum of Goods.\u201D '
        'American Economic Review, 67(5): 823\u2013839.',

        'Arkolakis, C., A. Costinot, and A. Rodr\u00EDguez-Clare. (2012). \u201CNew Trade '
        'Models, Same Old Gains?\u201D American Economic Review, 102(1): 94\u2013130.',

        'van der Ploeg, F. (2011). \u201CNatural Resources: Curse or Blessing?\u201D '
        'Journal of Economic Literature, 49(2): 366\u2013420.',

        'Ohlin, B. (1933). Interregional and International Trade. '
        'Cambridge, MA: Harvard University Press.',

        'Biglaiser, G., J. Cr\u00E9mer, and A. Mantovani. (2024). \u201CThe Economics of the Cloud.\u201D '
        'Toulouse School of Economics Working Paper No. 24-1520.',

        'Stojkoski, V., P. Coll-Ruiz, N. Mar\u00E9chal, and C. Requier-Desjardins. (2024). '
        '\u201CTrade in Cloud Computing and AI Services.\u201D WTO Staff Working Paper ERSD-2024-03.',

        'World Bank. (2025). Digital Progress and Trends Report 2025: '
        'Strengthening AI Foundations. Washington, DC: World Bank.',
    ]

    ref_txts = sorted(new_refs, key=lambda x: x.lower())
    for el in ref_els:
        body.remove(el)

    # Build reverse map: reference text prefix -> key
    def find_ref_key(ref_text):
        """Find the citation key for a reference text."""
        for key, prefix in REF_KEY_MAP.items():
            if ref_text.startswith(prefix):
                return key
        return None

    bm_id_refs = [500]  # bookmark IDs for references
    cur = refs
    for rt in ref_txts:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(12)
        italic_portion = find_italic_portion(rt)
        key = find_ref_key(rt)
        if key:
            # Add bookmark target for in-text citation links
            p._element.append(make_bookmark(bm_id_refs[0], key))
            # Author portion as hyperlink back to in-text citation
            split_pos = rt.find('\u201C')
            if split_pos < 0:
                split_pos = rt.find('(', 3)
            if split_pos < 0:
                split_pos = len(rt)
            author_part = rt[:split_pos]
            remaining = rt[split_pos:]
            p._element.append(make_hyperlink(f'{key}txt', author_part))
            if remaining:
                _write_ref_segments(p, remaining, italic_portion)
            p._element.append(make_bookmark_end(bm_id_refs[0]))
            bm_id_refs[0] += 1
        else:
            _write_ref_segments(p, rt, italic_portion)
        el = p._element
        body.remove(el)
        cur.addnext(el)
        cur = el
    print(f"  {len(ref_txts)} references")
    return cur  # last reference element


def link_citations(body):
    print("Linking citations...")
    bm_id_cite = [200]
    passes = 0
    while True:
        n = link_citations_pass(body, CITE_MAP, bm_id_cite)
        passes += 1
        if n == 0 or passes > 10:
            break
    print(f"  {bm_id_cite[0] - 200} citation links created in {passes} passes")


def link_equations(body):
    """Link 'equation (N)' mentions in text to their display equation bookmarks."""
    print("Linking equation references...")
    import re
    count = 0
    bm_id_eq = [900]
    eq_pattern = re.compile(r'equation \((\d+)\)')
    for p_el in list(body.findall(qn('w:p'))):
        for child in list(p_el):
            if child.tag != qn('w:r'):
                continue
            t_el = child.find(qn('w:t'))
            if t_el is None or not t_el.text:
                continue
            text = t_el.text
            m = eq_pattern.search(text)
            if not m:
                continue
            eq_num = m.group(1)
            anchor = f'Eq{eq_num}'
            match_text = m.group(0)
            idx = m.start()
            before = text[:idx]
            after = text[idx + len(match_text):]
            rPr_orig = child.find(qn('w:rPr'))
            t_el.text = before
            t_el.set(XML_SPACE, SPACE_PRESERVE)
            ins = child
            bm_start = make_bookmark(bm_id_eq[0], f'{anchor}txt{bm_id_eq[0]}')
            ins.addnext(bm_start)
            ins = bm_start
            hyperlink = make_hyperlink(anchor, match_text, rPr_orig)
            ins.addnext(hyperlink)
            ins = hyperlink
            bm_end = make_bookmark_end(bm_id_eq[0])
            ins.addnext(bm_end)
            ins = bm_end
            bm_id_eq[0] += 1
            if after:
                ra = OxmlElement('w:r')
                if rPr_orig is not None:
                    ra.append(copy.deepcopy(rPr_orig))
                ta = OxmlElement('w:t')
                ta.set(XML_SPACE, SPACE_PRESERVE)
                ta.text = after
                ra.append(ta)
                ins.addnext(ra)
            count += 1
    print(f"  {count} equation links created")


def fix_orphan_backlinks(body, refs):
    """Remove hyperlink wrappers in references whose back-link targets don't exist in the body."""
    # Collect all bookmark names in the document
    all_bookmarks = set()
    for el in body:
        for bm in el.findall(f'.//{{{W_NS}}}bookmarkStart'):
            name = bm.get(f'{{{W_NS}}}name', '')
            if name:
                all_bookmarks.add(name)

    # Scan reference paragraphs for hyperlinks with missing targets
    refs_idx = list(body).index(refs)
    fixed = 0
    for el in list(body)[refs_idx + 1:]:
        if el.tag == f'{{{W_NS}}}sectPr':
            break
        # Stop at headings (e.g. Appendix) that follow references
        if el.tag == qn('w:p'):
            pPr = el.find(qn('w:pPr'))
            if pPr is not None:
                pS = pPr.find(qn('w:pStyle'))
                if pS is not None and 'Heading' in pS.get(qn('w:val'), ''):
                    break
                if pPr.find(f'{{{W_NS}}}sectPr') is not None:
                    break
        for hl in el.findall(f'.//{{{W_NS}}}hyperlink'):
            anchor = hl.get(f'{{{W_NS}}}anchor', '')
            if anchor and anchor not in all_bookmarks:
                # Replace hyperlink element with its child runs (keep text, drop link)
                parent = hl.getparent()
                idx = list(parent).index(hl)
                children = list(hl)
                for child in children:
                    hl.remove(child)
                    parent.insert(idx, child)
                    idx += 1
                parent.remove(hl)
                fixed += 1
    if fixed:
        print(f"  Fixed {fixed} orphan back-link(s) in references")


def apply_formatting(doc, body, refs, title_el, author_el, ver_el, abs_text_el):
    print("Applying formatting...")
    # Set Normal style defaults
    normal = doc.styles['Normal']
    normal.font.name = TIMES_NEW_ROMAN
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5

    # Identify reference paragraphs to protect their spacing
    refs_idx = list(body).index(refs)
    ref_elements = set()
    for el in list(body)[refs_idx + 1:]:
        if el.tag == qn('w:sectPr'):
            break
        if el.tag == qn('w:p'):
            pPr = el.find(qn('w:pPr'))
            if pPr is not None:
                pS = pPr.find(qn('w:pStyle'))
                if pS is not None and 'Heading' in pS.get(qn('w:val'), ''):
                    break
                if pPr.find(f'{{{W_NS}}}sectPr') is not None:
                    break
            ref_elements.add(el)

    # Paragraphs to protect from global formatting (centered title page elements)
    _protected = {title_el, author_el, ver_el, abs_text_el}

    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        # Heading 1: Times New Roman, blue, 14pt, bold
        if style == 'Heading 1':
            for run in p.runs:
                run.font.color.rgb = HEADING_BLUE
                run.font.name = TIMES_NEW_ROMAN
                run.font.size = Pt(14)
                run.bold = True
            continue
        # Heading 2: Times New Roman, blue, 12pt, italic, no bold
        if style == 'Heading 2':
            for run in p.runs:
                run.font.color.rgb = HEADING_BLUE
                run.font.name = TIMES_NEW_ROMAN
                run.font.size = Pt(12)
                run.italic = True
                run.bold = False
            continue
        if 'Heading' not in style and p.text.strip():
            # Skip title page elements (centered)
            if p._element not in _protected:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if p.paragraph_format.first_line_indent is None or p.paragraph_format.first_line_indent > 0:
                    p.paragraph_format.first_line_indent = Inches(0)
            # Subtitle runs: italic first run ending with ". " → font 12, TNR, not bold
            runs = [r for r in p.runs if r.text.strip()]
            if runs and runs[0].italic and runs[0].text.rstrip().endswith('.'):
                runs[0].font.size = Pt(12)
                runs[0].font.name = TIMES_NEW_ROMAN
                runs[0].bold = False
            # Preserve reference formatting (hanging indent + 4pt spacing)
            if p._element in ref_elements:
                continue
            # Preserve title page spacing
            if p._element in _protected:
                continue
            p.paragraph_format.space_before = Pt(0)
            # Preserve Pt(2) spacing on paragraphs immediately before equations
            if p.paragraph_format.space_after is None or p.paragraph_format.space_after >= Pt(8):
                p.paragraph_format.space_after = Pt(8)


def add_page_numbers_and_break(doc, body, kw_el):
    print("Adding page numbers...")
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.different_first_page_header_footer = True

    # Default footer: right-aligned page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.clear()
    # Insert PAGE field: w:fldSimple or fldChar sequence
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), ' PAGE ')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '1'
    r.append(t)
    fld.append(r)
    fp._element.append(fld)

    # First page footer: empty (no page number on title page)
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    if first_footer.paragraphs:
        first_footer.paragraphs[0].clear()

    # ═══════════════════════════════════════════════════════════════════════
    # STEP 17: PAGE BREAK AFTER ABSTRACT (title on separate page)
    # ═══════════════════════════════════════════════════════════════════════
    print("Adding page break after keywords...")
    add_page_break(doc, body, kw_el)


def main():
    # ═══════════════════════════════════════════════════════════════════════
    # LOAD DATA (v3)
    # ═══════════════════════════════════════════════════════════════════════

    print("Loading data...")
    cal = []
    with open(DATA / "calibration_results_v3.csv", encoding="utf-8") as f:
        cal = list(csv.DictReader(f))
    reg = {}
    with open(DATA / "calibration_regimes_v3.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            reg[row["iso3"]] = row
    # World Bank operational ECA region (developing Europe & Central Asia)
    eca = {
        'ALB', 'ARM', 'AZE', 'BLR', 'BIH', 'BGR', 'HRV', 'CZE', 'EST',
        'GEO', 'HUN', 'KAZ', 'XKX', 'KGZ', 'LVA', 'LTU', 'MDA', 'MNE',
        'MKD', 'POL', 'ROU', 'RUS', 'SRB', 'SVK', 'SVN', 'TJK', 'TUR',
        'TKM', 'UKR', 'UZB',
    }

    eca_cal = [row for row in cal if row["iso3"] in eca]
    non_eca_cal = [row for row in cal if row["iso3"] not in eca]
    n_eca = len(eca_cal)
    n_total = len(cal)

    # Reliability index ξ_j ∈ (0, 1]
    xi = {}
    with open(DATA / "reliability_index.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            xi[row["iso3"]] = float(row["xi_reliability"])

    _reg_init = {"full import": 0, "import training + build inference": 0,
                 "full domestic": 0, "build training + import inference": 0}
    all_reg = dict(_reg_init)
    all_sov = dict(_reg_init)
    for row in cal:
        iso = row["iso3"]
        if iso in reg:
            rr = reg[iso]["regime"]
            rs = reg[iso]["regime_with_sovereignty"]
            if rr in all_reg:
                all_reg[rr] += 1
            if rs in all_sov:
                all_sov[rs] += 1

    print(f"  Total: {n_total}, ECA: {n_eca}")
    print(f"  All regimes: {dict((k, v) for k, v in all_reg.items() if v)}")

    # ═══════════════════════════════════════════════════════════════════════
    # DEMAND CALIBRATION (MW-capacity-based shares)
    # ═══════════════════════════════════════════════════════════════════════
    print("Loading data center capacity estimates...")
    dc_counts = {}
    dc_capacity = {}
    dc_sources = {}
    with open(DATA / "dc_capacity_estimates.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            dc_counts[row["iso3"]] = int(row["n_datacenters"])
            dc_capacity[row["iso3"]] = float(row["capacity_mw"])
            dc_sources[row["iso3"]] = row["source"]

    # Capacity for each calibration country
    dc_k = {}
    for row in cal:
        iso = row["iso3"]
        if iso in dc_capacity:
            dc_k[iso] = dc_capacity[iso]
        else:
            dc_k[iso] = 5.0  # minimum 5 MW for countries with no data
    total_dc = sum(dc_k.values())
    omega = {iso: d / total_dc for iso, d in dc_k.items()}

    # Top demand centers
    sorted_omega = sorted(omega.items(), key=lambda x: -x[1])
    top5_labels = []
    for iso, w in sorted_omega[:5]:
        co = next(r["country"] for r in cal if r["iso3"] == iso)
        top5_labels.append((iso, co, w))
    top5_share = sum(w for _, _, w in top5_labels)

    # Training export revenue shares
    train_revenue = {}
    for iso in dc_k:
        if iso in reg:
            src = reg[iso]["best_train_source"]
            train_revenue[src] = train_revenue.get(src, 0) + omega[iso]

    # Inference export revenue shares
    inf_revenue = {}
    for iso in dc_k:
        if iso in reg:
            src = reg[iso]["best_inf_source"]
            inf_revenue[src] = inf_revenue.get(src, 0) + omega[iso]

    # HHI
    hhi_t = sum(s**2 for s in train_revenue.values())
    hhi_i = sum(s**2 for s in inf_revenue.values())

    # Lambda* for each country
    costs_dict = {row["iso3"]: float(row["c_j_total"]) + ETA for row in cal}
    lambda_star = {}
    for iso, c_k in costs_dict.items():
        min_foreign = min(c for i, c in costs_dict.items() if i != iso)
        lambda_star[iso] = c_k / min_foreign - 1

    # Welfare cost of sovereignty
    welfare_train = 0
    welfare_inf = 0
    for iso in dc_k:
        if iso in reg and iso in costs_dict:
            c_k = costs_dict[iso]
            best_train = float(reg[iso]["best_train_cost"])
            best_inf = float(reg[iso]["best_inf_cost"])
            c_k_inf = float(reg[iso]["P_I_domestic"])
            welfare_train += omega[iso] * max(0, c_k - best_train)
            welfare_inf += omega[iso] * max(0, c_k_inf - best_inf)
    welfare_total = welfare_train + welfare_inf
    weighted_avg_cost = sum(omega[iso] * costs_dict[iso]
                            for iso in dc_k if iso in costs_dict)
    welfare_pct = welfare_total / weighted_avg_cost * 100

    # Counterfactual: doubling sovereignty to 20%
    min_cost = min(costs_dict.values())
    count_dom_10 = sum(
        1 for iso in dc_k
        if iso in costs_dict and costs_dict[iso] <= 1.10 * min_cost)
    count_dom_20 = sum(
        1 for iso in dc_k
        if iso in costs_dict and costs_dict[iso] <= 1.20 * min_cost)
    extra_dom = count_dom_20 - count_dom_10
    export_share_10 = sum(
        omega[iso] for iso in dc_k
        if iso in costs_dict and costs_dict[iso] > 1.10 * min_cost)
    export_share_20 = sum(
        omega[iso] for iso in dc_k
        if iso in costs_dict and costs_dict[iso] > 1.20 * min_cost)

    sanctioned = {'IRN'}

    # Kyrgyzstan inference clients
    kgz_inf_clients = []
    for iso in dc_k:
        if iso in reg and reg[iso]["best_inf_source"] == "KGZ":
            co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
            kgz_inf_clients.append((iso, co, omega[iso] * 100))

    # Build demand_data dict for passing to write functions
    demand_data = {
        "omega": omega, "sorted_omega": sorted_omega,
        "top5_labels": top5_labels, "top5_share": top5_share,
        "train_revenue": train_revenue, "inf_revenue": inf_revenue,
        "hhi_t": hhi_t, "hhi_i": hhi_i,
        "lambda_star": lambda_star, "costs_dict": costs_dict,
        "welfare_total": welfare_total, "welfare_pct": welfare_pct,
        "welfare_train": welfare_train, "welfare_inf": welfare_inf,
        "weighted_avg_cost": weighted_avg_cost,
        "count_dom_10": count_dom_10, "count_dom_20": count_dom_20,
        "extra_dom": extra_dom,
        "export_share_10": export_share_10, "export_share_20": export_share_20,
        "kgz_inf_clients": kgz_inf_clients,
        "dc_k": dc_k, "dc_counts": dc_counts, "dc_sources": dc_sources,
    }
    print(f"  DC data for {len(dc_k)} countries, HHI_T={hhi_t:.4f}, HHI_I={hhi_i:.4f}")

    # ═══════════════════════════════════════════════════════════════════════
    # CAPACITY-CONSTRAINED EQUILIBRIUM
    # ═══════════════════════════════════════════════════════════════════════
    print("Computing capacity-constrained equilibrium...")

    # Load grid capacity data (apply scale correction)
    k_bar = {}
    with open(DATA / "grid_capacity_estimates.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            k_bar[row["iso3"]] = float(row["K_bar_gpu_hours"]) * K_BAR_SCALE

    # Training supply stack: rank countries by c_j, compute cumulative capacity
    supply_stack = sorted(
        [(iso, costs_dict[iso], k_bar.get(iso, 1e12))
         for iso in costs_dict if iso in k_bar],
        key=lambda x: x[1]
    )

    def solve_capacity_equilibrium(lam, label):
        """Solve for capacity-constrained training equilibrium at given lambda."""
        p_T = supply_stack[0][1]  # start with cheapest
        m_T = 0
        Q_TX = 0
        for _ in range(30):
            Q_TX = 0
            for iso in dc_k:
                if iso in costs_dict:
                    c_k = costs_dict[iso]
                    if c_k > (1 + lam) * p_T:
                        Q_TX += ALPHA * omega.get(iso, 0) * Q_TOTAL
            cum_cap = 0
            found = False
            for idx, (iso_j, c_j, k_j) in enumerate(supply_stack):
                if iso_j in sanctioned:
                    continue
                cap_available = k_j * ALPHA
                cum_cap += cap_available
                if cum_cap >= Q_TX and Q_TX > 0:
                    p_T_new = c_j
                    m_T = idx
                    found = True
                    break
            if found and abs(p_T_new - p_T) < 0.0001:
                p_T = p_T_new
                break
            if found:
                p_T = p_T_new

        # Compute shares
        shares = {}
        remaining = Q_TX
        for iso_j, c_j, k_j in supply_stack:
            if iso_j in sanctioned:
                continue
            if c_j > p_T:
                break
            ca = min(k_j * ALPHA, remaining)
            if ca > 0:
                shares[iso_j] = ca
                remaining -= ca
            if remaining <= 0:
                break
        total_exp = sum(shares.values())
        hhi = sum((s / total_exp) ** 2 for s in shares.values()) if total_exp > 0 else 1.0
        # Shadow values
        mu = {}
        for iso_j, c_j, k_j in supply_stack:
            if iso_j in sanctioned:
                continue
            if c_j < p_T:
                allocated = shares.get(iso_j, 0)
                if allocated >= k_j * ALPHA * 0.99:
                    mu[iso_j] = p_T - c_j
        # Lambda_star under capacity constraints
        ls_cap = {iso: c_k / p_T - 1 for iso, c_k in costs_dict.items()}

        print(f"  [{label}] p_T = ${p_T:.3f}/hr, {len(shares)} exporters, "
              f"HHI_T = {hhi:.4f}, {len(mu)} constrained")
        for iso_m, mu_v in sorted(mu.items(), key=lambda x: -x[1])[:5]:
            co = next((r["country"] for r in cal if r["iso3"] == iso_m), iso_m)
            print(f"    {co}: \u03bc = ${mu_v:.3f}/hr")
        return p_T, m_T, shares, hhi, mu, ls_cap, len(shares)

    # Pass 1: pure cost minimization (lambda=0) — main capacity result
    (p_T_0, _, _, cap_hhi_0, mu_0, ls_0, n_exp_0
     ) = solve_capacity_equilibrium(0.0, "\u03bb=0")

    # Pass 2: with sovereignty (lambda=LAMBDA)
    (p_T_sov, _, _, cap_hhi_sov, _, _, n_exp_sov
     ) = solve_capacity_equilibrium(LAMBDA, f"\u03bb={LAMBDA}")

    # Store both sets of results
    demand_data["p_T"] = p_T_0                       # pure-cost training price
    demand_data["p_T_sov"] = p_T_sov                 # sovereignty training price
    demand_data["cap_hhi_t"] = cap_hhi_0             # pure-cost HHI
    demand_data["cap_hhi_t_sov"] = cap_hhi_sov       # sovereignty HHI
    demand_data["n_train_exporters"] = n_exp_0
    demand_data["n_train_exporters_sov"] = n_exp_sov
    demand_data["mu_j"] = mu_0

    # ═══════════════════════════════════════════════════════════════════════
    # COST-RECOVERY ADJUSTMENT (PREFERRED BASELINE)
    # ═══════════════════════════════════════════════════════════════════════
    print("Computing cost-recovery adjustment...")

    # Load latency data for inference recomputation
    latency_data = {}
    with open(DATA / "country_pair_latency.csv", encoding="utf-8") as f:
        for lrow in csv.DictReader(f):
            latency_data[(lrow["iso3_from"], lrow["iso3_to"])] = float(lrow["avg_ms"])
    DOMESTIC_LATENCY_DEFAULT = 5.0

    def _get_latency(j, k):
        if j == k:
            return latency_data.get((j, k), DOMESTIC_LATENCY_DEFAULT)
        if (j, k) in latency_data:
            return latency_data[(j, k)]
        if (k, j) in latency_data:
            return latency_data[(k, j)]
        return None

    adj_costs = dict(costs_dict)  # copy baseline
    adj_changes = {}
    for iso, p_E_adj in SUBSIDY_ADJ.items():
        if iso not in adj_costs:
            continue
        row = next(r for r in cal if r["iso3"] == iso)
        p_E_orig = float(row["p_E_usd_kwh"])
        pue = float(row["pue"])
        delta_elec = pue * GAMMA * (p_E_adj - p_E_orig)
        adj_costs[iso] = costs_dict[iso] + delta_elec
        subsidy_gap = p_E_adj - p_E_orig  # $/kWh gap
        # Fiscal transfer for a hypothetical 100 MW data center ($/year)
        fiscal_transfer_100mw = subsidy_gap * 1000 * 100 * H_YR  # kWh/yr * $/kWh
        adj_changes[iso] = {
            "country": row["country"],
            "p_E_orig": p_E_orig, "p_E_adj": p_E_adj,
            "c_j_orig": costs_dict[iso], "c_j_adj": adj_costs[iso],
            "subsidy_gap": subsidy_gap,
            "fiscal_transfer_100mw": fiscal_transfer_100mw,
        }

    adj_ranked = sorted(adj_costs.items(), key=lambda x: x[1])
    adj_rank_map = {iso: rank for rank, (iso, _) in enumerate(adj_ranked, 1)}

    # Count regime changes under adjusted costs
    regime_changes = 0
    for iso_k in dc_k:
        if iso_k not in reg:
            continue
        orig_regime = reg[iso_k]["regime"]
        # Recompute regime under adjusted costs
        c_k_adj = adj_costs.get(iso_k, costs_dict.get(iso_k))
        if c_k_adj is None:
            continue
        adj_cheapest_train = adj_ranked[0][0]
        adj_train_cost = adj_costs[adj_cheapest_train]
        is_dom_train = (adj_train_cost >= c_k_adj)
        # Inference: find best source under adjusted costs
        l_kk = _get_latency(iso_k, iso_k)
        P_I_dom = (1 + TAU * l_kk) * c_k_adj if l_kk is not None else c_k_adj
        best_inf_adj = P_I_dom
        best_inf_src = iso_k
        for j, c_j in adj_costs.items():
            if j == iso_k:
                continue
            l_jk = _get_latency(j, iso_k)
            if l_jk is None:
                continue
            cost = (1 + TAU * l_jk) * c_j
            if cost < best_inf_adj:
                best_inf_adj = cost
                best_inf_src = j
        is_dom_inf = (best_inf_src == iso_k)
        if is_dom_train and is_dom_inf:
            new_regime = "full domestic"
        elif is_dom_train:
            new_regime = "build training + import inference"
        elif is_dom_inf:
            new_regime = "import training + build inference"
        else:
            new_regime = "full import"
        if new_regime != orig_regime:
            regime_changes += 1

    adj_cheapest = adj_ranked[0][0]
    # Top 5 adjusted ranking
    adj_top5 = []
    for iso, c in adj_ranked[:5]:
        co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
        adj_top5.append((iso, co, c))

    # Subsidy gap statistics
    gaps = [v["subsidy_gap"] * 1000 for v in adj_changes.values()]  # $/MWh
    max_gap_iso = max(adj_changes, key=lambda x: adj_changes[x]["subsidy_gap"])
    max_gap_entry = adj_changes[max_gap_iso]

    demand_data["adj_top5"] = adj_top5
    demand_data["adj_cheapest_name"] = next(
        (r["country"] for r in cal if r["iso3"] == adj_cheapest), adj_cheapest)
    demand_data["adj_rank_map"] = adj_rank_map
    demand_data["adj_costs"] = adj_costs
    demand_data["n_adjusted"] = len(adj_changes)
    demand_data["regime_changes"] = regime_changes
    demand_data["max_gap_country"] = max_gap_entry["country"]
    demand_data["max_fiscal_transfer"] = max_gap_entry["fiscal_transfer_100mw"]
    demand_data["min_gap_mwh"] = min(gaps)
    demand_data["max_gap_mwh_val"] = max(gaps)

    print(f"  Adjusted {len(adj_changes)} countries; new cheapest: "
          f"{demand_data['adj_cheapest_name']} (${adj_costs[adj_cheapest]:.3f}/hr)")
    print(f"  Regime changes: {regime_changes}")
    print(f"  Subsidy gap range: ${min(gaps):.0f}\u2013${max(gaps):.0f}/MWh")
    print(f"  Max fiscal transfer (100 MW): {max_gap_entry['country']} "
          f"${max_gap_entry['fiscal_transfer_100mw'] / 1e6:.0f}M/yr")
    for iso, co, c in adj_top5:
        flag = " *" if iso in adj_changes else ""
        print(f"    {adj_rank_map[iso]:>2}. {co:<24} ${c:.3f}/hr{flag}")

    # ═══════════════════════════════════════════════════════════════════════
    # RE-COMPUTE EQUILIBRIUM ON COST-RECOVERY BASELINE
    # ═══════════════════════════════════════════════════════════════════════
    print("Re-computing equilibrium on cost-recovery baseline...")

    # Build cost-recovery supply stack and override closure variables
    adj_supply_stack = sorted(
        [(iso, adj_costs[iso], k_bar.get(iso, 1e12))
         for iso in adj_costs if iso in k_bar],
        key=lambda x: x[1]
    )
    supply_stack = adj_supply_stack  # noqa: F841
    costs_dict = adj_costs

    # Re-run capacity equilibrium on cost-recovery costs
    (p_T_0, _, _, cap_hhi_0, mu_0, ls_0, n_exp_0
     ) = solve_capacity_equilibrium(0.0, "\u03bb=0 cost-recovery")
    (p_T_sov, _, _, cap_hhi_sov, _, _, n_exp_sov
     ) = solve_capacity_equilibrium(LAMBDA, f"\u03bb={LAMBDA} cost-recovery")

    demand_data["p_T"] = p_T_0
    demand_data["p_T_sov"] = p_T_sov
    demand_data["cap_hhi_t"] = cap_hhi_0
    demand_data["cap_hhi_t_sov"] = cap_hhi_sov
    demand_data["n_train_exporters"] = n_exp_0
    demand_data["n_train_exporters_sov"] = n_exp_sov
    demand_data["mu_j"] = mu_0
    demand_data["lambda_star"] = ls_0

    # Recompute inference sourcing under cost-recovery costs
    adj_reg = {}
    for iso_k in dc_k:
        c_k = adj_costs.get(iso_k)
        if c_k is None:
            continue
        l_kk = _get_latency(iso_k, iso_k)
        P_I_dom = (1 + TAU * (l_kk or 0)) * c_k
        best_inf_cost = P_I_dom
        best_inf_src = iso_k
        for iso_j, c_j in adj_costs.items():
            if iso_j == iso_k:
                continue
            l_jk = _get_latency(iso_j, iso_k)
            if l_jk is None:
                continue
            cost_del = (1 + TAU * l_jk) * c_j
            if cost_del < best_inf_cost:
                best_inf_cost = cost_del
                best_inf_src = iso_j
        adj_reg[iso_k] = {
            'best_inf_source': best_inf_src,
            'best_inf_cost': f'{best_inf_cost:.4f}',
            'P_I_domestic': f'{P_I_dom:.4f}',
        }

    # Recompute inference revenue shares
    adj_inf_revenue = {}
    for iso in dc_k:
        if iso in adj_reg:
            src = adj_reg[iso]['best_inf_source']
            adj_inf_revenue[src] = adj_inf_revenue.get(src, 0) + omega.get(iso, 0)
    adj_hhi_i = sum(s**2 for s in adj_inf_revenue.values())
    demand_data["inf_revenue"] = adj_inf_revenue
    demand_data["hhi_i"] = adj_hhi_i

    # Recompute welfare
    adj_welfare_train = 0
    adj_welfare_inf = 0
    for iso in dc_k:
        if iso in adj_reg and iso in adj_costs:
            c_k = adj_costs[iso]
            min_foreign = min(
                c for i, c in adj_costs.items()
                if i != iso and i not in sanctioned)
            adj_welfare_train += omega.get(iso, 0) * max(0, c_k - min_foreign)
            best_inf = float(adj_reg[iso]["best_inf_cost"])
            P_I_dom = float(adj_reg[iso]["P_I_domestic"])
            adj_welfare_inf += omega.get(iso, 0) * max(0, P_I_dom - best_inf)
    adj_welfare_total = adj_welfare_train + adj_welfare_inf
    adj_weighted_avg = sum(
        omega.get(iso, 0) * adj_costs[iso]
        for iso in dc_k if iso in adj_costs)
    adj_welfare_pct = (adj_welfare_total / adj_weighted_avg * 100
                       if adj_weighted_avg > 0 else 0)
    demand_data["welfare_total"] = adj_welfare_total
    demand_data["welfare_pct"] = adj_welfare_pct
    demand_data["welfare_train"] = adj_welfare_train
    demand_data["welfare_inf"] = adj_welfare_inf
    demand_data["weighted_avg_cost"] = adj_weighted_avg

    # Recompute counterfactual
    adj_min_cost = min(adj_costs.values())
    adj_count_dom_10 = sum(
        1 for iso in dc_k
        if iso in adj_costs and adj_costs[iso] <= 1.10 * adj_min_cost)
    adj_count_dom_20 = sum(
        1 for iso in dc_k
        if iso in adj_costs and adj_costs[iso] <= 1.20 * adj_min_cost)
    demand_data["extra_dom"] = adj_count_dom_20 - adj_count_dom_10
    demand_data["export_share_10"] = sum(
        omega.get(iso, 0) for iso in dc_k
        if iso in adj_costs and adj_costs[iso] > 1.10 * adj_min_cost)
    demand_data["export_share_20"] = sum(
        omega.get(iso, 0) for iso in dc_k
        if iso in adj_costs and adj_costs[iso] > 1.20 * adj_min_cost)

    # Recompute KGZ inference clients
    adj_kgz_clients = []
    for iso in dc_k:
        if iso in adj_reg and adj_reg[iso]["best_inf_source"] == "KGZ":
            co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
            adj_kgz_clients.append((iso, co, omega.get(iso, 0) * 100))
    demand_data["kgz_inf_clients"] = adj_kgz_clients

    # Store adj_reg and adj_costs for write functions
    demand_data["adj_reg"] = adj_reg
    demand_data["costs_dict"] = adj_costs

    # Print summary
    print(f"  Cost-recovery inference HHI_I = {adj_hhi_i:.4f}")
    adj_inf_top5 = sorted(adj_inf_revenue.items(), key=lambda x: -x[1])[:5]
    for iso, share in adj_inf_top5:
        co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
        print(f"    {co}: {share * 100:.1f}%")

    # ═══════════════════════════════════════════════════════════════════════
    # SENSITIVITY ANALYSIS
    # ═══════════════════════════════════════════════════════════════════════
    print("\nRunning sensitivity analysis...")
    sens_results = run_sensitivity(cal, omega, dc_k, k_bar, sanctioned)
    demand_data["sensitivity"] = sens_results

    # ═══════════════════════════════════════════════════════════════════════
    # RELIABILITY-ADJUSTED COST RANKINGS
    # ═══════════════════════════════════════════════════════════════════════
    print("Computing reliability-adjusted rankings...")
    # Use cost-recovery adjusted costs (preferred baseline)
    xi_costs = {}
    for row in cal:
        iso = row["iso3"]
        c_j = float(row["c_j_total"]) + ETA
        xi_j = xi.get(iso, 1.0)
        # Apply cost-recovery adjustment if applicable
        if iso in SUBSIDY_ADJ:
            p_E_cr = SUBSIDY_ADJ[iso]
            p_E_raw = float(row["p_E_usd_kwh"])
            c_j = c_j + (p_E_cr - p_E_raw) * float(row["pue"]) * GAMMA
        xi_costs[iso] = c_j / xi_j

    # Baseline (no xi adjustment) costs for comparison
    baseline_costs = {}
    for row in cal:
        iso = row["iso3"]
        c_j = float(row["c_j_total"]) + ETA
        if iso in SUBSIDY_ADJ:
            p_E_cr = SUBSIDY_ADJ[iso]
            p_E_raw = float(row["p_E_usd_kwh"])
            c_j = c_j + (p_E_cr - p_E_raw) * float(row["pue"]) * GAMMA
        baseline_costs[iso] = c_j

    # Rank both
    baseline_rank = sorted(baseline_costs.items(), key=lambda x: x[1])
    xi_rank = sorted(xi_costs.items(), key=lambda x: x[1])
    baseline_order = [iso for iso, _ in baseline_rank]
    xi_order = [iso for iso, _ in xi_rank]

    # Top 5 with names
    xi_top5 = []
    for iso, cost in xi_rank[:5]:
        co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
        xi_top5.append((co, cost))

    # Spearman rank correlation
    n_r = len(baseline_order)
    rank_base = {iso: i for i, iso in enumerate(baseline_order)}
    rank_xi = {iso: i for i, iso in enumerate(xi_order)}
    d_sq = sum((rank_base[iso] - rank_xi[iso]) ** 2
               for iso in baseline_order if iso in rank_xi)
    spearman = 1 - 6 * d_sq / (n_r * (n_r ** 2 - 1))

    # Count how many top-10 baseline producers fall out of top-10
    base_top10 = set(baseline_order[:10])
    xi_top10 = set(xi_order[:10])
    n_changed_top10 = len(base_top10 - xi_top10)

    demand_data["xi_adjusted"] = {
        "top5": xi_top5,
        "rank_corr": spearman,
        "n_changed_top10": n_changed_top10,
        "xi_order": xi_order[:10],
        "baseline_rank_map": {iso: i for i, iso in enumerate(baseline_order)},
        "xi_rank_map": {iso: i for i, iso in enumerate(xi_order)},
    }
    demand_data["xi"] = xi
    # Country name map for figure labels
    demand_data["iso_country"] = {r["iso3"]: r["country"] for r in cal}
    print(f"  Reliability-adjusted top 5: {[f'{co} (${c:.2f})' for co, c in xi_top5]}")
    print(f"  Spearman rank corr: {spearman:.4f}, top-10 changes: {n_changed_top10}")

    # ═══════════════════════════════════════════════════════════════════════
    # LOAD v8 AND INDEX HEADINGS
    # ═══════════════════════════════════════════════════════════════════════

    print("\nLoading v8...")
    doc = Document(str(DOCS / "flop_trade_model_v8.docx"))
    body = doc.element.body
    all_el = list(body)
    init_footnotes(doc)

    hmap = {}
    for el in all_el:
        if el.tag == qn('w:p'):
            pPr = el.find(qn('w:pPr'))
            if pPr is not None:
                pS = pPr.find(qn('w:pStyle'))
                if pS is not None and 'Heading' in pS.get(qn('w:val'), ''):
                    ft = "".join(r.text or "" for r in el.findall(f'.//{qn("w:t")}'))
                    if '1.2' in ft:
                        hmap['1.2'] = el
                    elif '1.1' in ft:
                        hmap['1.1'] = el
                    elif '1.' in ft and 'Model' in ft:
                        hmap['1'] = el
                    elif '2.' in ft and 'Comp' in ft:
                        hmap['2'] = el
                    elif '3.' in ft and 'Make' in ft:
                        hmap['3'] = el
                    elif '4.' in ft and 'Calib' in ft:
                        hmap['4'] = el
                    elif '5.' in ft and 'Conc' in ft:
                        hmap['5'] = el
                    elif ft.strip() == 'References':
                        hmap['refs'] = el
                    elif ft.strip() == 'Abstract':
                        hmap['abs'] = el

    # ═══════════════════════════════════════════════════════════════════════
    # STEPS
    # ═══════════════════════════════════════════════════════════════════════

    title_el, author_el, ver_el, abs_text_el, kw_el = write_title_and_abstract(doc, body, all_el, hmap)
    write_introduction(doc, body, hmap)
    write_literature(doc, body, hmap)
    write_production_technology(doc, body, hmap)
    write_trade_costs(doc, body, hmap)
    renumber_sections(hmap)
    # New Section 3 subsections (3.3, 3.4) — inserted before the renumbered Section 4 heading
    write_demand(doc, body, hmap, demand_data)
    write_sourcing_and_equilibrium(doc, body, hmap, demand_data)
    # Section 4: Equilibrium Properties (replaces old Comparative Advantage + Make-or-Buy)
    write_equilibrium_properties(doc, body, hmap, demand_data)
    write_data_section(doc, body, hmap, demand_data)
    write_calibration(doc, body, hmap, cal, reg, n_eca, n_total, all_reg, all_sov, demand_data)
    write_conclusion(doc, body, hmap, demand_data)

    refs = hmap['refs']
    last_ref = write_references(doc, body, refs)
    last_fig4b = write_figure4b(doc, body, last_ref, demand_data)
    last_table1 = write_table1(doc, body, last_fig4b, demand_data)
    last_app_note = write_appendix(doc, body, last_table1, eca_cal, non_eca_cal, reg, demand_data)
    last_model_app = write_model_appendix(doc, body, last_app_note)
    last_sens_app = write_sensitivity_appendix(doc, body, last_model_app, demand_data)
    last_dcf_app = write_kyrgyzstan_appendix(doc, body, last_sens_app)
    write_construction_regression_appendix(doc, body, last_dcf_app)
    link_citations(body)
    link_equations(body)
    fix_orphan_backlinks(body, refs)
    apply_formatting(doc, body, refs, title_el, author_el, ver_el, abs_text_el)
    add_page_numbers_and_break(doc, body, kw_el)

    # ═══════════════════════════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════════════════════════
    flush_footnotes()
    doc.core_properties.author = 'Michael Lokshin'
    out = DOCS / "flop_trade_model_v21.docx"
    for _attempt in range(30):
        try:
            doc.save(str(out))
            break
        except PermissionError:
            if _attempt == 0:
                print(f"\nFile locked — waiting for Word to release {out.name}...")
            import time
            time.sleep(2)
    else:
        raise PermissionError(f"Could not save {out} after 60 seconds. Close Word and retry.")
    print(f"\nSaved {out}")


if __name__ == '__main__':
    main()
