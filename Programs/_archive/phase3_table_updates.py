"""Phase 3: Update table values with new xi and costs (no floor)."""
import csv
import copy
import os
import zipfile
from decimal import Decimal, ROUND_HALF_UP
from lxml import etree
from docx import Document

DOCS = r"F:\onedrive\__documents\papers\FLOPsExport\Documents"
DATA = r"F:\onedrive\__documents\papers\FLOPsExport\Data"
SRC = os.path.join(DOCS, "flop_trade_model_v29_phase2.docx")
OUT = os.path.join(DOCS, "flop_trade_model_v29.docx")

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

HIGH_INCOME = {
    'Norway', 'Finland', 'Canada', 'Sweden', 'Iceland', 'New Zealand',
    'United States of America', 'Australia', 'United Kingdom', 'Denmark',
    'France', 'Luxembourg', 'Netherlands', 'Belgium', 'Estonia', 'Latvia',
    'Switzerland', 'Portugal', 'Japan', 'Germany', 'Austria', 'Ireland',
    'Singapore', 'South Korea', 'Czechia', 'Czech Republic',
    'Slovenia', 'Lithuania', 'Spain', 'Italy', 'Greece', 'Poland',
    'Hungary', 'Croatia', 'Slovakia', 'Chile', 'Uruguay', 'Romania',
    'Bulgaria', 'United Arab Emirates', 'Saudi Arabia', 'Qatar', 'Bahrain',
    'Kuwait', 'Oman', 'Brunei',
}


def rhup(value, dp=2):
    """Round half-up."""
    return float(Decimal(str(value)).quantize(Decimal(10) ** -dp,
                                              rounding=ROUND_HALF_UP))


# ── Load Phase 1 results ──
print("Loading country_results_v29.csv...")
results = {}
with open(os.path.join(DATA, "country_results_v29.csv"), encoding="utf-8") as f:
    for row in csv.DictReader(f):
        results[row["country"]] = {
            "xi_new": float(row["xi_new"]),
            "c_adj_new": float(row["c_adj_new"]),
            "rank_new": int(row["rank_new"]),
            "c_cr": float(row["c_cr"]),
            "rank_cr": int(row["rank_cr"]),
        }

# Short name aliases
NAME_ALIASES = {
    "United States of A.": "United States of America",
    "United States of Am": "United States of America",
    "United Arab Emirate": "United Arab Emirates",
    "Bosnia and Herz.": "Bosnia and Herz.",
    "North Macedonia": "North Macedonia",
    "South Korea": "South Korea",
}


def lookup(country_name):
    """Find country in results, trying aliases."""
    if country_name in results:
        return results[country_name]
    # Try alias
    if country_name in NAME_ALIASES:
        return results.get(NAME_ALIASES[country_name])
    # Try partial match
    for name in results:
        if name.startswith(country_name[:15]) or country_name.startswith(name[:15]):
            return results[name]
    return None


# ── Load Table A3 results ──
print("Loading tableA3_v29.csv...")
a3_data = {}
with open(os.path.join(DATA, "tableA3_v29.csv"), encoding="utf-8") as f:
    for row in csv.DictReader(f):
        a3_data[row["Scenario"]] = row

# ── Open Phase 2 docx ──
print(f"Opening {SRC}...")
doc = Document(SRC)

# ═══════════════════════════════════════════════════════════════════════
# TABLE A1 (index 11): 85 countries, update cols 7 (ξ) and 8 (c_j)
# ═══════════════════════════════════════════════════════════════════════
print("\n=== Table A1 (index 11) ===")
tbl_a1 = doc.tables[11]
a1_rows = tbl_a1.rows
print(f"  Rows: {len(a1_rows)} (header + {len(a1_rows)-1} data)")

# Update values
updated_a1 = 0
a1_data_for_sort = []  # (c_adj_new, row_xml_element)
for ri in range(1, len(a1_rows)):
    cells = a1_rows[ri].cells
    country = cells[0].text.strip()
    r = lookup(country)
    if r is None:
        print(f"  WARNING: No data for '{country}'")
        continue

    # Col 7: ξ_new
    xi_str = f"{rhup(r['xi_new']):.2f}"
    # Col 8: c_adj_new
    cost_str = f"${rhup(r['c_adj_new']):.2f}"

    # Preserve formatting: get existing run properties
    for col_idx, new_val in [(7, xi_str), (8, cost_str)]:
        cell = cells[col_idx]
        para = cell.paragraphs[0]
        runs = para.runs
        if runs:
            # Copy rPr from first run
            old_rpr = copy.deepcopy(runs[0]._r.find(f'{{{W_NS}}}rPr'))
            # Clear all runs
            for run in runs:
                run.text = ""
            runs[0].text = new_val
            # Restore formatting
            if old_rpr is not None:
                new_rpr = runs[0]._r.find(f'{{{W_NS}}}rPr')
                if new_rpr is not None:
                    runs[0]._r.remove(new_rpr)
                runs[0]._r.insert(0, old_rpr)
        else:
            # No existing runs, create one
            para.text = new_val

    a1_data_for_sort.append((r['c_adj_new'], a1_rows[ri]._tr))
    updated_a1 += 1

print(f"  Updated {updated_a1} rows")

# Re-sort by c_adj ascending
print("  Sorting by c_adj ascending...")
tbl_el = tbl_a1._tbl
header_tr = tbl_el.findall(f'{{{W_NS}}}tr')[0]

# Remove all data rows
for _, tr in a1_data_for_sort:
    tbl_el.remove(tr)

# Re-add sorted
a1_data_for_sort.sort(key=lambda x: x[0])
for _, tr in a1_data_for_sort:
    tbl_el.append(tr)

print(f"  Sorted {len(a1_data_for_sort)} data rows")

# Spot check
print("  Spot checks:")
for ri in range(1, min(4, len(tbl_a1.rows))):
    cells = tbl_a1.rows[ri].cells
    print(f"    Row {ri}: {cells[0].text.strip():20s}  xi={cells[7].text.strip()}  c={cells[8].text.strip()}")
# Last row
lr = len(tbl_a1.rows) - 1
cells = tbl_a1.rows[lr].cells
print(f"    Row {lr}: {cells[0].text.strip():20s}  xi={cells[7].text.strip()}  c={cells[8].text.strip()}")


# ═══════════════════════════════════════════════════════════════════════
# TABLE 3a (index 9): ~22 countries, update cols 7-10
# ═══════════════════════════════════════════════════════════════════════
print("\n=== Table 3a (index 9) ===")
tbl_3a = doc.tables[9]
print(f"  Rows: {len(tbl_3a.rows)}")

updated_3a = 0
for ri in range(2, len(tbl_3a.rows)):  # Skip header rows (0 and 1)
    cells = tbl_3a.rows[ri].cells
    country = cells[0].text.strip()
    if not country or country == "Country":
        continue
    r = lookup(country)
    if r is None:
        print(f"  WARNING: No data for '{country}' in row {ri}")
        continue

    dev_type = "D" if country not in HIGH_INCOME else "A"

    updates = [
        (7, f"${rhup(r['c_adj_new']):.2f}"),
        (8, f"{rhup(r['xi_new']):.2f}"),
        (9, str(r['rank_new'])),
        (10, dev_type),
    ]

    for col_idx, new_val in updates:
        if col_idx < len(cells):
            cell = cells[col_idx]
            para = cell.paragraphs[0]
            runs = para.runs
            if runs:
                old_rpr = copy.deepcopy(runs[0]._r.find(f'{{{W_NS}}}rPr'))
                for run in runs:
                    run.text = ""
                runs[0].text = new_val
                if old_rpr is not None:
                    new_rpr = runs[0]._r.find(f'{{{W_NS}}}rPr')
                    if new_rpr is not None:
                        runs[0]._r.remove(new_rpr)
                    runs[0]._r.insert(0, old_rpr)
            else:
                para.text = new_val
    updated_3a += 1

print(f"  Updated {updated_3a} rows")
# Spot check
for ri in [2, 3, 4]:
    cells = tbl_3a.rows[ri].cells
    print(f"    Row {ri}: {cells[0].text.strip():20s}  c_adj={cells[7].text.strip()}  "
          f"xi={cells[8].text.strip()}  rank={cells[9].text.strip()}  type={cells[10].text.strip()}")


# ═══════════════════════════════════════════════════════════════════════
# TABLE A2 (index 12): 85 countries, same cols 7-10
# ═══════════════════════════════════════════════════════════════════════
print("\n=== Table A2 (index 12) ===")
tbl_a2 = doc.tables[12]
print(f"  Rows: {len(tbl_a2.rows)}")

updated_a2 = 0
for ri in range(2, len(tbl_a2.rows)):  # Skip 2 header rows
    cells = tbl_a2.rows[ri].cells
    country = cells[0].text.strip()
    if not country or country == "Country":
        continue
    r = lookup(country)
    if r is None:
        print(f"  WARNING: No data for '{country}' in row {ri}")
        continue

    dev_type = "D" if country not in HIGH_INCOME else "A"

    updates = [
        (7, f"${rhup(r['c_adj_new']):.2f}"),
        (8, f"{rhup(r['xi_new']):.2f}"),
        (9, str(r['rank_new'])),
        (10, dev_type),
    ]

    for col_idx, new_val in updates:
        if col_idx < len(cells):
            cell = cells[col_idx]
            para = cell.paragraphs[0]
            runs = para.runs
            if runs:
                old_rpr = copy.deepcopy(runs[0]._r.find(f'{{{W_NS}}}rPr'))
                for run in runs:
                    run.text = ""
                runs[0].text = new_val
                if old_rpr is not None:
                    new_rpr = runs[0]._r.find(f'{{{W_NS}}}rPr')
                    if new_rpr is not None:
                        runs[0]._r.remove(new_rpr)
                    runs[0]._r.insert(0, old_rpr)
            else:
                para.text = new_val
    updated_a2 += 1

print(f"  Updated {updated_a2} rows")


# ═══════════════════════════════════════════════════════════════════════
# TABLE A3 (index 18): 5 data rows, update cols 2-5
# ═══════════════════════════════════════════════════════════════════════
print("\n=== Table A3 (index 18) ===")
tbl_a3 = doc.tables[18]
print(f"  Rows: {len(tbl_a3.rows)}")

# Map scenario names to CSV data
SCENARIO_MAP = {
    "Baseline": "Baseline (omega=0.50, rho=$1.36, Form B)",
    "Form A": "Form A (v26 specification)",
    "High governance": "High governance weight (omega=0.85)",
    "Low hardware": "Low hardware share (rho=$1.30)",
    "High hardware": "High hardware share (rho=$1.42)",
}

updated_a3 = 0
for ri in range(1, len(tbl_a3.rows)):
    cells = tbl_a3.rows[ri].cells
    scenario = cells[0].text.strip()

    # Find matching CSV row
    csv_row = None
    for key, csv_name in SCENARIO_MAP.items():
        if key.lower() in scenario.lower():
            csv_row = a3_data.get(csv_name)
            break

    if csv_row is None:
        print(f"  WARNING: No A3 data for '{scenario}'")
        continue

    updates = [
        (2, csv_row["Dev_top15"]),
        (3, csv_row["Max_markup"]),
        (4, csv_row["Spearman_rho"]),
        (5, csv_row["Top5"]),
    ]

    for col_idx, new_val in updates:
        if col_idx < len(cells):
            cell = cells[col_idx]
            para = cell.paragraphs[0]
            runs = para.runs
            if runs:
                old_rpr = copy.deepcopy(runs[0]._r.find(f'{{{W_NS}}}rPr'))
                for run in runs:
                    run.text = ""
                runs[0].text = new_val
                if old_rpr is not None:
                    new_rpr = runs[0]._r.find(f'{{{W_NS}}}rPr')
                    if new_rpr is not None:
                        runs[0]._r.remove(new_rpr)
                    runs[0]._r.insert(0, old_rpr)
            else:
                para.text = new_val
    updated_a3 += 1
    print(f"  Row {ri}: {scenario[:40]:40s} -> Dev={csv_row['Dev_top15']}, "
          f"Markup={csv_row['Max_markup']}, Spearman={csv_row['Spearman_rho']}")

print(f"  Updated {updated_a3} rows")


# ═══════════════════════════════════════════════════════════════════════
# SET METADATA
# ═══════════════════════════════════════════════════════════════════════
doc.core_properties.author = 'Michael Lokshin'

# ═══════════════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════════════
print(f"\nSaving {OUT}...")
doc.save(OUT)
print("Saved.")

# ═══════════════════════════════════════════════════════════════════════
# FINAL VERIFICATION
# ═══════════════════════════════════════════════════════════════════════
print("\n=== FINAL VERIFICATION ===")
doc2 = Document(OUT)

# 1. Floor check
floor_p = sum(1 for p in doc2.paragraphs if 'floor' in p.text.lower())
floor_t = 0
for t in doc2.tables:
    for row in t.rows:
        for cell in row.cells:
            if 'floor' in cell.text.lower():
                floor_t += 1
print(f"1. Floor in paragraphs: {floor_p}")
print(f"2. Floor in tables: {floor_t}")

# 3. Table A1 checks
ta1 = doc2.tables[11]
print(f"3. Table A1: {len(ta1.rows)-1} data rows")
# Spot check Norway and Kyrgyzstan
for ri in range(1, len(ta1.rows)):
    cells = ta1.rows[ri].cells
    name = cells[0].text.strip()
    if name in ("Norway", "Kyrgyzstan"):
        xi = cells[7].text.strip()
        cost = cells[8].text.strip()
        print(f"   {name}: xi={xi} c={cost}")

# 4. Table A3
ta3 = doc2.tables[18]
print(f"4. Table A3: {len(ta3.rows)} rows (expect 6)")
for ri in range(len(ta3.rows)):
    cells = ta3.rows[ri].cells
    print(f"   Row {ri}: {cells[0].text.strip()[:40]:40s} | {cells[2].text.strip():>3s} | "
          f"{cells[3].text.strip():>6s} | {cells[4].text.strip():>5s}")

# 5. Table 3a spot check
t3a = doc2.tables[9]
print(f"5. Table 3a: {len(t3a.rows)} rows")
for ri in [2, 3]:
    cells = t3a.rows[ri].cells
    print(f"   Row {ri}: {cells[0].text.strip():20s}  c_adj={cells[7].text.strip()} "
          f"xi={cells[8].text.strip()} rank={cells[9].text.strip()}")

print("\n=== Phase 3 complete ===")
