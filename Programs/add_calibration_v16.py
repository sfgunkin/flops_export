"""
Produce flop_trade_model_v18.docx from v8.docx.

v18 changes (15 accumulated comments + referee R2 Major Issues 2 & 3):
  - Sovereignty premium λ in all equations
  - Proposition 1 as buyer's optimization equation
  - Proposition 2 with country-specific F_j
  - Country taxonomy from model
  - Structure around exporters/importers
  - New Data section (Section 6)
  - Table 1: construction column filled, no c_j header, regime defined
  - 81 countries (20 new non-ECA)
  - Define $/hr costs, explain hardware amortization
  - All data source references
  - Helpman-Melitz reference
  - Major Issue 2: demand calibration (q_k = ω_k · Q via GDP shares)
  - Major Issue 3: Corollaries 1-3 (substantive comparative statics)
  - AI language cleanup
"""

import csv
import pathlib
import sys
import io
import copy
from datetime import datetime
from lxml import etree
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

DOCS = pathlib.Path(r"F:\onedrive\__documents\papers\FLOPsExport\Documents")
DATA = pathlib.Path(r"F:\onedrive\__documents\papers\FLOPsExport\Data")

TAU = 0.0008
LAMBDA = 0.10
PHI = 1.08
DELTA_PUE = 0.015
THETA_REF = 15.0
GAMMA = 0.700
GPU_PRICE = 25_000
GPU_LIFE = 3
GPU_UTIL = 0.90
DC_LIFE = 15
H_YR = 365.25 * 24
RHO = GPU_PRICE / (GPU_LIFE * H_YR * GPU_UTIL)

# Cost-reflective electricity prices for subsidy robustness check ($/kWh)
# Replacement = estimated LRMC of dominant generation at opportunity-cost fuel price
SUBSIDY_ADJ = {
    'IRN': 0.050,  # Gas CCGT at export-parity fuel cost (~$5/MMBtu)
    'TKM': 0.050,  # Gas CCGT at export-parity fuel cost
    'DZA': 0.065,  # Gas at near-export parity
    'EGY': 0.075,  # Gas/oil, Egyptian subsidy reform target
    'UZB': 0.085,  # Gas, moderate subsidy
}

# Fonts
CAMBRIA_MATH = 'Cambria Math'
TIMES_NEW_ROMAN = 'Times New Roman'

# Colors
HEADING_BLUE = RGBColor(0x2F, 0x54, 0x96)
LINK_COLOR = '1F3864'

# XML namespaces
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'
SPACE_PRESERVE = 'preserve'

# Table formatting
TABLE_WIDTH_PCT = '5000'

# ═══════════════════════════════════════════════════════════════════════
# OMML HELPERS
# ═══════════════════════════════════════════════════════════════════════


def _mr(text, italic=True):
    r = OxmlElement('m:r')
    rPr = OxmlElement('m:rPr')
    sty = OxmlElement('m:sty')
    sty.set(qn('m:val'), 'i' if italic else 'p')
    rPr.append(sty)
    r.append(rPr)
    wrPr = OxmlElement('w:rPr')
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:ascii'), CAMBRIA_MATH)
    rF.set(qn('w:hAnsi'), CAMBRIA_MATH)
    wrPr.append(rF)
    r.append(wrPr)
    t = OxmlElement('m:t')
    t.set(XML_SPACE, SPACE_PRESERVE)
    t.text = text
    r.append(t)
    return r


def _v(text):
    return _mr(text, True)


def _t(text):
    return _mr(text, False)


def _msub(base, sub, base_italic=True, sub_italic=True):
    el = OxmlElement('m:sSub')
    el.append(OxmlElement('m:sSubPr'))
    e = OxmlElement('m:e')
    e.append(_mr(base, base_italic))
    el.append(e)
    s = OxmlElement('m:sub')
    s.append(_mr(sub, sub_italic))
    el.append(s)
    return el


def _msup(base, sup, base_italic=True, sup_italic=True):
    el = OxmlElement('m:sSup')
    el.append(OxmlElement('m:sSupPr'))
    e = OxmlElement('m:e')
    e.append(_mr(base, base_italic))
    el.append(e)
    s = OxmlElement('m:sup')
    s.append(_mr(sup, sup_italic))
    el.append(s)
    return el


def _msubsup(base, sub, sup):
    """Subscript-superscript combo."""
    el = OxmlElement('m:sSubSup')
    el.append(OxmlElement('m:sSubSupPr'))
    e = OxmlElement('m:e')
    e.append(_mr(base))
    el.append(e)
    s = OxmlElement('m:sub')
    s.append(_mr(sub))
    el.append(s)
    u = OxmlElement('m:sup')
    u.append(_mr(sup))
    el.append(u)
    return el


def _nary(char, sub_parts, sup_parts, e_parts):
    """N-ary operator (summation, product) with sub/sup limits."""
    el = OxmlElement('m:nary')
    pr = OxmlElement('m:naryPr')
    ch = OxmlElement('m:chr')
    ch.set(qn('m:val'), char)
    pr.append(ch)
    if not sup_parts:
        supHide = OxmlElement('m:supHide')
        supHide.set(qn('m:val'), '1')
        pr.append(supHide)
    el.append(pr)
    sub = OxmlElement('m:sub')
    for p in sub_parts:
        sub.append(p)
    el.append(sub)
    sup = OxmlElement('m:sup')
    for p in sup_parts:
        sup.append(p)
    el.append(sup)
    e = OxmlElement('m:e')
    for p in e_parts:
        e.append(p)
    el.append(e)
    return el


def _limlow(e_parts, lim_parts):
    """Lower limit (argmin, min, lim) with limit underneath."""
    el = OxmlElement('m:limLow')
    el.append(OxmlElement('m:limLowPr'))
    e = OxmlElement('m:e')
    for p in e_parts:
        e.append(p)
    el.append(e)
    lim = OxmlElement('m:lim')
    for p in lim_parts:
        lim.append(p)
    el.append(lim)
    return el


def omath(p, parts):
    om = OxmlElement('m:oMath')
    for part in parts:
        om.append(part)
    p._element.append(om)


def omath_display(doc, body, cursor, parts, eq_num=None):
    """Display equation in a borderless 2-column table: centered equation + right-aligned number."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove all borders
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'none')
        e.set(qn('w:sz'), '0')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), 'auto')
        borders.append(e)
    tblPr.append(borders)
    # Full-width table
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), TABLE_WIDTH_PCT)
    tblW.set(qn('w:type'), 'pct')
    old_w = tblPr.find(qn('w:tblW'))
    if old_w is not None:
        tblPr.remove(old_w)
    tblPr.append(tblW)
    # Column widths: equation 85%, number 15%
    for j, w in enumerate([8100, 1400]):
        tc = tbl.cell(0, j)._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(w))
        tcW.set(qn('w:type'), 'dxa')
        old = tcPr.find(qn('w:tcW'))
        if old is not None:
            tcPr.remove(old)
        tcPr.append(tcW)
    # Equation in first cell, centered
    p0 = tbl.cell(0, 0).paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr0 = p0._element.get_or_add_pPr()
    sp0 = OxmlElement('w:spacing')
    sp0.set(qn('w:before'), '60')
    sp0.set(qn('w:after'), '60')
    pPr0.append(sp0)
    om = OxmlElement('m:oMath')
    for part in parts:
        om.append(part)
    p0._element.append(om)
    # Number in second cell, right-aligned
    p1 = tbl.cell(0, 1).paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr1 = p1._element.get_or_add_pPr()
    sp1 = OxmlElement('w:spacing')
    sp1.set(qn('w:before'), '60')
    sp1.set(qn('w:after'), '60')
    pPr1.append(sp1)
    if eq_num:
        p1.add_run(f'({eq_num})')
    # Vertical center cell content
    for j in range(2):
        tc = tbl.cell(0, j)._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
    tbl_el = tbl._tbl
    body.remove(tbl_el)
    cursor.addnext(tbl_el)
    return None, tbl_el

# ═══════════════════════════════════════════════════════════════════════
# PARAGRAPH HELPERS
# ═══════════════════════════════════════════════════════════════════════


def mkp(doc, body, cursor, space_before=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(space_before if space_before is not None else 0)
    p.paragraph_format.space_after = Pt(8)
    el = p._element
    body.remove(el)
    cursor.addnext(el)
    return p, el


def replace_p(doc, body, old_el):
    prev = old_el.getprevious()
    body.remove(old_el)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    el = p._element
    body.remove(el)
    prev.addnext(el)
    return p, el


def mkh(doc, body, cursor, text, level=1):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    el = p._element
    body.remove(el)
    cursor.addnext(el)
    return el


def add_italic(p, text):
    """Add an italic run to paragraph p."""
    r = p.add_run(text)
    r.italic = True
    return r


def add_page_break(doc, body, after_el):
    """Insert a page break paragraph after after_el. Returns the new element."""
    pb_p = doc.add_paragraph()
    pb_p.paragraph_format.space_before = Pt(0)
    pb_p.paragraph_format.space_after = Pt(0)
    pb_run = pb_p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    pb_run._element.append(br)
    pb_el = pb_p._element
    body.remove(pb_el)
    after_el.addnext(pb_el)
    return pb_el


def make_bookmark(bm_id, name):
    """Create a w:bookmarkStart element."""
    bs = OxmlElement('w:bookmarkStart')
    bs.set(qn('w:id'), str(bm_id))
    bs.set(qn('w:name'), name)
    return bs


def make_bookmark_end(bm_id):
    """Create a w:bookmarkEnd element."""
    be = OxmlElement('w:bookmarkEnd')
    be.set(qn('w:id'), str(bm_id))
    return be


def make_hyperlink(anchor, text, rPr_orig=None, color=LINK_COLOR):
    """Create a w:hyperlink element with blue underlined text."""
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), anchor)
    hl.set(qn('w:history'), '1')
    r = OxmlElement('w:r')
    if rPr_orig is not None:
        new_rPr = copy.deepcopy(rPr_orig)
    else:
        new_rPr = OxmlElement('w:rPr')
    clr = OxmlElement('w:color')
    clr.set(qn('w:val'), color)
    uu = OxmlElement('w:u')
    uu.set(qn('w:val'), 'single')
    new_rPr.append(clr)
    new_rPr.append(uu)
    r.append(new_rPr)
    t = OxmlElement('w:t')
    t.set(XML_SPACE, SPACE_PRESERVE)
    t.text = text
    r.append(t)
    hl.append(r)
    return hl


# ═══════════════════════════════════════════════════════════════════════
# FOOTNOTE HELPER
# ═══════════════════════════════════════════════════════════════════════
_fn_xml = [None]   # cached parsed footnotes XML
_fn_part = [None]  # cached footnotes Part

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def init_footnotes(doc):
    """Parse the footnotes part from the document package and remove old content footnotes."""
    for rel in doc.part.rels.values():
        if 'footnotes' in rel.reltype:
            _fn_part[0] = rel.target_part
            _fn_xml[0] = etree.fromstring(_fn_part[0].blob)
            # Remove existing content footnotes (keep IDs 0 and -1 = Word separators)
            for fn in list(_fn_xml[0]):
                fid = fn.get(f'{{{W_NS}}}id', '')
                if fid not in ('0', '-1', ''):
                    _fn_xml[0].remove(fn)
            return
    print("  Warning: no footnotes part found in template")


def make_footnote(p, fn_text, fn_id):
    """Add a footnote at the end of paragraph p."""
    if _fn_xml[0] is None:
        return
    fn_el = etree.SubElement(_fn_xml[0], f'{{{W_NS}}}footnote')
    fn_el.set(f'{{{W_NS}}}id', str(fn_id))
    fn_p = etree.SubElement(fn_el, f'{{{W_NS}}}p')
    fn_pPr = etree.SubElement(fn_p, f'{{{W_NS}}}pPr')
    fn_pStyle = etree.SubElement(fn_pPr, f'{{{W_NS}}}pStyle')
    fn_pStyle.set(f'{{{W_NS}}}val', 'FootnoteText')
    # Auto-number mark
    fn_r1 = etree.SubElement(fn_p, f'{{{W_NS}}}r')
    fn_rPr1 = etree.SubElement(fn_r1, f'{{{W_NS}}}rPr')
    fn_rStyle1 = etree.SubElement(fn_rPr1, f'{{{W_NS}}}rStyle')
    fn_rStyle1.set(f'{{{W_NS}}}val', 'FootnoteReference')
    etree.SubElement(fn_r1, f'{{{W_NS}}}footnoteRef')
    # Footnote text
    fn_r2 = etree.SubElement(fn_p, f'{{{W_NS}}}r')
    fn_t2 = etree.SubElement(fn_r2, f'{{{W_NS}}}t')
    fn_t2.set(XML_SPACE, SPACE_PRESERVE)
    fn_t2.text = ' ' + fn_text
    # Footnote reference in main text
    fn_ref_r = OxmlElement('w:r')
    fn_ref_rPr = OxmlElement('w:rPr')
    fn_ref_rStyle = OxmlElement('w:rStyle')
    fn_ref_rStyle.set(qn('w:val'), 'FootnoteReference')
    fn_ref_rPr.append(fn_ref_rStyle)
    fn_ref_r.append(fn_ref_rPr)
    fn_ref_el = OxmlElement('w:footnoteReference')
    fn_ref_el.set(qn('w:id'), str(fn_id))
    fn_ref_r.append(fn_ref_el)
    p._element.append(fn_ref_r)


def flush_footnotes():
    """Write cached footnotes XML back to the part."""
    if _fn_part[0] is not None and _fn_xml[0] is not None:
        _fn_part[0]._blob = etree.tostring(
            _fn_xml[0], xml_declaration=True, encoding='UTF-8', standalone=True)

# ═══════════════════════════════════════════════════════════════════════
# CITATION CROSS-REFERENCE SYSTEM
# ═══════════════════════════════════════════════════════════════════════


# Map: citation text as it appears in-text -> bookmark key
# (cite_author, year, key, ref_anchor)
# cite_author: author string as it appears in inline citations
# year: publication year
# key: unique bookmark key
# ref_anchor: string to match in the reference list for back-linking
CITATIONS = [
    ('Epoch AI', '2024', 'EpochAI2024', 'Epoch AI. (2024)'),
    ('Deloitte', '2025', 'Deloitte2025', 'Deloitte. (2025)'),
    ('IEA', '2025', 'IEA2025', 'IEA. (2025)'),
    ('Goldman Sachs', '2024', 'GoldmanSachs2024', 'Goldman Sachs. (2024)'),
    ('EPRI', '2024', 'EPRI2024', 'EPRI. (2024)'),
    ('Hausmann, Hwang, and Rodrik', '2007', 'Hausmann2007', 'Hausmann, R.'),
    ('Uptime Institute', '2024', 'UptimeInstitute2024',
     'Uptime Institute. (2024)'),
    ('Firebird', '2026', 'Firebird2026', 'Firebird. (2026)'),
    ('Flucker, Tozer, and Whitehead', '2013', 'Flucker2013', 'Flucker, S.'),
    ('Oltmanns, Krcmarik, and Gatti', '2021', 'Oltmanns2021', 'Oltmanns, J.'),
    ('Liu et al.', '2023', 'Liu2023', 'Liu, Z.'),
    ('Goldfarb and Trefler', '2018', 'Goldfarb2018', 'Goldfarb, A.'),
    ('Korinek and Stiglitz', '2021', 'Korinek2021', 'Korinek, A.'),
    ('UNCTAD', '2025', 'UNCTAD2025', 'UNCTAD. (2025)'),
    ('Grossman and Rossi-Hansberg', '2008', 'Grossman2008', 'Grossman, G.'),
    ('Hummels and Schaur', '2013', 'Hummels2013', 'Hummels, D.'),
    ('Brainard', '1997', 'Brainard1997', 'Brainard, S.'),
    ('Helpman, Melitz, and Yeaple', '2004', 'HMY2004', 'Helpman, E.'),
    ('Lim\u00E3o and Venables', '2001', 'Limao2001', 'Lim\u00E3o, N.'),
    ('Eurostat', '2025', 'Eurostat2025', 'Eurostat. (2025)'),
    ('EIA', '2025', 'EIA2025', 'EIA. (2025)'),
    ('Hersbach et al.', '2020', 'Hersbach2020', 'Hersbach, H.'),
    ('Turner & Townsend', '2025', 'TurnerTownsend2025',
     'Turner & Townsend. (2025)'),
    ('WonderNetwork', '2024', 'WonderNetwork2024', 'WonderNetwork. (2024)'),
    ('NVIDIA', '2024', 'NVIDIA2024', 'NVIDIA. (2024)'),
    ('Krugman', '1991', 'Krugman1991', 'Krugman, P.'),
    ('World Bank', '2024', 'WorldBank2024', 'World Bank. (2024)'),
    ('GlobalPetrolPrices', '2025', 'GlobalPetrolPrices2025',
     'GlobalPetrolPrices. (2025)'),
    ('Anderson and van Wincoop', '2003', 'Anderson2003', 'Anderson, J.'),
    ('Sastry, Heim, et al.', '2024', 'Sastry2024', 'Sastry, G.'),
    ('Lehdonvirta, Wu, and Hawkins', '2024', 'Lehdonvirta2024',
     'Lehdonvirta, V.'),
    ('Pilz, Mahmood, and Heim', '2025', 'Pilz2025', 'Pilz, K.'),
]

# Auto-generate CITE_MAP: both "Author (Year)" and "Author Year" forms
CITE_MAP = {}
for _auth, _yr, _key, _ in CITATIONS:
    CITE_MAP[f'{_auth} ({_yr})'] = _key   # narrative: Author (Year)
    CITE_MAP[f'{_auth} {_yr}'] = _key      # parenthetical: Author Year
CITE_MAP['World Bank'] = 'WorldBank2024'   # bare mention without year

# Auto-generate REF_KEY_MAP for back-linking from reference list
REF_KEY_MAP = {_key: _anchor for _, _, _key, _anchor in CITATIONS}


def link_citations_pass(body, cite_map, bm_id):
    """Single pass: find citation text in runs and replace with bookmark+hyperlink. Returns count."""
    count = 0
    # Sort by length descending so longer citations match first
    sorted_cites = sorted(cite_map.items(), key=lambda x: -len(x[0]))
    for p_el in list(body.findall(qn('w:p'))):
        for child in list(p_el):
            if child.tag != qn('w:r'):
                continue
            t_el = child.find(qn('w:t'))
            if t_el is None or not t_el.text:
                continue
            text = t_el.text
            for cite_text, key in sorted_cites:
                if cite_text not in text:
                    continue
                idx = text.index(cite_text)
                before = text[:idx]
                after = text[idx + len(cite_text):]
                rPr_orig = child.find(qn('w:rPr'))
                # Modify current run to "before" text only
                t_el.text = before
                t_el.set(XML_SPACE, SPACE_PRESERVE)
                ins = child
                # bookmarkStart
                bm_start = make_bookmark(bm_id[0], f'{key}txt')
                ins.addnext(bm_start)
                ins = bm_start
                # hyperlink (blue underline)
                hyperlink = make_hyperlink(key, cite_text, rPr_orig)
                ins.addnext(hyperlink)
                ins = hyperlink
                # bookmarkEnd
                bm_end = make_bookmark_end(bm_id[0])
                ins.addnext(bm_end)
                ins = bm_end
                bm_id[0] += 1
                # after text
                if after:
                    ra = OxmlElement('w:r')
                    if rPr_orig is not None:
                        ra.append(copy.deepcopy(rPr_orig))
                    ta = OxmlElement('w:t')
                    ta.set(XML_SPACE, SPACE_PRESERVE)
                    ta.text = after
                    ra.append(ta)
                    ins.addnext(ra)
                count += 1
                break  # one per run per pass
    return count


# ═══════════════════════════════════════════════════════════════════════
# ITALIC JOURNAL / BOOK TITLES IN REFERENCES
# ═══════════════════════════════════════════════════════════════════════
ITALIC_IN_REFS = {
    'Anderson': 'American Economic Review',
    'Antr\u00E0s': 'Journal of Political Economy',
    'Brainard': 'American Economic Review',
    'Deloitte': 'Deloitte Insights',
    'EIA': 'Electric Power Monthly',
    'Eurostat': 'Electricity Prices for Non-Household Consumers (nrg_pc_205)',
    'Flucker': 'Building Services Engineering Research and Technology',
    'GlobalPetrolPrices': 'Electricity Prices Around the World',
    'Goldfarb': 'The Economics of Artificial Intelligence',
    'Goldman': 'Goldman Sachs Research',
    'Grossman': 'American Economic Review',
    'Hausmann': 'Journal of Economic Growth',
    'Helpman': 'American Economic Review',
    'Hersbach': 'Quarterly Journal of the Royal Meteorological Society',
    'Hummels': 'American Economic Review',
    'Korinek': 'NBER Working Paper',
    'Krugman': 'Journal of Political Economy',
    'Lim\u00E3o': 'World Bank Economic Review',
    'Liu': 'Proceedings of ACM e-Energy',
    'NVIDIA': 'NVIDIA H100 Tensor Core GPU Datasheet',
    'Oltmanns': 'Journal of Property Investment & Finance',
    'Turner': 'Data Centre Construction Cost Index 2025',
    'Samuelson': 'Economic Journal',
    'UNCTAD': 'Technology and Innovation Report 2025',
    'WonderNetwork': 'Global Ping Statistics',
    'World Bank': 'World Development Indicators',
    'Lehdonvirta': 'Proceedings of the AAAI/ACM Conference on AI, Ethics, and Society',
    'Pilz': 'AI\u2019s Power Requirements Under Exponential Growth',
}


def find_italic_portion(full_rt):
    """Find the journal/book title that should be italicized in a reference."""
    for author_start, italic_text in ITALIC_IN_REFS.items():
        if full_rt.startswith(author_start) and italic_text in full_rt:
            return italic_text
    return None


def _write_ref_segments(p, text, italic_portion):
    """Write reference text with italic journal/book title."""
    if italic_portion and italic_portion in text:
        idx = text.index(italic_portion)
        if idx > 0:
            p.add_run(text[:idx])
        r = p.add_run(italic_portion)
        r.italic = True
        after = text[idx + len(italic_portion):]
        if after:
            p.add_run(after)
    else:
        p.add_run(text)


def add_table(doc, body, after_el, headers, rows, col_widths=None, title=None):
    if title:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after = Pt(3)
        tp.paragraph_format.first_line_indent = Inches(0)
        run = tp.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
        tbl_el = tp._element
        body.remove(tbl_el)
        after_el.addnext(tbl_el)
        after_el = tbl_el
    nr = len(rows) + 1
    nc = len(headers)
    table = doc.add_table(rows=nr, cols=nc)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    # Remove all table-level borders
    tblPr = table._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    old_borders = tblPr.find(qn('w:tblBorders'))
    if old_borders is not None:
        tblPr.remove(old_borders)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    # AutoFit to window: 100% page width
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), TABLE_WIDTH_PCT)
    tblW.set(qn('w:type'), 'pct')
    # Academic-style horizontal rules on header row (top + bottom)

    def _cell_border(tc, sides):
        tcPr = tc.get_or_add_tcPr()
        tcB = OxmlElement('w:tcBorders')
        for s in sides:
            b = OxmlElement(f'w:{s}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), '4')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            tcB.append(b)
        tcPr.append(tcB)
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ""
        pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pp.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        _cell_border(c._tc, ['top', 'bottom'])
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = table.cell(i + 1, j)
            c.text = ""
            pp = c.paragraphs[0]
            if 2 <= j <= 6:
                pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = pp.add_run(str(val))
            run.font.size = Pt(8)
            # Bottom border on last data row
            if i == len(rows) - 1:
                _cell_border(c._tc, ['bottom'])
    if col_widths:
        for j, w in enumerate(col_widths):
            for i in range(nr):
                c = table.cell(i, j)
                tcPr = c._tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(w))
                tcW.set(qn('w:type'), 'dxa')
                old = tcPr.find(qn('w:tcW'))
                if old is not None:
                    tcPr.remove(old)
                tcPr.append(tcW)
    for row in table.rows:
        for cell in row.cells:
            for pp in cell.paragraphs:
                pPr = pp._element.get_or_add_pPr()
                sp = OxmlElement('w:spacing')
                sp.set(qn('w:before'), '10')
                sp.set(qn('w:after'), '10')
                pPr.append(sp)
    tbl_el = table._tbl
    body.remove(tbl_el)
    after_el.addnext(tbl_el)
    return tbl_el


def write_title_and_abstract(doc, body, all_el, hmap):
    print("Rewriting title and abstract...")
    # Replace title (first element — no previous, so clear and rewrite in place)
    title_el = all_el[0]
    for child in list(title_el):
        if child.tag != qn('w:pPr'):
            title_el.remove(child)
    title_p = doc.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.first_line_indent = Inches(0)
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(0)
    r_title = title_p.add_run('Selling FLOPs:\na New Export Industry for Developing Countries')
    r_title.bold = False
    r_title.font.size = Pt(16)
    r_title.font.name = TIMES_NEW_ROMAN

    # Add author name
    author_p, author_el = mkp(doc, body, title_el, space_before=12)
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.paragraph_format.space_after = Pt(12)
    r_author = author_p.add_run('Michael Lokshin')
    r_author.italic = True
    make_footnote(author_p,
                  'This paper\u2019s findings, interpretations, and conclusions are entirely those of the '
                  'authors and do not necessarily represent the views of their employers including the '
                  'World Bank, its Executive Directors, or the countries they represent. '
                  'Michael Lokshin: mlokshin@worldbank.org', 1)

    # Version stamp
    ver_p, ver_el = mkp(doc, body, author_el, space_before=2)
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_p.paragraph_format.space_after = Pt(12)
    r_ver = ver_p.add_run(f'v18  \u2014  {datetime.now().strftime("%B %d, %Y  %H:%M")}')
    r_ver.font.size = Pt(9)
    r_ver.font.color.rgb = RGBColor(128, 128, 128)
    r_ver.font.name = TIMES_NEW_ROMAN

    # Replace Abstract heading + text with single paragraph
    # Remove old Abstract heading
    abs_heading = hmap['abs']
    abs_text = all_el[2]
    body.remove(abs_heading)
    body.remove(abs_text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    r_abs_label = p.add_run('Abstract')
    r_abs_label.bold = True
    p.add_run(
        ': The rapid growth of artificial intelligence is generating surging global demand '
        'for computational resources, yet the cost of producing a unit of computation varies '
        'by a factor of two across countries. The paper develops a trade model in which countries '
        'produce and export computing services (FLOPs), with costs determined by electricity '
        'prices, climate, and construction costs. The model distinguishes two services: '
        'AI training, which is latency-insensitive and can be offshored to the cheapest '
        'producer, and AI inference, which degrades with distance and favors proximity to users. '
        'A sovereignty premium captures governments\u2019 preference for domestic data processing. '
        'Calibrating the model for 86 countries, we find that many cheap-energy '
        'economies, including several low-income countries, could serve the world\u2019s training '
        'needs, while regional inference hubs emerge around major demand centers. For '
        'developing countries with cheap energy but narrow export baskets, compute '
        'exporting provides a direct link between a natural resource endowment and the '
        'fastest-growing segment of global electricity demand, projected to more than '
        'double from 415 TWh in 2024 to 945 TWh by 2030.'
    )
    el = p._element
    body.remove(el)
    ver_el.addnext(el)
    abs_text_el = el

    # JEL classification and keywords after abstract
    p_jel, jel_el = mkp(doc, body, abs_text_el, space_before=12)
    r_jel_label = p_jel.add_run('JEL Classification: ')
    r_jel_label.bold = True
    p_jel.add_run('F14, F18, L86, O14, O33, Q40')

    p_kw, kw_el = mkp(doc, body, jel_el, space_before=2)
    r_kw_label = p_kw.add_run('Keywords: ')
    r_kw_label.bold = True
    p_kw.add_run(
        'compute trade, FLOPs, artificial intelligence, data centers, '
        'comparative advantage, electricity costs, developing countries'
    )

    return title_el, author_el, abs_text_el, kw_el


def write_introduction(doc, body, hmap):
    print("Inserting Section 1: Introduction...")
    cur = mkh(doc, body, hmap['1'].getprevious(), '1. Introduction', level=1)

    # Para 1: AI compute demand + electricity footprint (consolidated)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The rapid expansion of artificial intelligence is creating a rapidly growing demand '
        'for computational resources. The computation used to train frontier AI '
        'models has been doubling every six months (Epoch AI 2024), '
        'with inference workloads expected to account for roughly two-thirds of all compute '
        'by 2026 (Deloitte 2025). '
        'The electricity footprint is enormous: data centers consumed approximately 415 TWh '
        'in 2024, accounting for 1.5% of global electricity demand, projected to reach 945 TWh by 2030 '
        '(IEA 2025), with U.S. data center electricity consumption expected to triple over '
        'that period (EPRI 2024).'
    )
    make_footnote(p, 'Based on Epoch AI\u2019s dataset of 225 notable training runs; '
                  'growth has accelerated since 2020.', 2)

    # Para 3: FLOP exporting as value chain upgrading
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'This surge in demand creates a new type of export opportunity. Countries with abundant, '
        'inexpensive electricity and land, from Central Asia to North Africa, can produce '
        'and export computational services, measured in floating-point operations per second '
        '(FLOP/s). We call this '
    )
    add_italic(p, 'FLOP exporting')
    p.add_run(
        ': the production of compute services in one country for consumption in another. '
        'FLOP exporting is a form of value chain upgrading. Rather than '
        'exporting raw energy resources (oil, natural gas, or coal) as primary commodities, '
        'countries can convert cheap electricity into a higher value-added digital service. '
        'Just as exporting refined petroleum products captures more value than exporting crude '
        'oil, exporting FLOPs captures more value than exporting the kilowatt-hours that power '
        'them. For energy-rich developing countries, FLOP exporting offers a route up the '
        'value chain without the heavy industrialization traditionally required for '
        'such upgrading (Hausmann, Hwang, and Rodrik 2007).'
    )

    # Para 4: Low labor + option value
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Data center operations require minimal labor: '
        'a typical hyperscale facility employs only about 50 permanent staff '
        '(Uptime Institute 2024).'
    )
    make_footnote(p, 'A 40 MW facility houses over 6,000 servers with approximately '
                  '53,000 GPUs. Staffing ranges from approximately 30 (Tier II) to 70 (Tier IV '
                  'with full redundancy); remote monitoring further reduces on-site requirements '
                  '(Uptime Institute 2024).', 3)
    p.add_run(
        ' The binding input is cheap electricity, so some human capital constraints that have '
        'historically limited export upgrading in developing countries (Hausmann, Hwang, and '
        'Rodrik 2007) are largely absent. Investing in FLOP production capacity today also '
        'has option value: export-oriented data centers can later serve the domestic '
        'market as local AI demand grows.'
    )

    # Para 5: ECA opportunity
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The opportunity is particularly relevant for the Europe and Central Asia (ECA) region. '
        'Several ECA countries, including Turkmenistan, Kyrgyzstan, and the countries of the '
        'South Caucasus, have among the world\u2019s lowest electricity prices but limited '
        'integration into the global digital economy. Building data centers in these locations '
        'and selling compute services to high-cost markets could generate export revenue, attract '
        'foreign investment, and accelerate digital infrastructure development.'
    )

    # Para 6: Real data center plans + profit estimate
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Large-scale investments confirm that FLOP exporting from developing countries is '
        'feasible. Armenia is deploying 50,000 GPUs in a $4 billion '
        'AI megaproject (Firebird 2026), while Kenya, Saudi Arabia, Morocco, Malaysia, and '
        'Indonesia have each attracted billion-dollar data center commitments.'
    )
    make_footnote(p, 'Microsoft and G42 announced a $1 billion geothermal-powered data center '
                  'in Kenya (2024); AWS committed $5.3 billion to a cloud region in Saudi Arabia (2024); '
                  'Morocco allocated $1.1 billion under its Digital Morocco 2030 strategy; '
                  'Microsoft ($2.2 billion) and Google ($2 billion) announced data centers in Malaysia (2024); '
                  'Microsoft committed $1.7 billion to cloud and AI infrastructure in Indonesia (2024).', 4)
    p.add_run(
        ' The economic stakes are substantial: a single 40 MW data center in Kyrgyzstan could '
        'generate annual revenue of $630 million\u2013$950 million at wholesale contract rates, '
        'adding over 15% to the country\u2019s export base (World Bank 2024).'
    )
    make_footnote(p, 'At $0.038/kWh electricity, a 40 MW facility houses approximately '
                  '53,000 GPUs with production costs of $453 million per year. Gross revenue depends '
                  'on pricing: at hyperscaler retail rates ($2.00\u20132.50/GPU-hour), annual revenue '
                  'would reach $830 million\u2013$1 billion, but a Kyrgyz operator would more likely '
                  'sell at wholesale or long-term contract rates, perhaps $0.80\u20131.20/GPU-hour, '
                  'yielding $630\u2013950 million and thinner margins. Even at the lower bound, '
                  'this exceeds 15% of Kyrgyzstan\u2019s $3.8 billion in goods exports (2024).', 5)

    # Para 9: First paper + contributions
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Recent work has examined compute governance and the geography of AI infrastructure '
        '(Sastry, Heim, et al. 2024, Lehdonvirta, Wu, and Hawkins 2024, '
        'Pilz, Mahmood, and Heim 2025), but no formal trade model of compute exists. '
        'We offer the first such model, treating FLOPs as commodities produced and exported '
        'according to Ricardian comparative advantage. '
        'We make three contributions. First, we develop a trade model of FLOP production and '
        'export that decomposes the cost of a FLOP into electricity, hardware, and construction '
        'components, and introduces an iceberg trade cost for inference that captures latency '
        'degradation, alongside a sovereignty premium for domestic production preference. '
        'Second, we calibrate the model for 86 countries using data on electricity prices, '
        'climate, data center construction costs, and inter-country network latency. '
        'Third, we characterize the resulting trade regimes (which countries export, which '
        'import, and which adopt hybrid strategies) and show how the sovereignty premium '
        'determines the boundary between domestic and foreign sourcing.'
    )

    # Para 10: Roadmap
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The remainder of the paper is organized as follows. Section 2 reviews the related '
        'literature. Section 3 develops the model, defining the production technology for FLOPs '
        'and the trade cost structure that distinguishes training from inference. Section 4 '
        'derives the comparative advantage results, presents the buyer\u2019s optimal sourcing '
        'problem, and develops a country taxonomy. Section 5 analyzes the make-or-buy decision '
        'under country-specific fixed entry costs and the sovereignty premium. Section 6 '
        'describes the data. Section 7 calibrates the model and discusses the results. '
        'Section 8 concludes.'
    )


def write_literature(doc, body, hmap):
    print("Inserting Section 2: Related Literature...")
    cur = mkh(doc, body, hmap['1'].getprevious(), '2. Related Literature', level=1)

    # Para 1: AI comparative advantage + value chain upgrading (merged)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Goldfarb and Trefler (2018) argue that AI shifts comparative advantage toward '
        'countries with data, human capital, and institutional capacity. Our model introduces '
        'a complementary channel: comparative advantage in compute '
    )
    add_italic(p, 'production')
    p.add_run(
        ' depends on electricity costs and climate, so resource-rich countries could become '
        'compute exporters without domestic AI research industries. Korinek and Stiglitz (2021) '
        'raise the possibility that developing countries could be left behind in the AI '
        'revolution; FLOP exporting offers a route by which energy-rich developing '
        'countries could participate. The concept of FLOP exporting as value chain upgrading '
        'connects to Hausmann, Hwang, and Rodrik (2007), who show that what a country exports '
        'matters for growth. Lim\u00E3o and Venables (2001) show that infrastructure quality '
        'determines trade costs. In our model, network infrastructure plays the analogous role '
        'for digital trade.'
    )

    # Para 2: Data center location literature (kept as-is)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'A growing literature examines the determinants of data center location. '
        'Flucker, Tozer, and Whitehead (2013) show that climate significantly affects data center '
        'cooling costs. Oltmanns, Krcmarik, and Gatti (2021) model data center siting as a '
        'function of electricity prices, climate, connectivity, and political stability. '
        'Liu et al. (2023) study data center placement under renewable energy constraints. '
        'These studies focus on where firms should build data centers; our contribution is to '
        'embed this location decision in a trade framework that endogenizes the sourcing of '
        'compute across countries.'
    )

    # Para 3: Compute governance literature
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Several recent papers address compute governance. Sastry, Heim, et al. (2024) '
        'argue that compute is well suited as a regulatory lever because it is detectable, '
        'excludable, and quantifiable, and because the chip supply chain is concentrated in '
        'a handful of firms. Lehdonvirta, Wu, and Hawkins (2024) '
        'map the global geography of cloud GPU infrastructure and find that training-capable '
        'hardware sits in roughly 30 countries (their \u201CCompute North\u201D), while '
        'a \u201CCompute South\u201D is limited to inference-grade chips, '
        'a geographic split that mirrors our model\u2019s training/inference exporter distinction. '
        'Pilz, Mahmood, and Heim (2025) project that AI data center power demand '
        'could reach 327 GW by 2030 and that domestic power shortages may push '
        'compute infrastructure abroad. These papers describe where compute is and who '
        'controls it. We add the economics of why it locates where it does and how trade '
        'in compute responds to costs and policy.'
    )

    # Para 4: Trade theory connections (trimmed, no Melitz)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Our model builds on the trade-in-tasks framework of Grossman and Rossi-Hansberg (2008), '
        'where tasks differ in their offshoring costs. Training and inference use the same input '
        '(FLOPs) but differ in latency sensitivity, generating task-specific sourcing patterns. '
        'The iceberg trade cost for inference connects to Hummels and Schaur (2013), who estimate '
        'that each day of shipping time is equivalent to a tariff; in our setting, milliseconds '
        'replace days and network latency replaces shipping time. The make-or-buy decision follows '
        'the proximity-concentration tradeoff of Brainard (1997) and the heterogeneous-firm trade '
        'model of Helpman, Melitz, and Yeaple (2004), where firms sort into exporting versus FDI '
        'based on productivity; in our setting, countries sort into importing versus domestic '
        'production based on their cost advantage.'
    )


def write_model_opening(doc, body, all_el):
    print("Rewriting Section 3 opening (FLOP production)...")

    p, cur = replace_p(doc, body, all_el[4])
    p.add_run(
        'A floating-point operation (FLOP) is a single arithmetic computation. Computing power '
        'is measured in petaFLOP/s (10\u00b9\u2075 FLOPs per second); a current-generation NVIDIA H100 '
        'GPU delivers approximately 1 petaFLOP/s at 16-bit precision. The production of compute '
        'at scale takes place in data centers, purpose-built facilities that house thousands of '
        'GPU-equipped servers.'
    )

    p2, cur = mkp(doc, body, cur)
    p2.add_run('The production process has three cost components. First, ')
    add_italic(p2, 'hardware')
    p2.add_run(
        ': a current-generation NVIDIA H100 costs approximately $25,000 and has a useful life '
        'of about three years at 90% utilization, yielding an amortized cost of approximately '
        '$1.06/hr.'
    )
    make_footnote(p2, 'Amortization: $25,000 / (3 years \u00d7 8,766 hours/year \u00d7 90% '
                  'utilization) \u2248 $1.06/hr. This assumes the NVIDIA list price; street prices '
                  'fluctuated between $25,000 and $40,000 during the 2023\u20132024 GPU shortage.', 6)
    p2.add_run(' Second, ')
    add_italic(p2, 'electricity')
    p2.add_run(
        ': each GPU draws approximately 700 watts, and a large data center may consume '
        '40\u2013100 MW. Electricity is the primary recurring cost and the main source of '
        'cross-country variation. Third, '
    )
    add_italic(p2, 'construction')
    p2.add_run(
        ': the physical facility, including power distribution, connectivity, and cooling. '
        'Cooling overhead is captured by the power usage effectiveness (PUE), which varies '
        'from 1.08 in cold climates to over 1.4 in hot ones '
        '(Flucker, Tozer, and Whitehead 2013).'
    )
    make_footnote(p2, 'Google reports a fleet-wide trailing twelve-month PUE of 1.10 (2024). '
                  'The PUE floor of 1.08 represents current best practice with free-air cooling '
                  'in Scandinavian or Icelandic climates.', 7)

    p3, cur = mkp(doc, body, cur)
    p3.add_run(
        'Because GPU prices are set on global markets and do not vary significantly across '
        'countries, while electricity prices and construction costs differ substantially, '
        'the cross-country variation in FLOP costs is driven primarily by energy and '
        'infrastructure. This is the basis of our model.'
    )


def write_production_technology(doc, body, hmap):
    print("Rewriting Section 3.1 (Production Technology)...")

    all_now = list(body)
    s11i = all_now.index(hmap['1.1'])
    s12i = all_now.index(hmap['1.2'])
    for el in all_now[s11i + 1:s12i]:
        body.remove(el)
    cur = hmap['1.1']

    p, cur = mkp(doc, body, cur)
    p.add_run('Consider ')
    omath(p, [_v('N')])
    p.add_run(
        ' countries, each capable of producing compute services (FLOPs). The unit cost of '
        'producing one GPU-hour of compute in country '
    )
    omath(p, [_v('j')])
    p.add_run(
        ' depends on three inputs: electricity, hardware, and data center construction. '
        'The key cost driver is energy intensity, the electrical power a GPU draws during '
        'operation. We denote energy intensity by '
    )
    omath(p, [_v('\u03B3')])
    p.add_run(', measured in kilowatts (kW). For the NVIDIA H100, ')
    omath(p, [_v('\u03B3'), _t(' = 0.700 kW')])
    p.add_run(
        ' (700 watts). The actual electricity consumed depends also on the '
    )
    add_italic(p, 'power usage effectiveness')
    p.add_run(' PUE(')
    omath(p, [_msub('\u03B8', 'j')])
    p.add_run(
        '), a dimensionless ratio of total facility energy to IT equipment energy '
        '(Flucker, Tozer, and Whitehead 2013). We model PUE as:'
    )
    p.paragraph_format.space_after = Pt(2)

    # PUE equation
    _, cur = omath_display(doc, body, cur, [
        _t('PUE('), _msub('\u03B8', 'j'), _t(') = '),
        _v('\u03C6'), _t(' + '),
        _v('\u03B4'), _t(' \u00b7 max(0, '),
        _msub('\u03B8', 'j'), _t(' \u2212 '),
        _v('\u03B8\u0304'), _t(')'),
    ], eq_num='1')

    p, cur = mkp(doc, body, cur)
    p.add_run('where ')
    omath(p, [_v('\u03C6'), _t(' = 1.08')])
    p.add_run(' is the baseline PUE in cold climates, ')
    omath(p, [_v('\u03B4'), _t(' = 0.015')])
    p.add_run(' is the PUE sensitivity per \u00b0C above the reference, and ')
    omath(p, [_v('\u03B8\u0304'), _t(' = 15\u00b0C')])
    p.add_run(
        ' is the reference temperature (Flucker, Tozer, and Whitehead 2013).'
    )
    make_footnote(p, 'The linear PUE model is a simplification. Modern liquid and immersion '
                  'cooling technologies can achieve PUE \u2248 1.2 even in hot climates, flattening the '
                  'temperature\u2013PUE relationship. Our specification thus overstates the cooling penalty '
                  'for countries that adopt advanced cooling. Capping PUE at 1.20 (simulating universal '
                  'liquid cooling) yields a Kendall rank correlation of 0.96 with the baseline '
                  'rankings; the top five countries are unchanged and the maximum rank shift is six '
                  'positions. Gulf states and North Africa gain the most (UAE moves from 26th to 20th, '
                  'Qatar from 15th to 11th), but the effect is small because electricity prices, not '
                  'cooling, dominate cross-country cost variation.', 8)

    # Hardware and construction
    p, cur = mkp(doc, body, cur)
    p.add_run('Hardware costs are captured by ')
    omath(p, [_v('\u03C1')])
    p.add_run(', the amortized cost of one GPU-hour: ')
    omath(p, [
        _v('\u03C1'), _t(' = '), _msub('P', 'GPU'),
        _t(' / ('), _v('L'), _t(' \u00b7 '), _v('H'),
        _t(' \u00b7 '), _v('\u03B2'), _t(')'),
    ])
    p.add_run(', where ')
    omath(p, [_msub('P', 'GPU')])
    p.add_run(' is the purchase price, ')
    omath(p, [_v('L')])
    p.add_run(' the lifetime in years, ')
    omath(p, [_v('H')])
    p.add_run(' = 8,766 hours per year, and ')
    omath(p, [_v('\u03B2')])
    p.add_run(' the utilization rate. Construction costs enter through ')
    omath(p, [_msub('p', 'L,j')])
    p.add_run(
        ', the cost of building one kilowatt of data center IT capacity in country '
    )
    omath(p, [_v('j')])
    p.add_run(' ($/kW), amortized over the facility lifetime ')
    omath(p, [_v('D')])
    p.add_run('.')

    # Equation (1)
    p, cur = mkp(doc, body, cur)
    p.add_run('Then, the total cost per GPU-hour in country ')
    omath(p, [_v('j')])
    p.add_run(' is:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msub('c', 'j'), _t(' = '),
        _t('PUE('), _msub('\u03B8', 'j'), _t(') \u00b7 '),
        _v('\u03B3'), _t(' \u00b7 '),
        _msub('p', 'E,j'), _t(' + '),
        _v('\u03C1'), _t(' + '),
        _v('\u03B3'), _t(' \u00b7 '),
        _msub('p', 'L,j'), _t(' / ('),
        _v('D'), _t(' \u00b7 '), _v('H'), _t(')'),
    ], eq_num='2')

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The first term is the electricity cost: the PUE-adjusted power draw times the '
        'electricity price '
    )
    omath(p, [_msub('p', 'E,j')])
    p.add_run(
        ' ($/kWh). The second term is amortized hardware. The third is amortized construction. '
        'Since GPU prices, lifetimes, and utilization rates are set by global markets and '
        'engineering constraints, '
    )
    omath(p, [_v('\u03C1')])
    p.add_run(
        ' and its components ('
    )
    omath(p, [_msub('P', 'GPU'), _t(', '), _v('L'), _t(', '), _v('\u03B2')])
    p.add_run(
        ') are common across countries. The main exception is countries subject to U.S. '
        'export controls on advanced GPUs (e.g., China, Russia, Iran), where grey-market '
        'procurement raises the effective '
    )
    omath(p, [_v('\u03C1')])
    p.add_run(
        '; we discuss this in Section 7. Cross-country variation in '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' is therefore driven primarily by electricity prices, climate (through PUE), '
        'and construction costs.'
    )
    make_footnote(p, 'China is developing an alternative domestic chip stack based on '
                  'Huawei\u2019s Ascend series (910B/910C) and other domestic accelerators. If these '
                  'achieve comparable FLOPs per watt at lower prices, China\u2019s effective \u03C1 could '
                  'diverge from the NVIDIA-based benchmark used here, potentially improving its '
                  'cost position despite export controls.', 9)


def write_trade_costs(doc, body, hmap):
    print("Rewriting Section 3.2 (Trade Costs)...")

    all_now = list(body)
    s12i = all_now.index(hmap['1.2'])
    s2i = all_now.index(hmap['2'])
    for el in all_now[s12i + 1:s2i]:
        body.remove(el)
    cur = hmap['1.2']

    # Redefine two service types
    p, cur = mkp(doc, body, cur)
    p.add_run('Countries produce and trade two types of compute services. ')
    add_italic(p, 'Training services')
    p.add_run(
        ' encompass batch workloads: model training, fine-tuning, and large-scale data '
        'processing. A frontier model training run typically takes weeks to months on '
        'thousands of GPUs. The client ships its data to where FLOPs are cheapest; the computation '
        'executes locally and the output is returned. Since neither input nor output is '
        'time-sensitive, network latency plays no role. '
    )
    add_italic(p, 'Inference services')
    p.add_run(
        ' encompass real-time workloads: chatbot responses, autonomous decisions, interactive '
        'agents. Each query must travel to the server and back within milliseconds, so the '
        'service degrades with delivery delay.'
    )

    # Latency definition
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Latency')
    p.add_run(', denoted ')
    omath(p, [_msub('l', 'jk')])
    p.add_run(
        ', is the round-trip time for a data packet to travel from server country '
    )
    omath(p, [_v('j')])
    p.add_run(' to demand center ')
    omath(p, [_v('k')])
    p.add_run(
        ' and back, measured in milliseconds (ms). '
        'Within a country, latency is typically 5\u201310 ms; across continents it can exceed '
        '150 ms. For training, the workload ships to the producer, so effective latency is zero.'
    )

    # Sovereignty premium definition
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Governments and firms may prefer to process data domestically for reasons of national '
        'security, regulatory compliance, or political preference. We capture this through a '
    )
    add_italic(p, 'sovereignty premium')
    p.add_run(' ')
    omath(p, [_v('\u03BB'), _t(' \u2265 0')])
    p.add_run(
        ', which acts as a proportional markup on the cost of foreign-sourced compute. '
        'When a country sources compute from a foreign producer, the effective cost is '
        'inflated by the factor '
    )
    omath(p, [_t('(1 + '), _v('\u03BB'), _t(')')])
    p.add_run(
        '. The sovereignty premium is zero for domestic production.'
    )

    # Equation (2) with sovereignty
    p, cur = mkp(doc, body, cur)
    p.add_run('The delivered cost of service ')
    omath(p, [_v('s'), _t(' \u2208 {'), _v('T'), _t(', '), _v('I'), _t('}')])
    p.add_run(' from producer ')
    omath(p, [_v('j')])
    p.add_run(' to demand center ')
    omath(p, [_v('k')])
    p.add_run(' is:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msub('P', 's'), _t('('), _v('j'), _t(', '), _v('k'),
        _t(') = (1 + '), _msub('\u03BB', 'jk'),
        _t(') \u00b7 (1 + '), _msub('\u03C4', 's'), _t(' \u00b7 '),
        _msub('l', 'jk'), _t(') \u00b7 '),
        _msub('c', 'j'),
    ], eq_num='3')

    p, cur = mkp(doc, body, cur)
    p.add_run('where ')
    omath(p, [_msub('\u03BB', 'jk'), _t(' = '), _v('\u03BB')])
    p.add_run(' if ')
    omath(p, [_v('j'), _t(' \u2260 '), _v('k')])
    p.add_run(' (foreign sourcing) and ')
    omath(p, [_msub('\u03BB', 'jk'), _t(' = 0')])
    p.add_run(' if ')
    omath(p, [_v('j'), _t(' = '), _v('k')])
    p.add_run(' (domestic); ')
    omath(p, [_msub('\u03C4', 'T'), _t(' = 0')])
    p.add_run(' (training has zero effective latency) and ')
    omath(p, [_msub('\u03C4', 'I'), _t(' = '), _v('\u03C4'), _t(' > 0')])
    p.add_run(
        ' (inference degrades with latency). The parameter '
    )
    omath(p, [_v('\u03C4')])
    p.add_run(' measures the rate of quality degradation per millisecond of round-trip latency.')

    # Implication
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'For training, equation (3) simplifies to '
    )
    omath(p, [_msub('P', 'T'), _t('('), _v('j'), _t(', '), _v('k'),
              _t(') = (1 + '), _msub('\u03BB', 'jk'), _t(') \u00b7 '), _msub('c', 'j')])
    p.add_run(
        ': the delivered cost depends only on the production cost and whether the source is '
        'foreign. For inference: '
    )
    omath(p, [_msub('P', 'I'), _t('('), _v('j'), _t(', '), _v('k'),
              _t(') = (1 + '), _msub('\u03BB', 'jk'), _t(') \u00b7 (1 + '), _v('\u03C4'),
              _t(' \u00b7 '), _msub('l', 'jk'), _t(') \u00b7 '), _msub('c', 'j')])
    p.add_run(
        ': the delivered cost depends on production cost, latency, and sovereignty. '
        'In practice, inference also faces a hard latency ceiling: beyond a threshold '
    )
    omath(p, [_v('\u0305l')])
    p.add_run(
        ' (typically 200\u2013300 ms for interactive applications), the service becomes '
        'unusable regardless of price. We model this as '
    )
    omath(p, [_msub('P', 'I'), _t('('), _v('j'), _t(', '), _v('k'),
              _t(') = \u221E')])
    p.add_run(' if ')
    omath(p, [_msub('l', 'jk'), _t(' > '), _v('\u0305l')])
    p.add_run(
        '. Training concentrates at the globally cheapest FLOP source '
        'regardless of distance, while inference disperses toward demand centers bounded '
        'by latency.'
    )


def renumber_sections(hmap):
    print("Renumbering sections...")
    # v8: 1→3, 1.1→3.1, 1.2→3.2, 2→4, 3→5, 4→7 (skip 6 for Data), 5→8
    renumber = [
        ('1.2', '1.2', '3.2'), ('1.1', '1.1', '3.1'), ('1', '1.', '3.'),
        ('2', '2.', '4.'), ('3', '3.', '5.'),
        ('4', '4.', '7.'), ('5', '5.', '8.'),
    ]
    for key, old, new in renumber:
        if key in hmap:
            el = hmap[key]
            for t in el.findall(f'.//{qn("w:t")}'):
                if t.text and old in t.text:
                    t.text = t.text.replace(old, new, 1)
                    break


def write_comparative_advantage(doc, body, hmap, demand_data):
    print("Rewriting Section 4 (Comparative Advantage)...")

    # Clear content between section 4 heading (was v8 "2") and section 5 heading (was v8 "3")
    all_now = list(body)
    s4 = hmap['2']
    s5 = hmap['3']
    s4i = all_now.index(s4)
    s5i = all_now.index(s5)
    for el in all_now[s4i + 1:s5i]:
        body.remove(el)
    cur = s4

    # Buyer's problem intro
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'We now characterize the optimal sourcing decision for a country that consumes compute '
        'services. The analysis is structured around two types of countries: '
    )
    add_italic(p, 'importers')
    p.add_run(' (countries that purchase FLOPs from abroad) and ')
    add_italic(p, 'exporters')
    p.add_run(' (countries that produce FLOPs for foreign consumption).')

    # Optimal sourcing rule
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The buyer\u2019s optimal sourcing rule follows directly from the cost structure. '
        'For each service type '
    )
    omath(p, [_v('s'), _t(' \u2208 {'), _v('T'), _t(', '), _v('I'), _t('}')])
    p.add_run(', demand country ')
    omath(p, [_v('k')])
    p.add_run(' chooses the source ')
    omath(p, [_msubsup('j', 's', '*')])
    p.add_run(' that minimizes the delivered cost:')
    p.paragraph_format.space_after = Pt(2)

    # Optimization equation — buyer's problem
    _, cur = omath_display(doc, body, cur, [
        _msubsup('j', 's', '*'), _t('('), _v('k'),
        _t(') = '),
        _limlow([_t('arg min')], [_v('j')]),
        _t(' '), _msub('P', 's'), _t('('), _v('j'), _t(', '), _v('k'), _t(')'),
    ], eq_num='4')

    # Training case
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Training. ')
    p.add_run('Since ')
    omath(p, [_msub('\u03C4', 'T'), _t(' = 0')])
    p.add_run(', the buyer\u2019s problem reduces to ')
    omath(p, [_msubsup('j', 'T', '*'), _t('('), _v('k'),
              _t(') = '),
              _limlow([_t('arg min')], [_v('j')]),
              _t(' '), _msub('P', 'T'), _t('('), _v('j'), _t(', '), _v('k'), _t(')')
              ])
    p.add_run(' = ')
    omath(p, [_limlow([_t('arg min')], [_v('j')]),
              _t(' (1 + '), _msub('\u03BB', 'jk'), _t(') \u00b7 '), _msub('c', 'j')])
    p.add_run(
        '. Country '
    )
    omath(p, [_v('k')])
    p.add_run(' sources training domestically if and only if ')
    omath(p, [_msub('c', 'k'), _t(' \u2264 (1 + '), _v('\u03BB'), _t(') \u00b7 '),
              _t('min'), _msub(' ', 'j\u2260k'), _t(' '), _msub('c', 'j')])
    p.add_run(
        '. Without the sovereignty premium, all countries import training from the single '
        'cheapest global producer. With a positive '
    )
    omath(p, [_v('\u03BB')])
    p.add_run(
        ', countries whose domestic cost falls within the sovereignty band produce '
        'domestically.'
    )

    # Critical sovereignty threshold
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'This condition yields a country-specific switching threshold. Country '
    )
    omath(p, [_v('k')])
    p.add_run(' sources training domestically if and only if ')
    omath(p, [_v('\u03BB'), _t(' \u2265 '),
              _msubsup('\u03BB', 'k', '*')])
    p.add_run(', where ')
    omath(p, [_msubsup('\u03BB', 'k', '*'), _t(' = '),
              _msub('c', 'k'), _t(' / '),
              _t('min'), _msub(' ', 'j\u2260k'), _t(' '),
              _msub('c', 'j'), _t(' \u2212 1')])
    p.add_run('.')

    p, cur = mkp(doc, body, cur)
    ls = demand_data["lambda_star"]
    p.add_run(
        'The threshold is country-specific and varies widely. Near-frontier '
        f'countries switch to domestic production with minimal sovereignty premia: '
        f'Kyrgyzstan requires only \u03BB* = {ls["KGZ"] * 100:.1f}%, '
        f'China {ls["CHN"] * 100:.1f}%, and the United States {ls["USA"] * 100:.1f}%. '
        f'High-cost countries require much larger premia: Germany needs '
        f'\u03BB* = {ls["DEU"] * 100:.1f}% and Japan {ls["JPN"] * 100:.1f}%. '
        'This predicts which countries are most likely to resist compute trade '
        'liberalization: those with high \u03BB* face large cost penalties from '
        'sovereignty-driven domestic sourcing.'
    )

    # Inference case
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Inference. ')
    p.add_run('Since ')
    omath(p, [_msub('\u03C4', 'I'), _t(' = '), _v('\u03C4'), _t(' > 0')])
    p.add_run(
        ', the buyer weighs both cost and latency: '
    )
    omath(p, [_msubsup('j', 'I', '*'), _t('('), _v('k'),
              _t(') = '),
              _limlow([_t('arg min')], [_v('j')]),
              _t(' '), _msub('P', 'I'), _t('('), _v('j'), _t(', '), _v('k'), _t(')')])
    p.add_run(
        '. A country with low production cost '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' but high latency to demand center '
    )
    omath(p, [_v('k')])
    p.add_run(
        ' may be uncompetitive for inference even if it dominates in training. This '
        'creates regional inference hubs: countries that combine moderate costs with '
        'geographic proximity to major markets.'
    )

    # Country taxonomy
    p, cur = mkp(doc, body, cur, space_before=6)
    r = p.add_run('Country Taxonomy. ')
    r.bold = True
    p.add_run(
        'The model generates a natural classification of countries based on their cost '
        'position and geographic proximity to demand centers:'
    )

    p, cur = mkp(doc, body, cur)
    p.add_run('Equation (3) generates four possible sourcing regimes for each buyer ')
    omath(p, [_v('k')])
    p.add_run(', depending on whether training and inference are sourced domestically or imported:')

    p, cur = mkp(doc, body, cur)
    add_italic(p, '(i) Full import ')
    omath(p, [_t('('), _msubsup('j', 'T', '*'), _t(' \u2260 k, '),
              _msubsup('j', 'I', '*'), _t(' \u2260 k)')])
    r = p.add_run('. ')
    p.add_run('Countries with high ')
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' that import both training and inference from abroad. '
        'These countries have no cost advantage and rely entirely '
        'on foreign compute. Examples: Ireland, Croatia, Greenland.'
    )

    p, cur = mkp(doc, body, cur)
    add_italic(p, '(ii) Import training, domestic inference ')
    omath(p, [_t('('), _msubsup('j', 'T', '*'), _t(' \u2260 k, '),
              _msubsup('j', 'I', '*'), _t(' = k)')])
    r = p.add_run('. ')
    p.add_run(
        'The most common pure-cost regime: the country cannot compete on global '
        'training costs but its domestic latency advantage makes local inference cheaper than '
        'importing. Examples: Norway, Canada, Brazil, Australia.'
    )

    p, cur = mkp(doc, body, cur)
    add_italic(p, '(iii) Full domestic ')
    omath(p, [_t('('), _msubsup('j', 'T', '*'), _t(' = k, '),
              _msubsup('j', 'I', '*'), _t(' = k)')])
    r = p.add_run('. ')
    p.add_run('Under pure cost minimization, only the globally cheapest producer (Iran) '
              'falls in this category. With a sovereignty premium ')
    omath(p, [_v('\u03BB'), _t(' > 0')])
    p.add_run(
        ', this becomes the dominant regime: countries with high production costs '
        'still produce domestically because the sovereignty markup makes imports more '
        'expensive. For instance, with '
    )
    omath(p, [_v('\u03BB'), _t(' = 10%')])
    p.add_run(
        ', countries like Japan ($1.25/hr) and Germany ($1.24/hr) '
        'produce both services domestically, even though their costs far exceed those of '
        'the cheapest producers.'
    )

    p, cur = mkp(doc, body, cur)
    add_italic(p, '(iv) Domestic training, import inference ')
    omath(p, [_t('('), _msubsup('j', 'T', '*'), _t(' = k, '),
              _msubsup('j', 'I', '*'), _t(' \u2260 k)')])
    r = p.add_run('. ')
    p.add_run(
        'Theoretically possible if a country is the cheapest globally (optimal for zero-latency '
        'training), but sufficiently remote that a nearby country offers cheaper inference. '
        'Empirically, this regime is empty: the cheapest producer\u2019s cost advantage extends to '
        'domestic inference as well.'
    )

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'From the supply side, two patterns of specialization emerge. '
    )
    add_italic(p, 'Training exporters ')
    p.add_run(
        'are countries with the lowest global costs, which meet training demand worldwide '
        'because training has zero effective latency cost. '
    )
    add_italic(p, 'Regional inference hubs ')
    p.add_run(
        'are countries with moderate costs and low latency to major demand centers that export '
        'inference to nearby high-cost countries. Their competitive radius is bounded by latency. '
        'Examples: Kosovo (Southeastern Europe), Finland (Baltics), Algeria (Mediterranean).'
    )

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Note that the production cost per GPU-hour is the same regardless of whether the GPU '
        'is used for training or inference: both services use the same hardware. The cost '
        'difference arises entirely from the trade cost: inference delivered over distance '
        'costs more because of the latency markup. From the buyer\u2019s perspective, an inference '
        'FLOP costs more than a training FLOP because some compute is effectively lost to latency.'
    )

    # Training exporters ⊂ inference exporters
    p, cur = mkp(doc, body, cur, space_before=6)
    p.add_run(
        'The cost structure also implies that every training exporter also exports inference '
        'to at least its own geographic neighborhood; the set of training exporters is a weak '
        'subset of inference exporters. A training exporter has the globally lowest '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ', and for inference to nearby countries this cost advantage dominates the latency '
        'markup, so the same country wins the inference competition for proximate demand '
        'centers. This explains why regime (iv), domestic training with imported '
        'inference, is empirically empty and predicts that training exporters hold '
        'asymmetric market power: they compete globally for training and regionally for '
        'inference.'
    )

    # Training more concentrated than inference
    p, cur = mkp(doc, body, cur, space_before=6)
    p.add_run('A related implication concerns market concentration. Let ')
    omath(p, [_msub('HHI', 's'), _t(' = '),
              _nary('\u2211', [_v('j')], [],
                    [_msup('(demand share served by j)', '2', False, False)])])
    p.add_run(' for service ')
    omath(p, [_v('s')])
    p.add_run('. Then ')
    omath(p, [_msub('HHI', 'T'), _t(' \u2265 '), _msub('HHI', 'I')])
    p.add_run(
        ': training has no geography dimension, so demand pools at the single '
        'cheapest producer ('
    )
    omath(p, [_msub('HHI', 'T'), _t(' \u2192 1')])
    p.add_run(
        '), while inference has latency costs that segment the market into regional '
        'hubs, dispersing market share ('
    )
    omath(p, [_msub('HHI', 'I'), _t(' < 1')])
    p.add_run(
        f'). In our calibration, the demand-weighted HHI for training is '
        f'{demand_data["hhi_t"]:.2f} (highly concentrated), versus '
        f'{demand_data["hhi_i"]:.2f} for inference (moderately concentrated).'
    )


def write_make_or_buy(doc, body, hmap, demand_data):
    print("Rewriting Section 5 (Make-or-Buy)...")

    all_now = list(body)
    s5 = hmap['3']
    s4_old = hmap['4']
    s5i = all_now.index(s5)
    s4i = all_now.index(s4_old)
    for el in all_now[s5i + 1:s4i]:
        body.remove(el)
    cur = s5

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'A country that wishes to produce FLOPs domestically must first build a data center, '
        'incurring a country-specific fixed cost '
    )
    omath(p, [_msub('F', 'j')])
    p.add_run(
        '. This cost depends on local construction prices, land costs, and regulatory '
        'requirements, and is proportional to the construction cost parameter '
    )
    omath(p, [_msub('p', 'L,j')])
    p.add_run(
        '. The entry decision trades off the fixed cost against the expected operating margin.'
    )

    # Demand specification: Equation (5)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'We close the model by specifying demand. Let '
    )
    omath(p, [_msub('q', 'k')])
    p.add_run(
        ' denote the volume of compute purchased by country '
    )
    omath(p, [_v('k')])
    p.add_run(
        '. Following the gravity-model tradition (Anderson and van Wincoop 2003), we proxy '
        'compute demand using GDP at purchasing power parity:'
    )
    p.paragraph_format.space_after = Pt(2)

    # Equation (5): q_k = ω_k · Q
    _, cur = omath_display(doc, body, cur, [
        _msub('q', 'k'), _t(' = '),
        _msub('\u03C9', 'k'), _t(' \u00b7 '), _v('Q'),
        _t(',     '),
        _msub('\u03C9', 'k'), _t(' = '),
        _msub('GDP', 'k'), _t(' / '),
        _nary('\u2211', [_v("k\u2032")], [],
              [_msub('GDP', "k\u2032")]),
    ], eq_num='5')

    p, cur = mkp(doc, body, cur)
    p.add_run('where ')
    omath(p, [_v('Q')])
    p.add_run(
        ' is total global compute spending and '
    )
    omath(p, [_msub('\u03C9', 'k')])
    p.add_run(
        ' is country '
    )
    omath(p, [_v('k')])
    p.add_run(
        '\u2019s share of global demand, measured by its share of GDP at purchasing power '
        'parity. This is a standard gravity-model proxy: larger economies demand more '
        'compute services. Since all results below depend only on demand '
    )
    add_italic(p, 'shares')
    p.add_run(
        ', not on the absolute level '
    )
    omath(p, [_v('Q')])
    p.add_run(
        ', the calibration does not require an estimate of total global compute spending.'
    )

    # Proposition 2
    p, cur = mkp(doc, body, cur, space_before=6)
    p.add_run('The entry decision has a natural characterization. Country ')
    omath(p, [_v('j')])
    p.add_run(' enters the FLOP export market (builds a data center) if and only if:')
    p.paragraph_format.space_after = Pt(2)

    # Entry condition equation
    _, cur = omath_display(doc, body, cur, [
        _nary('\u2211',
              [_v('k'), _t(' : '), _msubsup('j', 's', '*'), _t('('), _v('k'), _t(') = '), _v('j')],
              [],
              [_msub('q', 'k'), _t(' \u00b7 ['), _msub('P', 's'),
               _t('('), _v('j'), _t(', '), _v('k'),
               _t(') \u2212 '), _msub('c', 'j'), _t(']')]
              ),
        _t(' \u2265 '), _msub('F', 'j'),
    ], eq_num='6')

    p, cur = mkp(doc, body, cur)
    p.add_run('where ')
    omath(p, [_msub('q', 'k')])
    p.add_run(
        ' is the volume of compute demanded by country '
    )
    omath(p, [_v('k')])
    p.add_run(', and the summation ranges over all demand centers ')
    omath(p, [_v('k')])
    p.add_run(' for which ')
    omath(p, [_v('j')])
    p.add_run(
        ' is the optimal source per equation (4). Expected volume is determined by the '
        'global market, not by country '
    )
    omath(p, [_v('j')])
    p.add_run(
        '\u2019s own size. A small country like Kyrgyzstan could, if it is the cheapest producer, '
        'attract training demand from the entire world. However, small cheap-energy countries '
        'may have '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' below the global average yet still fail the entry condition if an even cheaper country '
        'exists, because the cheapest producer captures all training volume. For inference, '
        'the competitive catchment area is limited by latency: a remote country, no matter how '
        'cheap, serves only a small geographic radius.'
    )

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Construction costs vary substantially across countries. A '
        'data center in Norway costs more to build than one in Uzbekistan, even if Norway '
        'has cheaper electricity. The entry condition thus depends on both the operating margin '
        '(driven by '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run(') and the fixed cost (driven by ')
    omath(p, [_msub('F', 'j')])
    p.add_run(
        '). As in the heterogeneous-firm trade model of Helpman, Melitz, and Yeaple (2004), '
        'only sufficiently productive units, here countries with sufficiently low cost, '
        'find it worthwhile to enter.'
    )

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The sovereignty premium '
    )
    omath(p, [_v('\u03BB')])
    p.add_run(
        ' interacts with the entry decision in two ways. First, it shifts some import demand '
        'toward domestic production, reducing the volume available to foreign exporters. Second, '
        'it raises the effective price that domestic producers can charge, improving the margin '
        'for domestic entry. The combined effect is that the sovereignty premium expands the set '
        'of countries that find it profitable to build data centers.'
    )

    # Welfare cost of sovereignty
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Welfare cost of sovereignty. ')
    p.add_run(
        'The sovereignty premium imposes a welfare cost by diverting demand from the cheapest '
        'global producer to more expensive domestic sources. We define the aggregate '
        'welfare cost as the demand-weighted extra spending from sovereignty-induced domestic '
        'sourcing: '
    )
    omath(p, [_nary('\u2211', [_v('k')], [],
              [_msub('\u03C9', 'k'), _t(' \u00b7 ('),
               _msub('c', 'k'), _t(' \u2212 '),
               _t('min'), _msub(' ', 'j'), _t(' '),
               _msub('P', 's'), _t('('), _v('j'), _t(', '), _v('k'), _t('))')
               ])])
    p.add_run(
        f', summed over countries that switch to domestic production. '
        f'In our calibration (Section 7), this cost amounts to '
        f'{demand_data["welfare_pct"]:.1f}% of demand-weighted average spending.'
    )

    # Transition to empirical implementation
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The model in Sections 3\u20135 yields predictions about which countries produce FLOPs '
        'and which import them, conditional on observable parameters: electricity prices, '
        'temperatures, construction costs, and bilateral latencies. We now calibrate the model '
        'using data for 86 countries. Section 6 describes the data sources, and Section 7 '
        'presents the calibration results and discusses the implied sourcing patterns for '
        'major demand centers.'
    )


def write_data_section(doc, body, hmap, demand_data):
    print("Inserting Section 6: Data...")

    sec7_heading = hmap['4']  # was v8 "4", now renumbered to "7"
    cur = mkh(doc, body, sec7_heading.getprevious(), '6. Data', level=1)

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Calibrating the model requires data on each component of equations (2) and (3): '
        'electricity prices, temperatures, construction costs, and bilateral latencies. '
    )
    add_italic(p, 'Electricity prices. ')
    p.add_run(
        'For European countries, we use Eurostat industrial electricity prices in the '
        '20,000\u201369,999 MWh consumption band (nrg_pc_205), which corresponds to large '
        'industrial consumers (Eurostat 2025). For non-European countries, we use national '
        'regulator tariff sheets and secondary sources: U.S. Energy Information Administration '
        '(EIA 2025) for the United States; KEPCO for South Korea; national utility tariffs for '
        'Central Asian countries (Barki Tojik, AERA, Ministry of Energy of Uzbekistan); and '
        'GlobalPetrolPrices (2025) for remaining countries. All prices are converted to USD/kWh '
        'at 2024 average exchange rates.'
    )

    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Temperature and construction. ')
    p.add_run(
        'Peak summer temperature for each country is computed from ERA5 reanalysis data '
        '(Hersbach et al. 2020) as the average monthly maximum in the three warmest months, '
        'aggregated across populated grid cells. '
    )
    p.add_run(
        'Data center construction costs per watt of IT capacity are from the Turner & Townsend '
        'Data Centre Construction Cost Index 2025, which reports actual costs ($/W) for 37 '
        'countries across 52 markets (Turner & Townsend 2025). For the remaining countries, '
        'we predict construction costs using a log-linear regression: ln($/W) = '
    )
    omath(p, [_v('a'), _t(' + '), _msub('b', '1'),
              _t(' \u00b7 ln(GDP per capita)')])
    p.add_run(' (')
    omath(p, [_msup('R', '2'), _t(' = 0.43')])
    p.add_run(
        '). The low explanatory power reflects the fact that construction costs depend on '
        'factors beyond income levels, such as labor markets, building codes, and imported-materials '
        'logistics, that GDP per capita captures imperfectly. Resource-rich Gulf states, '
        'for instance, may have lower construction costs than their GDP would predict due to '
        'imported labor and streamlined permitting. '
        'Since construction is only 3\u20136% of total FLOP costs, imputation '
        'error matters less than it might seem: 95% prediction intervals for imputed '
        'countries span about \u00b1$3.50/W, which translates to '
        '\u00b1$0.02/hr in total cost (1.5\u20132% of the mean). '
        'Costs are amortized over 15 years. In equation (2), construction costs '
        'are expressed per kilowatt ($/kW); the raw data in $/W is converted by multiplying by 1,000.'
    )
    make_footnote(p, 'The 37 DCCI countries span 52 markets: Australia, Austria, Brazil, '
                  'Canada, Chile, China, Colombia, Denmark, Finland, France, Germany, Greece, India, '
                  'Indonesia, Ireland, Italy, Japan, Kenya, Malaysia, Mexico, Netherlands, New Zealand, '
                  'Nigeria, Norway, Poland, Portugal, Saudi Arabia, Singapore, South Africa, South Korea, '
                  'Spain, Sweden, Switzerland, UAE, UK, Uruguay, and USA.', 10)

    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Latency. ')
    p.add_run(
        'Inter-country round-trip latency is measured using WonderNetwork\u2019s global ping dataset '
        '(WonderNetwork 2024). For each country pair, we use the median round-trip time (RTT) '
        'in milliseconds. '
        'Domestic latency defaults to 5 ms where no intra-country measurement is available. '
        'These measurements reflect today\u2019s network infrastructure. New undersea cables, '
        'terrestrial fiber, and CDN expansions could cut bilateral latencies enough to '
        'redraw inference trade patterns, opening distant low-cost producers to markets '
        'they cannot currently reach. '
    )
    add_italic(p, 'Hardware. ')
    p.add_run(
        'We use the NVIDIA H100 SXM GPU as the reference hardware platform: list price '
        '$25,000, thermal design power 700W, economic lifetime 3 years, utilization rate 90% '
        '(NVIDIA 2024). This yields an amortized hardware cost '
    )
    omath(p, [_v('\u03C1'), _t(f' = ${RHO:.3f}')])
    p.add_run(
        '/hr. GPU prices are assumed uniform across countries; in practice, export controls '
        'raise effective prices for some countries (see Section 7).'
    )

    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Other parameters. ')
    p.add_run('The latency degradation parameter is set at ')
    omath(p, [_v('\u03C4'), _t(f' = {TAU}'), _t(' per ms')])
    p.add_run(
        f', implying that 100 ms of latency inflates inference cost by {TAU * 100:.0%}. '
        'This value is calibrated to match the observed willingness of cloud providers to '
        'invest in regional points of presence: at '
    )
    omath(p, [_v('\u03C4'), _t(f' = {TAU}'), _t(' per ms')])
    p.add_run(
        ', a latency difference of 100 ms (roughly the intercontinental round-trip between '
        'Europe and East Asia) imposes an 8% cost penalty, consistent with industry evidence '
        'that web-service revenue declines by approximately 1% per 100 ms of additional '
        'latency. '
        'The sovereignty premium is '
    )
    omath(p, [_v('\u03BB'), _t(f' = {LAMBDA:.0%}')])
    p.add_run('. Sensitivity to both parameters is explored in Section 7.')
    make_footnote(p, 'The 10% sovereignty premium is conservative. Survey evidence on data '
                  'localization suggests enterprises pay 15\u201330% more for guaranteed domestic '
                  'data residency (UNCTAD 2025).', 11)

    # Demand data paragraph
    p, cur = mkp(doc, body, cur)
    top5 = demand_data["top5_labels"]
    add_italic(p, 'Demand. ')
    p.add_run(
        'We proxy compute demand '
    )
    omath(p, [_msub('q', 'k')])
    p.add_run(
        ' with total GDP at purchasing power parity (World Bank 2024, '
        'Anderson and van Wincoop 2003), as specified in equation (5). '
        f'The five largest demand centers\u2014'
        f'{top5[0][1]} ({top5[0][2] * 100:.0f}%), '
        f'{top5[1][1]} ({top5[1][2] * 100:.0f}%), '
        f'{top5[2][1]} ({top5[2][2] * 100:.1f}%), '
        f'{top5[3][1]}, and '
        f'{top5[4][1]}\u2014'
        f'account for {demand_data["top5_share"] * 100:.0f}% of global demand. '
        'GDP is a rough proxy. Actual compute demand is probably more concentrated, since '
        'the United States and China have outsized technology sectors. Cloud revenue shares or '
        'installed data center capacity would be better demand measures, but neither is '
        'available at the country level. If anything, our GDP-based shares understate '
        'concentration, which would push training-market HHI even higher.'
    )


def write_calibration(doc, body, hmap, cal, reg, n_eca, n_total, all_reg, all_sov, demand_data):
    print("Replacing Section 7 (Calibration)...")

    sec7 = hmap['4']
    sec8 = hmap['5']
    all_now = list(body)
    s7i = all_now.index(sec7)
    s8i = all_now.index(sec8)
    for el in all_now[s7i + 1:s8i]:
        body.remove(el)
    cur = sec7

    # Introductory paragraph with explanation of costs
    p, cur = mkp(doc, body, cur)
    p.add_run('We calibrate the model for ')
    omath(p, [_v('N'), _t(f' = {n_total}')])
    p.add_run(f' countries ({n_eca} in ECA, {n_total - n_eca} non-ECA comparators). ')
    p.add_run('The unit cost ')
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' represents the total hourly cost of operating one GPU in country '
    )
    omath(p, [_v('j')])
    p.add_run(
        ', measured in dollars per GPU-hour ($/hr). It is the sum of hourly electricity '
        'cost, amortized hardware cost ($1.06/hr), and amortized construction cost.'
    )

    # Define regime column
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The regime column reports the optimal sourcing strategy for each country as a demand '
        'center, derived from the sourcing rule in equation (4) under pure cost minimization '
        '(without the sovereignty premium). '
    )
    add_italic(p, 'Import')
    p.add_run(' means the country imports both training and inference. ')
    add_italic(p, 'Hybrid')
    p.add_run(
        ' means the country imports training from the cheapest global source but produces '
        'inference domestically or from a nearby hub. '
    )
    add_italic(p, 'Domestic')
    p.add_run(
        ' means the country produces both services domestically because it has the lowest '
        'cost including latency. '
    )
    p._element.append(make_bookmark(100, 'TableA1txt'))
    p._element.append(make_hyperlink('TableA1', 'Table A1'))
    p._element.append(make_bookmark_end(100))
    p.add_run(' in the Appendix reports the full results for all ')
    omath(p, [_v('N'), _t(f' = {n_total}')])
    p.add_run(' countries, sorted by total unit cost.')

    # Key findings
    p, cur = mkp(doc, body, cur, space_before=6)
    cheapest = cal[0]
    p.add_run(
        f'The cheapest FLOP producer globally is {cheapest["country"]} '
        f'(${float(cheapest["c_j_total"]):.2f}/hr), benefiting from heavily subsidized '
        f'electricity at ${float(cheapest["p_E_usd_kwh"]):.3f}/kWh. '
        'Among ECA countries, the five cheapest are Turkmenistan ($1.11/hr, subsidized '
        'electricity), Kyrgyzstan ($1.13/hr, hydropower), Russia ($1.14/hr), '
        'Kosovo ($1.14/hr), and Ukraine ($1.15/hr). The Nordics benefit from low PUE '
        '(1.08\u20131.10). At the expensive end, Ireland ($1.28/hr) and Greenland ($1.32/hr) '
        'face high electricity prices. '
        'Construction costs account for 3\u20136% of total costs, ranging from $0.033/hr '
        '(China) to $0.078/hr (Japan, Singapore).'
    )

    # Regime results
    pct_imp = round(100 * all_reg.get("full import", 0) / n_total)
    pct_hyb = round(100 * all_reg.get("import training + build inference", 0) / n_total)
    pct_dom = round(100 * all_reg.get("full domestic", 0) / n_total)
    pct_sov_dom = round(100 * all_sov.get("full domestic", 0) / n_total)
    pct_sov_hyb = round(100 * all_sov.get("import training + build inference", 0) / n_total)
    pct_sov_imp = round(100 * all_sov.get("full import", 0) / n_total)

    p, cur = mkp(doc, body, cur)
    p.add_run(
        f'Under pure cost minimization (without sovereignty), '
        f'{all_reg.get("full import", 0)} of {n_total} countries ({pct_imp}%) are in '
    )
    add_italic(p, 'full import')
    p.add_run(f', {all_reg.get("import training + build inference", 0)} ({pct_hyb}%) in ')
    add_italic(p, 'hybrid')
    p.add_run(f', and {all_reg.get("full domestic", 0)} ({pct_dom}%) in ')
    add_italic(p, 'domestic')
    p.add_run(
        '. Training concentrates entirely at the cheapest global producer. '
        'Inference organizes into regional hubs: '
        'Kosovo serves Southeastern Europe, Algeria the Mediterranean, '
        'Finland the Baltics, Canada North America. '
        'A sovereignty premium '
    )
    omath(p, [_v('\u03BB'), _t(f' = {LAMBDA:.0%}')])
    p.add_run(
        f' shifts {all_sov.get("full domestic", 0)} countries ({pct_sov_dom}%) to '
    )
    add_italic(p, 'domestic')
    p.add_run(f', {all_sov.get("import training + build inference", 0)} ({pct_sov_hyb}%) to ')
    add_italic(p, 'hybrid')
    p.add_run(f', and only {all_sov.get("full import", 0)} ({pct_sov_imp}%) remain ')
    add_italic(p, 'full import')
    p.add_run(
        '. The sovereignty premium is particularly powerful for inference: since the latency '
        'markup within Europe is moderate (10\u201340 ms, adding 1\u20133%), even a small domestic '
        'preference tips the decision away from importing.'
    )

    # Lambda calibration
    p, cur = mkp(doc, body, cur)
    p.add_run('In practice, ')
    omath(p, [_v('\u03BB')])
    p.add_run(
        ' is bilateral and heterogeneous. Between allies with mutual data adequacy '
        'agreements (e.g., EU member states), the effective sovereignty premium may be near zero. '
        'Between geopolitical adversaries, it is effectively infinite: the United States '
        'would not source training from Iran regardless of cost, and current sanctions make '
        'such transactions illegal. Our uniform '
    )
    omath(p, [_v('\u03BB'), _t(' = 10%')])
    p.add_run(
        ' should therefore be understood as an average over non-adversarial country pairs. '
        'Although Iran has the lowest production cost in our calibration, this reflects heavily '
        'subsidized electricity ($0.005/kWh) and is largely irrelevant for trade because '
        'international sanctions exclude it from most markets. In a model with bilateral '
    )
    omath(p, [_msub('\u03BB', 'jk')])
    p.add_run(
        ', sanctioned countries would be excluded from serving most demand centers, '
        'and the effective cheapest suppliers become Turkmenistan ($1.11/hr) and '
        'Kyrgyzstan ($1.13/hr).'
    )

    # Demand-weighted trade flows
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Demand-weighted trade flows. ')
    tr_ex = demand_data["train_revenue_ex"]
    ir = demand_data["inf_revenue"]
    # Top training exporter (excluding sanctioned)
    top_train_ex = sorted(tr_ex.items(), key=lambda x: -x[1])
    top_train_iso, top_train_share = top_train_ex[0]
    top_train_co = next(r["country"] for r in cal if r["iso3"] == top_train_iso)
    p.add_run(
        f'Weighting the sourcing patterns by demand shares (equation 5), '
        f'training concentrates at the single cheapest producer. '
        f'Since Iran\u2019s low electricity price ($0.005/kWh) reflects heavy subsidies and the '
        f'country faces broad international sanctions, we report training flows excluding '
        f'sanctioned producers: {top_train_co} '
        f'(${demand_data["costs_dict"][top_train_iso]:.2f}/hr) captures '
        f'{top_train_share * 100:.1f}% of global training demand. '
    )
    # Top inference exporters
    top_inf = sorted(ir.items(), key=lambda x: -x[1])
    top5_inf = top_inf[:5]
    inf_labels = []
    for iso, share in top5_inf:
        co = next(r["country"] for r in cal if r["iso3"] == iso)
        inf_labels.append(f'{co} ({share * 100:.0f}%)')
    p.add_run(
        f'Inference is far more dispersed: the top five inference suppliers are '
        f'{", ".join(inf_labels)}, together serving '
        f'{sum(s for _, s in top5_inf) * 100:.0f}% of global inference demand. '
        f'The demand-weighted HHI for training is '
        f'{demand_data["hhi_t"]:.2f}, versus {demand_data["hhi_i"]:.2f} for inference, '
        'confirming that training is far more concentrated than inference.'
    )

    # Training concentration: limitation + bounded-capacity extension
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'That all training demand pools at one producer is an artifact of the model\u2019s '
        'assumptions\u2014homogeneous FLOPs, zero transport costs, and no capacity limits. '
        'In practice, training would be spread across suppliers by power-grid constraints '
        '(no single country can absorb unlimited demand), risk diversification '
        '(firms hedge against outages by splitting workloads), differences in regulatory and '
        'data-governance regimes, and long-term contracts. '
        'A simple bounded-capacity extension makes the point: if each country can supply at most '
        'a fixed share of global demand, proportional to its power generation capacity, '
        'buyers fill the cheapest supplier first, then spill over to the next cheapest, and '
        'so on. Training demand then spreads across the low-cost tier rather '
        'than collapsing to a single country, and '
    )
    omath(p, [_msub('HHI', 'T')])
    p.add_run(
        ' falls sharply. Inference, already dispersed by latency, is less affected. '
        'We keep the unconstrained formulation for tractability. The cost rankings and '
    )
    omath(p, [_msub('HHI', 'T'), _t(' \u2265 '), _msub('HHI', 'I')])
    p.add_run(' survive under bounded capacity.')

    # Revenue for developing countries
    p, cur = mkp(doc, body, cur)
    kgz_clients = demand_data["kgz_inf_clients"]
    kgz_total = sum(w for _, _, w in kgz_clients)
    kgz_client_names = [co for _, co, _ in sorted(kgz_clients, key=lambda x: -x[2]) if co != "Kyrgyzstan"]
    names = kgz_client_names[:3]
    if len(names) <= 2:
        kgz_list = " and ".join(names)
    else:
        kgz_list = f'{", ".join(names[:-1])}, and {names[-1]}'
    p.add_run(
        f'Among developing countries, Kyrgyzstan captures {kgz_total:.0f}% of global '
        f'inference demand by serving {kgz_list}\u2014a striking '
        'result for a country with GDP under $15 billion. '
        'Algeria serves as the inference hub for Western Europe, capturing '
        f'{ir.get("DZA", 0) * 100:.0f}% of global inference demand from 14 European '
        'countries including Germany, France, the United Kingdom, and Italy. '
        'These patterns show how cheap-energy developing countries can earn export '
        'revenue from much larger economies.'
    )

    # Counterfactual
    p, cur = mkp(doc, body, cur)
    p.add_run(
        f'Doubling the sovereignty premium to 20% shifts {demand_data["extra_dom"]} '
        f'additional countries to domestic training production, reducing '
        f'the share of global training demand available to foreign producers '
        f'from {demand_data["export_share_10"] * 100:.0f}% to '
        f'{demand_data["export_share_20"] * 100:.0f}%. '
        'At 20%, nearly all countries produce training domestically, and the '
        'export market for training effectively disappears. '
        'Inference export revenue is more resilient to sovereignty premia because '
        'the latency advantage of proximity partially insulates regional hubs.'
    )

    # Major consumer markets
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Major demand centers. ')
    usa_inf = reg.get('USA', {}).get('best_inf_source', 'CAN')
    usa_inf_cost = reg.get('USA', {}).get('best_inf_cost', '1.190')
    deu_inf = reg.get('DEU', {}).get('best_inf_source', 'DZA')
    deu_inf_cost = reg.get('DEU', {}).get('best_inf_cost', '1.180')
    gbr_inf_cost = reg.get('GBR', {}).get('best_inf_cost', '1.176')
    fra_inf_cost = reg.get('FRA', {}).get('best_inf_cost', '1.174')
    chn_inf = reg.get('CHN', {}).get('best_inf_source', 'KGZ')
    chn_inf_cost = reg.get('CHN', {}).get('best_inf_cost', '1.161')
    rus_inf_cost = reg.get('RUS', {}).get('best_inf_cost', '1.179')
    sau_inf_cost = reg.get('SAU', {}).get('best_inf_cost', '1.168')
    are_inf_cost = reg.get('ARE', {}).get('best_inf_cost', '1.170')
    p.add_run(
        'The model\u2019s predictions vary across major AI demand centers because '
        'each faces a different latency geography. '
        f'For the United States, the pure-cost optimum sources training from the cheapest '
        f'available producer and inference from {usa_inf} (${float(usa_inf_cost):.2f}/hr), reflecting Canada\u2019s '
        'combination of low cost and minimal cross-border latency. '
        f'For major European demand centers, inference is sourced from {deu_inf}: '
        f'Germany at ${float(deu_inf_cost):.2f}/hr, '
        f'the United Kingdom at ${float(gbr_inf_cost):.2f}/hr, '
        f'and France at ${float(fra_inf_cost):.2f}/hr; '
        'Algeria\u2019s subsidized electricity and moderate Mediterranean latency make it the '
        'European inference hub. '
        f'For China, the cheapest inference source is {chn_inf} '
        f'(${float(chn_inf_cost):.2f}/hr), a bordering country with hydropower-based electricity. '
        f'Russia produces inference domestically even under pure cost minimization '
        f'(${float(rus_inf_cost):.2f}/hr), as its low energy costs offset moderate latency. '
        f'The Gulf states present another pattern: Saudi Arabia (${float(sau_inf_cost):.2f}/hr) '
        f'and the UAE (${float(are_inf_cost):.2f}/hr) also produce inference domestically, '
        'benefiting from subsidized energy and expanding datacenter capacity. '
        'These patterns illustrate the model\u2019s core prediction: inference organizes around '
        'latency-bounded regional hubs, and each major market has a distinct optimal supplier '
        'determined by geography. '
        f'With a sovereignty premium of {LAMBDA:.0%}, the United States, China, Russia, and '
        'the Gulf states all shift to full domestic production, while Germany still imports '
        'training but produces inference domestically.'
    )

    # Governance and Political Economy discussion (condensed to 3 paragraphs)
    # Para 1: Caveats + Institutional factors
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Governance and political economy. ')
    p.add_run(
        'These cost rankings assume globally uniform hardware prices and abstract from '
        'institutional heterogeneity. In practice, GPU export controls raise effective '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' for Iran, Russia, and Belarus; grid reliability varies widely; the EU\u2019s GDPR and '
        'AI Act create hard barriers for inference on personal data; and agglomeration '
        'economies (Krugman 1991) favor established hubs despite higher costs. '
        'Data center investments are large, long-lived, and immobile, so the viability of a '
        'country as a compute exporter depends on institutional factors not captured by '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' alone. Several of the cheapest producers in our calibration (Iran, Turkmenistan, '
        'Uzbekistan) rank poorly on property rights and rule of law indices, and subsidized '
        'electricity prices may be politically fragile. The fixed entry cost '
    )
    omath(p, [_msub('F', 'j')])
    p.add_run(
        ' should therefore be interpreted as including an institutional premium required '
        'to attract foreign capital.'
    )

    # Para 2: Regulatory, geopolitical barriers, and infrastructure
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The EU\u2019s GDPR and AI Act segment the compute market along regulatory lines, '
        'reinforcing the sovereignty premium '
    )
    omath(p, [_v('\u03BB')])
    p.add_run(
        ' as a structural feature. U.S. export controls on advanced GPUs raise the effective '
        'hardware cost '
    )
    omath(p, [_v('\u03C1')])
    p.add_run(
        ' for sanctioned countries, potentially offsetting any electricity cost advantage and '
        'discouraging long-term investment. '
        'Grid reliability further narrows the set of viable exporters: countries with frequent '
        'outages face backup-generation costs not reflected in headline electricity prices, so '
        'effective costs should be understood as reliability-adjusted. Taken together, these '
        'governance factors suggest that viable compute exporters are a strict subset of low-cost '
        'producers: those that combine cheap energy with adequate institutional quality, '
        'such as the Nordic countries, Canada, and parts of the Gulf and Central Asia. '
        'Water is another constraint. Evaporative cooling consumes large volumes, and several '
        'of the cheapest producers (Iran, Turkmenistan, Egypt, Saudi Arabia) are water-scarce. '
        'Liquid cooling reduces water needs but does not eliminate them.'
    )

    # Subsidy sustainability
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'A deeper issue concerns the sustainability of subsidized electricity prices at an export '
        'scale. Several of the cheapest producers in our calibration, including Iran '
        '($0.005/kWh), Turkmenistan ($0.01/kWh), and Egypt ($0.038/kWh), benefit from domestic '
        'energy subsidies. At a small scale, hosting a data center is a net gain: the facility '
        'pays the subsidized tariff and generates export revenue, employment, and tax receipts. '
        'At export scale, however, the fiscal arithmetic reverses. A country that becomes a '
        'major training hub would channel hundreds of megawatts of subsidized electricity into '
        'serving foreign demand, effectively transferring the subsidy to foreign AI companies. '
        'The implicit fiscal cost, the gap between the subsidized tariff and the long-run '
        'marginal cost of generation, would scale linearly with export volume. Governments '
        'would face a choice between raising electricity prices for data centers (eroding the cost '
        'advantage), maintaining subsidies at growing fiscal cost, or capping data center capacity. '
        'This suggests that the lowest-cost producers in our calibration are unlikely to serve '
        'global training demand at the volumes implied by pure cost minimization. The '
        'sustainable export frontier is defined not by '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' alone but by the price at which electricity can be supplied without fiscal subsidy.'
    )

    # Robustness: cost-reflective prices
    adj_top5 = demand_data["adj_top5"]
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'How sensitive are the rankings to subsidized electricity? We replace observed tariffs '
        f'in {demand_data["n_adjusted"]} subsidy-dependent countries '
        '(Iran, Turkmenistan, Algeria, Egypt, and Uzbekistan) with estimated '
        'long-run marginal costs of generation at opportunity-cost fuel prices, '
        'specifically $0.05/kWh for gas exporters (the approximate cost of '
        'combined-cycle generation at export-parity fuel prices) and '
        '$0.075\u2013$0.085/kWh for gas importers with subsidies. '
        f'The five cheapest producers become '
        f'{adj_top5[0][1]} (${adj_top5[0][2]:.2f}/hr), '
        f'{adj_top5[1][1]} (${adj_top5[1][2]:.2f}/hr), '
        f'{adj_top5[2][1]} (${adj_top5[2][2]:.2f}/hr), '
        f'{adj_top5[3][1]} (${adj_top5[3][2]:.2f}/hr), '
        f'and {adj_top5[4][1]} (${adj_top5[4][2]:.2f}/hr). '
        'Hydropower countries (Kyrgyzstan, Canada, Norway) and cold-climate '
        'producers (Russia, Kosovo) move up, and Iran drops from first '
        f'to {demand_data["adj_rank_map"]["IRN"]}th. '
        'The model\u2019s structure is unchanged: training still goes to the cheapest '
        'producer, inference still clusters in regional hubs, and '
        'the sovereignty premium still pushes countries toward domestic production. '
        'What changes is who sits at the top: hydropower replaces '
        'subsidized gas. Cheap electricity still confers comparative advantage in compute, '
        'but the advantage has to be real rather than fiscal.'
    )

    # Model extensions (condensed to 1 paragraph)
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Model extensions. ')
    p.add_run(
        'The model can be extended in several directions. (i) bounded capacity per country, '
        'so that grid limits spread training across multiple low-cost producers as discussed '
        'above, (ii) stochastic disruptions (grid outages, political instability) that '
        'give buyers a reason to split workloads across providers, '
        '(iii) heterogeneous demand segmented by latency tolerance, (iv) endogenous electricity '
        'prices with upward-sloping supply curves, (v) carbon pricing that introduces a '
        '\u201Cgreen premium\u201D for hydropower-rich countries, (vi) strategic interaction among '
        'oligopolistic providers, and (vii) governance as a multiplicative cost shifter on '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run('.')


def write_conclusion(doc, body, hmap, demand_data):
    print("Rewriting Section 8 (Conclusion)...")
    sec8 = hmap['5']
    refs = hmap['refs']
    all_now = list(body)
    c8i = all_now.index(sec8)
    ri = all_now.index(refs)
    for el in all_now[c8i + 1:ri]:
        body.remove(el)

    p, cur_concl = mkp(doc, body, sec8)
    p.add_run(
        'We develop a trade model for computing services (FLOPs) in which countries '
        'produce and export computing capacity based on their electricity prices, climate, and '
        'construction costs. We distinguish two service types, latency-insensitive training '
        'and latency-sensitive inference, and introduce a sovereignty premium to capture '
        'governments\u2019 preference for domestic data processing. We derive buyers\u2019 optimal '
        'sourcing decisions and producers\u2019 entry conditions, '
        'and establish three comparative static results: country-specific sovereignty '
        'thresholds, the nesting of training exporters within inference exporters, '
        'and greater concentration in training than in inference markets. '
        'We calibrate the model for 86 countries using data on electricity '
        'prices, temperatures, construction costs, bilateral latencies, and GDP-based demand shares.'
    )

    p, cur_concl = mkp(doc, body, cur_concl)
    p.add_run(
        'Across 86 countries, cheap-energy peripheries serve as FLOP exporters for training, '
        'while inference organizes into regional hubs bounded by latency. '
        'The sovereignty premium rationalizes widespread domestic investment, shifting '
        'the majority of countries from import to domestic production, '
        f'at a demand-weighted welfare cost of {demand_data["welfare_pct"]:.0f}% of '
        'average compute spending. '
        'The model generates a country taxonomy (full importers, training exporters, '
        'inference hubs, and hybrid regimes) that maps onto observed investment patterns.'
    )

    p, cur_concl = mkp(doc, body, cur_concl)
    p.add_run(
        'For developing countries, the results point to a new avenue for economic participation '
        'in the global economy. Countries like Kyrgyzstan, Uzbekistan, and Egypt, which rank '
        'among the cheapest FLOP producers in our calibration, could convert cheap electricity '
        'into a high-value digital export without building a domestic AI research ecosystem. '
        'FLOP exporting is the digital equivalent of resource-based industrialization, but '
        'with the advantage that the \u201Cresource\u201D (electricity) is renewable and the product '
        '(compute) serves the fastest-growing sector of the world economy.'
    )

    p, _ = mkp(doc, body, cur_concl)
    p.add_run(
        'The policy implications are asymmetric. Restricting training imports raises costs '
        'without a proximity benefit, since training is latency-insensitive. Supporting domestic '
        'inference has a genuine latency rationale, but is less justified for countries close '
        'to low-cost neighbors. For developing countries seeking to enter the compute export '
        'market, the binding constraints are not technological but institutional, namely reliable power '
        'grids, political stability, data governance frameworks, and international connectivity '
        'determine whether cost advantages translate into actual exports.'
    )


def write_appendix(doc, body, refs, eca_cal, non_eca_cal, reg):
    print("Inserting Appendix (Table A1)...")

    # Page break before Appendix
    pb_el = add_page_break(doc, body, refs.getprevious())

    # Insert Appendix heading before References
    cur_app = mkh(doc, body, pb_el, 'Appendix', level=1)

    headers = ["Rank", "Country", "Total Cost\n($/hr)", "Electricity\nCost ($/hr)",
               "Construction\nCost ($/hr)", "PUE", "Elec. Price\n($/kWh)", "Pure-Cost\nRegime"]
    rows = []
    for r_row in eca_cal:
        co = r_row["country"]
        if len(co) > 22:
            co = co[:21] + "."
        regime = reg.get(r_row["iso3"], {}).get("regime", "n/a")
        rs = {"full import": "import", "import training + build inference": "hybrid",
              "full domestic": "domestic"}.get(regime, regime)
        rows.append((r_row["rank"], co,
                     f'${float(r_row["c_j_total"]):.2f}', f'${float(r_row["c_j_electricity"]):.3f}',
                     f'${float(r_row["c_j_construction"]):.3f}', f'{float(r_row["pue"]):.2f}',
                     f'${float(r_row["p_E_usd_kwh"]):.3f}', rs))

    rows.append(("", "Non-ECA comparators", "", "", "", "", "", ""))
    for r_row in non_eca_cal:
        co = r_row["country"]
        if len(co) > 22:
            co = co[:21] + "."
        regime = reg.get(r_row["iso3"], {}).get("regime", "n/a")
        rs = {"full import": "import", "import training + build inference": "hybrid",
              "full domestic": "domestic"}.get(regime, regime)
        rows.append((r_row["rank"], co,
                     f'${float(r_row["c_j_total"]):.2f}', f'${float(r_row["c_j_electricity"]):.3f}',
                     f'${float(r_row["c_j_construction"]):.3f}', f'{float(r_row["pue"]):.2f}',
                     f'${float(r_row["p_E_usd_kwh"]):.3f}', rs))

    cw = [500, 1800, 950, 950, 950, 550, 950, 950]
    # Table title with bookmark + back-link to in-text mention
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(6)
    tp.paragraph_format.space_after = Pt(3)
    tp.paragraph_format.first_line_indent = Inches(0)
    tp._element.append(make_bookmark(101, 'TableA1'))
    # Custom hyperlink with bold + size formatting for table title
    hl_t = OxmlElement('w:hyperlink')
    hl_t.set(qn('w:anchor'), 'TableA1txt')
    hl_t.set(qn('w:history'), '1')
    r_t = OxmlElement('w:r')
    rPr_t = OxmlElement('w:rPr')
    b_t = OxmlElement('w:b')
    rPr_t.append(b_t)
    sz_t = OxmlElement('w:sz')
    sz_t.set(qn('w:val'), '20')
    rPr_t.append(sz_t)
    clr_t = OxmlElement('w:color')
    clr_t.set(qn('w:val'), LINK_COLOR)
    uu_t = OxmlElement('w:u')
    uu_t.set(qn('w:val'), 'single')
    rPr_t.append(clr_t)
    rPr_t.append(uu_t)
    r_t.append(rPr_t)
    t_t = OxmlElement('w:t')
    t_t.text = 'Table A1'
    r_t.append(t_t)
    hl_t.append(r_t)
    tp._element.append(hl_t)
    tp._element.append(make_bookmark_end(101))
    run_tt = tp.add_run('. Unit cost of FLOP production by country (H100 GPU-hour)')
    run_tt.bold = True
    run_tt.font.size = Pt(10)
    tbl_el = tp._element
    body.remove(tbl_el)
    cur_app.addnext(tbl_el)
    add_table(doc, body, tbl_el, headers, rows, cw)


def write_references(doc, body, refs):
    print("Updating references...")

    # Page break before References
    add_page_break(doc, body, refs.getprevious())

    all_now = list(body)
    ri = all_now.index(refs)
    ref_els = []
    ref_txts = []
    for i in range(ri + 1, len(all_now)):
        el = all_now[i]
        if el.tag == qn('w:p'):
            t = "".join(r.text or "" for r in el.findall(f'.//{qn("w:t")}'))
            if t.strip():
                ref_txts.append(t.strip())
                ref_els.append(el)
        elif el.tag == qn('w:sectPr'):
            break

    new_refs = [
        'Anderson, J. E., and E. van Wincoop. (2003). \u201CGravity with Gravitas: '
        'A Solution to the Border Puzzle.\u201D '
        'American Economic Review, 93(1): 170\u2013192.',

        'Antr\u00E0s, P., and E. Helpman. (2004). \u201CGlobal Sourcing.\u201D '
        'Journal of Political Economy, 112(3): 552\u2013580.',

        'Brainard, S. L. (1997). \u201CAn Empirical Assessment of the Proximity-Concentration '
        'Trade-off.\u201D American Economic Review, 87(4): 520\u2013544.',

        'Deloitte. (2025). \u201CTechnology, Media, and Telecommunications Predictions 2026.\u201D '
        'Deloitte Insights.',

        'EIA. (2025). Electric Power Monthly. U.S. Energy Information Administration.',

        'Epoch AI. (2024). \u201CThe Training Compute of Notable AI Models.\u201D epochai.org.',

        'EPRI. (2024). \u201CPowering Intelligence: Analyzing AI and Data Center Energy '
        'Consumption.\u201D Electric Power Research Institute.',

        'Eurostat. (2025). Electricity Prices for Non-Household Consumers '
        '(nrg_pc_205). Luxembourg: Eurostat.',

        'Firebird. (2026). \u201CPhase 2 of Armenia AI Megaproject, Scaling to $4 Billion '
        'and 50,000 GPUs.\u201D Press release, January 2026.',

        'Flucker, S., R. Tozer, and R. Whitehead. (2013). \u201CData Centre Energy Efficiency '
        'Analysis.\u201D Building Services Engineering Research and Technology, 34(1): 103\u2013117.',

        'GlobalPetrolPrices. (2025). Electricity Prices Around the World. '
        'globalpetrolprices.com.',

        'Goldfarb, A., and D. Trefler. (2018). \u201CAI and International Trade.\u201D '
        'In The Economics of Artificial Intelligence. Chicago: Univ. of Chicago Press, '
        'pp. 463\u2013492.',

        'Goldman Sachs. (2024). \u201CAI Is Poised to Drive 165% Increase in Data Center '
        'Power Demand.\u201D Goldman Sachs Research.',

        'Grossman, G. M., and E. Rossi-Hansberg. (2008). \u201CTrading Tasks: A Simple Theory '
        'of Offshoring.\u201D American Economic Review, 98(5): 1978\u20131997.',

        'Hausmann, R., J. Hwang, and D. Rodrik. (2007). \u201CWhat You Export Matters.\u201D '
        'Journal of Economic Growth, 12(1): 1\u201325.',

        'Helpman, E., Melitz, M. J., and S. R. Yeaple. (2004). \u201CExport Versus FDI with '
        'Heterogeneous Firms.\u201D American Economic Review, 94(1): 300\u2013316.',

        'Hersbach, H., et al. (2020). \u201CThe ERA5 Global Reanalysis.\u201D '
        'Quarterly Journal of the Royal Meteorological Society, 146(730): 1999\u20132049.',

        'Hummels, D., and G. Schaur. (2013). \u201CTime as a Trade Barrier.\u201D '
        'American Economic Review, 103(7): 2935\u20132959.',

        'IEA. (2025). \u201CEnergy Demand from AI.\u201D Published online at iea.org.',

        'Korinek, A., and J. Stiglitz. (2021). \u201CAI, Globalization, and Strategies for '
        'Economic Development.\u201D NBER Working Paper No. 28453.',

        'Krugman, P. (1991). \u201CIncreasing Returns and Economic Geography.\u201D '
        'Journal of Political Economy, 99(3): 483\u2013499.',

        'Lim\u00E3o, N., and A. J. Venables. (2001). \u201CInfrastructure, Geographical '
        'Disadvantage, Transport Costs, and Trade.\u201D '
        'World Bank Economic Review, 15(3): 451\u2013479.',

        'Liu, Z., A. Wierman, Y. Chen, B. Raber, and J. Moriarty. (2023). '
        '\u201CSustainability of Data Center Digital Twins.\u201D '
        'Proceedings of ACM e-Energy, pp. 178\u2013189.',

        'NVIDIA. (2024). NVIDIA H100 Tensor Core GPU Datasheet. nvidia.com.',

        'Oltmanns, J., D. Krcmarik, and R. Gatti. (2021). \u201CData Centre Site Selection.\u201D '
        'Journal of Property Investment & Finance, 39(1): 55\u201372.',

        'Samuelson, P. (1954). \u201CThe Transfer Problem and Transport Costs, II.\u201D '
        'Economic Journal, 64(254): 264\u2013289.',

        'Turner & Townsend. (2025). Data Centre Construction Cost Index 2025. '
        'turnerandtownsend.com.',

        'UNCTAD. (2025). Technology and Innovation Report 2025. Geneva: United Nations.',

        'Uptime Institute. (2024). \u201CData Center Staffing: Trends and Best Practices.\u201D',

        'WonderNetwork. (2024). Global Ping Statistics. wondernetwork.com.',

        'World Bank. (2024). World Development Indicators. Washington, DC.',

        'Lehdonvirta, V., B. Wu, and Z. Hawkins. (2024). \u201CCompute North vs. Compute South: '
        'The Uneven Possibilities of Compute-Based AI Governance Around the Globe.\u201D '
        'Proceedings of the AAAI/ACM Conference on AI, Ethics, and Society, 7(1): 828\u2013838.',

        'Pilz, K. F., Y. Mahmood, and L. Heim. (2025). AI\u2019s Power Requirements Under '
        'Exponential Growth. Santa Monica, CA: RAND Corporation, RR-A3572-1.',

        'Sastry, G., L. Heim, et al. (2024). \u201CComputing Power and the Governance of '
        'Artificial Intelligence.\u201D arXiv:2402.08797.',
    ]

    existing_auth = {r.split(',')[0].split('.')[0].strip().lower() for r in ref_txts}
    for r in new_refs:
        a = r.split(',')[0].split('.')[0].strip().lower()
        if a not in existing_auth:
            ref_txts.append(r)
    ref_txts.sort(key=lambda x: x.lower())
    for el in ref_els:
        body.remove(el)

    # Build reverse map: reference text prefix -> key
    def find_ref_key(ref_text):
        """Find the citation key for a reference text."""
        for key, prefix in REF_KEY_MAP.items():
            if ref_text.startswith(prefix):
                return key
        return None

    bm_id_refs = [500]  # bookmark IDs for references
    cur = refs
    for rt in ref_txts:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.0
        italic_portion = find_italic_portion(rt)
        key = find_ref_key(rt)
        if key:
            # Add bookmark target for in-text citation links
            p._element.append(make_bookmark(bm_id_refs[0], key))
            # Author portion as hyperlink back to in-text citation
            split_pos = rt.find('\u201C')
            if split_pos < 0:
                split_pos = rt.find('(', 10)
            if split_pos < 0:
                split_pos = len(rt)
            author_part = rt[:split_pos]
            remaining = rt[split_pos:]
            p._element.append(make_hyperlink(f'{key}txt', author_part))
            if remaining:
                _write_ref_segments(p, remaining, italic_portion)
            p._element.append(make_bookmark_end(bm_id_refs[0]))
            bm_id_refs[0] += 1
        else:
            _write_ref_segments(p, rt, italic_portion)
        el = p._element
        body.remove(el)
        cur.addnext(el)
        cur = el
    print(f"  {len(ref_txts)} references")


def link_citations(body):
    print("Linking citations...")
    bm_id_cite = [200]
    passes = 0
    while True:
        n = link_citations_pass(body, CITE_MAP, bm_id_cite)
        passes += 1
        if n == 0 or passes > 10:
            break
    print(f"  {bm_id_cite[0] - 200} citation links created in {passes} passes")


def apply_formatting(doc, body, refs, title_el, author_el, abs_text_el):
    print("Applying formatting...")
    # Identify reference paragraphs to protect their spacing
    refs_idx = list(body).index(refs)
    ref_elements = set()
    for el in list(body)[refs_idx + 1:]:
        if el.tag == qn('w:sectPr'):
            break
        if el.tag == qn('w:p'):
            ref_elements.add(el)

    # Paragraphs to protect from global formatting (centered title page elements)
    _protected = {title_el, author_el, abs_text_el}

    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        # Heading 1: blue, Times New Roman
        if style == 'Heading 1':
            for run in p.runs:
                run.font.color.rgb = HEADING_BLUE
                run.font.name = TIMES_NEW_ROMAN
            continue
        # Heading 2: blue italic, Times New Roman
        if style == 'Heading 2':
            for run in p.runs:
                run.font.color.rgb = HEADING_BLUE
                run.italic = True
                run.font.name = TIMES_NEW_ROMAN
            continue
        if 'Heading' not in style and p.text.strip():
            # Skip title page elements (centered)
            if p._element not in _protected:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if p.paragraph_format.first_line_indent is None or p.paragraph_format.first_line_indent > 0:
                    p.paragraph_format.first_line_indent = Inches(0)
            # Subtitle runs: italic first run ending with ". " → font 12, TNR, not bold
            runs = [r for r in p.runs if r.text.strip()]
            if runs and runs[0].italic and runs[0].text.rstrip().endswith('.'):
                runs[0].font.size = Pt(12)
                runs[0].font.name = TIMES_NEW_ROMAN
                runs[0].bold = False
            # Preserve reference formatting (hanging indent + 4pt spacing)
            if p._element in ref_elements:
                continue
            # Preserve title page spacing
            if p._element in _protected:
                continue
            p.paragraph_format.space_before = Pt(0)
            # Preserve Pt(2) spacing on paragraphs immediately before equations
            if p.paragraph_format.space_after is None or p.paragraph_format.space_after >= Pt(8):
                p.paragraph_format.space_after = Pt(8)


def add_page_numbers_and_break(doc, body, kw_el):
    print("Adding page numbers...")
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    # Default footer: right-aligned page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.clear()
    # Insert PAGE field: w:fldSimple or fldChar sequence
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), ' PAGE ')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '1'
    r.append(t)
    fld.append(r)
    fp._element.append(fld)

    # First page footer: empty (no page number on title page)
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    if first_footer.paragraphs:
        first_footer.paragraphs[0].clear()

    # ═══════════════════════════════════════════════════════════════════════
    # STEP 17: PAGE BREAK AFTER ABSTRACT (title on separate page)
    # ═══════════════════════════════════════════════════════════════════════
    print("Adding page break after keywords...")
    add_page_break(doc, body, kw_el)


def main():
    # ═══════════════════════════════════════════════════════════════════════
    # LOAD DATA (v3)
    # ═══════════════════════════════════════════════════════════════════════

    print("Loading data...")
    cal = []
    with open(DATA / "calibration_results_v3.csv", encoding="utf-8") as f:
        cal = list(csv.DictReader(f))
    reg = {}
    with open(DATA / "calibration_regimes_v3.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            reg[row["iso3"]] = row
    # World Bank operational ECA region (developing Europe & Central Asia)
    eca = {
        'ALB', 'ARM', 'AZE', 'BLR', 'BIH', 'BGR', 'HRV', 'CZE', 'EST',
        'GEO', 'HUN', 'KAZ', 'XKX', 'KGZ', 'LVA', 'LTU', 'MDA', 'MNE',
        'MKD', 'POL', 'ROU', 'RUS', 'SRB', 'SVK', 'SVN', 'TJK', 'TUR',
        'TKM', 'UKR', 'UZB',
    }

    eca_cal = [row for row in cal if row["iso3"] in eca]
    non_eca_cal = [row for row in cal if row["iso3"] not in eca]
    n_eca = len(eca_cal)
    n_total = len(cal)

    eca_reg = {"full import": 0, "import training + build inference": 0,
               "full domestic": 0, "build training + import inference": 0}
    eca_sov = dict(eca_reg)
    all_reg = dict(eca_reg)
    all_sov = dict(eca_reg)
    for row in cal:
        iso = row["iso3"]
        if iso in reg:
            rr = reg[iso]["regime"]
            rs = reg[iso]["regime_with_sovereignty"]
            if rr in all_reg:
                all_reg[rr] += 1
            if rs in all_sov:
                all_sov[rs] += 1
            if iso in eca:
                if rr in eca_reg:
                    eca_reg[rr] += 1
                if rs in eca_sov:
                    eca_sov[rs] += 1

    print(f"  Total: {n_total}, ECA: {n_eca}")
    print(f"  All regimes: {dict((k, v) for k, v in all_reg.items() if v)}")

    # ═══════════════════════════════════════════════════════════════════════
    # DEMAND CALIBRATION (GDP-based shares)
    # ═══════════════════════════════════════════════════════════════════════
    print("Loading GDP and population data...")
    gdp_pcap = {}
    with open(DATA / "wb_gdp_per_capita_ppp_2023.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            gdp_pcap[row["iso3"]] = float(row["gdp_pcap_ppp_2023"])
    pop_data = {}
    with open(DATA / "wb_population_2023.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            pop_data[row["iso3"]] = int(row["population_2023"])

    # GDP_k for each calibration country
    gdp_k = {}
    for row in cal:
        iso = row["iso3"]
        if iso in gdp_pcap and iso in pop_data:
            gdp_k[iso] = gdp_pcap[iso] * pop_data[iso]
    total_gdp = sum(gdp_k.values())
    omega = {iso: g / total_gdp for iso, g in gdp_k.items()}

    # Top demand centers
    sorted_omega = sorted(omega.items(), key=lambda x: -x[1])
    top5_labels = []
    for iso, w in sorted_omega[:5]:
        co = next(r["country"] for r in cal if r["iso3"] == iso)
        top5_labels.append((iso, co, w))
    top5_share = sum(w for _, _, w in top5_labels)

    # Training export revenue shares
    train_revenue = {}
    for iso in gdp_k:
        if iso in reg:
            src = reg[iso]["best_train_source"]
            train_revenue[src] = train_revenue.get(src, 0) + omega[iso]

    # Inference export revenue shares
    inf_revenue = {}
    for iso in gdp_k:
        if iso in reg:
            src = reg[iso]["best_inf_source"]
            inf_revenue[src] = inf_revenue.get(src, 0) + omega[iso]

    # HHI
    hhi_t = sum(s**2 for s in train_revenue.values())
    hhi_i = sum(s**2 for s in inf_revenue.values())

    # Lambda* for each country
    costs_dict = {row["iso3"]: float(row["c_j_total"]) for row in cal}
    lambda_star = {}
    for iso, c_k in costs_dict.items():
        min_foreign = min(c for i, c in costs_dict.items() if i != iso)
        lambda_star[iso] = c_k / min_foreign - 1

    # Welfare cost of sovereignty
    welfare_train = 0
    welfare_inf = 0
    for iso in gdp_k:
        if iso in reg and iso in costs_dict:
            c_k = costs_dict[iso]
            best_train = float(reg[iso]["best_train_cost"])
            best_inf = float(reg[iso]["best_inf_cost"])
            c_k_inf = float(reg[iso]["P_I_domestic"])
            welfare_train += omega[iso] * max(0, c_k - best_train)
            welfare_inf += omega[iso] * max(0, c_k_inf - best_inf)
    welfare_total = welfare_train + welfare_inf
    weighted_avg_cost = sum(omega[iso] * costs_dict[iso]
                            for iso in gdp_k if iso in costs_dict)
    welfare_pct = welfare_total / weighted_avg_cost * 100

    # Counterfactual: doubling sovereignty to 20%
    min_cost = min(costs_dict.values())
    count_dom_10 = sum(
        1 for iso in gdp_k
        if iso in costs_dict and costs_dict[iso] <= 1.10 * min_cost)
    count_dom_20 = sum(
        1 for iso in gdp_k
        if iso in costs_dict and costs_dict[iso] <= 1.20 * min_cost)
    extra_dom = count_dom_20 - count_dom_10
    export_share_10 = sum(
        omega[iso] for iso in gdp_k
        if iso in costs_dict and costs_dict[iso] > 1.10 * min_cost)
    export_share_20 = sum(
        omega[iso] for iso in gdp_k
        if iso in costs_dict and costs_dict[iso] > 1.20 * min_cost)
    export_reduction_pp = (export_share_10 - export_share_20) * 100

    # Kyrgyzstan inference clients
    kgz_inf_clients = []
    for iso in gdp_k:
        if iso in reg and reg[iso]["best_inf_source"] == "KGZ":
            co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
            kgz_inf_clients.append((iso, co, omega[iso] * 100))

    # Training revenue excluding sanctioned countries (IRN)
    sanctioned = {'IRN'}
    train_revenue_ex = {}
    for iso in gdp_k:
        if iso in reg:
            best_src = min((j for j in costs_dict if j not in sanctioned),
                           key=lambda j: costs_dict[j])
            train_revenue_ex[best_src] = train_revenue_ex.get(best_src, 0) + omega[iso]

    # Build demand_data dict for passing to write functions
    demand_data = {
        "omega": omega, "sorted_omega": sorted_omega,
        "top5_labels": top5_labels, "top5_share": top5_share,
        "train_revenue": train_revenue, "inf_revenue": inf_revenue,
        "hhi_t": hhi_t, "hhi_i": hhi_i,
        "lambda_star": lambda_star, "costs_dict": costs_dict,
        "welfare_total": welfare_total, "welfare_pct": welfare_pct,
        "welfare_train": welfare_train, "welfare_inf": welfare_inf,
        "weighted_avg_cost": weighted_avg_cost,
        "count_dom_10": count_dom_10, "count_dom_20": count_dom_20,
        "extra_dom": extra_dom,
        "export_share_10": export_share_10, "export_share_20": export_share_20,
        "export_reduction_pp": export_reduction_pp,
        "kgz_inf_clients": kgz_inf_clients,
        "train_revenue_ex": train_revenue_ex,
    }
    print(f"  GDP data for {len(gdp_k)} countries, HHI_T={hhi_t:.4f}, HHI_I={hhi_i:.4f}")

    # ═══════════════════════════════════════════════════════════════════════
    # ROBUSTNESS: COST-REFLECTIVE ELECTRICITY PRICES
    # ═══════════════════════════════════════════════════════════════════════
    print("Computing subsidy-adjusted robustness check...")
    adj_costs = dict(costs_dict)  # copy baseline
    adj_changes = {}
    for iso, p_E_adj in SUBSIDY_ADJ.items():
        if iso not in adj_costs:
            continue
        row = next(r for r in cal if r["iso3"] == iso)
        p_E_orig = float(row["p_E_usd_kwh"])
        pue = float(row["pue"])
        delta_elec = pue * GAMMA * (p_E_adj - p_E_orig)
        adj_costs[iso] = costs_dict[iso] + delta_elec
        adj_changes[iso] = {
            "country": row["country"],
            "p_E_orig": p_E_orig, "p_E_adj": p_E_adj,
            "c_j_orig": costs_dict[iso], "c_j_adj": adj_costs[iso],
        }

    adj_ranked = sorted(adj_costs.items(), key=lambda x: x[1])
    adj_rank_map = {iso: rank for rank, (iso, _) in enumerate(adj_ranked, 1)}

    # Re-do training allocation under adjusted costs
    adj_cheapest = adj_ranked[0][0]
    adj_train_rev = {}
    for iso in gdp_k:
        adj_train_rev[adj_cheapest] = adj_train_rev.get(adj_cheapest, 0) + omega.get(iso, 0)
    adj_hhi_t = sum(s**2 for s in adj_train_rev.values())

    # Re-do inference allocation: need to re-find best inf source with adjusted costs
    adj_inf_rev = {}
    for iso_k in gdp_k:
        if iso_k not in reg:
            continue
        # Inference assignments depend on latency (unchanged by price adjustments).
        # Keep original inference sourcing.
        orig_src = reg[iso_k]["best_inf_source"]
        adj_inf_rev[orig_src] = adj_inf_rev.get(orig_src, 0) + omega.get(iso_k, 0)
    adj_hhi_i = sum(s**2 for s in adj_inf_rev.values())

    # Top 5 adjusted ranking
    adj_top5 = []
    for iso, c in adj_ranked[:5]:
        co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
        adj_top5.append((iso, co, c))

    demand_data["adj_top5"] = adj_top5
    demand_data["adj_hhi_t"] = adj_hhi_t
    demand_data["adj_hhi_i"] = adj_hhi_i
    demand_data["adj_cheapest"] = adj_cheapest
    demand_data["adj_cheapest_name"] = next(
        (r["country"] for r in cal if r["iso3"] == adj_cheapest), adj_cheapest)
    demand_data["adj_cheapest_cost"] = adj_costs[adj_cheapest]
    demand_data["adj_changes"] = adj_changes
    demand_data["adj_rank_map"] = adj_rank_map
    demand_data["n_adjusted"] = len(adj_changes)

    print(f"  Adjusted {len(adj_changes)} countries; new cheapest: "
          f"{demand_data['adj_cheapest_name']} (${adj_costs[adj_cheapest]:.3f}/hr)")
    for iso, co, c in adj_top5:
        flag = " *" if iso in adj_changes else ""
        print(f"    {adj_rank_map[iso]:>2}. {co:<24} ${c:.3f}/hr{flag}")

    # ═══════════════════════════════════════════════════════════════════════
    # LOAD v8 AND INDEX HEADINGS
    # ═══════════════════════════════════════════════════════════════════════

    print("\nLoading v8...")
    doc = Document(str(DOCS / "flop_trade_model_v8.docx"))
    body = doc.element.body
    all_el = list(body)
    init_footnotes(doc)

    hmap = {}
    for el in all_el:
        if el.tag == qn('w:p'):
            pPr = el.find(qn('w:pPr'))
            if pPr is not None:
                pS = pPr.find(qn('w:pStyle'))
                if pS is not None and 'Heading' in pS.get(qn('w:val'), ''):
                    ft = "".join(r.text or "" for r in el.findall(f'.//{qn("w:t")}'))
                    if '1.2' in ft:
                        hmap['1.2'] = el
                    elif '1.1' in ft:
                        hmap['1.1'] = el
                    elif '1.' in ft and 'Model' in ft:
                        hmap['1'] = el
                    elif '2.' in ft and 'Comp' in ft:
                        hmap['2'] = el
                    elif '3.' in ft and 'Make' in ft:
                        hmap['3'] = el
                    elif '4.' in ft and 'Calib' in ft:
                        hmap['4'] = el
                    elif '5.' in ft and 'Conc' in ft:
                        hmap['5'] = el
                    elif ft.strip() == 'References':
                        hmap['refs'] = el
                    elif ft.strip() == 'Abstract':
                        hmap['abs'] = el

    # ═══════════════════════════════════════════════════════════════════════
    # STEPS
    # ═══════════════════════════════════════════════════════════════════════

    title_el, author_el, abs_text_el, kw_el = write_title_and_abstract(doc, body, all_el, hmap)
    write_introduction(doc, body, hmap)
    write_literature(doc, body, hmap)
    write_model_opening(doc, body, all_el)
    write_production_technology(doc, body, hmap)
    write_trade_costs(doc, body, hmap)
    renumber_sections(hmap)
    write_comparative_advantage(doc, body, hmap, demand_data)
    write_make_or_buy(doc, body, hmap, demand_data)
    write_data_section(doc, body, hmap, demand_data)
    write_calibration(doc, body, hmap, cal, reg, n_eca, n_total, all_reg, all_sov, demand_data)
    write_conclusion(doc, body, hmap, demand_data)

    refs = hmap['refs']
    write_appendix(doc, body, refs, eca_cal, non_eca_cal, reg)
    write_references(doc, body, refs)
    link_citations(body)
    apply_formatting(doc, body, refs, title_el, author_el, abs_text_el)
    add_page_numbers_and_break(doc, body, kw_el)

    # ═══════════════════════════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════════════════════════
    flush_footnotes()
    doc.core_properties.author = 'Michael Lokshin'
    out = DOCS / "flop_trade_model_v18.docx"
    doc.save(str(out))
    print(f"\nSaved {out}")


if __name__ == '__main__':
    main()
