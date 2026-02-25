"""
Produce flop_trade_model_v28.docx from v8.docx.

v28: Hyperscaler FDI trust channel — equation (2'), new equilibrium, Table 3b column (7):
  - compute_fdi_lambda(): bilateral premium with hyperscaler intermediation
  - FDI equilibrium: λ^FDI replaces λ_{ij}, 14 exporters, 7 developing
  - Table 3b column (7): regime assignments under FDI specification
  - Equation (2'): λ_{jk}^FDI variant in Section 3.2
  - FDI subsection in Section 6.2 (~280 words)
  - Revised abstract, conclusion, sovereignty editorial
  - AI language fixes throughout

v27: Form B upgrade — efficiency adjustment on non-hardware costs only:
  - c_adj = ρ + (c_cr − ρ) / ξ_eff (Form B replaces Form A: c_adj = c_cr / ξ_eff)
  - ω = 0.50 (equal weight governance/grid; was 0.85)
  - ξ_floor = 0.30 (was 0.50)
  - G source: raw WGI Rule of Law from xi_scenarios.xlsx (was reliability_index.csv)
  - Table A3: 7 robustness scenarios from form_b_simulations.xlsx
  - New top 5: Canada, Finland, Norway, China, Kyrgyzstan
  - Abstract rewritten with developing-country opportunity framing

v25: Based on v24 with the following changes:
  - ξ now weighted geometric mean: governance^ω × grid^(1−ω), ω = 0.85 (enclave argument)
  - P33/P73 rewritten with weighted formula and enclave justification
  - ω robustness paragraph added to Section 7.1
  - FN11 (HHI definition) deleted
  - P72 merged into P71 (sovereignty premium paragraph)
  - P92, P93, P95 condensed
  - All ξ values recomputed (Kyrgyzstan: 0.52, Uzbekistan: 0.48, Ethiopia: 0.42)

v24: Based on v23 with the following changes:
  - Decomposed ξ: now ξ_j^{eff} = grid reliability × operational risk only (no sanctions)
  - Bilateral λ_{ij} = α₁·G_{ij} + α₂·(1-R_{ij}) + α₃·S_{ij}
    (geopolitical distance, regulatory compatibility, sanctions)
  - Demand tiering: Tier 1 (sovereign, 10%), Tier 2 (regulated, 20%), Tier 3 (commercial, 70%)
  - New display equation (2) for λ_{ij}; all equations renumbered (+1 from old eq 2 onward)
  - Table 3 split into Table 3a (cost specs) and Table 3b (sovereignty specs)
  - Bailey, Strezhnev, and Voeten (2017) reference added
  - Full recomputation of equilibrium under bilateral sovereignty

v23: Based on v22 with the following changes:
  - Table 3 (country rankings under alternative pricing assumptions)
  - Section 6.2 restructured around Table 3 columns
  - Appendix tables renumbered (A2→A3 through A7→A8)
  - Letter codes (EE, IE, ID, DD, II) in Proposition 1 and Table 1
  - All Table/Figure references hyperlinked

v22: Based on v21 with the following changes:
  - 5-type regime classification (Proposition 1) replacing old 4-type scheme

v21: Based on v20 with the following changes:
  - Reliability-adjusted results lead Section 6 (preferred specification up front)
  - Observed-tariff ranking (Iran cheapest) demoted to illustrative comparison
  - Governance section trimmed (no redundant reliability repeat)
  - λ* subscript fix (_msup instead of _msubsup with empty sub)
  - Eq (5) inlined, equations renumbered (6→5, 7→6)
  - USD/kWh → $/kWh consistency
  - "Hardware generation" renamed to "GPU vintage"
  - Appendix E: construction cost regression table
  - Results preview paragraph added to Introduction
  - Manual edits integrated: "AI compute", "compute facilities", active voice,
    staffing paragraph removed, citations shortened to et al., "Our model",
    Section 3.1 opening trimmed, Costinot reference removed from Section 6

v20: Capacity-Constrained Ricardian Model restructuring.
  - Merged Sections 3+4+5 into unified Model (Section 3)
    3.1 Production Technology, 3.2 Trade Costs, 3.3 Demand,
    3.4 Sourcing and Market Equilibrium
  - New Section 4: Equilibrium Properties (Propositions 1-5, purely theoretical)
  - Renumbered: 5=Data, 6=Calibration, 7=Robustness, 8=Conclusion
  - Capacity constraints K_bar_j: training supply stack, market-clearing p_T,
    Ricardian rents, shadow values mu_j (inline), HHI_T < 1
  - New Appendix B: model derivation (B.1-B.6, 5 display equations)
  - Training/inference demand split: alpha parameter
  - 6 main-text display equations (p_T inlined), 5 propositions
"""

import copy
import csv
import io
import pathlib
import sys
from datetime import datetime
from decimal import Decimal, ROUND_HALF_UP
from functools import partial

import matplotlib

matplotlib.use('Agg')
import matplotlib.pyplot as plt  # noqa: E402
from docx import Document  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.oxml import OxmlElement  # noqa: E402
from docx.oxml.ns import qn  # noqa: E402
from docx.shared import Inches, Pt, RGBColor  # noqa: E402
from lxml import etree  # noqa: E402

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

DOCS = pathlib.Path(r"F:\onedrive\__documents\papers\FLOPsExport\Documents")
DATA = pathlib.Path(r"F:\onedrive\__documents\papers\FLOPsExport\Data")

TAU = 0.0008
LAMBDA = 0.10
PHI = 1.08
DELTA_PUE = 0.015
THETA_REF = 15.0
GAMMA = 0.700
GPU_PRICE = 25_000
GPU_LIFE = 3
GPU_UTIL = 0.70
DC_LIFE = 15
H_YR = 365.25 * 24
RHO = GPU_PRICE / (GPU_LIFE * H_YR * GPU_UTIL)
GPU_TDP_KW = 0.700  # GPU thermal design power in kW (same as GAMMA)
ETA = 0.15          # Amortized networking cost per GPU-hour ($/hr)

# Training share of compute demand (industry estimate)
ALPHA = 0.50

# ξ weighting: governance^ω × grid^(1−ω) — equal weight to governance and reliability
OMEGA_XI = 0.50

# Institutional enclave floor for ξ — minimum operational quality achievable
# through investor-side risk mitigation (arbitration, dedicated PPA, SEZ)
XI_FLOOR = 0.30

# Total global compute demand in GPU-hours per period
# Calibrated to ~10 million data center GPUs at ~70% utilization (~8,766 hrs/yr)
Q_TOTAL = 60_000_000_000   # 60 billion GPU-hours

# Grid capacity CSV unit correction (kWh→GWh used /1e9 instead of /1e6)
K_BAR_SCALE = 1000

# OECD members (38) for computing ξ range in body text
OECD_ISOS = {
    'AUS', 'AUT', 'BEL', 'CAN', 'CHL', 'COL', 'CRI', 'CZE', 'DNK', 'EST',
    'FIN', 'FRA', 'DEU', 'GRC', 'HUN', 'ISL', 'IRL', 'ISR', 'ITA', 'JPN',
    'KOR', 'LVA', 'LTU', 'LUX', 'MEX', 'NLD', 'NZL', 'NOR', 'POL', 'PRT',
    'SVK', 'SVN', 'ESP', 'SWE', 'CHE', 'TUR', 'GBR', 'USA',
}


def rhup(value, dp=2):
    """Round half-up (0.985 → 0.99, 0.765 → 0.77). Avoids banker's rounding."""
    return float(Decimal(str(value)).quantize(Decimal(10) ** -dp,
                                              rounding=ROUND_HALF_UP))


# Cost-reflective electricity prices for cost-recovery adjustment ($/kWh)
# Replacement = estimated LRMC of dominant generation at opportunity-cost fuel price
# Sources: IMF (2025) explicit fossil-fuel subsidy estimates; Lazard (2025) LCOE v17
SUBSIDY_ADJ = {
    'IRN': 0.085,  # Gas CCGT at export-parity fuel cost; IMF 2025, Lazard CCGT
    'TKM': 0.070,  # Gas CCGT at export-parity fuel cost; IMF 2025
    'DZA': 0.065,  # Gas CCGT at near-export parity; IMF 2025
    'EGY': 0.080,  # Gas/oil, reform target; IMF 2025, Egypt subsidy reform
    'UZB': 0.090,  # Gas, WB reform target; World Bank 2024
    'QAT': 0.100,  # Gas CCGT opportunity cost of LNG export; MEI GCC study
    'SAU': 0.100,  # Gas/oil opportunity cost; MEI GCC study
    'ARE': 0.095,  # Gas CCGT; MEI GCC study
    'RUS': 0.065,  # Gas/coal/nuclear mix; IMF 2025
    'KAZ': 0.085,  # Coal at cost recovery; IMF 2025
    'NGA': 0.080,  # Gas/hydro mix; IMF 2025
    'ZAF': 0.095,  # Coal, Eskom cost recovery; NERSA 2025 tariff path
    'ETH': 0.050,  # Hydro cost-recovery target; IMF 2025
}

# ═══════════════════════════════════════════════════════════════════════
# v24: BILATERAL SOVEREIGNTY PREMIUM λ_{ij}
# ═══════════════════════════════════════════════════════════════════════

# Bilateral λ_{ij} coefficients: λ_{ij} = α₁·G_{ij} + α₂·(1-R_{ij}) + α₃·S_{ij}
ALPHA_GEO = 0.08     # α₁: geopolitical distance weight
ALPHA_REG = 0.04     # α₂: regulatory incompatibility weight
# α₃ = ∞ (sanctions → trade prohibited; handled by exclusion)

# Uniform λ retained ONLY for robustness comparison (old specification)
LAMBDA_UNIFORM = 0.10

# Demand tier weights (Deloitte 2025 estimates)
W_TIER1 = 0.10   # Sovereign: military, intelligence, critical infrastructure
W_TIER2 = 0.20   # Regulated: health, financial, GDPR-covered personal data
W_TIER3 = 0.70   # Commercial: routine training, commercial inference, non-personal

# Countries under comprehensive sanctions (S_{ij} = 1 for Western buyers)
SANCTIONED = {'IRN', 'RUS', 'BLR', 'PRK', 'SYR', 'TKM'}

# Geopolitical blocs for G_{ij} computation
# Bloc assignments based on UN General Assembly voting patterns
# (Bailey, Strezhnev, and Voeten 2017)
BLOC_WESTERN = {
    'USA', 'CAN', 'GBR', 'FRA', 'DEU', 'ITA', 'ESP', 'PRT', 'NLD', 'BEL',
    'LUX', 'AUT', 'CHE', 'IRL', 'DNK', 'NOR', 'SWE', 'FIN', 'ISL', 'GRC',
    'CZE', 'POL', 'HUN', 'SVK', 'SVN', 'EST', 'LVA', 'LTU', 'HRV', 'BGR',
    'ROU', 'CYP', 'MLT', 'JPN', 'KOR', 'AUS', 'NZL', 'ISR', 'TWN',
}
BLOC_CHINA_ALIGNED = {
    'CHN', 'RUS', 'BLR', 'PRK', 'SYR', 'IRN',
    'VEN', 'CUB', 'NIC', 'MMR',
}
BLOC_NON_ALIGNED = None  # Default: any country not in WESTERN or CHINA_ALIGNED

# Inter-bloc geopolitical distance G_{ij} ∈ [0, 1]
# Rows/cols: W=Western, C=China-aligned, N=Non-aligned
BLOC_DISTANCE = {
    ('W', 'W'): 0.00,
    ('W', 'C'): 0.95,
    ('W', 'N'): 0.40,
    ('C', 'W'): 0.95,
    ('C', 'C'): 0.00,
    ('C', 'N'): 0.55,
    ('N', 'W'): 0.40,
    ('N', 'C'): 0.55,
    ('N', 'N'): 0.20,
}

# EU member states (mutual data adequacy → R_{ij} = 1 for all pairs)
EU_MEMBERS = {
    'AUT', 'BEL', 'BGR', 'HRV', 'CYP', 'CZE', 'DNK', 'EST', 'FIN', 'FRA',
    'DEU', 'GRC', 'HUN', 'IRL', 'ITA', 'LVA', 'LTU', 'LUX', 'MLT', 'NLD',
    'POL', 'PRT', 'ROU', 'SVK', 'SVN', 'ESP', 'SWE',
}

# APEC Cross-Border Privacy Rules (CBPR) members → R_{ij} = 1 for pairs
APEC_CBPR = {
    'AUS', 'CAN', 'JPN', 'KOR', 'MEX', 'PHL', 'SGP', 'TWN', 'USA',
}

# Additional bilateral regulatory agreements
# USMCA digital trade: USA-CAN-MEX (covered by APEC_CBPR)
# DEPA: SGP-CHL-NZL
DEPA_MEMBERS = {'SGP', 'CHL', 'NZL'}


def _get_bloc(iso):
    """Return bloc code ('W', 'C', 'N') for a country."""
    if iso in BLOC_WESTERN:
        return 'W'
    if iso in BLOC_CHINA_ALIGNED:
        return 'C'
    return 'N'


def compute_geo_distance(iso_i, iso_j):
    """Compute geopolitical distance G_{ij} ∈ [0, 1] from bloc assignments."""
    if iso_i == iso_j:
        return 0.0
    bi, bj = _get_bloc(iso_i), _get_bloc(iso_j)
    return BLOC_DISTANCE.get((bi, bj), 0.40)


def compute_reg_compat(iso_i, iso_j):
    """Compute regulatory compatibility R_{ij} ∈ {0, 1}."""
    if iso_i == iso_j:
        return 1
    # EU mutual adequacy
    if iso_i in EU_MEMBERS and iso_j in EU_MEMBERS:
        return 1
    # APEC CBPR
    if iso_i in APEC_CBPR and iso_j in APEC_CBPR:
        return 1
    # DEPA
    if iso_i in DEPA_MEMBERS and iso_j in DEPA_MEMBERS:
        return 1
    return 0


def compute_bilateral_lambda(iso_i, iso_j):
    """Compute bilateral sovereignty premium λ_{ij}.

    Returns float (the premium) or float('inf') for sanctioned pairs.
    λ_{ii} = 0 by definition (domestic production).
    """
    if iso_i == iso_j:
        return 0.0
    # Sanctions: either country sanctioned vis-à-vis the other
    if iso_i in SANCTIONED or iso_j in SANCTIONED:
        # Sanctioned-to-sanctioned pairs within same bloc can still trade
        if iso_i in SANCTIONED and iso_j in SANCTIONED:
            bi, bj = _get_bloc(iso_i), _get_bloc(iso_j)
            if bi == bj:
                return ALPHA_GEO * 0.0 + ALPHA_REG * (1 - 0)
        return float('inf')
    G_ij = compute_geo_distance(iso_i, iso_j)
    R_ij = compute_reg_compat(iso_i, iso_j)
    return ALPHA_GEO * G_ij + ALPHA_REG * (1 - R_ij)


# v28: GPU export controls — partial restriction on training hardware for China
# α₃ is normally ∞ (full ban) for SANCTIONED countries.  China is NOT sanctioned
# in the comprehensive sense, but faces GPU export controls on training-grade hardware.
# We model this as a partial sanctions indicator: S(CHN, buyer) = 0.5 for training only.
GPU_EXPORT_CONTROLLED = {'CHN'}
GPU_CONTROL_ALPHA3 = 0.10  # partial α₃ for GPU-controlled (not full sanction)


def compute_fdi_lambda(host_j, buyer_k, hyperscaler_h='USA'):
    """Compute FDI sovereignty premium λ_{jk}^FDI (equation 2').

    When a hyperscaler headquartered in h operates a facility in host j,
    the buyer k evaluates trust against the operator h, not the host j:
      λ_{jk}^FDI = α₁·G(h,k) + α₂·(1 − R(h,k)) + α₃·S(j,k)

    G and R terms use the hyperscaler's home country h (trust attaches to operator).
    S term uses the host country j (GPUs can't ship to sanctioned hosts).

    Returns float (premium) or float('inf') for sanctioned hosts.
    """
    if host_j == buyer_k:
        return 0.0
    # Sanctions on HOST: GPUs cannot ship to sanctioned countries regardless of operator
    if host_j in SANCTIONED:
        return float('inf')
    # GPU export controls on host (China): partial restriction
    s_jk = 0.0
    if host_j in GPU_EXPORT_CONTROLLED:
        s_jk = 0.5  # partial — training hardware restricted, inference less so
    # G and R terms use hyperscaler home country h, not host j
    G_hk = compute_geo_distance(hyperscaler_h, buyer_k)
    R_hk = compute_reg_compat(hyperscaler_h, buyer_k)
    return ALPHA_GEO * G_hk + ALPHA_REG * (1 - R_hk) + GPU_CONTROL_ALPHA3 * s_jk


def compute_xi_eff(governance_score, grid_score, xi_floor=XI_FLOOR):
    """Compute production-efficiency index ξ_j^{eff}.

    Weighted geometric mean with institutional floor:
      ξ_raw = governance^ω × grid^(1−ω)
      ξ_eff = ξ_floor + (1 − ξ_floor) × ξ_raw

    ω = 0.50 assigns equal weight to governance and grid reliability.
    ξ_floor = 0.30 reflects minimum operational quality via investor-side
    risk mitigation (arbitration, dedicated PPA, SEZ provisions).
    """
    xi_raw = (governance_score ** OMEGA_XI) * (grid_score ** (1 - OMEGA_XI))
    return xi_floor + (1 - xi_floor) * xi_raw


def recompute_costs(cal, gpu_price=None, gpu_util=None,
                    p_E_delta=0.0, pue_cap=None, subsidy_adj=None):
    """Re-derive c_j from CSV primitives with parameter overrides."""
    gp = gpu_price or GPU_PRICE
    gu = gpu_util or GPU_UTIL
    rho = gp / (GPU_LIFE * H_YR * gu)
    costs = {}
    for row in cal:
        iso = row["iso3"]
        p_E = float(row["p_E_usd_kwh"])
        if subsidy_adj and iso in subsidy_adj:
            p_E = subsidy_adj[iso]
        p_E += p_E_delta
        theta = float(row["theta_summer_C"])
        pue = PHI + DELTA_PUE * max(0, theta - THETA_REF)
        if pue_cap is not None:
            pue = min(pue, pue_cap)
        c_elec = pue * GAMMA * p_E
        c_constr = float(row["c_j_construction"])
        costs[iso] = c_elec + rho + ETA + c_constr
    return costs


def run_sensitivity(cal, omega, dc_k, k_bar, sanctioned, xi_vals, xi_raw_vals):
    """Load pre-computed sensitivity results from form_b_simulations.xlsx.

    v27: 7 robustness scenarios (Form B with parameter variations + Form A reference).
    Values are read from the Summary sheet of form_b_simulations.xlsx.

    Returns list of scenario result dicts for Table A3.
    """
    import openpyxl
    wb = openpyxl.load_workbook(DATA / "form_b_simulations.xlsx", read_only=True)
    ws = wb['Summary']
    hdr = [c.value for c in next(ws.iter_rows(max_row=1))]
    sim_rows = {}
    for row in ws.iter_rows(min_row=2, values_only=True):
        d = dict(zip(hdr, row))
        sim_rows[d['Config']] = d
    wb.close()

    # 7 scenarios for Table A3, in order
    scenario_defs = [
        ('C2',    'Baseline (\u03C9=0.50, floor=0.30, \u03C1=$1.36, Form B)',
         '\u03C9=0.50, floor=0.30, \u03C1=$1.36'),
        ('REF_A', 'Form A (v26 specification)',
         'c_adj = c_cr / \u03BE_eff, \u03C9=0.85, floor=0.50'),
        ('C1',    'No floor (\u03C9=0.50, floor=0.00)',
         '\u03C9=0.50, floor=0.00'),
        ('C3',    'High floor (\u03C9=0.50, floor=0.50)',
         '\u03C9=0.50, floor=0.50'),
        ('A2',    'High governance weight (\u03C9=0.85)',
         '\u03C9=0.85, floor=0.30'),
        ('H1',    'Low hardware share (\u03C1=$1.30)',
         '\u03C1=$1.30 (\u22124.4%)'),
        ('H4',    'High hardware share (\u03C1=$1.42)',
         '\u03C1=$1.42 (+4.4%)'),
    ]

    results = []
    baseline_top5 = None

    for config_key, label, param_change in scenario_defs:
        s = sim_rows[config_key]
        dev_top15 = int(s.get('Dev top15', 0))
        max_markup = float(s.get('Max markup%', 0)) * 100  # xlsx stores as fraction
        spearman = float(s.get('Spearman vs cr', 0))
        top5_str = str(s.get('Top 5 countries', ''))
        # Clean star markers
        top5_clean = top5_str.replace(' \u2605', '').replace('\u2605', '')

        if baseline_top5 is None:
            baseline_top5 = top5_clean

        result = {
            "label": label,
            "param_change": param_change,
            "dev_top15": dev_top15,
            "max_markup": max_markup,
            "rank_corr": spearman,
            "top5_str": top5_clean,
            "top5_unchanged": (top5_clean == baseline_top5),
            "kwargs": {},
            # Legacy fields for compatibility
            "p_T": 0, "n_exporters": 0, "hhi_T": 0, "top5": [],
        }
        results.append(result)
        print(f"  Sensitivity [{label}]: dev_top15={dev_top15}, "
              f"max_markup={max_markup:.1f}%, \u03C1_cr={spearman:.2f}, "
              f"top5={top5_clean}")

    return results


def _ordinal(n):
    """Return ordinal string for integer n (e.g. 1 -> '1st', 23 -> '23rd')."""
    s = ('th', 'st', 'nd', 'rd') + ('th',) * 6
    if 11 <= (n % 100) <= 13:
        return f'{n}th'
    return f'{n}{s[n % 10]}'


def _num_word(n):
    """Return English word for small integers, digit string otherwise."""
    words = {0: 'zero', 1: 'one', 2: 'two', 3: 'three', 4: 'four',
             5: 'five', 6: 'six', 7: 'seven', 8: 'eight', 9: 'nine',
             10: 'ten', 11: 'eleven', 12: 'twelve', 13: 'thirteen',
             14: 'fourteen', 15: 'fifteen', 16: 'sixteen', 17: 'seventeen',
             18: 'eighteen', 19: 'nineteen', 20: 'twenty'}
    return words.get(n, str(n))


# Fonts
CAMBRIA_MATH = 'Cambria Math'
TIMES_NEW_ROMAN = 'Times New Roman'

# Colors
HEADING_BLUE = RGBColor(0x2F, 0x54, 0x96)
LINK_COLOR = '1F3864'

# XML namespaces
XML_SPACE = '{http://www.w3.org/XML/1998/namespace}space'
SPACE_PRESERVE = 'preserve'

# Table formatting
TABLE_WIDTH_PCT = '5000'

# ═══════════════════════════════════════════════════════════════════════
# OMML HELPERS
# ═══════════════════════════════════════════════════════════════════════


def _mr(text, italic=True):
    r = OxmlElement('m:r')
    rPr = OxmlElement('m:rPr')
    sty = OxmlElement('m:sty')
    sty.set(qn('m:val'), 'i' if italic else 'p')
    rPr.append(sty)
    r.append(rPr)
    wrPr = OxmlElement('w:rPr')
    rF = OxmlElement('w:rFonts')
    rF.set(qn('w:ascii'), CAMBRIA_MATH)
    rF.set(qn('w:hAnsi'), CAMBRIA_MATH)
    wrPr.append(rF)
    r.append(wrPr)
    t = OxmlElement('m:t')
    t.set(XML_SPACE, SPACE_PRESERVE)
    t.text = text
    r.append(t)
    return r


def _v(text):
    return _mr(text, True)


def _t(text):
    return _mr(text, False)


def _msub(base, sub, base_italic=True, sub_italic=True):
    el = OxmlElement('m:sSub')
    el.append(OxmlElement('m:sSubPr'))
    e = OxmlElement('m:e')
    e.append(_mr(base, base_italic))
    el.append(e)
    s = OxmlElement('m:sub')
    s.append(_mr(sub, sub_italic))
    el.append(s)
    return el


def _msup(base, sup, base_italic=True, sup_italic=True):
    el = OxmlElement('m:sSup')
    el.append(OxmlElement('m:sSupPr'))
    e = OxmlElement('m:e')
    e.append(_mr(base, base_italic))
    el.append(e)
    s = OxmlElement('m:sup')
    s.append(_mr(sup, sup_italic))
    el.append(s)
    return el


def _mbar(base, base_italic=True):
    """Overbar accent using OMML <m:bar> element (renders better than combining macron)."""
    el = OxmlElement('m:bar')
    barPr = OxmlElement('m:barPr')
    pos = OxmlElement('m:pos')
    pos.set(qn('m:val'), 'top')
    barPr.append(pos)
    el.append(barPr)
    e = OxmlElement('m:e')
    e.append(_mr(base, base_italic))
    el.append(e)
    return el


def _mbar_sub(base, sub, base_italic=True, sub_italic=True):
    """Barred base with subscript: properly nested as sSub(bar(base), sub)."""
    el = OxmlElement('m:sSub')
    el.append(OxmlElement('m:sSubPr'))
    e = OxmlElement('m:e')
    e.append(_mbar(base, base_italic))
    el.append(e)
    s = OxmlElement('m:sub')
    s.append(_mr(sub, sub_italic))
    el.append(s)
    return el


def _mfrac(num_parts, den_parts):
    """OMML fraction: numerator / denominator, each a list of OMML elements."""
    el = OxmlElement('m:f')
    fPr = OxmlElement('m:fPr')
    el.append(fPr)
    num = OxmlElement('m:num')
    for p in num_parts:
        num.append(p)
    el.append(num)
    den = OxmlElement('m:den')
    for p in den_parts:
        den.append(p)
    el.append(den)
    return el


def _msubsup(base, sub, sup):
    """Subscript-superscript combo."""
    el = OxmlElement('m:sSubSup')
    el.append(OxmlElement('m:sSubSupPr'))
    e = OxmlElement('m:e')
    e.append(_mr(base))
    el.append(e)
    s = OxmlElement('m:sub')
    s.append(_mr(sub))
    el.append(s)
    u = OxmlElement('m:sup')
    u.append(_mr(sup))
    el.append(u)
    return el


def _nary(char, sub_parts, sup_parts, e_parts):
    """N-ary operator (summation, product) with sub/sup limits."""
    el = OxmlElement('m:nary')
    pr = OxmlElement('m:naryPr')
    ch = OxmlElement('m:chr')
    ch.set(qn('m:val'), char)
    pr.append(ch)
    if not sup_parts:
        supHide = OxmlElement('m:supHide')
        supHide.set(qn('m:val'), '1')
        pr.append(supHide)
    el.append(pr)
    sub = OxmlElement('m:sub')
    for p in sub_parts:
        sub.append(p)
    el.append(sub)
    sup = OxmlElement('m:sup')
    for p in sup_parts:
        sup.append(p)
    el.append(sup)
    e = OxmlElement('m:e')
    for p in e_parts:
        e.append(p)
    el.append(e)
    return el


def _limlow(e_parts, lim_parts):
    """Lower limit (argmin, min, lim) with limit underneath."""
    el = OxmlElement('m:limLow')
    el.append(OxmlElement('m:limLowPr'))
    e = OxmlElement('m:e')
    for p in e_parts:
        e.append(p)
    el.append(e)
    lim = OxmlElement('m:lim')
    for p in lim_parts:
        lim.append(p)
    el.append(lim)
    return el


def omath(p, parts):
    om = OxmlElement('m:oMath')
    for part in parts:
        om.append(part)
    p._element.append(om)


def omath_display(doc, body, cursor, parts, eq_num=None):
    """Display equation in a borderless 2-column table: centered equation + right-aligned number."""
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove all borders
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        e = OxmlElement(f'w:{edge}')
        e.set(qn('w:val'), 'none')
        e.set(qn('w:sz'), '0')
        e.set(qn('w:space'), '0')
        e.set(qn('w:color'), 'auto')
        borders.append(e)
    tblPr.append(borders)
    # Full-width table
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'), TABLE_WIDTH_PCT)
    tblW.set(qn('w:type'), 'pct')
    old_w = tblPr.find(qn('w:tblW'))
    if old_w is not None:
        tblPr.remove(old_w)
    tblPr.append(tblW)
    # Column widths: equation 85%, number 15%
    for j, w in enumerate([8100, 1400]):
        tc = tbl.cell(0, j)._tc
        tcPr = tc.get_or_add_tcPr()
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(w))
        tcW.set(qn('w:type'), 'dxa')
        old = tcPr.find(qn('w:tcW'))
        if old is not None:
            tcPr.remove(old)
        tcPr.append(tcW)
    # Equation in first cell, centered
    p0 = tbl.cell(0, 0).paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pPr0 = p0._element.get_or_add_pPr()
    sp0 = OxmlElement('w:spacing')
    sp0.set(qn('w:before'), '60')
    sp0.set(qn('w:after'), '60')
    pPr0.append(sp0)
    om = OxmlElement('m:oMath')
    for part in parts:
        om.append(part)
    p0._element.append(om)
    # Number in second cell, right-aligned
    p1 = tbl.cell(0, 1).paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr1 = p1._element.get_or_add_pPr()
    sp1 = OxmlElement('w:spacing')
    sp1.set(qn('w:before'), '60')
    sp1.set(qn('w:after'), '60')
    pPr1.append(sp1)
    if eq_num:
        # Add bookmark target so in-text "equation (N)" mentions can link here
        bm_name = f'Eq{eq_num}'
        _eq_clean = eq_num.replace('.', '').replace('B', '90').replace('a', '01').replace('b', '02').replace('\u2032', '9')
        bm_id_val = 800 + int(_eq_clean)
        p1._element.append(make_bookmark(bm_id_val, bm_name))
        p1.add_run(f'({eq_num})')
        p1._element.append(make_bookmark_end(bm_id_val))
    # Vertical center cell content
    for j in range(2):
        tc = tbl.cell(0, j)._tc
        tcPr = tc.get_or_add_tcPr()
        vAlign = OxmlElement('w:vAlign')
        vAlign.set(qn('w:val'), 'center')
        tcPr.append(vAlign)
    tbl_el = tbl._tbl
    body.remove(tbl_el)
    cursor.addnext(tbl_el)
    return None, tbl_el

# ═══════════════════════════════════════════════════════════════════════
# PARAGRAPH HELPERS
# ═══════════════════════════════════════════════════════════════════════


def mkp(doc, body, cursor, space_before=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(space_before if space_before is not None else 0)
    p.paragraph_format.space_after = Pt(8)
    el = p._element
    body.remove(el)
    cursor.addnext(el)
    return p, el


def mkh(doc, body, cursor, text, level=1):
    p = doc.add_paragraph(text, style=f'Heading {level}')
    el = p._element
    body.remove(el)
    cursor.addnext(el)
    return el


def add_italic(p, text):
    """Add an italic run to paragraph p."""
    r = p.add_run(text)
    r.italic = True
    return r


def add_page_break(doc, body, after_el):
    """Insert a page break paragraph after after_el. Returns the new element."""
    pb_p = doc.add_paragraph()
    pb_p.paragraph_format.space_before = Pt(0)
    pb_p.paragraph_format.space_after = Pt(0)
    pb_run = pb_p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    pb_run._element.append(br)
    pb_el = pb_p._element
    body.remove(pb_el)
    after_el.addnext(pb_el)
    return pb_el


def make_bookmark(bm_id, name):
    """Create a w:bookmarkStart element."""
    bs = OxmlElement('w:bookmarkStart')
    bs.set(qn('w:id'), str(bm_id))
    bs.set(qn('w:name'), name)
    return bs


def make_bookmark_end(bm_id):
    """Create a w:bookmarkEnd element."""
    be = OxmlElement('w:bookmarkEnd')
    be.set(qn('w:id'), str(bm_id))
    return be


def make_hyperlink(anchor, text, rPr_orig=None, color=LINK_COLOR):
    """Create a w:hyperlink element with blue underlined text."""
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), anchor)
    hl.set(qn('w:history'), '1')
    r = OxmlElement('w:r')
    new_rPr = copy.deepcopy(rPr_orig) if rPr_orig is not None else OxmlElement('w:rPr')
    clr = OxmlElement('w:color')
    clr.set(qn('w:val'), color)
    uu = OxmlElement('w:u')
    uu.set(qn('w:val'), 'single')
    new_rPr.append(clr)
    new_rPr.append(uu)
    r.append(new_rPr)
    t = OxmlElement('w:t')
    t.set(XML_SPACE, SPACE_PRESERVE)
    t.text = text
    r.append(t)
    hl.append(r)
    return hl


def _rPr_pt(pt_size):
    """Return a w:rPr element with the given font size (in points)."""
    rPr = OxmlElement('w:rPr')
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(pt_size * 2))  # half-points
    rPr.append(sz)
    return rPr


# ═══════════════════════════════════════════════════════════════════════
# TABLE CELL HELPERS (shared across Table 3a, 3b, A2)
# ═══════════════════════════════════════════════════════════════════════

def _tbl_border(tc, sides, sz='4', style='single'):
    """Add borders to a table cell element."""
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement('w:tcBorders')
    for s in sides:
        b = OxmlElement(f'w:{s}')
        b.set(qn('w:val'), style)
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tcB.append(b)
    tcPr.append(tcB)


def _tbl_set(tbl, row_i, col_j, text, bold=False, align='center', font_size=9):
    """Set text content and formatting of a table cell."""
    cell = tbl.cell(row_i, col_j)
    cell.text = ''
    pp = cell.paragraphs[0]
    if align == 'left':
        pp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif align == 'right':
        pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(text)
    rr.font.size = Pt(font_size)
    rr.font.name = 'Times New Roman'
    if bold:
        rr.bold = True
    return cell


def _tbl_merge(tbl, row_i, col_start, col_end):
    """Merge cells in a row from col_start to col_end (inclusive)."""
    tbl.cell(row_i, col_start).merge(tbl.cell(row_i, col_end))


def _tbl_clear_borders(tbl):
    """Remove all table-level borders and set width to 100%."""
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    old_bdr = tblPr.find(qn('w:tblBorders'))
    if old_bdr is not None:
        tblPr.remove(old_bdr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        tblBorders.append(b)
    tblPr.append(tblBorders)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = OxmlElement('w:tblW')
        tblPr.append(tblW)
    tblW.set(qn('w:w'), '5000')
    tblW.set(qn('w:type'), 'pct')


def _tbl_col_widths(tbl, col_widths):
    """Set column widths (in twips) for all rows."""
    n_rows = len(tbl.rows)
    for j, w in enumerate(col_widths):
        for i in range(n_rows):
            try:
                c = tbl.cell(i, j)
            except IndexError:
                continue
            tcPr = c._tc.get_or_add_tcPr()
            tcW = OxmlElement('w:tcW')
            tcW.set(qn('w:w'), str(w))
            tcW.set(qn('w:type'), 'dxa')
            old = tcPr.find(qn('w:tcW'))
            if old is not None:
                tcPr.remove(old)
            tcPr.append(tcW)


def _tbl_cell_spacing(tbl, before='10', after='10'):
    """Set paragraph spacing inside all table cells."""
    for row in tbl.rows:
        for cell in row.cells:
            for pp in cell.paragraphs:
                pPr = pp._element.get_or_add_pPr()
                sp = OxmlElement('w:spacing')
                sp.set(qn('w:before'), before)
                sp.set(qn('w:after'), after)
                pPr.append(sp)


# ═══════════════════════════════════════════════════════════════════════
# FOOTNOTE HELPER
# ═══════════════════════════════════════════════════════════════════════
_fn_xml = [None]   # cached parsed footnotes XML
_fn_part = [None]  # cached footnotes Part

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'


def init_footnotes(doc):
    """Parse the footnotes part from the document package and remove old content footnotes."""
    for rel in doc.part.rels.values():
        if 'footnotes' in rel.reltype:
            _fn_part[0] = rel.target_part
            _fn_xml[0] = etree.fromstring(_fn_part[0].blob)
            # Remove existing content footnotes (keep IDs 0 and -1 = Word separators)
            for fn in list(_fn_xml[0]):
                fid = fn.get(f'{{{W_NS}}}id', '')
                if fid not in ('0', '-1', ''):
                    _fn_xml[0].remove(fn)
            return
    print("  Warning: no footnotes part found in template")


def make_footnote(p, fn_text, fn_id):
    """Add a footnote at the end of paragraph p."""
    if _fn_xml[0] is None:
        return
    fn_el = etree.SubElement(_fn_xml[0], f'{{{W_NS}}}footnote')
    fn_el.set(f'{{{W_NS}}}id', str(fn_id))
    fn_p = etree.SubElement(fn_el, f'{{{W_NS}}}p')
    fn_pPr = etree.SubElement(fn_p, f'{{{W_NS}}}pPr')
    fn_pStyle = etree.SubElement(fn_pPr, f'{{{W_NS}}}pStyle')
    fn_pStyle.set(f'{{{W_NS}}}val', 'FootnoteText')
    # Auto-number mark
    fn_r1 = etree.SubElement(fn_p, f'{{{W_NS}}}r')
    fn_rPr1 = etree.SubElement(fn_r1, f'{{{W_NS}}}rPr')
    fn_rStyle1 = etree.SubElement(fn_rPr1, f'{{{W_NS}}}rStyle')
    fn_rStyle1.set(f'{{{W_NS}}}val', 'FootnoteReference')
    etree.SubElement(fn_r1, f'{{{W_NS}}}footnoteRef')
    # Footnote text
    fn_r2 = etree.SubElement(fn_p, f'{{{W_NS}}}r')
    fn_t2 = etree.SubElement(fn_r2, f'{{{W_NS}}}t')
    fn_t2.set(XML_SPACE, SPACE_PRESERVE)
    fn_t2.text = ' ' + fn_text
    # Footnote reference in main text
    fn_ref_r = OxmlElement('w:r')
    fn_ref_rPr = OxmlElement('w:rPr')
    fn_ref_rStyle = OxmlElement('w:rStyle')
    fn_ref_rStyle.set(qn('w:val'), 'FootnoteReference')
    fn_ref_rPr.append(fn_ref_rStyle)
    fn_ref_r.append(fn_ref_rPr)
    fn_ref_el = OxmlElement('w:footnoteReference')
    fn_ref_el.set(qn('w:id'), str(fn_id))
    fn_ref_r.append(fn_ref_el)
    p._element.append(fn_ref_r)


def flush_footnotes():
    """Write cached footnotes XML back to the part."""
    if _fn_part[0] is not None and _fn_xml[0] is not None:
        _fn_part[0]._blob = etree.tostring(
            _fn_xml[0], xml_declaration=True, encoding='UTF-8', standalone=True)

# ═══════════════════════════════════════════════════════════════════════
# CITATION CROSS-REFERENCE SYSTEM
# ═══════════════════════════════════════════════════════════════════════


# Map: citation text as it appears in-text -> bookmark key
# (cite_author, year, key, ref_anchor)
# cite_author: author string as it appears in inline citations
# year: publication year
# key: unique bookmark key
# ref_anchor: string to match in the reference list for back-linking
CITATIONS = [
    ('Epoch AI', '2024', 'EpochAI2024', 'Epoch AI. (2024)'),
    ('Deloitte', '2025', 'Deloitte2025', 'Deloitte. (2025)'),
    ('Deloitte', '2020', 'Deloitte2020', 'Deloitte and Google. (2020)'),
    ('IEA', '2025', 'IEA2025', 'IEA. (2025)'),
    ('Ohlin', '1933', 'Ohlin1933', 'Ohlin, B.'),
    ('Biglaiser et al.', '2024', 'Biglaiser2024', 'Biglaiser, G.'),
    ('Blinder', '2006', 'Blinder2006', 'Blinder, A.'),
    ('Stojkoski et al.', '2024', 'Stojkoski2024', 'Stojkoski, V.'),
    ('World Bank', '2025', 'WorldBank2025', 'World Bank. (2025). Digital'),
    ('Hausmann et al.', '2007', 'Hausmann2007', 'Hausmann, R.'),
    ('Uptime Institute', '2024', 'UptimeInstitute2024',
     'Uptime Institute. (2024)'),
    ('Firebird', '2026', 'Firebird2026', 'Firebird. (2026)'),
    ('Flucker et al.', '2013', 'Flucker2013', 'Flucker, S.'),
    ('Eaton and Kortum', '2002', 'EatonKortum2002', 'Eaton, J.'),
    ('Liu et al.', '2023', 'Liu2023', 'Liu, Z.'),
    ('Goldfarb and Trefler', '2018', 'Goldfarb2018', 'Goldfarb, A.'),
    ('Korinek and Stiglitz', '2021', 'Korinek2021', 'Korinek, A.'),
    ('UNCTAD', '2025', 'UNCTAD2025', 'UNCTAD. (2025)'),
    ('Google', '2024', 'Google2024', 'Google. (2024)'),

    ('Brainard', '1997', 'Brainard1997', 'Brainard, S.'),
    ('Helpman et al.', '2004', 'HMY2004', 'Helpman, E.'),
    ('Lim\u00E3o and Venables', '2001', 'Limao2001', 'Lim\u00E3o, N.'),
    ('Eurostat', '2025', 'Eurostat2025', 'Eurostat. (2025)'),
    ('EIA', '2025', 'EIA2025', 'EIA. (2025)'),
    ('Hersbach et al.', '2020', 'Hersbach2020', 'Hersbach, H.'),
    ('Turner & Townsend', '2025', 'TurnerTownsend2025',
     'Turner & Townsend. (2025)'),
    ('WonderNetwork', '2024', 'WonderNetwork2024', 'WonderNetwork. (2024)'),
    ('NVIDIA', '2024', 'NVIDIA2024', 'NVIDIA. (2024)'),
    ('Krugman', '1991', 'Krugman1991', 'Krugman, P.'),
    ('World Bank', '2024', 'WorldBank2024', 'World Bank. (2024)'),
    ('GlobalPetrolPrices', '2025', 'GlobalPetrolPrices2025',
     'GlobalPetrolPrices. (2025)'),
    ('Cloudscene', '2025', 'Cloudscene2025', 'Cloudscene. (2025)'),
    ('Sastry, Heim, et al.', '2024', 'Sastry2024', 'Sastry, G.'),
    ('Lehdonvirta et al.', '2024', 'Lehdonvirta2024',
     'Lehdonvirta, V.'),
    ('Pilz et al.', '2025', 'Pilz2025', 'Pilz, K.'),
    ('Turner Lee and West', '2025', 'TurnerLee2025', 'Turner Lee, N.'),
    ('IMF', '2025', 'IMF2025', 'IMF. (2025)'),
    ('Lazard', '2025', 'Lazard2025', 'Lazard. (2025)'),
    ('Arkolakis et al.', '2012', 'ACR2012',
     'Arkolakis, C.'),
    ('van der Ploeg', '2011', 'vanderPloeg2011', 'van der Ploeg, F.'),
    ('Barroso et al.', '2018', 'Barroso2018', 'Barroso, L.'),
    ('ABD', '2020', 'ABD2020', 'Asian Development Bank. (2020)'),
    ('Calcaterra et al.', '2024', 'Calcaterra2024', 'Calcaterra, M.'),
    # v24: Bailey et al. for UN General Assembly ideal-point data
    ('Bailey, Strezhnev, and Voeten', '2017', 'BaileyEtAl2017', 'Bailey, M.'),
    # v26: SEZ/institutional floor references
    ('Farole', '2011', 'Farole2011', 'Farole, T.'),
    ('Frick, Rodr\u00EDguez-Pose, and Wong', '2019', 'Frick2019', 'Frick, S.'),
    ('World Bank', '2017', 'WorldBank2017', 'World Bank. (2017)'),
    # v27: Enterprise Surveys reference for grid reliability data
    ('World Bank Enterprise Surveys', '2025', 'WBES2025',
     'World Bank. (2025). Enterprise'),
]

# Auto-generate CITE_MAP: both "Author (Year)" and "Author Year" forms
CITE_MAP = {}
for _auth, _yr, _key, _ in CITATIONS:
    CITE_MAP[f'{_auth} ({_yr})'] = _key   # narrative: Author (Year)
    CITE_MAP[f'{_auth} {_yr}'] = _key      # parenthetical: Author Year
CITE_MAP['World Bank'] = 'WorldBank2024'   # bare mention without year
CITE_MAP['Cloudscene'] = 'Cloudscene2025'  # bare mention without year
# Abbreviated citation form (CITATIONS has "Sastry, Heim, et al.")
CITE_MAP['Sastry et al. (2024)'] = 'Sastry2024'
CITE_MAP['Sastry et al. 2024'] = 'Sastry2024'

# Auto-generate REF_KEY_MAP for back-linking from reference list
REF_KEY_MAP = {_key: _anchor for _, _, _key, _anchor in CITATIONS}


def link_citations_pass(body, cite_map, bm_id):
    """Single pass: find citation text in runs and replace with bookmark+hyperlink. Returns count."""
    count = 0
    # Sort by length descending so longer citations match first
    sorted_cites = sorted(cite_map.items(), key=lambda x: -len(x[0]))
    for p_el in list(body.findall(qn('w:p'))):
        for child in list(p_el):
            if child.tag != qn('w:r'):
                continue
            t_el = child.find(qn('w:t'))
            if t_el is None or not t_el.text:
                continue
            text = t_el.text
            for cite_text, key in sorted_cites:
                if cite_text not in text:
                    continue
                idx = text.index(cite_text)
                before = text[:idx]
                after = text[idx + len(cite_text):]
                rPr_orig = child.find(qn('w:rPr'))
                # Modify current run to "before" text only
                t_el.text = before
                t_el.set(XML_SPACE, SPACE_PRESERVE)
                ins = child
                # bookmarkStart
                bm_start = make_bookmark(bm_id[0], f'{key}txt')
                ins.addnext(bm_start)
                ins = bm_start
                # hyperlink (blue underline)
                hyperlink = make_hyperlink(key, cite_text, rPr_orig)
                ins.addnext(hyperlink)
                ins = hyperlink
                # bookmarkEnd
                bm_end = make_bookmark_end(bm_id[0])
                ins.addnext(bm_end)
                ins = bm_end
                bm_id[0] += 1
                # after text
                if after:
                    ra = OxmlElement('w:r')
                    if rPr_orig is not None:
                        ra.append(copy.deepcopy(rPr_orig))
                    ta = OxmlElement('w:t')
                    ta.set(XML_SPACE, SPACE_PRESERVE)
                    ta.text = after
                    ra.append(ta)
                    ins.addnext(ra)
                count += 1
                break  # one per run per pass
    return count


# ═══════════════════════════════════════════════════════════════════════
# ITALIC JOURNAL / BOOK TITLES IN REFERENCES
# ═══════════════════════════════════════════════════════════════════════
ITALIC_IN_REFS = {
    'Brainard': 'American Economic Review',
    'Calcaterra': 'Nature Energy',
    'Cloudscene': 'Global Data Center Directory',
    'Deloitte': 'Deloitte Insights',
    'EIA': 'Electric Power Monthly',
    'Eurostat': 'Electricity Prices for Non-Household Consumers (nrg_pc_205)',
    'Flucker': 'Building Services Engineering Research and Technology',
    'GlobalPetrolPrices': 'Electricity Prices Around the World',
    'Goldfarb': 'The Economics of Artificial Intelligence',
    'Google': '2024 Environmental Report',
    'Hausmann': 'Journal of Economic Growth',
    'Helpman': 'American Economic Review',
    'Hersbach': 'Quarterly Journal of the Royal Meteorological Society',

    'Korinek': 'NBER Working Paper',
    'Krugman': 'Journal of Political Economy',
    'Lim\u00E3o': 'World Bank Economic Review',
    'Liu': 'Proceedings of ACM e-Energy',
    'NVIDIA': 'NVIDIA H100 Tensor Core GPU Datasheet',
    'Eaton': 'Econometrica',
    'Turner': 'Data Centre Construction Cost Index 2025',
    'UNCTAD': 'Technology and Innovation Report 2025',
    'Uptime Institute': 'Global Data Center Survey Results 2024',
    'WonderNetwork': 'Global Ping Statistics',
    'World Bank': 'World Development Indicators',
    'Lehdonvirta': 'Proceedings of the AAAI/ACM Conference on AI, Ethics, and Society',
    'Pilz': 'AI\u2019s Power Requirements Under Exponential Growth',
    'IMF': 'IMF Working Paper',
    'Lazard': 'Lazard\u2019s Levelized Cost of Energy Analysis, Version 17.0',
    'U.S. Department': 'Horizontal Merger Guidelines',
    'Arkolakis': 'American Economic Review',
    'Bailey': 'Journal of Conflict Resolution',
    'van der Ploeg': 'Journal of Economic Literature',
    'Barroso': 'The Datacenter as a Computer',
    'Ohlin': 'Interregional and International Trade',
    'Biglaiser': 'Toulouse School of Economics Working Paper',
    'Blinder': 'Foreign Affairs',
    'Stojkoski': 'Nature Communications',
    'World Bank. (2025). Digital': 'Digital Progress and Trends Report 2025',
    'World Bank. (2025). Enterprise': 'Enterprise Surveys',
    'World Bank. (2017)': 'Special Economic Zones: An Operational Review of Their Impacts',
    'Farole': 'Special Economic Zones in Africa: Comparing Performance and Learning from Global Experiences',
    'Frick': 'Economic Geography',
}


def find_italic_portion(full_rt):
    """Find the journal/book title that should be italicized in a reference."""
    for author_start, italic_text in ITALIC_IN_REFS.items():
        if full_rt.startswith(author_start) and italic_text in full_rt:
            return italic_text
    return None


def _write_ref_segments(p, text, italic_portion):
    """Write reference text with italic journal/book title."""
    if italic_portion and italic_portion in text:
        idx = text.index(italic_portion)
        if idx > 0:
            p.add_run(text[:idx])
        r = p.add_run(italic_portion)
        r.italic = True
        after = text[idx + len(italic_portion):]
        if after:
            p.add_run(after)
    else:
        p.add_run(text)


def add_table(doc, body, after_el, headers, rows, col_widths=None, title=None,
              center_cols=None, bookmark_id=None, bookmark_name=None,
              backlink_name=None):
    if title:
        tp = doc.add_paragraph()
        tp.paragraph_format.space_before = Pt(6)
        tp.paragraph_format.space_after = Pt(3)
        tp.paragraph_format.first_line_indent = Inches(0)
        if bookmark_id and bookmark_name:
            # Split title into table number + rest (e.g. "Table A3" + ". Sensitivity...")
            dot_pos = title.find('. ')
            tbl_num = title[:dot_pos] if dot_pos > 0 else title
            tbl_rest = title[dot_pos:] if dot_pos > 0 else ''
            tp._element.append(make_bookmark(bookmark_id, bookmark_name))
            if backlink_name:
                hl = OxmlElement('w:hyperlink')
                hl.set(qn('w:anchor'), backlink_name)
                hl.set(qn('w:history'), '1')
                r_hl = OxmlElement('w:r')
                rPr_hl = OxmlElement('w:rPr')
                b_hl = OxmlElement('w:b')
                rPr_hl.append(b_hl)
                sz_hl = OxmlElement('w:sz')
                sz_hl.set(qn('w:val'), '20')
                rPr_hl.append(sz_hl)
                clr_hl = OxmlElement('w:color')
                clr_hl.set(qn('w:val'), LINK_COLOR)
                uu_hl = OxmlElement('w:u')
                uu_hl.set(qn('w:val'), 'single')
                rPr_hl.append(clr_hl)
                rPr_hl.append(uu_hl)
                r_hl.append(rPr_hl)
                t_hl = OxmlElement('w:t')
                t_hl.text = tbl_num
                r_hl.append(t_hl)
                hl.append(r_hl)
                tp._element.append(hl)
            else:
                run_num = tp.add_run(tbl_num)
                run_num.bold = True
                run_num.font.size = Pt(10)
            tp._element.append(make_bookmark_end(bookmark_id))
            if tbl_rest:
                run_rest = tp.add_run(tbl_rest)
                run_rest.bold = True
                run_rest.font.size = Pt(10)
        else:
            run = tp.add_run(title)
            run.bold = True
            run.font.size = Pt(10)
        tbl_el = tp._element
        body.remove(tbl_el)
        after_el.addnext(tbl_el)
        after_el = tbl_el
    nr = len(rows) + 1
    nc = len(headers)
    table = doc.add_table(rows=nr, cols=nc)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    _tbl_clear_borders(table)
    # Academic-style horizontal rules on header row (top + bottom)
    for j, h in enumerate(headers):
        c = table.cell(0, j)
        c.text = ""
        pp = c.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pp.add_run(h)
        run.bold = True
        run.font.size = Pt(8)
        _tbl_border(c._tc, ['top', 'bottom'])
    _center_set = set(center_cols) if center_cols else set()
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = table.cell(i + 1, j)
            c.text = ""
            pp = c.paragraphs[0]
            if j in _center_set:
                pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif j >= 2:
                pp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = pp.add_run(str(val))
            run.font.size = Pt(8)
            # Double bottom border on last data row
            if i == len(rows) - 1:
                _tbl_border(c._tc, ['bottom'], style='double')
    if col_widths:
        for j, w in enumerate(col_widths):
            for i in range(nr):
                c = table.cell(i, j)
                tcPr = c._tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(w))
                tcW.set(qn('w:type'), 'dxa')
                old = tcPr.find(qn('w:tcW'))
                if old is not None:
                    tcPr.remove(old)
                tcPr.append(tcW)
    for row in table.rows:
        for cell in row.cells:
            for pp in cell.paragraphs:
                pPr = pp._element.get_or_add_pPr()
                sp = OxmlElement('w:spacing')
                sp.set(qn('w:before'), '10')
                sp.set(qn('w:after'), '10')
                pPr.append(sp)
    tbl_el = table._tbl
    body.remove(tbl_el)
    after_el.addnext(tbl_el)
    return tbl_el


def write_title_and_abstract(doc, body, all_el, hmap, demand_data=None):
    print("Rewriting title and abstract...")
    # Replace title (first element — no previous, so clear and rewrite in place)
    title_el = all_el[0]
    for child in list(title_el):
        if child.tag != qn('w:pPr'):
            title_el.remove(child)
    title_p = doc.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.first_line_indent = Inches(0)
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after = Pt(0)
    r_title = title_p.add_run('Cheap Energy Might Not Be Enough:\nA Trade Model of AI Compute Services')
    r_title.bold = False
    r_title.font.size = Pt(16)
    r_title.font.name = TIMES_NEW_ROMAN

    # Add author name
    author_p, author_el = mkp(doc, body, title_el, space_before=12)
    author_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_p.paragraph_format.space_after = Pt(12)
    r_author = author_p.add_run('Michael Lokshin')
    r_author.italic = True
    make_footnote(author_p,
                  'This paper\u2019s findings, interpretations, and conclusions are entirely those of the '
                  'author and do not necessarily represent the views of the author\u2019s employer, the '
                  'World Bank, its Executive Directors, or the countries they represent. '
                  'Michael Lokshin: mlokshin@worldbank.org', 1)

    # Version stamp
    ver_p, ver_el = mkp(doc, body, author_el, space_before=2)
    ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ver_p.paragraph_format.space_after = Pt(12)
    r_ver = ver_p.add_run(f'v28  \u2014  {datetime.now().strftime("%B %d, %Y  %H:%M")}')
    r_ver.font.size = Pt(9)
    r_ver.font.color.rgb = RGBColor(128, 128, 128)
    r_ver.font.name = TIMES_NEW_ROMAN

    # Replace Abstract heading + text with single paragraph
    # Remove old Abstract heading
    abs_heading = hmap['abs']
    abs_text = all_el[2]
    body.remove(abs_heading)
    body.remove(abs_text)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    r_abs_label = p.add_run('Abstract')
    r_abs_label.bold = True
    p.add_run(
        ': Energy-rich developing countries could convert cheap electricity into high-value '
        'AI compute exports. This paper develops the first trade model of compute services '
        '(FLOPs) to assess which countries have comparative advantage in this emerging sector. '
        'Latency-insensitive AI training can be offshored '
        'to the lowest-cost producers, while latency-sensitive inference favors proximity '
        'to users; a bilateral sovereignty premium captures governments\u2019 preference for domestic data '
        'processing, varying with geopolitical alignment, regulatory compatibility, and sanctions. '
        'Calibration across 85 countries shows that six developing countries rank among the '
        'fifteen cheapest producers after governance adjustment, but bilateral trust deficits '
        'eliminate all developing-country exports under the standard sovereignty specification. '
        'When a hyperscaler intermediates the transaction, the buyer\u2019s trust attaches to '
        'the operator rather than the host country, restoring developing-country '
        f'competitiveness: {_num_word(demand_data.get("n_dev_fdi_exporters", 7) if demand_data else 7)} '
        'developing economies re-enter as potential exporters. '
        'The cross-country cost spread is only 12\u201320 percent, making compute both the '
        'easiest tradable sector for developing countries to enter on cost grounds and the '
        'one most vulnerable to small policy-induced frictions. For energy-rich developing '
        'countries, the binding constraint is not electricity cost but the institutional '
        'credibility needed to attract hyperscaler investment.'
    )
    el = p._element
    body.remove(el)
    ver_el.addnext(el)
    abs_text_el = el

    # Two blank lines after abstract (explicit empty paragraphs)
    p_blank1, blank1_el = mkp(doc, body, abs_text_el)
    p_blank1.paragraph_format.space_before = Pt(0)
    p_blank1.paragraph_format.space_after = Pt(0)
    p_blank1.paragraph_format.line_spacing = 1.0
    p_blank1.add_run(' ')
    p_blank2, blank2_el = mkp(doc, body, blank1_el)
    p_blank2.paragraph_format.space_before = Pt(0)
    p_blank2.paragraph_format.space_after = Pt(0)
    p_blank2.paragraph_format.line_spacing = 1.0
    p_blank2.add_run(' ')

    # JEL classification and keywords after abstract
    p_jel, jel_el = mkp(doc, body, blank2_el, space_before=0)
    p_jel.paragraph_format.left_indent = Inches(0.5)
    p_jel.paragraph_format.right_indent = Inches(0.5)
    p_jel.paragraph_format.line_spacing = 1.0
    r_jel_label = p_jel.add_run('JEL Classification: ')
    r_jel_label.bold = True
    p_jel.add_run('F14, F18, L86, O14, O33, Q40')

    p_kw, kw_el = mkp(doc, body, jel_el, space_before=2)
    p_kw.paragraph_format.left_indent = Inches(0.5)
    p_kw.paragraph_format.right_indent = Inches(0.5)
    p_kw.paragraph_format.line_spacing = 1.0
    r_kw_label = p_kw.add_run('Keywords: ')
    r_kw_label.bold = True
    p_kw.add_run(
        'compute trade, FLOPs, artificial intelligence, data centers, '
        'comparative advantage, electricity costs, developing countries'
    )

    return title_el, author_el, ver_el, abs_text_el, kw_el


def write_introduction(doc, body, hmap):
    print("Inserting Section 1: Introduction...")
    cur = mkh(doc, body, hmap['1'].getprevious(), '1. Introduction', level=1)

    # Para 1: AI compute demand + electricity footprint (consolidated)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The expansion of artificial intelligence drives the demand for computational '
        'resources. The compute used to train the largest AI '
        'models has been doubling every six months since 2010 (Epoch AI 2024). '
        'Data centers accounted for approximately 1.5% of global electricity demand in '
        '2024\u2014more than the electricity consumption of France\u2014a share projected '
        'to more than double by 2030 '
        '(IEA 2025). '
        'AI-oriented facilities are qualitatively different from traditional cloud or enterprise '
        'data centers. They deploy thousands of graphic processing units (GPUs) at power '
        'densities of 40\u2013100 kW per rack '
        '(versus 5\u201310 kW in conventional facilities), and can consume over 500,000 gallons of cooling '
        'water per day (Turner Lee and West 2025).'
    )
    # footnote 2 removed (unclear)

    # Para 3: FLOP exporting (v28: Task 5 AI language fixes)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'This surge in demand creates an export opportunity. '
        'We refer to the production of compute services in one country '
        'for consumption in another as '
    )
    add_italic(p, 'FLOP exporting')
    p.add_run(
        '. Rather than '
        'exporting raw energy resources as primary commodities, '
        'countries can use cheap electricity to produce a higher value-added digital service. '
        'For resource-rich developing countries, FLOP exporting could offer a route up the '
        'value chain, bypassing the manufacturing stage traditionally required for '
        'such upgrading (Hausmann et al. 2007).'
    )

    # v28: Task 3 — ECA paragraph DELETED

    # Para 6: Real data center plans + profit estimate (v28: Task 4 compressed)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Recent megaprojects confirm that FLOP exporting is technically and commercially '
        'feasible. Armenia is deploying 50,000 GPUs in a $4 billion '
        'investment (Firebird 2026), Kenya, Saudi Arabia, and Malaysia have each attracted '
        'billion-dollar data center commitments, and cloud computing exports already exceed '
        '$9 billion annually (World Bank 2025). '
        'A 40 MW data center in Kyrgyzstan could '
        'generate annual revenue of $630\u2013950 million at wholesale contract rates, '
        'equivalent to over 15% of Kyrgyzstan\u2019s $3.8 billion in goods exports (World Bank 2024).'
    )
    make_footnote(p, 'At $0.038/kWh electricity, a 40 MW facility houses approximately '
                  '53,000 GPUs with production costs of $453 million per year. A Kyrgyz operator '
                  'would most likely sell at wholesale or long-term contract rates of roughly '
                  '$0.80\u20131.20/GPU-hour, yielding gross revenue of $630\u2013950 million. '
                  'Hyperscaler retail rates ($2.00\u20132.50/GPU-hour) represent an upper bound that '
                  'is unlikely for a new market entrant. Even at the wholesale lower bound, '
                  'this exceeds 15% of Kyrgyzstan\u2019s $3.8 billion in goods exports (2024).', 5)

    # Para 9: First paper + contributions
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Recent work examines compute governance and the geography of AI infrastructure '
        '(Sastry et al. 2024, Lehdonvirta et al. 2024, '
        'Pilz et al. 2025), but no formal trade model of compute exists. '
        'We offer the first such model, treating FLOPs as commodities produced and exported '
        'according to Ricardian comparative advantage. '
        'We make three contributions. First, we develop a capacity-constrained '
        'Ricardian model in which countries produce and export compute services. '
        'An iceberg trade cost captures latency degradation for inference, and a bilateral '
        'sovereignty premium captures geopolitical and regulatory frictions. '
        'Capacity ceilings generate scarcity rents and predictions about concentration '
        'and trade patterns. '
        'Second, we calibrate the model for 85 countries using data on electricity '
        'prices, climate, data center construction costs, and inter-country network latency, '
        'correcting for energy subsidies that distort headline cost rankings. '
        'Third, we characterize the resulting trade regimes\u2014which countries export, which '
        'import, and which adopt hybrid strategies\u2014and quantify the welfare cost of '
        'the sovereignty premium.'
    )

    # Para 11: Calibration findings preview
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The calibration reveals that cheap electricity and favorable cooling conditions '
        'are necessary but not sufficient for competitive compute exporting. Hardware '
        'amortization accounts for approximately 90 percent of the compute unit cost and '
        'is identical across countries, compressing the total cost spread to roughly '
        '12 percent under raw electricity tariffs (20 percent after efficiency adjustment). '
        'Once production-efficiency penalties for weak governance '
        'are applied, the cost ranking changes substantially. '
        'Kyrgyzstan, the cheapest producer under cost-recovery pricing, drops from 1st '
        'to 5th after governance adjustment\u2009\u2014\u2009but countries combining cheap energy '
        'with adequate institutions, including China, Kosovo, and Vietnam, remain in the top 15. '
        'Durable comparative advantage requires credible '
        'production efficiency alongside low energy costs, while market access depends on '
        'bilateral trust and regulatory alignment.'
    )



def write_literature(doc, body, hmap):
    print("Inserting Section 2: Related Literature...")
    cur = mkh(doc, body, hmap['1'].getprevious(), '2. Related Literature', level=1)

    # Para 1: AI comparative advantage + value chain upgrading (merged)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Goldfarb and Trefler (2018) argue that AI shifts comparative advantage toward '
        'countries with data, human capital, and institutional capacity. Our model introduces '
        'a complementary mechanism where comparative advantage in compute '
    )
    add_italic(p, 'production')
    p.add_run(
        ' depends on electricity costs and climate, so resource-rich countries could become '
        'compute exporters without domestic AI research industries. Korinek and Stiglitz (2021) '
        'suggest that developing countries could be left behind in the AI '
        'revolution. FLOP exporting would allow energy-rich developing '
        'countries to participate in that revolution. The concept of FLOP exporting as value chain upgrading '
        'connects to Hausmann et al. (2007), who show that what a country exports '
        'matters for growth. Lim\u00E3o and Venables (2001) demonstrate that infrastructure quality '
        'determines trade costs. In our model, network infrastructure plays the analogous role '
        'for digital trade.'
    )

    # IT-offshoring contrast (ChatGPT review)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Unlike the offshoring of IT services, which is labor-intensive and skill-biased '
        '(Blinder 2006), FLOP exporting is capital- and energy-intensive: the binding input '
        'is cheap electricity, not cheap labor, so the set of potential exporters is '
        'fundamentally different.'
    )

    # Para 2: Data center location literature
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Several studies examine the determinants of data center location. '
        'Flucker et al. (2013) show that climate affects data center '
        'cooling costs. '
        'Liu et al. (2023) study data center placement under renewable energy constraints. '
        'These studies focus on where firms should build data centers. '
        'In international trade theory, Brainard (1997) formalizes the proximity-concentration '
        'trade-off between serving a market locally and concentrating production abroad, '
        'and Helpman et al. (2004) extend this to heterogeneous firms choosing between '
        'exporting and FDI.'
    )

    # Para 3: Compute governance literature
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'On the governance side, Sastry et al. (2024) '
        'argue that compute is well-suited for regulation because governments can track '
        'the number of chips in circulation, restrict who can buy them, and measure how much computation they '
        'perform. Lehdonvirta et al. (2024) '
        'map the global geography of cloud GPU infrastructure, distinguishing a '
        '\u201CCompute North\u201D with training-capable hardware from a '
        '\u201CCompute South\u201D limited to inference-grade chips. '
        'Pilz et al. (2025) project that AI data center power demand '
        'could reach 327 GW by 2030 and that domestic power shortages may push '
        'compute infrastructure abroad. The World Bank (2025) documents the resulting '
        'global compute divide: high-income countries hold 77% of colocation data center '
        'capacity and account for 87% of cloud computing exports (Stojkoski et al. 2024), '
        'but the report offers descriptive evidence without a formal framework linking '
        'production costs to trade patterns. On the industrial organization side, '
        'Biglaiser et al. (2024) survey the economics of cloud '
        'markets, including switching costs, egress fees, and platform competition among '
        'hyperscalers, but the supply-side question of where compute is produced and '
        'whether developing countries can become competitive exporters has not been addressed. '
        'Stojkoski et al. (2024) document the geography of cloud exports but treat cloud '
        'services as homogeneous; the present model adds supply-side cost structure and a '
        'training\u2013inference distinction that generates location-specific comparative advantage. '
        'The present model provides the formal framework these studies lack.'
    )


def write_production_technology(doc, body, hmap):
    print("Rewriting Section 3.1 (Production Technology, merged with Section 3 opening)...")

    # Clear everything between Section 3 heading and Section 3.2 heading,
    # preserving the Section 3.1 heading element
    all_now = list(body)
    s1i = all_now.index(hmap['1'])
    s12i = all_now.index(hmap['1.2'])
    s11_el = hmap['1.1']
    for el in all_now[s1i + 1:s12i]:
        if el is not s11_el:
            body.remove(el)
    cur = hmap['1']  # start after Section 3 heading

    # Para 1: linking paragraph from lit review to model (before 3.1 subtitle)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'This section models compute as a tradable good with '
        'country-specific production costs, a delivery cost that depends on whether the '
        'workload is training (latency-insensitive) or inference (latency-sensitive), and '
        'a sovereignty premium reflecting governments\u2019 preference for domestic production.'
    )

    cur = hmap['1.1']  # continue after 3.1 subtitle

    # Para 2: formal setup
    p, cur = mkp(doc, body, cur)
    p.add_run('Consider ')
    omath(p, [_v('N')])
    p.add_run(
        ' countries, each capable of producing compute services. The cost of '
        'producing a unit of compute in country '
    )
    omath(p, [_v('j')])
    p.add_run(
        ' depends on three inputs: electricity, hardware, and data center construction.'
    )

    # PUE inlined (no display equation) — merged with equation lead-in
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'A data center consumes electricity not only for its GPUs but also for '
        'cooling, power distribution, and lighting. '
        'This overhead is measured by the '
    )
    add_italic(p, 'power usage effectiveness')
    p.add_run(
        ' (PUE), the ratio of total facility energy consumption to IT equipment energy consumption. '
        'The cost per GPU-hour in country '
    )
    omath(p, [_v('j')])
    p.add_run(' is:')
    make_footnote(p, 'The linear PUE model is a simplification. Modern liquid and immersion '
                  'cooling technologies can achieve PUE \u2248 1.2 even in hot climates, flattening the '
                  'temperature\u2013PUE relationship. The robustness check in Section 6 confirms that '
                  'the results are insensitive to this specification. '
                  'Google (2024) reports a fleet-wide trailing '
                  'twelve-month PUE of 1.10.', 7)
    p.paragraph_format.space_after = Pt(2)

    # Equation (2): cost function (with networking η)
    _, cur = omath_display(doc, body, cur, [
        _msub('c', 'j'), _t(' = '),
        _t('PUE('), _msub('\u03B8', 'j'), _t(') \u00b7 '),
        _v('\u03B3'), _t(' \u00b7 '),
        _msub('p', 'E,j'), _t(' + '),
        _v('\u03C1'), _t(' + '),
        _v('\u03B7'), _t(' + '),
        _msub('p', 'L,j'), _t(' / ('),
        _v('D'), _t(' \u00b7 '), _v('H'), _t('),'),
    ], eq_num='1')

    # Equation explanation (streamlined — no "first term/second term" redundancy)
    p, cur = mkp(doc, body, cur)
    p.add_run('where ')
    omath(p, [_v('\u03B3')])
    p.add_run(
        ' is GPU power draw (kW), '
    )
    omath(p, [_msub('p', 'E,j')])
    p.add_run(' is the electricity price ($/kWh), ')
    omath(p, [_v('\u03C1'), _t(' = '), _msub('P', 'GPU'),
              _t(' / ('), _v('L'), _t(' \u00b7 '), _v('H'),
              _t(' \u00b7 '), _v('\u03B2'), _t(')')])
    p.add_run(
        ' is amortized hardware cost per GPU-hour '
        '('
    )
    omath(p, [_msub('P', 'GPU')])
    p.add_run(' = purchase price, ')
    omath(p, [_v('L')])
    p.add_run(' = lifetime in years, ')
    omath(p, [_v('H')])
    p.add_run(' = 8,766 hours per year, ')
    omath(p, [_v('\u03B2')])
    p.add_run(' = utilization rate),')
    make_footnote(p, 'For the NVIDIA H100: $25,000 / (3 years \u00d7 8,766 hours/year \u00d7 70% '
                  'utilization) \u2248 $1.36/hr. Street prices have fallen to $18,000\u2013$22,000 '
                  'as of late 2025. Each GPU draws approximately 700 watts.', 9)
    p.add_run(' ')
    omath(p, [_v('\u03B7')])
    p.add_run(
        ' is amortized networking cost (high-speed interconnect such as InfiniBand), '
        'and the last term amortizes construction costs '
    )
    omath(p, [_msub('p', 'L,j')])
    p.add_run(
        ' ($/W of IT capacity) over the facility\u2019s lifetime '
    )
    omath(p, [_v('D')])
    p.add_run('. Both ')
    omath(p, [_v('\u03C1')])
    p.add_run(' and ')
    omath(p, [_v('\u03B7')])
    p.add_run(
        ' are determined in global hardware markets and are common across countries.'
    )
    make_footnote(p, 'China is developing an alternative domestic chip stack based on '
                  'Huawei\u2019s Ascend series (910B/910C) and other domestic accelerators. If these '
                  'achieve comparable FLOPs per watt at lower prices, China\u2019s effective \u03C1 could '
                  'diverge from the NVIDIA-based benchmark used here, potentially improving its '
                  'cost position despite export controls.', 10)
    p.add_run(
        ' Cross-country variation in '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' is therefore driven by electricity prices, climate, '
        'and construction costs.'
    )

    # Endowment paragraph
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Countries export goods intensive in their '
        'abundant factors (Ohlin 1933). For compute production, the relevant endowment is not electricity '
        'per se but the natural resources that generate it\u2014hydropower '
        '(Kyrgyzstan, Ethiopia, Georgia), oil and gas (Iran, Turkmenistan, Qatar), solar '
        'irradiance (North Africa, the Gulf), and geothermal energy (Kenya, Iceland). '
        'The electricity price '
    )
    omath(p, [_msub('p', 'E,j')])
    p.add_run(
        ' in equation (1) is therefore a reduced-form expression for country '
    )
    omath(p, [_v('j')])
    p.add_run('\u2019s energy resource endowment.')


def write_trade_costs(doc, body, hmap):
    print("Rewriting Section 3.2 (Trade Costs)...")

    all_now = list(body)
    s12i = all_now.index(hmap['1.2'])
    s2i = all_now.index(hmap['2'])
    for el in all_now[s12i + 1:s2i]:
        body.remove(el)
    cur = hmap['1.2']

    # Redefine two service types (merged with latency definition)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Countries produce and trade two types of compute services that differ in their '
        'offshoring costs. '
    )
    add_italic(p, 'Training services')
    p.add_run(
        ' (denoted T) encompass batch workloads such as model training, fine-tuning, '
        'and large-scale data '
        'processing. Training a state-of-the-art AI model can take weeks to months across '
        'thousands of GPUs. The client ships its data to a data center, the computation '
        'executes locally, and the output is returned to the client. Since neither input '
        'nor output is time-sensitive, network latency plays no role. '
    )
    add_italic(p, 'Inference services')
    p.add_run(
        ' (denoted I) encompass real-time workloads such as chatbot responses, autonomous '
        'decisions, and interactive '
        'agents. Each query must travel to the server and back within milliseconds, so the '
        'service quality degrades as delivery delays (latency) increase. '
    )
    add_italic(p, 'Latency')
    p.add_run(', denoted ')
    omath(p, [_msub('l', 'jk')])
    p.add_run(
        ', is the round-trip time for a data packet to travel from seller '
    )
    omath(p, [_v('j')])
    p.add_run(' to buyer ')
    omath(p, [_v('k')])
    p.add_run(
        ', typically 5\u201310 ms within a country and over 150 ms across continents '
        '(Appendix F summarizes workload types and their latency sensitivity).'
    )

    # v24: Bilateral sovereignty premium definition
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Governments and firms may prefer to process data domestically for national '
        'security, regulatory compliance, or political reasons. This is captured by a '
        'bilateral '
    )
    add_italic(p, 'sovereignty premium')
    p.add_run(' ')
    omath(p, [_msub('\u03BB', 'ij'), _t(' \u2265 0')])
    p.add_run(
        ', which acts as a proportional markup on the cost of compute sourced from seller '
    )
    omath(p, [_v('i')])
    p.add_run(' by buyer ')
    omath(p, [_v('j')])
    p.add_run(
        '. When buyer '
    )
    omath(p, [_v('j')])
    p.add_run(' sources compute from seller ')
    omath(p, [_v('i'), _t(' \u2260 '), _v('j')])
    p.add_run(', the effective cost is inflated by the factor ')
    omath(p, [_t('(1 + '), _msub('\u03BB', 'ij'), _t(')')])
    p.add_run(
        '. The sovereignty premium is zero for domestic production ('
    )
    omath(p, [_msub('\u03BB', 'ii'), _t(' = 0')])
    p.add_run(
        '). The bilateral premium is modeled as a function of three observables: '
        'geopolitical alignment '
    )
    omath(p, [_msub('G', 'ij')])
    p.add_run(
        ', measured by UN General Assembly ideal-point distance '
        '(Bailey, Strezhnev, and Voeten 2017); regulatory compatibility '
    )
    omath(p, [_msub('R', 'ij')])
    p.add_run(
        ', coded as 1 for country pairs covered by a mutual data-adequacy agreement '
        'and 0 otherwise; and a sanctions indicator '
    )
    omath(p, [_msub('S', 'ij')])
    p.add_run(
        ' equal to 1 if either country maintains comprehensive trade sanctions against the other:'
    )
    p.paragraph_format.space_after = Pt(2)

    # v24: NEW Equation (2) — bilateral λ_{ij}
    _, cur = omath_display(doc, body, cur, [
        _msub('\u03BB', 'ij'), _t(' = '),
        _msub('\u03B1', '1'), _t(' \u00b7 '),
        _msub('G', 'ij'), _t(' + '),
        _msub('\u03B1', '2'), _t(' \u00b7 (1 \u2212 '),
        _msub('R', 'ij'), _t(') + '),
        _msub('\u03B1', '3'), _t(' \u00b7 '),
        _msub('S', 'ij'), _t('.'),
    ], eq_num='2')

    p, cur = mkp(doc, body, cur)
    p.add_run('The calibration of these coefficients is described in Section 6.1.')

    # v28: Equation (2') — Hyperscaler FDI trust specification
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'This specification treats the host country as the relevant trust counterparty. '
        'When a hyperscaler intermediates the transaction, the relevant counterparty is '
        'the operator, not the host. For a facility in host country '
    )
    omath(p, [_v('j')])
    p.add_run(' operated by a hyperscaler headquartered in country ')
    omath(p, [_v('h')])
    p.add_run(', the effective premium becomes:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msubsup('\u03BB', 'jk', 'FDI'), _t(' = '),
        _msub('\u03B1', '1'), _t(' \u00b7 '),
        _v('G'), _t('('), _v('h'), _t(', '), _v('k'), _t(') + '),
        _msub('\u03B1', '2'), _t(' \u00b7 (1 \u2212 '),
        _v('R'), _t('('), _v('h'), _t(', '), _v('k'), _t(')) + '),
        _msub('\u03B1', '3'), _t(' \u00b7 '),
        _v('S'), _t('('), _v('j'), _t(', '), _v('k'), _t(').'),
    ], eq_num='2\u2032')

    p, cur = mkp(doc, body, cur)
    p.add_run('Equation (2\u2032) nests equation (2): when ')
    omath(p, [_v('h'), _t(' = '), _v('j')])
    p.add_run(
        ' (domestically owned facility), the two coincide. '
        'Section 6.2 reports calibration results under both specifications.'
    )

    # Equation (3): delivered cost with ξ_j^{eff} (was eq 2)
    p, cur = mkp(doc, body, cur)
    p.add_run('The delivered cost of service ')
    omath(p, [_v('s'), _t(' \u2208 {'), _v('T'), _t(', '), _v('I'), _t('}')])
    p.add_run(' from seller ')
    omath(p, [_v('j')])
    p.add_run(' to buyer ')
    omath(p, [_v('k')])
    p.add_run(' is:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msub('P', 's'), _t('('), _v('j'), _t(', '), _v('k'),
        _t(') = (1 + '), _msub('\u03BB', 'jk'),
        _t(') \u00b7 (1 + '), _msub('\u03C4', 's'), _t(' \u00b7 '),
        _msub('l', 'jk'), _t(') \u00b7 ['),
        _v('\u03C1'), _t(' + '),
        _mfrac([_msub('c', 'j'), _t(' \u2212 '), _v('\u03C1')],
               [_msubsup('\u03BE', 'j', 'eff')]),
        _t('],'),
    ], eq_num='3')

    p, cur = mkp(doc, body, cur)
    p.add_run('where ')
    omath(p, [_msubsup('\u03BE', 'j', 'eff'), _t(' \u2208 (0, 1]')])
    p.add_run(
        ' is a production-efficiency index that captures factors reducing the real output '
        'of a GPU-hour in country '
    )
    omath(p, [_v('j')])
    p.add_run(
        '. The effective index incorporates an institutional floor: '
    )
    omath(p, [_msubsup('\u03BE', 'j', 'eff'), _t(' = '),
              _msub('\u03BE', 'floor'), _t(' + (1 \u2212 '),
              _msub('\u03BE', 'floor'), _t(') \u00b7 '),
              _msubsup('G', 'j', '\u03C9'), _t(' \u00b7 '),
              _msubsup('R', 'j', '1\u2212\u03C9')])
    p.add_run(', where ')
    omath(p, [_v('G')])
    p.add_run(
        ' is the Rule of Law percentile (rescaled to [0,\u20091]) '
        '(World Bank 2024)'
    )
    make_footnote(p,
        'The Rule of Law percentile is a country-level composite capturing perceptions '
        'of contract enforcement, property rights, courts, and police across the entire '
        'economy. The governance relevant to a data center investment is narrower: '
        'credibility of long-term power purchase agreements, stability of the tax regime, '
        'and enforceability of international arbitration clauses. Sector-specific measures '
        'such as bilateral investment treaty coverage or ICSID arbitration caseload would '
        'be preferable if available across all 85 countries.',
        23)
    p.add_run(', ')
    omath(p, [_v('R')])
    p.add_run(
        ' is grid reliability '
        '(World Bank Enterprise Surveys 2025), also rescaled to [0,\u20091], and '
    )
    omath(p, [_v('\u03C9'), _t(' = 0.50')])
    p.add_run(
        ' assigns equal weight to governance and grid reliability. '
        'The floor '
    )
    omath(p, [_msub('\u03BE', 'floor')])
    p.add_run(
        ' reflects the minimum operational quality achievable through investor-side '
        'risk mitigation, calibrated to evidence on zone-level governance in '
        'special economic zones (Section 6). '
        'The efficiency adjustment applies only to the non-hardware cost component: '
        'since hardware is priced on global markets, governance penalties operate on '
        'local costs rather than on the full unit '
        'cost (see equation 3). The baseline sets '
    )
    omath(p, [_msub('\u03BE', 'floor'), _t(' = 0.30')])
    p.add_run('. '
        'The index approaches one for countries with strong institutions and '
        'reliable power, and falls below one for countries where institutional weakness '
        'or operational risk reduces effective delivery. '
        'Sanctions, data-protection trust, and regulatory compatibility '
        'do not enter '
    )
    omath(p, [_msubsup('\u03BE', 'j', 'eff')])
    p.add_run(
        ' because they affect trade willingness between specific country pairs, not the '
        'physical cost of producing a GPU-hour; these frictions are captured by the '
        'bilateral sovereignty premium '
    )
    omath(p, [_msub('\u03BB', 'ij')])
    p.add_run('.')

    # New paragraph: τ latency degradation (split from ξ paragraph)
    p, cur = mkp(doc, body, cur)
    p.add_run('The parameter ')
    omath(p, [_v('\u03C4')])
    p.add_run(
        ' measures the rate of quality degradation per millisecond of round-trip latency, with '
    )
    omath(p, [_msub('\u03C4', 'T'), _t(' = 0')])
    p.add_run(' and ')
    omath(p, [_msub('\u03C4', 'I'), _t(' = '), _v('\u03C4'), _t(' > 0')])
    p.add_run(
        '. For training ('
    )
    omath(p, [_msub('\u03C4', 'T'), _t(' = 0')])
    p.add_run(
        '), the delivered cost is the sum of the production cost and the sovereignty '
        'markup. '
        'For inference, the delivered cost increases with latency at rate '
    )
    omath(p, [_msub('\u03C4', 'I')])
    p.add_run('. Beyond a threshold ')
    omath(p, [_mbar('l')])
    p.add_run(
        ' (typically 200\u2013300 ms for interactive applications), the service becomes '
        'unusable regardless of price, modeled as '
    )
    omath(p, [_msub('P', 'I'), _t('('), _v('j'), _t(', '), _v('k'),
              _t(') = \u221E')])
    p.add_run(' if ')
    omath(p, [_msub('l', 'jk'), _t(' > '), _mbar('l')])
    p.add_run('.')


def renumber_sections(hmap):
    print("Renumbering sections...")
    # v20 structure: 1=Intro, 2=Lit, 3=Model(3.1,3.2), 4=Equil Props, 5=Data, 6=Calib, 7=Robustness, 8=Concl
    # v8 headings: 1→3 (Model), 1.1→3.1, 1.2→3.2, 2→4 (Equil Props), 4→6 (Calib), 5→8 (Concl)
    # Section 3 (Make-or-Buy) content will be absorbed; heading removed by write functions
    renumber = [
        ('1.2', '1.2', '3.2'), ('1.1', '1.1', '3.1'), ('1', '1.', '3.'),
        ('2', '2.', '4.'),
        ('4', '4.', '6.'), ('5', '5.', '8.'),
    ]
    for key, old, new in renumber:
        if key in hmap:
            el = hmap[key]
            for t in el.findall(f'.//{qn("w:t")}'):
                if t.text and old in t.text:
                    t.text = t.text.replace(old, new, 1)
                    break
    # Rename Section 6 heading: "Discussion" → "Results"
    if '4' in hmap:
        for t in hmap['4'].findall(f'.//{qn("w:t")}'):
            if t.text and 'Discussion' in t.text:
                t.text = t.text.replace('Discussion', 'Results')
                break
    # Rename Section 3.1 heading
    if '1.1' in hmap:
        for t in hmap['1.1'].findall(f'.//{qn("w:t")}'):
            if t.text and 'Production Technology' in t.text:
                t.text = t.text.replace('Production Technology',
                                        'Production Technology and Cost Structure')
                break


def write_demand(doc, body, hmap, demand_data):
    """Section 3.3 Demand — moved from old Section 5 (Make-or-Buy)."""
    print("Inserting Section 3.3 (Demand)...")

    # Insert 3.3 heading after the end of Section 3.2 content (before old Section 2 heading)
    s2 = hmap['2']  # v8 heading "2", becomes Section 4
    cur = mkh(doc, body, s2.getprevious(), '3.3 Global Compute Demand', level=2)

    # Demand specification: Equation (4)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The model is closed by specifying demand for compute services. Let '
    )
    omath(p, [_msub('q', 'k')])
    p.add_run(
        ' denote the volume of compute purchased by buyer '
    )
    omath(p, [_v('k')])
    p.add_run(
        '. We measure compute demand '
        'using installed data center capacity in megawatts (MW):'
    )
    p.paragraph_format.space_after = Pt(2)

    # Equation (4): q_k = ω_k · Q
    _, cur = omath_display(doc, body, cur, [
        _msub('q', 'k'), _t(' = '),
        _msub('\u03C9', 'k'), _t(' \u00b7 '), _v('Q'),
        _t(',     '),
        _msub('\u03C9', 'k'), _t(' = '),
        _msub('M', 'k'), _t(' / '),
        _nary('\u2211', [_v("k\u2032")], [],
              [_msub('M', "k\u2032")]), _t(','),
    ], eq_num='4')

    p, cur = mkp(doc, body, cur)
    p.add_run('where ')
    omath(p, [_v('Q')])
    p.add_run(
        ' is the total global compute spending and '
    )
    omath(p, [_msub('\u03C9', 'k')])
    p.add_run(
        ' is country '
    )
    omath(p, [_v('k')])
    p.add_run(
        '\u2019s share of global demand, measured by its share of installed data center '
        'capacity (MW).'
    )
    make_footnote(p,
                  'Installed capacity is preferable to GDP as a demand proxy because '
                  'compute consumption is driven by data center infrastructure, not aggregate '
                  'income. Ireland and the Netherlands, for example, host far more capacity '
                  'per capita than their GDP shares would predict, while large economies like '
                  'India and Brazil account for modest shares of global data center power.',
                  20)

    # Training/inference split
    p, cur = mkp(doc, body, cur)
    p.add_run('Demand splits between training and inference. Training demand is ')
    omath(p, [_msub('q', 'Tk'), _t(' = '), _v('\u03B1'),
              _t(' \u00b7 '), _msub('q', 'k')])
    p.add_run(' and inference demand is ')
    omath(p, [_msub('q', 'Ik'), _t(' = (1 \u2212 '), _v('\u03B1'),
              _t(') \u00b7 '), _msub('q', 'k')])
    p.add_run(', where ')
    omath(p, [_v('\u03B1'), _t(' \u2208 (0, 1)')])
    p.add_run(
        ' is the exogenous training share. '
        'The parameter '
    )
    omath(p, [_v('\u03B1')])
    p.add_run(
        ' should be interpreted as the share of compute that is fully '
        'latency-insensitive and freely offshorable; the effective offshorable share may '
        'be smaller as intermediate workloads (agentic inference, fine-tuning) grow.'
    )
    make_footnote(p,
                  'Emerging workload categories, notably agentic inference (long-running, multi-step '
                  'reasoning tasks) and fine-tuning (rapid iterative retraining on proprietary data), '
                  'occupy a middle ground, tolerating moderate latency but requiring sustained GPU '
                  'allocation and proximity to data. Using installed capacity to proxy demand is a '
                  'static assumption; endogenizing demand, for instance, proportional to GDP or digital '
                  'adoption, is a natural extension.', 21)


def write_sourcing_and_equilibrium(doc, body, hmap, demand_data):
    """Section 3.4 Sourcing and Market Equilibrium — simplified (3 display equations)."""
    print("Inserting Section 3.4 (Sourcing and Market Equilibrium)...")

    s2 = hmap['2']
    cur = mkh(doc, body, s2.getprevious(), '3.4 Sourcing and Market Equilibrium', level=2)

    # Sourcing rule: Equation (5)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'For each service type '
    )
    omath(p, [_v('s'), _t(' \u2208 {'), _v('T'), _t(', '), _v('I'), _t('}')])
    p.add_run(', each buyer ')
    omath(p, [_v('k')])
    p.add_run(' chooses the source that minimizes the delivered cost:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msubsup('j', 's', '*'), _t('('), _v('k'),
        _t(') = '),
        _limlow([_t('arg min')], [_v('j')]),
        _t(' '), _msub('P', 's'), _t('('), _v('j'), _t(', '), _v('k'), _t(').'),
    ], eq_num='5')

    # Capacity ceiling — defined before training market (which references it)
    p, cur = mkp(doc, body, cur)
    p.add_run('Each country ')
    omath(p, [_v('j')])
    p.add_run(' is characterized by a capacity ceiling ')
    omath(p, [_mbar_sub('K', 'j')])
    p.add_run(
        ', measured in GPU-hours per period, representing the maximum volume of compute '
        'the country can supply. This ceiling reflects the joint constraint of grid '
        'electricity availability, institutional capacity for data center permitting and '
        'construction, and access to GPU financing.'
    )

    # Training market
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Training market. ')
    p.add_run('Since ')
    omath(p, [_msub('\u03C4', 'T'), _t(' = 0')])
    p.add_run(
        ', training is a homogeneous good with no distance-related quality degradation. '
        'Country '
    )
    omath(p, [_v('k')])
    p.add_run(' imports training whenever the world price, after adding the bilateral sovereignty '
              'premium, is lower than producing domestically: ')
    omath(p, [_t('(1 + '), _msub('\u03BB', 'jk'), _t(') \u00b7 '),
              _msub('p', 'T'), _t(' < '), _msub('c', 'k')])
    p.add_run(
        ', where '
    )
    omath(p, [_msub('p', 'T')])
    p.add_run(
        ' is the competitive world training price. '
        'In the capacity-constrained equilibrium, the cheapest producer supplies up to its capacity, '
        'then the next cheapest enters, and so on until demand is met. '
        'The marginal training exporter '
    )
    omath(p, [_msub('m', 'T')])
    p.add_run(
        ' is the producer whose entry just satisfies total export demand. '
        'The equilibrium training price equals the marginal exporter\u2019s cost: '
    )
    # p_T = c_{(m_T)} — inline (was display Eq 5)
    # Build subscript manually: c with subscript containing "(m_T)"
    c_sub = OxmlElement('m:sSub')
    c_sub.append(OxmlElement('m:sSubPr'))
    c_e = OxmlElement('m:e')
    c_e.append(_mr('c', True))
    c_sub.append(c_e)
    c_sub_content = OxmlElement('m:sub')
    c_sub_content.append(_mr('(', False))
    mt_sub = OxmlElement('m:sSub')
    mt_sub.append(OxmlElement('m:sSubPr'))
    mt_e = OxmlElement('m:e')
    mt_e.append(_mr('m', True))
    mt_sub.append(mt_e)
    mt_s = OxmlElement('m:sub')
    mt_s.append(_mr('T', True))
    mt_sub.append(mt_s)
    c_sub_content.append(mt_sub)
    c_sub_content.append(_mr(')', False))
    c_sub.append(c_sub_content)
    omath(p, [_msub('p', 'T'), _t(' = '), c_sub])
    p.add_run('.')

    # Rents and shadow value (K̄_j already defined before training market)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Without capacity constraints, '
    )
    omath(p, [_msub('m', 'T'), _t(' = 1')])
    p.add_run(
        ' and the cheapest country serves all demand at its own cost, earning zero rent. '
        'With binding capacity constraints, '
    )
    omath(p, [_msub('m', 'T'), _t(' > 1')])
    p.add_run(
        ', the price rises to the cost of the marginal entrant, and all infra-marginal '
        'exporters earn positive rents: '
    )
    omath(p, [_msub('\u03C0', 'Tj'), _t(' = ('), _msub('p', 'T'),
              _t(' \u2212 '), _msub('c', 'j'), _t(') \u00b7 '),
              _msub('K', 'Tj')])
    p.add_run(
        '. For a capacity-constrained exporter, the shadow value '
    )
    omath(p, [_msub('\u03BC', 'j')])
    p.add_run(
        ' of the capacity constraint measures how much one additional GPU-hour would be worth. '
        'It equals the difference between the market price and country\u2019s production price.'
    )

    # Inference: Equation (6)  [was display Eq 7, renumbered after inlining p_T]
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Inference market. ')
    p.add_run('Since ')
    omath(p, [_msub('\u03C4', 'I'), _t(' = '), _v('\u03C4'), _t(' > 0')])
    p.add_run(
        ', inference suffers distance-dependent quality degradation. '
        'The inference market for buyer '
    )
    omath(p, [_v('k')])
    p.add_run(
        ' is localized, as only countries with latency '
    )
    omath(p, [_msub('l', 'jk'), _t(' \u2264 '), _mbar('l')])
    p.add_run(
        ' can participate, and each faces a different delivered cost. '
        'The delivered inference price for buyer '
    )
    omath(p, [_v('k')])
    p.add_run(' is:')
    p.paragraph_format.space_after = Pt(2)

    # Build l_{m_I(k), k} and c_{m_I(k)} with (k) INSIDE the subscript
    # l subscripted with "m_I(k), k"
    l_sub = OxmlElement('m:sSub')
    l_sub.append(OxmlElement('m:sSubPr'))
    l_e = OxmlElement('m:e')
    l_e.append(_mr('l', True))
    l_sub.append(l_e)
    l_s = OxmlElement('m:sub')
    l_s.append(_msub('m', 'I'))
    l_s.append(_mr('(', False))
    l_s.append(_mr('k', True))
    l_s.append(_mr('),\u2009', False))
    l_s.append(_mr('k', True))
    l_sub.append(l_s)

    # c subscripted with "m_I(k)"
    c_sub2 = OxmlElement('m:sSub')
    c_sub2.append(OxmlElement('m:sSubPr'))
    c_e2 = OxmlElement('m:e')
    c_e2.append(_mr('c', True))
    c_sub2.append(c_e2)
    c_s2 = OxmlElement('m:sub')
    c_s2.append(_msub('m', 'I'))
    c_s2.append(_mr('(', False))
    c_s2.append(_mr('k', True))
    c_s2.append(_mr(')', False))
    c_sub2.append(c_s2)

    _, cur = omath_display(doc, body, cur, [
        _msubsup('p', 'I', 'f'), _t('('), _v('k'), _t(') = (1 + '),
        _v('\u03C4'), _t(' \u00b7 '), l_sub,
        _t(') \u00b7 '), c_sub2, _t(','),
    ], eq_num='6')

    p, cur = mkp(doc, body, cur)
    p.add_run('where ')
    omath(p, [_msub('m', 'I'), _t('('), _v('k'), _t(')')])
    p.add_run(
        ' is the marginal inference supplier to '
    )
    omath(p, [_v('k')])
    p.add_run(
        ', determined by the capacity-constrained supply stack for '
    )
    omath(p, [_v('k')])
    p.add_run(
        '\u2019s inference market. '
        'Each GPU-hour of capacity is allocated to its highest-margin use, whether '
        'training exports, inference exports to various destinations, or domestic supply.'
    )

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The cost-based equilibrium identifies the set of countries that could profitably '
        'produce and export compute. Which countries from this set actually attract '
        'investment depends additionally on agglomeration economies, hyperscaler market '
        'structure, and network connectivity (Krugman 1991). The cloud compute market is '
        'dominated by three firms (AWS, Azure, Google Cloud) whose location decisions '
        'reflect scale economies and self-reinforcing colocation dynamics that the '
        'competitive framework abstracts from. The capacity ceilings '
    )
    omath(p, [_msub('K\u0304', 'j')])
    p.add_run(
        ' partially capture the resulting gap between cost-based potential and realized '
        'investment; Section 7 discusses these limitations further.'
    )


def write_equilibrium_properties(doc, body, hmap, demand_data):
    print("Rewriting Section 4 (Equilibrium Properties)...")

    # Clear content between section 4 heading (was v8 "2") and section 5 heading (was v8 "3")
    # Also remove the old Make-or-Buy heading and its content
    all_now = list(body)
    s4 = hmap['2']
    s4_old_next = hmap['4']  # v8 heading "4", becomes Section 6
    s4i = all_now.index(s4)
    s4_next_i = all_now.index(s4_old_next)
    # Remove everything from after Section 4 heading to before old Calibration heading
    # This removes both old Section 4 content AND old Section 5 (Make-or-Buy) heading+content
    for el in all_now[s4i + 1:s4_next_i]:
        body.remove(el)
    cur = s4

    # Also rename the heading text from "Comparative Advantage" to "Equilibrium Properties"
    for t in s4.findall(f'.//{qn("w:t")}'):
        if t.text and 'Comp' in t.text:
            t.text = t.text.replace('Comparative Advantage', 'Equilibrium Properties').replace(' Results', '')
            break

    # Introduction
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'This section derives the formal properties of the capacity-constrained '
        'equilibrium defined in Section 3. Full derivations appear in Appendix B.'
    )

    # Proposition 1: Country taxonomy (5-regime version)
    p, cur = mkp(doc, body, cur, space_before=6)
    r = p.add_run('Proposition 1 (Country Taxonomy). ')
    r.bold = True
    r.italic = True
    p.add_run(
        'With two service types (training with '
    )
    omath(p, [_msub('\u03C4', 'T'), _t(' = 0')])
    p.add_run(', and inference with ')
    omath(p, [_msub('\u03C4', 'I'), _t(' > 0')])
    p.add_run(
        ') and three possible statuses (export, domestic production, import), '
        'there are nine potential regime combinations. '
        'In equilibrium, only five are realized. '
        'Each regime is coded by two letters: the first for training status, '
        'the second for inference (E\u2009=\u2009export, D\u2009=\u2009domestic production, '
        'I\u2009=\u2009import).'
    )

    # (i) EE — now its own paragraph (B7b split)
    p, cur = mkp(doc, body, cur)
    p.add_run('(i) ')
    r_code = p.add_run('EE')
    r_code.bold = True
    p.add_run(
        '. Training exporter + inference exporter. '
        'The cheapest producers, with '
    )
    omath(p, [_msub('c', 'j'), _t(' < '),
              _msubsup('p', 'T', '*')])
    p.add_run(
        ', supply training globally and inference to nearby demand centers.'
    )

    # (ii)
    p, cur = mkp(doc, body, cur)
    p.add_run('(ii) ')
    r_code = p.add_run('IE')
    r_code.bold = True
    p.add_run(
        '. Training importer + inference exporter. '
        'Countries with '
    )
    omath(p, [_msub('c', 'j'), _t(' > '),
              _msubsup('p', 'T', '*')])
    p.add_run(
        ' that are not cheap enough to compete in the global training market '
        'but serve as regional inference hubs due to low costs and proximity '
        'to demand centers ('
    )
    omath(p, [_msub('d', 'ij')])
    p.add_run(' below the latency threshold ')
    omath(p, [_mbar('d')])
    p.add_run(').')

    # (iii)
    p, cur = mkp(doc, body, cur)
    p.add_run('(iii) ')
    r_code = p.add_run('ID')
    r_code.bold = True
    p.add_run(
        '. Training importer + inference domestic producer. '
        'Countries that import training but produce inference domestically, '
        'because the bilateral sovereignty premium '
    )
    omath(p, [_msub('\u03BB', 'ij')])
    p.add_run(
        ' or geographic isolation makes all foreign inference sources '
        'more expensive than domestic production.'
    )

    # (iv) and (v) in same paragraph
    p, cur = mkp(doc, body, cur)
    p.add_run('(iv) ')
    r_code = p.add_run('DD')
    r_code.bold = True
    p.add_run(
        '. Domestic producer of both. Countries where '
    )
    omath(p, [_msub('c', 'k'), _t(' \u2264 (1 + '),
              _msub('\u03BB', 'jk'), _t(') \u00b7 '),
              _msubsup('p', 'T', '*')])
    p.add_run(
        ', so the bilateral sovereignty premium is large enough to justify '
        'domestic production of both training and inference.'
    )

    p, cur = mkp(doc, body, cur)
    p.add_run('(v) ')
    r_code = p.add_run('II')
    r_code.bold = True
    p.add_run(
        '. Importer of both. High-cost countries with '
    )
    omath(p, [_msub('c', 'k'), _t(' > (1 + '),
              _msub('\u03BB', 'jk'), _t(') \u00b7 '),
              _msubsup('p', 'T', '*')])
    p.add_run(
        ' that import both training and inference from cheaper or closer suppliers.'
    )

    # Ruling out the remaining four
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The remaining four combinations cannot arise in equilibrium. '
        'A training exporter cannot simultaneously produce inference domestically '
        'or import it: a country cheap enough to win a global training competition ('
    )
    omath(p, [_msub('c', 'j'), _t(' < '),
              _msubsup('p', 'T', '*')])
    p.add_run(
        ') is necessarily cheap enough to export inference to nearby buyers '
        '(Proposition 4). A country that produces training domestically cannot '
        'import inference, because the sovereignty premium that justifies domestic '
        'training production also justifies domestic inference production, which faces '
        'latency degradation when imported. A domestic training producer that '
        'exports inference would require simultaneously high sovereignty preference '
        'and low costs, a combination ruled out by the model\u2019s cost ordering. '
    )
    p._element.append(make_bookmark(130, 'Table1txt'))
    p._element.append(make_hyperlink('Table1', 'Table 1'))
    p._element.append(make_bookmark_end(130))
    p.add_run(' summarizes this taxonomy.')

    # Proposition 2: Concentration
    p, cur = mkp(doc, body, cur, space_before=6)
    r = p.add_run('Proposition 2 (Capacity Constraints Reduce Concentration). ')
    r.bold = True
    r.italic = True
    p.add_run(
        'Define the Herfindahl\u2013Hirschman Index (HHI), '
        'a standard measure of market concentration that is equal to the sum of squared market shares'
        ', for training market concentration as '
    )
    omath(p, [_msub('HHI', 'T'), _t(' = '),
              _nary('\u2211', [_v('j')], [],
                    [_msup('(', '2', False, False)]),
              _msub('K', 'Tj'), _t('/'),
              _msubsup('Q', 'T', 'X'),
              _msup(')', '2', False, False)])
    p.add_run(', where ')
    omath(p, [_msubsup('Q', 'T', 'X')])
    p.add_run(
        ' denote total training export demand, the sum of training demand '
        'across all importing countries. '
        'Without capacity constraints, the cheapest producer captures all training demand '
        'and '
    )
    omath(p, [_msub('HHI', 'T'), _t(' = 1')])
    p.add_run(
        '. With binding capacity constraints on the cheapest producers, '
    )
    omath(p, [_msub('HHI', 'T'), _t(' < 1')])
    p.add_run(
        ', and '
    )
    omath(p, [_msub('HHI', 'T')])
    p.add_run(
        ' is strictly decreasing in the number of capacity-constrained infra-marginal '
        'exporters. Intuitively, when cheap producers reach capacity '
        'limits, residual demand spills over to more expensive suppliers, spreading market shares '
        'more evenly.'
    )

    # Proposition 3: Sovereignty threshold
    p, cur = mkp(doc, body, cur, space_before=6)
    r = p.add_run('Proposition 3 (Sovereignty Switching Threshold). ')
    r.bold = True
    r.italic = True
    p.add_run(
        'A country will bear the additional cost of domestic AI training only if its sovereignty '
        'premium is large enough to justify the price premium over cheaper foreign producers. '
        'Formally, country '
    )
    omath(p, [_v('k')])
    p.add_run(' produces training domestically if and only if ')
    omath(p, [_msub('\u03BB', 'jk'), _t(' \u2265 '),
              _msubsup('\u03BB', 'k', '*'), _t(' = '),
              _msub('c', 'k'), _t('/'), _msub('p', 'T'), _t(' \u2212 1')])
    p.add_run(
        ' for every potential supplier '
    )
    omath(p, [_v('j')])
    p.add_run(' with ')
    omath(p, [_msub('c', 'j'), _t(' < '), _msub('c', 'k')])
    p.add_run(
        '; equivalently, '
    )
    omath(p, [_v('k')])
    p.add_run(' imports whenever there exists a supplier ')
    omath(p, [_v('j')])
    p.add_run(' for which ')
    omath(p, [_msub('\u03BB', 'jk'), _t(' < '),
              _msubsup('\u03BB', 'k', '*')])
    p.add_run(
        '. The threshold is increasing in '
    )
    omath(p, [_msub('c', 'k')])
    p.add_run(' and decreasing in ')
    omath(p, [_msub('p', 'T')])
    p.add_run(
        '. Under capacity constraints, '
    )
    omath(p, [_msub('p', 'T'), _t(' > '), _msub('c', '(1)')])
    p.add_run(
        ', so the threshold is lower than in the unconstrained model. '
        'Capacity constraints reduce the sovereignty premium required for domestic production '
        'because higher world prices make imports more expensive.'
    )
    make_footnote(p,
        'For large-demand countries (e.g., the United States with 43% of global compute '
        'demand), the observed domestic production may partly reflect scale economies '
        'rather than sovereignty preferences. The switching threshold \u03bb\u2096* conflates '
        'the home market effect with the sovereignty premium; disentangling the two would '
        'require a model with increasing returns, which lies outside the present framework.',
        22)

    # Corollary
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Corollary. ')
    p.add_run(
        'Capacity constraints reduce the welfare cost of sovereignty because '
        'the higher equilibrium price narrows the gap between domestic and import costs.'
    )

    # Proposition 4: Nesting
    p, cur = mkp(doc, body, cur, space_before=6)
    r = p.add_run('Proposition 4 (Training Exporters Nest Within Inference Exporters). ')
    r.bold = True
    r.italic = True
    p.add_run(
        'If a country is cheap enough to export training (which can be done from anywhere), '
        'it is also cheap enough to export inference to nearby demand centers. '
        'The set of training exporters is therefore a subset of the inference exporters for demand centers '
        'within the latency threshold '
    )
    omath(p, [_mbar('l')])
    p.add_run(
        '. A training exporter has the globally lowest '
    )
    omath(p, [_msub('c', 'j')])
    p.add_run(
        '. For inference to proximate demand centers, this cost advantage dominates '
        'the latency markup, so the same country wins the inference competition. '
        'Since training has no distance penalty while inference does, every country '
        'that exports training is also competitive in inference within its geographic '
        'neighborhood, but not vice versa.'
    )


def write_data_section(doc, body, hmap, demand_data):
    print("Inserting Section 5: Data...")

    sec6_heading = hmap['4']  # v8 heading "4", becomes Section 6
    cur = mkh(doc, body, sec6_heading.getprevious(), '5. Data', level=1)

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The propositions above generate testable predictions that depend on '
        'country-specific costs and bilateral frictions. '
        'Calibrating the production-cost and trade-cost parameters in equations (1)\u2013(4) requires data on '
        'electricity prices, temperatures, construction costs, bilateral latencies, and bilateral '
        'sovereignty frictions (geopolitical alignment, regulatory compatibility, sanctions). '
    )

    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Electricity prices. ')
    p.add_run(
        'For European countries, we use prices from Eurostat '
        '(industrial band, 20,000\u201369,999) (Eurostat 2025). '
        'For non-European countries, the prices are obtained from national '
        'regulator tariff sheets and secondary sources, including the U.S. Energy Information Administration '
        '(EIA 2025) for the United States, KEPCO for South Korea,'
    )
    make_footnote(p,
        'Korea Electric Power Corporation (KEPCO), Electricity Rate Table, accessed January 2025.', 24)
    p.add_run(
        ' national utility tariffs for '
        'Central Asian countries, and '
        'GlobalPetrolPrices (2025) for the remaining countries. All prices are converted to $/kWh '
        'at 2024 exchange rates.'
    )

    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Temperature and construction. ')
    p.add_run(
        'Peak summer temperature is derived from ERA5 data '
        '(Hersbach et al. 2020) as the average monthly maximum in the three warmest months, '
        'aggregated across populated grid cells. '
    )
    p.add_run(
        'Construction costs per watt of IT capacity are from the Turner & Townsend '
        'Data Centre Construction Cost Index 2025 (Turner & Townsend 2025), for 37 '
        'countries. For the remaining countries, costs are predicted '
        '(Appendix\u2009E). Since construction is only 3\u20136% of total per-GPU-hour '
        'costs, imputation error has a limited impact on cost rankings.'
    )

    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Latency. ')
    p.add_run(
        'Inter-country round-trip latency is measured using WonderNetwork\u2019s global ping dataset '
        '(WonderNetwork 2024). For each country pair, the median round-trip time (RTT) '
        'in milliseconds is used. '
        'Domestic latency defaults to 5 ms where no intra-country measurement is available.'
    )

    # Demand data paragraph
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Demand. ')
    p.add_run(
        'Compute demand '
    )
    omath(p, [_msub('q', 'k')])
    p.add_run(
        ' is proxied by installed data center capacity in MW, '
        'as specified in equation (4). '
        'For the top 15 markets, capacity estimates are based on industry reports '
        '(Synergy Research, Cushman & Wakefield, CBRE, Mordor Intelligence).'
    )
    make_footnote(p,
        'Synergy Research Group, Global Data Center Market Share (Q4 2024); '
        'Cushman & Wakefield, Global Data Center Market Comparison (2024); '
        'CBRE, Global Data Center Trends (2025); '
        'Mordor Intelligence, Data Center Market Forecast (2025).', 25)
    p.add_run(
        ' '
        'For smaller markets, capacity is estimated from facility counts (Cloudscene 2025) and regional averages. '
        'Since the results below depend only on demand shares, not on the absolute level '
    )
    omath(p, [_v('Q')])
    p.add_run(
        ', the calibration does not require an estimate of total global compute spending.'
    )

    # Calibration approach note
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Because the compute export market is still emerging and bilateral trade-flow data '
        'do not yet exist, the paper calibrates the model using engineering cost parameters '
        'rather than estimating it from observed trade. The calibration identifies the cost '
        'structure under which FLOP exporting becomes viable and provides a framework that '
        'can be taken to gravity-style estimation as transaction-level data emerge.'
    )


def write_calibration(doc, body, hmap, cal, reg, n_eca, n_total, demand_data):
    print("Replacing Section 6 (Calibration)...")

    sec7 = hmap['4']
    sec8 = hmap['5']
    all_now = list(body)
    s7i = all_now.index(sec7)
    s8i = all_now.index(sec8)
    for el in all_now[s7i + 1:s8i]:
        body.remove(el)
    cur = sec7

    # 6.1 Parameter calibration
    cur = mkh(doc, body, cur, '6.1 Parameter calibration', level=2)

    # PUE calibration
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'PUE. ')
    p.add_run('The baseline ')
    omath(p, [_v('\u03C6'), _t(' = 1.08')])
    p.add_run(
        ' matches Google\u2019s reported PUE for facilities with free-air cooling '
        'in cold climates (Uptime Institute 2024). The sensitivity coefficient '
    )
    omath(p, [_v('\u03B4'), _t(' = 0.015')])
    p.add_run(
        ' per \u00b0C is estimated from cross-sectional variation in PUE across data center '
        'locations with different cooling loads (Liu et al. 2023). The threshold '
    )
    omath(p, [_mbar('\u03B8'), _t(' = 15\u00b0C')])
    p.add_run(
        ' is the approximate outdoor temperature above which mechanical cooling is needed, '
        'below which free-air cooling suffices. '
        'Together, these yield PUE values from 1.08 (Iceland, Scandinavia) to 1.41 (UAE), '
        'consistent with the industry average of 1.56 (Uptime Institute 2024).'
    )
    make_footnote(p,
                  'Capping PUE at 1.20 (simulating universal liquid cooling) '
                  'yields a Kendall rank correlation of 0.96 with the baseline rankings. The top five '
                  'countries are unchanged, and the maximum rank shift is six positions. Gulf states and '
                  'North Africa gain the most (UAE moves from 26th to 20th, Qatar from 15th to 11th), '
                  'but the effect is small because electricity prices, not cooling, dominate '
                  'cross-country cost variation.', 13)

    # Hardware parameters
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Hardware. ')
    p.add_run(
        'The calibration uses the NVIDIA H100 SXM GPU as the reference hardware platform, with a list price '
        'of $25,000, a power of 700W, an economic lifetime of 3 years, and a utilization rate '
        '70% (Barroso et al. 2018). '
        'This yields an amortized hardware cost '
    )
    omath(p, [_v('\u03C1'), _t(f' = ${RHO:.3f}')])
    p.add_run(
        '/hr. '
        'Networking costs are calibrated at '
    )
    omath(p, [_v('\u03B7'), _t(f' = ${ETA:.2f}')])
    p.add_run(
        '/hr, based on the same three-year horizon (Barroso et al. 2018). '
        'GPU and networking equipment prices are assumed to be uniform across countries.'
    )

    # Latency degradation (split from old "Other parameters")
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Latency degradation. ')
    p.add_run('The parameter ')
    omath(p, [_v('\u03C4')])
    p.add_run(' is set at ')
    omath(p, [_v('\u03C4'), _t(f' = {TAU}'), _t(' per ms')])
    p.add_run(
        ', implying that 100 ms of round-trip latency '
        '(roughly the intercontinental round-trip between '
        'Europe and East Asia) inflates inference cost by 8%, '
        'consistent with the finding in Deloitte (2020) for e-commerce. '
    )

    # Sovereignty premium (split from old "Other parameters")
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Sovereignty premium. ')
    p.add_run(
        'The bilateral sovereignty premium '
    )
    omath(p, [_msub('\u03BB', 'ij')])
    p.add_run(
        ' is constructed from geopolitical distance, regulatory compatibility, and sanctions '
        'exposure, (eq. (2)), with coefficient weights '
    )
    omath(p, [_msub('\u03B1', '1'), _t(' = 0.08')])
    p.add_run(' (geopolitical distance), ')
    omath(p, [_msub('\u03B1', '2'), _t(' = 0.04')])
    p.add_run(' (regulatory incompatibility), and ')
    omath(p, [_msub('\u03B1', '3'), _t(' = 0.10')])
    p.add_run(' (sanctions). For sanctioned pairs, ')
    omath(p, [_msub('\u03BB', 'ij'), _t(' = \u221E')])
    p.add_run(' (trade is prohibited). For allies with mutual data-adequacy agreements '
              '(e.g., EU member states), ')
    omath(p, [_msub('\u03BB', 'ij'), _t(' \u2248 0')])
    p.add_run('. For non-adversarial pairs without regulatory agreements, ')
    omath(p, [_msub('\u03BB', 'ij')])
    p.add_run(' falls in the range 0.05\u20130.10.')
    make_footnote(p, 'The bilateral sovereignty coefficients are calibrated to match observed '
                  'patterns of data localization policy. Survey evidence suggests enterprises pay '
                  '15\u201330% more for guaranteed domestic data residency (UNCTAD 2025). '
                  'The uniform 10% premium serves as a robustness benchmark.', 15)
    p.add_run(' As a robustness check, we also report results under a uniform premium ')
    omath(p, [_v('\u03BB'), _t(f' = {LAMBDA:.0%}')])
    p.add_run('. The training share of compute demand is ')
    omath(p, [_v('\u03B1'), _t(' = 0.50')])
    p.add_run(', within the industry range of 0.4\u20130.6 (Deloitte 2025). '
              'Construction costs are amortized over 15 years.')

    # Production-efficiency index (ω = 0.50, ξ_floor = 0.30, Form B)
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Production efficiency. ')
    p.add_run('The production-efficiency index ')
    omath(p, [_msubsup('\u03BE', 'j', 'eff')])
    p.add_run(
        ' (eq. (3)) combines a weighted geometric mean of governance quality and grid '
        'reliability (weight '
    )
    omath(p, [_v('\u03C9'), _t(' = 0.50')])
    p.add_run(
        ' on governance) with an institutional floor '
    )
    omath(p, [_msub('\u03BE', 'floor'), _t(' = 0.30')])
    p.add_run(
        '. The equal weighting reflects the enclave character of '
        'data center operations: hyperscale facilities maintain independent backup power, '
        'making general grid reliability a weak predictor of data center uptime. '
        'Governance captures contract '
        'enforcement, expropriation risk, and regulatory stability, none of which '
        'can be mitigated through private infrastructure investment. '
        'The institutional floor reflects the minimum operational quality achievable '
        'through special economic zone provisions and international arbitration '
        '(Farole 2011; World Bank 2017; Frick, Rodr\u00EDguez-Pose, and Wong 2019). '
        'The floor is a political parameter, not a technological one: it depends on the '
        'durability of institutional commitments \u2014 power purchase agreements, SEZ provisions, '
        'arbitration clauses \u2014 that can be revoked if the political economy shifts. '
        'The efficiency adjustment operates only on the '
        'country-specific cost component (equation 3), since hardware is globally priced. '
        'For OECD countries, '
    )
    omath(p, [_msubsup('\u03BE', 'j', 'eff'), _t(' \u2248 1')])
    p.add_run(
        '. Developing countries face moderate penalties '
        f'(e.g., Kyrgyzstan:\u2009{rhup(demand_data["xi"].get("KGZ", 0.50)):.2f}, '
        f'Uzbekistan:\u2009{rhup(demand_data["xi"].get("UZB", 0.58)):.2f}, '
        f'Ethiopia:\u2009{rhup(demand_data["xi"].get("ETH", 0.57)):.2f}). '
        'Sensitivity to both '
    )
    omath(p, [_v('\u03C9')])
    p.add_run(' and ')
    omath(p, [_msub('\u03BE', 'floor')])
    p.add_run(' is examined in Section 7.')

    # Table 2 reference (end of parameter calibration subsection)
    p, cur = mkp(doc, body, cur)
    p._element.append(make_bookmark(103, 'Table2txt'))
    p._element.append(make_hyperlink('Table2', 'Table 2'))
    p._element.append(make_bookmark_end(103))
    p.add_run(
        ' reports all model parameters. Country-specific values, including electricity prices, '
        'temperatures, construction costs, and resulting unit costs, are reported in '
    )
    p._element.append(make_bookmark(100, 'TableA1txt'))
    p._element.append(make_hyperlink('TableA1', 'Table A1'))
    p._element.append(make_bookmark_end(100))
    p.add_run('.')

    # 6.2 Cost Rankings and Trade Patterns
    cur = mkh(doc, body, cur, '6.2 Cost Rankings and Trade Patterns', level=2)

    # ══════════════════════════════════════════════════════════════════════
    # 6.2  Cost Rankings and Trade Patterns
    # ══════════════════════════════════════════════════════════════════════
    _xi_adj = demand_data.get("xi_adjusted", {})
    _xi_top5 = _xi_adj["top5"] if _xi_adj else []
    _xi_n_changed = _xi_adj.get("n_changed_top10", 0)
    _t3 = demand_data["table3"]
    adj_top5 = demand_data["adj_top5"]
    cheapest = cal[0]
    max_gap_country = demand_data["max_gap_country"]
    max_fiscal_m = demand_data["max_fiscal_transfer"] / 1e6

    # ── A0. Figure 1 calibration strategy intro (NEW) ──
    p, cur = mkp(doc, body, cur)
    p._element.append(make_bookmark(150, 'Figure1calstxt'))
    p._element.append(make_hyperlink('Figure1cals', 'Figure 1'))
    p._element.append(make_bookmark_end(150))
    p.add_run(
        ' summarizes the calibration strategy. '
        'The analysis proceeds first by adjusting production costs, '
        'then applying trade frictions.'
    )

    # ── A1. Raw tariff contrast — Table 3 col (1) ──
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Under observed electricity tariffs and without efficiency '
        'adjustment, the cheapest producer in our sample of 85 countries is '
        f'{cheapest["country"]}, '
        f'followed by {cal[1]["country"]} '
        f'and {cal[2]["country"]} '
        '(Column\u2009(1) of '
    )
    p._element.append(make_bookmark(112, 'Table3txt'))
    p._element.append(make_hyperlink('Table3', 'Table 3a'))
    p._element.append(make_bookmark_end(112))
    p.add_run(
        '). But this ranking is misleading. '
        f'{cheapest["country"]}\u2019s headline electricity cost rests on one of the '
        'world\u2019s largest fossil fuel subsidies. Turkmenistan, Algeria, Qatar, and '
        'several other low-cost producers face similar distortions. '
    )

    # ── A2. Cost-recovery adjustment — merged ──
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'We replace subsidized tariffs with cost-recovery prices to distinguish '
        'real comparative advantage from fiscal artifact. '
        'The cost-recovery prices are defined as the long-run marginal cost (LRMC) of the dominant '
        'generation technology at opportunity-cost fuel prices (IMF 2025, Lazard 2025). '
        f'We apply this adjustment to {demand_data["n_adjusted"]} countries '
        'whose retail electricity prices fall below the estimated LRMC.'
    )
    make_footnote(p,
                  'For gas exporters (Iran, Turkmenistan, Algeria, Qatar), the calibration uses '
                  'combined-cycle gas generation at export-parity fuel prices '
                  '($0.065\u2013$0.100/kWh). For the Gulf states, it uses the opportunity cost '
                  'of domestic gas combustion relative to LNG exports. For coal-dependent producers '
                  '(Kazakhstan, South Africa), the calibration uses the Eskom-style cost-recovery '
                  'tariff. For Ethiopia, it uses the IMF\u2019s hydro cost-recovery target ($0.050/kWh). '
                  'The IMF estimates global fossil fuel subsidies at $6.7 trillion in 2024. '
                  'Explicit subsidies (below-cost pricing) account for 8%; the remainder reflects '
                  'unpriced environmental costs. The calibration uses only the explicit component.',
                  16)
    p.add_run(
        ' The subsidy gap ranges from '
        f'${demand_data["min_gap_mwh"] / 1000:.3f} to '
        f'${demand_data["max_gap_mwh_val"] / 1000:.3f}/kWh. '
        f'For {max_gap_country}, a 100\u2009MW IT-load data center would receive roughly '
        f'${max_fiscal_m:.0f}\u2009million per year in implicit fiscal transfer. '
        'At hyperscale scale, such subsidies are fiscally unsustainable. '
        'The resulting cost-recovery ranking (column\u2009(2) of '
    )
    p._element.append(make_hyperlink('Table3', 'Table 3a'))
    p.add_run(
        ') shifts the top of the ranking toward hydropower-rich countries: '
        f'{adj_top5[0][1]}, '
        f'{adj_top5[1][1]}, '
        f'{adj_top5[2][1]}, '
        f'{adj_top5[3][1]}, '
        f'and {adj_top5[4][1]}. '
        f'{cheapest["country"]} drops from first to '
        f'{_ordinal(demand_data["adj_rank_map"]["IRN"])}. '
        f'{_num_word(demand_data["regime_changes"]).capitalize()} '
        f'{"country changes" if demand_data["regime_changes"] == 1 else "countries change"} '
        'their trade regimes.'
    )

    # ── A3. Efficiency adjustment — moved, references Figure 2 ──
    # v27: Form B narrative — developing countries in top 15, rank changes
    _tkm_delta = next((d["delta"] for d in _t3 if d["iso"] == "TKM"), 0)
    _tjk_delta = next((d["delta"] for d in _t3 if d["iso"] == "TJK"), 0)
    _dnk_delta = next((d["delta"] for d in _t3 if d["iso"] == "DNK"), 0)
    _isl_delta = next((d["delta"] for d in _t3 if d["iso"] == "ISL"), 0)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Cost advantages must be weighed against institutional quality. '
        'Private infrastructure mitigates grid unreliability, but governance weaknesses '
        'lie outside the investor\u2019s control. '
        'Efficiency adjustment applies only '
        'to the country-specific cost component (energy procurement, cooling '
        'infrastructure, construction, and grid connection), assigning equal weight '
        'to governance quality and grid reliability ('
    )
    omath(p, [_v('\u03C9'), _t(' = 0.50')])
    p.add_run('). Column\u2009(3) of ')
    p._element.append(make_hyperlink('Table3', 'Table 3a'))
    p.add_run(
        ' reports the efficiency-adjusted ranking ('
    )
    p._element.append(make_bookmark(141, 'TableA2txt'))
    p._element.append(make_hyperlink('TableA2', 'Table A2'))
    p._element.append(make_bookmark_end(141))
    p.add_run(
        ' extends it to all countries). '
        'Under cost-recovery pricing with efficiency adjustment, the five cheapest '
        'producers are '
    )
    if _xi_top5:
        p.add_run(
            f'{_xi_top5[0][0]}, '
            f'{_xi_top5[1][0]}, '
            f'{_xi_top5[2][0]}, '
            f'{_xi_top5[3][0]}, '
            f'and {_xi_top5[4][0]}. '
        )
    p.add_run(
        f'Six developing countries appear in the top fifteen: '
        'China (4th), Kyrgyzstan (5th), Kosovo (7th), Montenegro (8th), Ethiopia (10th), '
        'and Vietnam (14th). '
        'The \u0394 column shows the rank change from raw tariffs to the preferred '
        'specification (column 1 to column 4): '
        f'Turkmenistan drops {abs(_tkm_delta)} places, '
        f'Tajikistan drops {abs(_tjk_delta)}, '
        f'while Denmark rises {abs(_dnk_delta)}, '
        f'and Iceland rises {abs(_isl_delta)}. '
        'Countries that combine cheap energy with adequate institutions retain their '
        'cost advantage after governance adjustment, while those with very weak '
        'governance (Turkmenistan, Tajikistan, Russia) lose ground. '
        'With hardware dominating unit costs, the remaining cross-country variation is '
        'narrow, and even modest governance penalties reshuffle adjacent ranks. '
        'Cheap energy is necessary but not sufficient for FLOP exporting. '
    )
    p._element.append(make_bookmark(121, 'Figure2txt'))
    p._element.append(make_hyperlink('Figure2', 'Figure 2'))
    p._element.append(make_bookmark_end(121))
    p.add_run(' illustrates the resulting rank reshuffling. ')

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The gap between cost-based comparative advantage and actual investment is wide. '
        'Among the twenty cheapest producers after efficiency adjustment, five have less '
        'than 100\u2009MW of installed data center capacity: Kyrgyzstan (rank 5, 5\u2009MW), '
        'Kosovo (rank 7, 5\u2009MW), Montenegro (rank 8, 5\u2009MW), Ethiopia '
        '(rank 10, 10\u2009MW), and Iceland (rank 11, 60\u2009MW). Meanwhile, several high-cost '
        'countries host large data center clusters: the Netherlands (rank 31, 1,800\u2009MW), '
        'Ireland (rank 71, 1,260\u2009MW), and Singapore (rank 66, 1,130\u2009MW). These patterns '
        'reflect agglomeration economies, hyperscaler location strategies, and network '
        'infrastructure that the cost-based ranking captures only through the capacity '
        'ceiling '
    )
    omath(p, [_msub('K\u0304', 'j')])
    p.add_run(
        '. The cost ranking should therefore be read as identifying the feasible set '
        'of exporters, not predicting which countries will attract investment absent the '
        'institutional and market-structure conditions discussed in Section 7.'
    )

    # ── A4. Bilateral sovereignty — Table 3b col (4) ──
    # Use table3 lam_k_star for inline values (consistency with table)
    _t3 = demand_data["table3"]
    _lks = {d["iso"]: d.get("lam_k_star", 0) for d in _t3}
    p, cur = mkp(doc, body, cur)
    p.add_run('The bilateral sovereignty premium ')
    omath(p, [_msub('\u03BB', 'ij')])
    p.add_run(
        ' from equation (2) reshapes trade patterns along geopolitical lines. '
        'Between allies with mutual data adequacy '
        'agreements (e.g., EU member states), the effective premium is near zero. '
        'Between geopolitical adversaries, it is effectively infinite\u2014the United States '
        'would not source training from Iran regardless of cost, and current sanctions '
        'make such transactions illegal. '
        'Column\u2009(4) of '
    )
    p._element.append(make_bookmark(115, 'Table3btxt'))
    p._element.append(make_hyperlink('Table3b', 'Table 3b'))
    p._element.append(make_bookmark_end(115))
    p.add_run(
        ' reports the bilateral specification. '
        'The last column reports the switching threshold '
    )
    omath(p, [_msubsup('\u03BB', 'k', '*')])
    p.add_run(
        ' from Proposition 3: negative values indicate countries cheap enough to export '
        '(e.g., Canada at '
    )
    omath(p, [_msubsup('\u03BB', 'k', '*'),
              _t(f' = \u2212{abs(_lks["CAN"]) * 100:.1f}%')])
    p.add_run(
        '); values closer to zero flag thinner cost advantages '
        'more easily erased by sovereignty frictions '
        '(e.g., Japan '
    )
    omath(p, [_msubsup('\u03BB', 'k', '*'),
              _t(f' = \u2212{abs(_lks["JPN"]) * 100:.1f}%')])
    p.add_run(', China ')
    omath(p, [_msubsup('\u03BB', 'k', '*'),
              _t(f' = \u2212{abs(_lks["CHN"]) * 100:.1f}%')])
    p.add_run(', Kyrgyzstan ')
    omath(p, [_msubsup('\u03BB', 'k', '*'),
              _t(f' = \u2212{abs(_lks["KGZ"]) * 100:.1f}%')])
    p.add_run(
        '). '
        'The bilateral sovereignty premium is particularly powerful for inference, '
        'since the latency markup within Europe is moderate (10\u201340\u2009ms, '
        'adding 1\u20133%), and even a small domestic preference can tip the decision '
        'away from importing.'
    )

    # ── A4. Trade flows under capacity constraints (KEEP P78) ──
    n_exp = demand_data.get("n_train_exporters", 1)
    cap_hhi = demand_data.get("cap_hhi_t", 1.0)
    p_T_val = demand_data.get("p_T", 1.10)
    mu_vals = demand_data.get("mu_j", {})
    n_exp_sov = demand_data.get("n_train_exporters_sov", 1)
    cap_hhi_sov = demand_data.get("cap_hhi_t_sov", 1.0)
    p_T_sov = demand_data.get("p_T_sov", p_T_val)
    ir = demand_data["inf_revenue"]

    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Trade flows under capacity constraints. ')
    p.add_run(
        'Weighting the sourcing patterns by demand shares from equation (4) and applying '
        'capacity constraints from Section 3.4, the equilibrium training price is '
    )
    omath(p, [_msub('p', 'T'), _t(f' = ${p_T_val:.2f}')])
    p.add_run(
        '/hr, set by the marginal exporter\u2019s cost. '
        f'Training demand is served by {n_exp} exporter{"s" if n_exp > 1 else ""} '
        f'(HHI = {cap_hhi:.2f}), confirming Proposition 2. '
    )
    if mu_vals:
        top_mu = sorted(mu_vals.items(), key=lambda x: -x[1])[:3]
        mu_labels = []
        for iso, mu in top_mu:
            co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
            mu_labels.append(f'{co} (${mu:.3f}/hr)')
        p.add_run(
            'The largest shadow values of grid capacity are in '
            f'{", ".join(mu_labels)}, '
            'indicating modest returns to capacity expansion. '
        )
    # Top inference exporters
    top_inf = sorted(ir.items(), key=lambda x: -x[1])
    top5_inf = top_inf[:5]
    inf_labels = []
    _the_countries = {'United Kingdom', 'United States', 'United Arab Emirates',
                      'Czech Republic', 'Netherlands', 'Philippines'}
    for iso, share in top5_inf:
        co = next(r["country"] for r in cal if r["iso3"] == iso)
        prefix = 'the ' if co in _the_countries else ''
        inf_labels.append(f'{prefix}{co} ({share * 100:.0f}%)')
    if len(inf_labels) > 1:
        inf_list = ', '.join(inf_labels[:-1]) + ', and ' + inf_labels[-1]
    else:
        inf_list = inf_labels[0] if inf_labels else ''
    p.add_run(
        'Under the bilateral sovereignty premium, most training demand '
        'shifts to domestic production, and the residual export market is '
        f'served by a single exporter at ${p_T_sov:.2f}/hr.'
    )

    # ── Inference dispersion ──
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Inference is more dispersed, with the top five suppliers being '
        f'{inf_list}, collectively accounting for '
        f'{sum(round(s * 100) for _, s in top5_inf):.0f}% of global inference demand '
        f'(HHI = {demand_data["hhi_i"]:.2f}). '
    )

    # ── A4. Developing countries (KEEP P79) ── (only create paragraph if content)
    _dev_para_created = False
    kgz_clients = demand_data["kgz_inf_clients"]
    kgz_total = sum(w for _, _, w in kgz_clients)
    if kgz_total > 0:
        if not _dev_para_created:
            p, cur = mkp(doc, body, cur)
            _dev_para_created = True
        kgz_client_names = [
            co for _, co, _ in sorted(kgz_clients, key=lambda x: -x[2])
            if co != "Kyrgyzstan"]
        names = kgz_client_names[:3]
        kgz_list = " and ".join(names) if len(names) <= 2 else f'{", ".join(names[:-1])}, and {names[-1]}'
        p.add_run(
            f'Among developing countries, Kyrgyzstan captures {kgz_total:.0f}% of global '
            f'inference demand by serving {kgz_list}, a large share '
            'for a country with a GDP of under $15 billion. '
        )
    _dev = {'DZA', 'KGZ', 'ETH', 'EGY', 'KOS', 'XKX', 'TKM', 'UZB', 'TJK',
            'ALB', 'MKD', 'GEO', 'ARM', 'MDA', 'UKR', 'BIH', 'SRB'}
    for _iso, _share in sorted(ir.items(), key=lambda x: -x[1]):
        if _iso in _dev and _iso != 'KGZ' and _share > 0.01:
            _co = next((r["country"] for r in cal if r["iso3"] == _iso), _iso)
            _n_served = sum(
                1 for i in demand_data.get("adj_reg", {})
                if demand_data["adj_reg"][i]["best_inf_source"] == _iso
                and i != _iso)
            if _n_served > 0:
                if not _dev_para_created:
                    p, cur = mkp(doc, body, cur)
                    _dev_para_created = True
                p.add_run(
                    f'{_co} serves as an inference hub for {_n_served} '
                    f'{"country" if _n_served == 1 else "countries"}, '
                    f'accounting for {_share * 100:.0f}% of global inference demand. '
                )
            break

    # ── A6. Major demand centers (KEEP P82) — moved before sovereignty ──
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Major demand centers. ')
    ar = demand_data.get("adj_reg", {})
    _iso_name = {r["iso3"]: r["country"] for r in cal}
    usa_inf = ar.get('USA', {}).get('best_inf_source', 'CAN')
    deu_inf = ar.get('DEU', {}).get('best_inf_source', 'KOS')
    gbr_inf = ar.get('GBR', {}).get('best_inf_source', 'GBR')
    fra_inf = ar.get('FRA', {}).get('best_inf_source', 'FRA')
    chn_inf = ar.get('CHN', {}).get('best_inf_source', 'KGZ')
    p.add_run(
        'The model\u2019s predictions vary across major AI demand centers because '
        'each faces a different latency geography. '
        'For the United States, the cost-recovery optimum sources training from the cheapest '
        'available producer, and inference from '
        f'{_iso_name.get(usa_inf, usa_inf)}. '
        'For Germany, inference is sourced from '
        f'{_iso_name.get(deu_inf, deu_inf)}, '
        f'for the United Kingdom {"domestically" if gbr_inf == "GBR" else "from " + _iso_name.get(gbr_inf, gbr_inf)}, '
        f'and for France {"domestically" if fra_inf == "FRA" else "from " + _iso_name.get(fra_inf, fra_inf)}. '
        f'For China, the cheapest source of inference is {_iso_name.get(chn_inf, chn_inf)}, '
        'a neighboring country with hydropower-based electricity. '
        'Inference supply thus concentrates around latency-bounded regional hubs, '
        'each major market sourcing from a distinct nearby producer.'
    )

    # ── A4. Sovereignty counterfactual (KEEP P80) ──
    p, cur = mkp(doc, body, cur)
    es10 = demand_data["export_share_10"]
    es20 = demand_data["export_share_20"]
    extra = demand_data["extra_dom"]
    if es10 < 0.005:
        p.add_run(
            'Under cost-recovery pricing, the narrow cost spread means that '
            'even a 10% sovereignty premium is sufficient to make domestic '
            'training viable for nearly all countries, leaving the share of '
            'global training demand available to foreign exporters negligible. '
            f'Raising the premium to 20% shifts {extra} additional '
            f'{"country" if extra == 1 else "countries"} to domestic '
            'production, but the marginal effect is small. '
            'Inference exports are more resilient to sovereignty premia because '
            'the latency advantage of proximity partially insulates regional hubs. '
            'The ease with which modest premia eliminate trade explains the '
            'welfare costs documented below.'
        )
    else:
        p.add_run(
            f'Doubling the sovereignty premium to 20% shifts {extra} '
            f'additional {"country" if extra == 1 else "countries"} '
            'to domestic training production, reducing '
            'the share of global training demand available to foreign producers '
            f'from {es10 * 100:.0f}% to '
            f'{es20 * 100:.0f}%. '
            'Inference exports are more resilient to sovereignty premia because '
            'the latency advantage of proximity partially insulates regional hubs. '
            'The ease with which modest premia eliminate trade explains the '
            'welfare costs documented below.'
        )

    # ── v28: Hyperscaler FDI and the trust channel ──
    n_dev_fdi = demand_data.get("n_dev_fdi_exporters", 7)
    # Build dynamic list of developing-country FDI exporters
    _fdi_regime = demand_data.get("regime_5_fdi", {})
    _DEVELOPING = demand_data.get("DEVELOPING", set())
    _iso_country = demand_data.get("iso_country", {})
    _dev_fdi_names = sorted(
        _iso_country.get(iso, iso)
        for iso, r in _fdi_regime.items()
        if iso in _DEVELOPING and r in ("T+I exporter", "inference hub"))
    _n_total_fdi_exp = sum(1 for r in _fdi_regime.values()
                           if r in ("T+I exporter", "inference hub"))
    _dev_fdi_str = ', '.join(_dev_fdi_names[:-1]) + ', and ' + _dev_fdi_names[-1] if len(_dev_fdi_names) > 1 else (_dev_fdi_names[0] if _dev_fdi_names else '')
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Hyperscaler FDI and the trust channel. ')
    p.add_run(
        'The bilateral specification assumes that buyers evaluate the host country\u2019s '
        'trustworthiness directly. In practice, the dominant delivery channel is hyperscaler '
        'FDI: AWS, Azure, or Google Cloud builds and operates the facility, and the buyer\u2019s '
        'contract is with the hyperscaler, not the host government. Under this arrangement, '
        'the effective sovereignty premium reflects the buyer\u2013operator relationship rather '
        'than the buyer\u2013host-country relationship, as specified in equation (2\u2032). '
        'For the three US-headquartered hyperscalers that control roughly 65 percent of the '
        'global cloud market, '
    )
    omath(p, [_v('G'), _t('('), _v('h'), _t(', '), _v('k'), _t(') \u2248 0')])
    p.add_run(' and ')
    omath(p, [_v('R'), _t('('), _v('h'), _t(', '), _v('k'), _t(') = 1')])
    p.add_run(
        ' for allied buyers, so the premium collapses to the sanctions indicator alone. '
        'Sanctions still bind on the host country: GPUs cannot be shipped to Iran, Russia, '
        'Belarus, or Turkmenistan regardless of who operates the facility. GPU export controls partially '
        'restrict China\u2019s access to training-grade hardware.'
    )

    p, cur = mkp(doc, body, cur)
    p.add_run('Column\u2009(7) of ')
    p._element.append(make_hyperlink('Table3b', 'Table 3b'))
    p.add_run(
        ' reports regime assignments under the FDI specification. '
        'Under bilateral sovereignty (column 4), only Canada exports; all other countries '
        'either produce domestically or import. Under hyperscaler FDI (column 7), '
        f'{_n_total_fdi_exp} countries become exporters, '
        f'{n_dev_fdi} of them developing economies: '
        f'{_dev_fdi_str}. '
        'The developing-country export opportunity identified in the cost rankings '
        '(column 3) is not eliminated by sovereignty \u2014 it is blocked by the absence of '
        'a trust intermediary and restored when one is present.'
    )

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'This result aligns with observed investment patterns. Countries actively '
        'attracting hyperscaler data center investment \u2014 India ($15 billion '
        'AdaniConneX/Google campus), Indonesia (120\u2009MW Jakarta), Malaysia (300\u2009MW '
        'Johor), Kenya ($1 billion Microsoft/G42 geothermal campus), and Armenia '
        '($4 billion Firebird project) \u2014 are developing economies that the bilateral '
        'specification assigns to DD or II but that the FDI specification identifies as '
        'potential exporters. The gap between columns (4) and (7) in '
    )
    p._element.append(make_hyperlink('Table3b', 'Table 3b'))
    p.add_run(' measures the value of hyperscaler intermediation as a trust mechanism.')

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Two caveats apply. First, the FDI specification assumes the hyperscaler absorbs '
        'all bilateral trust. In practice, data still physically resides in the host '
        'country\u2019s jurisdiction, and some buyers \u2014 particularly governments and '
        'regulated industries \u2014 may retain a residual sovereignty premium even when a '
        'trusted operator is present. Second, whether a hyperscaler chooses to build in a '
        'given country depends on agglomeration economies, network infrastructure, and '
        'market size that the cost ranking does not capture. The FDI column identifies the '
        'cost-feasible set of host countries; which ones actually attract investment is a '
        'separate question.'
    )

    # ── A4. Welfare cost of sovereignty (KEEP P81) ──
    p, cur = mkp(doc, body, cur)
    add_italic(p, 'Welfare cost of sovereignty. ')
    p.add_run(
        'The bilateral sovereignty premium imposes a welfare cost with two components, '
        'an import markup (importers pay '
    )
    omath(p, [_msub('\u03BB', 'ij'), _t(' \u00b7 '), _msub('p', 'T')])
    p.add_run(
        ' per unit above the competitive price) and an allocative inefficiency '
        '(countries with '
    )
    omath(p, [_msub('p', 'T'), _t(' < '), _msub('c', 'k'),
              _t(' \u2264 (1 + '), _msub('\u03BB', 'ij'), _t(') \u00b7 '), _msub('p', 'T')])
    p.add_run(
        ' produce domestically at above-world-price costs). '
        'Under capacity constraints, both components are smaller than in the unconstrained '
        'model because the higher world price narrows the gap between domestic and import costs. '
        f'The demand-weighted welfare cost is {demand_data["welfare_pct"]:.1f}% of '
        'average compute spending, comparable to the 0.2\u201310% welfare losses from trade barriers '
        'estimated for goods trade (Eaton and Kortum 2002, Arkolakis et al. 2012).'
        ' At current demand levels (approximately 6\u2009\u00d7\u200910\u00b9\u2070 GPU-hours at '
        '$1.50/hr), this amounts to roughly $1.3\u2009billion per year. This cost is small '
        'enough that governments with legitimate data-sovereignty objectives \u2014 military '
        'applications, health records, national statistical systems \u2014 may rationally '
        'prefer domestic production even at the efficiency loss the model documents.'
    )

    # ── A5. Sovereignty policy discussion (v28: Task 26 compressed) ──
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The welfare implications of the sovereignty premium depend on whether it reflects '
        'genuine security externalities or regulatory overreach. Some domestic processing '
        'preference is justified for genuinely confidential data, but much of the current '
        'policy push extends the sovereignty logic to routine commercial computation. '
        'The welfare cost is nontrivial: the bilateral premium shifts most countries toward '
        'domestic production, forgoing the cost savings from specialization. '
        'The premium is also partly endogenous to the absence of credible international data '
        'governance: where enforceable data-protection agreements exist (as within the EU), '
        'the premium falls; where they do not, even commercially motivated buyers face rational '
        'reasons to prefer domestic processing. '
        'Developing countries '
        'that adopt broad data localization requirements risk foreclosing both import savings and '
        'regional export opportunities \u2014 the very specialization gains the model predicts. '
        'The World Bank (2025) frames this as a core policy trade-off: building domestic '
        'capacity versus securing affordable access to international cloud services.'
    )

    # ══════════════════════════════════════════════════════════════════════
    # 7.  Robustness and Extensions
    # ══════════════════════════════════════════════════════════════════════
    cur = mkh(doc, body, cur, '7. Robustness and Extensions', level=1)

    # 7.1 Robustness and Caveats
    cur = mkh(doc, body, cur, '7.1 Robustness and Caveats', level=2)

    # ── B0. Caveats (consolidated: GPU controls, water, fiscal, hardware) ──
    p, cur = mkp(doc, body, cur, space_before=6)
    p.add_run(
        'The calibration adjusts for energy subsidies and institutional '
        'quality, but several constraints remain outside the model and would further narrow '
        'the set of viable exporters. GPU export controls bar Iran, Russia, and Belarus from '
        'acquiring current-generation hardware. Water scarcity constrains cooling in the Middle '
        'East and North Africa. Fiscal sustainability is a concern: regulated tariffs in many '
        'developing countries cover operating expenses but not full capital cost, so exporting '
        'compute at scale while the domestic energy sector cannot maintain its capital stock may '
        'prove politically unsustainable. Finally, export controls, logistics costs, and local '
        'distribution markups can raise effective GPU prices by 5\u201315% in developing countries, '
        'substantially eroding the thin cost advantages documented in '
    )
    p._element.append(make_hyperlink('TableA1', 'Table A1'))
    p.add_run(
        '. These omitted constraints all work against developing-country competitiveness; '
        'the calibration results should therefore be read as upper bounds.'
    )

    # ── B2. Endogenous electricity prices (v28: Task 27a compressed) ──
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Endogenous electricity prices. ')
    p.add_run(
        'The model treats electricity prices as exogenous. '
        'For small, cheap-energy countries, a hyperscale facility can be large relative to '
        'the host grid \u2014 a 100\u2009MW data center would consume roughly 3% of '
        'Kyrgyzstan\u2019s 3,800\u2009MW national output. '
        'At multi-facility scale, data centers would compete with residential heating, '
        'likely triggering regulatory intervention. The capacity ceiling '
    )
    omath(p, [_mbar_sub('K', 'j')])
    p.add_run(
        ' partially addresses this, '
        'but the fixed-price assumption means cost advantages in '
    )
    p._element.append(make_hyperlink('TableA1', 'Table A1'))
    p.add_run(
        ' are upper bounds. An extension '
        'with upward-sloping supply curves would compress these advantages.'
    )

    # ── B3. Cost of capital (v28: Task 27b compressed) ──
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Cost of capital. ')
    p.add_run(
        'The calibration assumes uniform financing via straight-line depreciation. '
        'In practice, the financing gap is large (Calcaterra et al. 2024). '
        'An OECD hyperscaler at 8% WACC faces hardware costs of $1.58/hr; '
        'a developing-country operator at 18% pays $1.87 \u2014 '
        'a $0.29 gap on hardware alone, roughly four times the electricity cost spread '
        'across the top 20 countries. '
        'If hyperscalers finance at their own WACC regardless of host country, '
        'baseline rankings hold. If locally financed, the same institutional weaknesses '
        'that lower \u03BE also raise the cost of capital.'
    )

    # ── B3b. GPU upgrades (v28: Task 27c compressed) ──
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'GPU upgrades. ')
    p.add_run(
        'Successor GPUs (B200, shipping 2025) deliver roughly four times the training '
        'throughput at about 1\u2009kW. Higher power draw widens the absolute electricity cost '
        'gap across countries, modestly strengthening developing-country comparative advantage. '
        'Countries operating older hardware must discount to compete, potentially eroding '
        'their cost advantage. Qualitative findings are robust to GPU generation choices.'
    )

    # ── B5. Sensitivity analysis (v27: references Table A3 with 7 scenarios) ──
    sens = demand_data.get("sensitivity", [])
    if sens:
        p, cur = mkp(doc, body, cur, space_before=6)
        add_italic(p, 'Sensitivity analysis. ')
        p.add_run(
            'Cost rankings are robust to substantial parameter variation. '
        )
        p._element.append(make_bookmark(143, 'TableA3txt'))
        p._element.append(make_hyperlink('TableA3', 'Table A3'))
        p._element.append(make_bookmark_end(143))
        p.add_run(
            ' reports results across seven robustness specifications, varying the '
            'governance weight, institutional floor, hardware cost share, and functional form.'
        )

    # Comparative statics intuition (ChatGPT review)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The sensitivity results can be read as comparative statics. A rise in global '
        'hardware costs increases the globally-priced cost share, compressing the '
        'locally-penalized component and muting governance penalties for '
        'developing-country exporters: the high-hardware scenario raises the '
        'developing-country count in the top fifteen to 10. Conversely, improved cooling '
        'technology that flattens the PUE\u2013temperature curve narrows the advantage of '
        'cold-climate countries but leaves energy-price differences intact. A reduction in '
        'sovereignty frictions shifts countries from domestic production to importing, '
        'expanding trade volumes but reducing rents for exporters. An increase in local '
        'energy prices \u2014 whether from subsidy removal or demand-driven grid strain \u2014 '
        'erodes the cost advantage that defines FLOP-exporting potential, as the endogenous '
        'electricity price extension above illustrates.'
    )

    # ── B5b. Efficiency parameters (ω + ξ_floor + ρ robustness) ──
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Efficiency parameters. ')
    p.add_run('The baseline assigns equal weight (')
    omath(p, [_v('\u03C9'), _t(' = 0.50')])
    p.add_run(') to governance and grid reliability, with an institutional floor ')
    omath(p, [_msub('\u03BE', 'floor'), _t(' = 0.30')])
    p.add_run('. ')
    p._element.append(make_hyperlink('TableA3', 'Table A3'))
    p.add_run(
        ' reports sensitivity along three dimensions. '
        'First, varying '
    )
    omath(p, [_v('\u03C9')])
    p.add_run(
        ' from 0.30 to 0.85 changes the number of developing countries in the '
        'top fifteen from 7 to 5, but leaves the top five exporters largely unchanged. '
        'Second, raising the floor from 0.00 to 0.50 increases developing-country '
        'representation from 3 to 9 in the top fifteen. '
        'Third, the results are most sensitive to the hardware cost share '
    )
    omath(p, [_v('\u03C1')])
    p.add_run(
        ': varying '
    )
    omath(p, [_v('\u03C1')])
    p.add_run(
        ' by \u00b14 percent changes the developing-country count by up to 7. '
        'The qualitative result, that energy-rich developing countries can compete in '
        'compute exports, follows from hardware being globally priced, not from any '
        'particular parameter choice. '
        'The qualitative finding is robust across all specifications.'
    )

    # v28: Task 28 — Comparative statics intuition paragraph
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The comparative statics have clear economic intuition. Raising the governance '
        'weight \u03C9 penalizes countries with weak rule of law, pushing Kyrgyzstan and '
        'Ethiopia down the ranking while rewarding Scandinavian producers. Raising the '
        'institutional floor '
    )
    omath(p, [_msub('\u03BE', 'floor')])
    p.add_run(
        ' compresses the governance penalty, allowing more developing countries into the '
        'top fifteen. The hardware cost share \u03C1 is the most consequential parameter: '
        'when hardware\u2019s share rises, the (identical) hardware cost dominates the unit cost, '
        'narrowing the cross-country spread and making governance differences less decisive '
        '\u2014 widening the door for developing-country entry.'
    )

    # ── B6. Uniform λ robustness check (reworded) ──
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Uniform sovereignty premium. ')
    p.add_run(
        'Column (6) of '
    )
    p._element.append(make_hyperlink('Table3b', 'Table 3b'))
    p.add_run(
        ' reports results under a uniform premium '
    )
    omath(p, [_v('\u03BB'), _t(' = 0.10')])
    p.add_run(
        '. '
        'The uniform premium produces broadly similar regime assignments: '
        'most countries produce domestically under either specification. '
        'The main difference is that the bilateral premium excludes sanctioned countries '
        'from serving any demand center, while the uniform premium treats all pairs '
        'identically. The bilateral specification is preferred because it captures '
        'observed heterogeneity in sovereignty preferences.'
    )

    # ── B7. Sovereignty tiers (reworded) ──
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Sovereignty tiers. ')
    p.add_run(
        'The model segments each country\u2019s demand into three tiers: '
        'sovereign workloads (10%, domestic only), '
        'regulated workloads (20%, higher regulatory compatibility weight), '
        'and commercial workloads (70%, geopolitical alignment only). '
        'Under calibrated parameters, tiering leaves regime assignments unchanged '
        'for all countries, hence columns (4) and (5) of '
    )
    p._element.append(make_hyperlink('Table3b', 'Table 3b'))
    p.add_run(
        ' are merged. The main impact is on inference sourcing: regulated workloads shift '
        'toward suppliers with strong data governance, favoring EU '
        'and APEC CBPR participants over closer but less regulated alternatives.'
    )
    make_footnote(p, 'The tier shares (10/20/70) are assumptions calibrated to match the '
                  'approximate composition of government, regulated-industry, and commercial '
                  'AI workloads (Deloitte 2025). Results are robust to moderate variation in '
                  'these shares.', 17)

    # ══════════════════════════════════════════════════════════════════════
    # 7.2  Model Extensions  (trimmed per Phase C)
    # ══════════════════════════════════════════════════════════════════════
    cur = mkh(doc, body, cur, '7.2 Model extensions', level=2)

    # v28: Task 29 — compressed extensions intro
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'The model can be extended to incorporate carbon pricing (\u201Cgreen premia\u201D), '
        'demand segmentation by latency tolerance, and stochastic disruptions that create '
        'incentives for supply diversification.'
    )

    # Agglomeration and market structure (KEEP P93)
    p, cur = mkp(doc, body, cur, space_before=6)
    add_italic(p, 'Agglomeration and market structure. ')
    p.add_run(
        'As noted in Section 3, the competitive framework abstracts from the industrial '
        'organization of the cloud market. The concentration of data centers in locations '
        'such as Northern Virginia \u2014 despite above-median electricity costs \u2014 reflects '
        'centripetal forces that the cost-based ranking cannot capture. The cost-competitive '
        'set identified in Section 6 should be interpreted as a necessary condition: a '
        'country must be cheap enough to attract investment, but whether investment '
        'materializes depends on agglomeration economies and network effects.'
    )


def write_conclusion(doc, body, hmap, demand_data):
    print("Rewriting Section 8 (Conclusion)...")
    sec8 = hmap['5']
    refs = hmap['refs']
    all_now = list(body)
    c8i = all_now.index(sec8)
    ri = all_now.index(refs)
    for el in all_now[c8i + 1:ri]:
        body.remove(el)

    p, cur_concl = mkp(doc, body, sec8)
    p.add_run(
        'This paper develops a capacity-constrained Ricardian model of compute trade, '
        'distinguishing latency-insensitive training from latency-sensitive inference and '
        'incorporating bilateral sovereignty premia. '
        'The calibration reveals a paradox: energy-rich developing countries hold a genuine '
        'cost advantage in compute production, but bilateral trust deficits eliminate that '
        'advantage for nearly all of them. '
        'Under the standard bilateral specification, only Canada exports; sovereignty premia '
        'shift every other country toward domestic production or importing, '
        f'at a demand-weighted welfare cost of {demand_data["welfare_pct"]:.1f}% of '
        'compute spending. '
        'The resolution lies in the hyperscaler FDI channel: when a trusted cloud provider '
        'operates the facility, the buyer\u2019s sovereignty concern attaches to the operator, '
        'not the host country. Under this specification, '
        f'{_num_word(demand_data.get("n_dev_fdi_exporters", 7))} developing countries re-enter '
        'as exporters, and the model\u2019s predictions align with observed hyperscaler investment '
        'in India, Kenya, Malaysia, and Southeast Asia. '
        'Because hardware accounts for roughly 90 percent of unit cost and is globally '
        'priced, the cross-country cost spread is only 12\u201320 percent \u2014 narrower than '
        'virtually any other tradable good. This makes compute both the easiest sector for '
        'developing countries to enter on cost grounds and the one most vulnerable to small '
        'frictions: a modest sovereignty premium, a slight governance penalty, or a higher '
        'cost of capital is sufficient to shift a country from exporter to importer. '
    )

    p, cur_concl = mkp(doc, body, cur_concl)
    p.add_run(
        'For developing countries, the results identify both the opportunity and the barrier. '
        'Countries like Kyrgyzstan, Kosovo, Ethiopia, and Vietnam rank among the fifteen '
        'cheapest FLOP producers after governance adjustment, and could use their energy '
        'resource endowments (hydropower, natural gas, and solar irradiance) to turn '
        'cheap power into exportable compute without building a domestic AI research ecosystem. '
        'But the bilateral specification shows that this opportunity is blocked by trust '
        'deficits \u2014 and restored only when a hyperscaler intermediates the transaction. '
        'The binding constraint is therefore not electricity cost but institutional credibility: '
        'a country must be non-sanctioned, offer credible power purchase agreements, maintain '
        'adequate network connectivity, and present a regulatory environment stable enough for '
        'a hyperscaler to commit capital over a 15-year horizon. '
        'These are achievable conditions. The countries currently attracting hyperscaler '
        'investment \u2014 India, Kenya, Malaysia, Indonesia \u2014 meet them. Those that do not, '
        'despite lower electricity costs, lack one or more of these prerequisites. '
        'The gap between columns (3) and (7) of Table 3b measures what institutional reform '
        'is worth. '
        'FLOP exporting resembles aluminum smelting near cheap hydropower \u2014 imported '
        'capital equipment transforms local electricity into an exportable product with '
        'minimal domestic labor \u2014 but electricity, unlike oil or minerals, is renewable '
        'where generated from hydro, solar, or geothermal sources, and compute demand is '
        'growing faster than demand for any physical commodity. '
        'The resource curse literature (van der Ploeg 2011) warns, however, that '
        'concentrated export revenues can lead to Dutch disease, institutional degradation, '
        'and exposure to demand cycles. '
        'Large-scale compute export revenues could also appreciate the real exchange rate, '
        'crowding out other tradable sectors, a channel particularly relevant for small, '
        'open economies where data center electricity consumption rivals existing industrial '
        'load. '
        'Whether FLOP exporting countries share these risks depends on the revenue-sharing '
        'model they adopt: a sovereign wealth fund approach (Norway) versus elite capture '
        '(Dutch disease). '
        'The share of surplus retained in the host country depends on ownership and fiscal '
        'structure: if the facility is foreign-owned, most operating surplus is repatriated, '
        'and the host retains only the electricity payment unless the government captures rent '
        'through taxation, equity participation, or resource royalties.'
    )

    p, _ = mkp(doc, body, cur_concl)
    p.add_run(
        'The policy implications are asymmetric across training and inference. '
        'Training workloads tolerate high latency, so restricting training imports '
        'raises costs without offsetting proximity gains. '
        'Inference, by contrast, is latency-sensitive, giving domestic production a genuine '
        'quality-of-service advantage \u2014 though this rationale weakens for countries close '
        'to low-cost neighbors. '
        'For developing countries, the five-column progression in Table 3b makes the policy '
        'challenge concrete: cheap energy gets a country into the cost-feasible set (column 3), '
        'but bilateral trust eliminates the opportunity (column 4), and only hyperscaler '
        'intermediation restores it (column 7). '
        'The binding constraints are not technological but institutional: non-sanctioned '
        'status, credible power purchase agreements, network connectivity, and a regulatory '
        'environment stable enough to justify a 15-year capital commitment.'
    )


def write_appendix(doc, body, last_ref_el, eca_cal, non_eca_cal, reg, demand_data):
    print("Inserting Appendix (Table A1)...")

    # Appendix heading — on the same landscape page as Table A1
    # (previous landscape sectPr already on Table 3 notes paragraph)
    cur_app = mkh(doc, body, last_ref_el, 'Appendix A', level=1)

    # ═══════════════════════════════════════════════════════════════════════
    # TABLE A1: COUNTRY-SPECIFIC CALIBRATION PARAMETERS (landscape)
    # ═══════════════════════════════════════════════════════════════════════
    print("Inserting Table A1 (Country parameters, landscape)...")

    # Table A1 title with bookmark + back-link (follows directly after A1 notes)
    tp2 = doc.add_paragraph()
    tp2.paragraph_format.space_before = Pt(6)
    tp2.paragraph_format.space_after = Pt(3)
    tp2.paragraph_format.first_line_indent = Inches(0)
    tp2._element.append(make_bookmark(104, 'TableA1'))
    hl_t = OxmlElement('w:hyperlink')
    hl_t.set(qn('w:anchor'), 'TableA1txt')
    hl_t.set(qn('w:history'), '1')
    r_t = OxmlElement('w:r')
    rPr_t = OxmlElement('w:rPr')
    b_t = OxmlElement('w:b')
    rPr_t.append(b_t)
    sz_t = OxmlElement('w:sz')
    sz_t.set(qn('w:val'), '20')
    rPr_t.append(sz_t)
    clr_t = OxmlElement('w:color')
    clr_t.set(qn('w:val'), LINK_COLOR)
    uu_t = OxmlElement('w:u')
    uu_t.set(qn('w:val'), 'single')
    rPr_t.append(clr_t)
    rPr_t.append(uu_t)
    r_t.append(rPr_t)
    t_t = OxmlElement('w:t')
    t_t.text = 'Table A1'
    r_t.append(t_t)
    hl_t.append(r_t)
    tp2._element.append(hl_t)
    tp2._element.append(make_bookmark_end(104))
    run_tt2 = tp2.add_run('. Country-specific calibration parameters')
    run_tt2.bold = True
    run_tt2.font.size = Pt(10)
    tp2_el = tp2._element
    body.remove(tp2_el)
    cur_app.addnext(tp2_el)

    # Gather all country data
    omega = demand_data["omega"]
    dc_k = demand_data.get("dc_k", {})
    xi = demand_data.get("xi", {})
    adj_rank_map = demand_data.get("adj_rank_map", {})
    # Sort by cost-recovery adjusted rank
    all_cal = sorted(eca_cal + non_eca_cal,
                     key=lambda r: adj_rank_map.get(r["iso3"], 999))

    a2_headers = ["Country", "p\u1d31\n($/kWh)", "\u03B8\u2c7c\n(\u00b0C)",
                  "PUE", "Constr.\n($/W)", "k\u0304\u2c7c\n(MW)",
                  "\u03C9\u2c7c\n(%)", "\u03BE\u2c7c\u1d49\u1da0\u1da0",
                  "c\u2c7c\n($/hr)", "Cost-Rec.\np\u1d31 ($/kWh)"]

    # Build row data; track which rows need bold in cost-recovery column
    a2_rows = []
    bold_cr_rows = []  # row indices (0-based) where cost-rec price is substituted
    for idx, r_row in enumerate(all_cal):
        iso = r_row["iso3"]
        co = r_row["country"]
        if len(co) > 20:
            co = co[:19] + "."
        # Cost-recovery price: substituted value for 13 countries, otherwise same as p_E
        p_E_raw = float(r_row["p_E_usd_kwh"])
        cr = SUBSIDY_ADJ.get(iso)
        cr_price = cr if cr is not None else p_E_raw
        cr_str = f'${cr_price:.3f}'
        if cr is not None:
            bold_cr_rows.append(idx)
        cap = dc_k.get(iso, 5.0)
        cap_str = f'{cap:,.0f}' if cap >= 10 else f'{cap:.0f}'
        share = omega.get(iso, 0)
        xi_j = xi.get(iso, 1.0)
        a2_rows.append((
            co,
            f'${p_E_raw:.3f}',
            f'{float(r_row["theta_summer_C"]):.1f}',
            f'{float(r_row["pue"]):.2f}',
            f'${float(r_row["p_L_usd_per_W"]):.2f}',
            cap_str,
            f'{share * 100:.1f}',
            f'{rhup(xi_j):.2f}',
            f'${float(r_row["c_j_total"]):.2f}',
            cr_str,
        ))

    # Column widths for landscape (9 inches usable)
    a2_cw = [1500, 650, 550, 500, 600, 650, 550, 450, 650, 700]
    last_a2_tbl = add_table(doc, body, tp2_el, a2_headers, a2_rows, a2_cw)

    # Post-process: bold the cost-recovery price cells for subsidized countries
    # Table rows: row 0 = header, data rows start at 1
    all_trs = last_a2_tbl.findall(f'{{{W_NS}}}tr')
    for row_idx in bold_cr_rows:
        tr = all_trs[row_idx + 1]  # skip header row
        # Column 9 = cost-recovery price (after removing Rank and Regime)
        tcs = tr.findall(f'{{{W_NS}}}tc')
        if len(tcs) > 9:
            tc = tcs[9]
            for r_el in tc.findall(f'.//{{{W_NS}}}r'):
                rPr = r_el.find(f'{{{W_NS}}}rPr')
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r_el.insert(0, rPr)
                b_el = OxmlElement('w:b')
                rPr.append(b_el)

    # Table A1 notes
    note_a2 = doc.add_paragraph()
    note_a2.paragraph_format.space_before = Pt(4)
    note_a2.paragraph_format.space_after = Pt(0)
    note_a2.paragraph_format.first_line_indent = Inches(0)
    note_a2.paragraph_format.line_spacing = 1.0
    rn = note_a2.add_run('Notes: ')
    rn.bold = True
    rn.font.size = Pt(10)
    rn = note_a2.add_run(
        'Countries sorted by cost-recovery adjusted rank (ascending). '
        'p\u1d31 = national electricity price for industrial/data center consumers ($/kWh). '
        '\u03B8\u2c7c = peak summer temperature (\u00b0C). '
        'PUE = Power Usage Effectiveness. '
        'Constr. = predicted data center construction cost ($/W of IT load). '
        'k\u0304\u2c7c = installed data center power capacity (MW). '
        '\u03C9\u2c7c = country share of global compute demand from equation (4). '
        '\u03BE\u2c7c\u1d49\u1da0\u1da0 = production-efficiency index (weighted geometric mean of '
        'governance quality and grid reliability, \u03C9 = 0.50; see equation 3). '
        'Sanctions captured by bilateral '
        '\u03BB\u1d62\u2c7c (equation 2). '
        'c\u2c7c = hourly cost of operating one H100 GPU (electricity + '
        'hardware at $1.36/hr + amortized construction; excludes networking \u03B7 = $0.15/hr, '
        'which is added in the equilibrium computations in Section 6). '
        'Cost-Rec. p\u1d31 = cost-recovery electricity price. '
        'For 13 countries with subsidized tariffs, this is the estimated long-run marginal cost '
        'of electricity generation (shown in bold). '
        'For all other countries, the cost-recovery price equals the observed tariff. '
        'The 37 DCCI countries span 52 markets: Australia, Austria, Brazil, Canada, Chile, '
        'China, Colombia, Denmark, Finland, France, Germany, Greece, India, Indonesia, Ireland, '
        'Italy, Japan, Kenya, Malaysia, Mexico, Netherlands, New Zealand, Nigeria, Norway, Poland, '
        'Portugal, Saudi Arabia, Singapore, South Africa, South Korea, Spain, Sweden, Switzerland, '
        'UAE, UK, Uruguay, and USA. '
        'The 95% prediction intervals for imputed countries span about \u00b1$3.50/W, '
        'which translates to \u00b1$0.02/hr in total cost (1.5\u20132% of the mean).'
    )
    rn.font.size = Pt(10)
    note_a2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_a2_el = note_a2._element
    body.remove(note_a2_el)
    last_a2_tbl.addnext(note_a2_el)

    # ─── Attach landscape sectPr to notes paragraph (no empty page) ───
    sect_pr = OxmlElement('w:sectPr')
    pg_sz = OxmlElement('w:pgSz')
    pg_sz.set(qn('w:w'), '15840')
    pg_sz.set(qn('w:h'), '12240')
    pg_sz.set(qn('w:orient'), 'landscape')
    sect_pr.append(pg_sz)
    pg_mar = OxmlElement('w:pgMar')
    pg_mar.set(qn('w:top'), '1440')
    pg_mar.set(qn('w:right'), '1440')
    pg_mar.set(qn('w:bottom'), '1440')
    pg_mar.set(qn('w:left'), '1440')
    pg_mar.set(qn('w:header'), '720')
    pg_mar.set(qn('w:footer'), '720')
    sect_pr.append(pg_mar)
    na2_pPr = note_a2_el.find(f'{{{W_NS}}}pPr')
    if na2_pPr is None:
        na2_pPr = etree.SubElement(note_a2_el, f'{{{W_NS}}}pPr')
    na2_pPr.append(sect_pr)

    return note_a2_el


def write_table_a2(doc, body, after_el, demand_data):
    """Table A2: Complete country rankings under alternative pricing assumptions (landscape)."""
    print("Inserting Table A2 (Complete country rankings, landscape)...")

    table3_data = demand_data["table3"]

    _short = {
        "United Arab Emirates": "UAE", "United Kingdom": "UK",
        "United States": "USA", "Bosnia and Herzegovina": "Bosnia & Herz.",
        "North Macedonia": "N. Macedonia", "Czech Republic": "Czechia",
    }

    def _sn(full):
        return _short.get(full, full[:18] + '.' if len(full) > 19 else full)

    # Title (placed directly after previous element — sectPr already on it)
    tp = doc.add_paragraph()
    tp.paragraph_format.space_before = Pt(6)
    tp.paragraph_format.space_after = Pt(3)
    tp.paragraph_format.first_line_indent = Inches(0)
    tp._element.append(make_bookmark(140, 'TableA2'))
    hl_a2t = OxmlElement('w:hyperlink')
    hl_a2t.set(qn('w:anchor'), 'TableA2txt')
    hl_a2t.set(qn('w:history'), '1')
    r_a2t = OxmlElement('w:r')
    rPr_a2t = OxmlElement('w:rPr')
    b_a2t = OxmlElement('w:b')
    rPr_a2t.append(b_a2t)
    sz_a2t = OxmlElement('w:sz')
    sz_a2t.set(qn('w:val'), '20')
    rPr_a2t.append(sz_a2t)
    clr_a2t = OxmlElement('w:color')
    clr_a2t.set(qn('w:val'), LINK_COLOR)
    uu_a2t = OxmlElement('w:u')
    uu_a2t.set(qn('w:val'), 'single')
    rPr_a2t.append(clr_a2t)
    rPr_a2t.append(uu_a2t)
    r_a2t.append(rPr_a2t)
    t_a2t = OxmlElement('w:t')
    t_a2t.text = 'Table A2'
    r_a2t.append(t_a2t)
    hl_a2t.append(r_a2t)
    tp._element.append(hl_a2t)
    tp._element.append(make_bookmark_end(140))
    run_t = tp.add_run('. Country rankings under alternative cost specifications (all countries)')
    run_t.bold = True
    run_t.font.size = Pt(10)
    tp_el = tp._element
    body.remove(tp_el)
    after_el.addnext(tp_el)

    # Sort all countries by spec (3) rank
    all_sorted = sorted(table3_data, key=lambda x: x["rank_eff"])

    n_rows = 2 + len(all_sorted)
    n_cols = 15
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    _tbl_clear_borders(tbl)

    _s = partial(_tbl_set, tbl, font_size=7)  # default fs=7 for A2

    # Row 0: group headers
    _s(0, 0, '')
    _tbl_merge(tbl, 0, 1, 3)
    _s(0, 1, '(1) Raw Electricity', bold=True)
    _tbl_merge(tbl, 0, 4, 6)
    _s(0, 4, '(2) Cost-Recovery', bold=True)
    _tbl_merge(tbl, 0, 7, 10)
    _s(0, 7, '(3) Efficiency-Adjusted', bold=True)
    _tbl_merge(tbl, 0, 11, 13)
    _s(0, 11, '(4) Bilateral \u03bb\u1d62\u2c7c', bold=True)
    _s(0, 14, '')
    for j in range(n_cols):
        _tbl_border(tbl.cell(0, j)._tc, ['top', 'bottom'])

    # Row 1: sub-headers
    sub_h = ['Country', 'c\u2c7c', 'Rank', 'Type',
             'c\u2c7c', 'Rank', 'Type',
             'c\u2c7c\u1d43\u1d48\u02b2', '\u03be\u2c7c\u1d49\u1da0\u1da0', 'Rank', 'Type',
             'p\u2c7c', 'Rank', 'Type', '\u0394']
    for j, h in enumerate(sub_h):
        _s(1, j, h, bold=True, align='left' if j == 0 else 'center')
        _tbl_border(tbl.cell(1, j)._tc, ['top', 'bottom'])

    # Data rows
    for i, d in enumerate(all_sorted):
        ri = i + 2
        _s(ri, 0, _sn(d["country"]), align='left')
        _s(ri, 1, f'${d["cj_raw"]:.2f}')
        _s(ri, 2, str(d["rank_raw"]))
        _s(ri, 3, d["type_raw"])
        _s(ri, 4, f'${d["cj_cr"]:.2f}')
        _s(ri, 5, str(d["rank_cr"]))
        _s(ri, 6, d["type_cr"])
        _s(ri, 7, f'${d["cj_eff"]:.2f}')
        _s(ri, 8, f'{rhup(d["xi"]):.2f}')
        _s(ri, 9, str(d["rank_eff"]))
        _s(ri, 10, d["type_eff"])
        _s(ri, 11, f'${d["cj_eff"]:.2f}')
        _s(ri, 12, str(d["rank_eff"]))
        _s(ri, 13, d.get("type_bilat", d.get("type_sov", "II")))
        dv = d["delta"]
        _s(ri, 14, f'+{dv}' if dv > 0 else str(dv))

    # Bottom border on last row
    for j in range(n_cols):
        _tbl_border(tbl.cell(n_rows - 1, j)._tc, ['bottom'], style='double')

    # Column widths and spacing
    _tbl_col_widths(tbl, [1800, 900, 540, 540, 900, 540, 540, 900, 450, 540, 540, 900, 540, 540, 540])
    _tbl_cell_spacing(tbl, before='5', after='5')

    tbl_el = tbl._tbl
    body.remove(tbl_el)
    tp_el.addnext(tbl_el)

    # Notes
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(2)
    note.paragraph_format.space_after = Pt(6)
    note.paragraph_format.first_line_indent = Inches(0)
    note.paragraph_format.line_spacing = 1.0
    rn = note.add_run('Notes: ')
    rn.bold = True
    rn.font.size = Pt(10)
    rn = note.add_run('See ')
    rn.font.size = Pt(10)
    note._element.append(make_hyperlink('Table3', 'Table 3a', rPr_orig=_rPr_pt(10)))
    rn = note.add_run(
        ' notes for column definitions. '
        'Countries sorted by specification (3) rank (ascending). '
        'Column (4) uses bilateral sovereignty premium \u03bb\u1d62\u2c7c from equation (2).'
    )
    rn.font.size = Pt(10)
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_el = note._element
    body.remove(note_el)
    tbl_el.addnext(note_el)

    # ─── Attach landscape sectPr to notes paragraph (no empty page) ───
    sect_pr = OxmlElement('w:sectPr')
    pg_sz = OxmlElement('w:pgSz')
    pg_sz.set(qn('w:w'), '15840')
    pg_sz.set(qn('w:h'), '12240')
    pg_sz.set(qn('w:orient'), 'landscape')
    sect_pr.append(pg_sz)
    pg_mar = OxmlElement('w:pgMar')
    pg_mar.set(qn('w:top'), '1080')
    pg_mar.set(qn('w:right'), '1080')
    pg_mar.set(qn('w:bottom'), '1080')
    pg_mar.set(qn('w:left'), '1080')
    pg_mar.set(qn('w:header'), '720')
    pg_mar.set(qn('w:footer'), '720')
    sect_pr.append(pg_mar)
    note_pPr = note_el.find(f'{{{W_NS}}}pPr')
    if note_pPr is None:
        note_pPr = etree.SubElement(note_el, f'{{{W_NS}}}pPr')
    note_pPr.append(sect_pr)

    return note_el


def write_model_appendix(doc, body, last_note):
    """Appendix B: Full model derivation from flops_capacity_model.md."""
    print("Inserting Appendix B (Model Derivation)...")

    # No page break needed — previous element has landscape sectPr which forces a new page
    cur = mkh(doc, body, last_note, 'Appendix B: Model Derivation', level=1)

    p, cur = mkp(doc, body, cur)
    p.add_run(
        'This appendix provides the full derivation of the capacity-constrained Ricardian '
        'model summarized in Sections 3\u20134.'
    )

    # B.1 Primitives
    cur = mkh(doc, body, cur, 'B.1 Primitives', level=2)
    p, cur = mkp(doc, body, cur)
    p.add_run('Each country ')
    omath(p, [_v('j')])
    p.add_run(' is endowed with a capacity ceiling ')
    omath(p, [_mbar_sub('K', 'j')])
    p.add_run(
        ' (GPU-hours per period), representing the maximum volume of compute it can supply. '
        'Country '
    )
    omath(p, [_v('j')])
    p.add_run(' faces unit production cost ')
    omath(p, [_msub('c', 'j')])
    p.add_run(
        ' from equation (1). On the demand side, total compute demand from country '
    )
    omath(p, [_v('k')])
    p.add_run(' is ')
    omath(p, [_msub('q', 'k')])
    p.add_run(' from equation (4). Training demand is ')
    omath(p, [_msub('q', 'Tk'), _t(' = '), _v('\u03B1'), _t(' \u00b7 '), _msub('q', 'k')])
    p.add_run(' and inference demand is ')
    omath(p, [_msub('q', 'Ik'), _t(' = (1 \u2212 '), _v('\u03B1'), _t(') \u00b7 '), _msub('q', 'k')])
    p.add_run('. Countries are ordered by cost: ')
    omath(p, [_msub('c', '(1)'), _t(' \u2264 '), _msub('c', '(2)'),
              _t(' \u2264 \u2026 \u2264 '), _msub('c', '(N)')])
    p.add_run('.')

    # B.2 Training Market
    cur = mkh(doc, body, cur, 'B.2 The Training Market', level=2)
    p, cur = mkp(doc, body, cur)
    p.add_run('Country ')
    omath(p, [_v('k')])
    p.add_run(' imports training if and only if ')
    omath(p, [_t('(1 + '), _msub('\u03BB', 'jk'), _t(') \u00b7 '), _msub('p', 'T'),
              _t(' < '), _msub('c', 'k')])
    p.add_run(
        ', where '
    )
    omath(p, [_msub('\u03BB', 'jk')])
    p.add_run(
        ' is the bilateral sovereignty premium from equation (2). '
        'The set of training importers is '
    )
    omath(p, [_msub('M', 'T'), _t(' = { '), _v('k'), _t(' : '),
              _msub('c', 'k'), _t(' > (1 + '), _msub('\u03BB', 'jk'), _t(') \u00b7 '),
              _msub('p', 'T'), _t(' }')])
    p.add_run(' and total training export demand is ')
    omath(p, [_msubsup('Q', 'T', 'X'), _t(' = '),
              _nary('\u2211', [_v('k'), _t(' \u2208 '), _msub('M', 'T')], [],
                    [_msub('q', 'Tk')])])
    p.add_run(
        '. The marginal training exporter '
    )
    omath(p, [_msub('m', 'T')])
    p.add_run(' is defined by:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msub('m', 'T'), _t(' = min { '), _v('m'),
        _t(' : '),
        _nary('\u2211', [_v('i'), _t(' = 1')], [_v('m')],
              [_msub('K', 'T,(i)')]),
        _t(' \u2265 '), _msubsup('Q', 'T', 'X'), _t(' }.'),
    ], eq_num='B.1')

    p, cur = mkp(doc, body, cur)
    p.add_run('The equilibrium training price is ')
    omath(p, [_msub('p', 'T'), _t(' = '), _msub('c', '('),
              _msub('m', 'T'), _t(')')])
    p.add_run('. Training rent for country ')
    omath(p, [_v('j')])
    p.add_run(' with ')
    omath(p, [_msub('c', 'j'), _t(' < '), _msub('p', 'T')])
    p.add_run(' is ')
    omath(p, [_msub('\u03C0', 'Tj'), _t(' = ('), _msub('p', 'T'),
              _t(' \u2212 '), _msub('c', 'j'), _t(') \u00b7 '),
              _msub('K', 'Tj')])
    p.add_run('.')

    # B.3 Inference Market
    cur = mkh(doc, body, cur, 'B.3 The Inference Market', level=2)
    p, cur = mkp(doc, body, cur)
    p.add_run('The feasible supplier set for demand center ')
    omath(p, [_v('k')])
    p.add_run(' is ')
    omath(p, [_v('S'), _t('('), _v('k'), _t(') = { '), _v('j'), _t(' : '),
              _msub('l', 'jk'), _t(' \u2264 '), _mbar('l'), _t(' }')])
    p.add_run('. The marginal cost of delivering one effective unit of inference from ')
    omath(p, [_v('j')])
    p.add_run(' to ')
    omath(p, [_v('k')])
    p.add_run(' is:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msub('MC', 'I'), _t('('), _v('j'), _t(', '), _v('k'),
        _t(') = (1 + '), _v('\u03C4'), _t(' \u00b7 '),
        _msub('l', 'jk'), _t(') \u00b7 '), _msub('c', 'j'), _t('.'),
    ], eq_num='B.2')

    p, cur = mkp(doc, body, cur)
    p.add_run('The inference rent per GPU-hour allocated to serving ')
    omath(p, [_v('k')])
    p.add_run(' is ')
    omath(p, [_msub('r', 'I'), _t('('), _v('j'), _t(', '), _v('k'),
              _t(') = '), _msubsup('p', 'I', 'f'), _t('('), _v('k'),
              _t(') / (1 + '), _v('\u03C4'), _t(' \u00b7 '),
              _msub('l', 'jk'), _t(') \u2212 '), _msub('c', 'j')])
    p.add_run('.')

    # B.4 Capacity Allocation
    cur = mkh(doc, body, cur, 'B.4 Capacity Allocation', level=2)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'Each GPU-hour is allocated to its highest-margin use. The margins per GPU-hour are: '
        'training exports '
    )
    omath(p, [_msub('r', 'T'), _t('('), _v('j'), _t(') = '),
              _msub('p', 'T'), _t(' \u2212 '), _msub('c', 'j')])
    p.add_run('; inference exports to ')
    omath(p, [_v('k')])
    p.add_run(': ')
    omath(p, [_msub('r', 'I'), _t('('), _v('j'), _t(', '), _v('k'),
              _t(') = '), _msubsup('p', 'I', 'f'), _t('('), _v('k'),
              _t(') / (1 + '), _v('\u03C4'), _t(' \u00b7 '),
              _msub('l', 'jk'), _t(') \u2212 '), _msub('c', 'j')])
    p.add_run(
        '. Total rent from operating '
    )
    omath(p, [_msub('K', 'j')])
    p.add_run(' GPU-hours is:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msub('\u03A0', 'j'), _t('('), _msub('K', 'j'),
        _t(') = '),
        _nary('\u2211', [_v('n'), _t(' = 1')], [_msub('K', 'j')],
              [_msubsup('r', 'j', '(n)')]), _t(','),
    ], eq_num='B.3')

    p, cur = mkp(doc, body, cur)
    p.add_run('which is concave and piecewise linear in ')
    omath(p, [_msub('K', 'j')])
    p.add_run('.')

    # B.5 Equilibrium and Existence
    cur = mkh(doc, body, cur, 'B.5 Equilibrium Definition and Existence', level=2)
    p, cur = mkp(doc, body, cur)
    p.add_run(
        'A competitive equilibrium consists of a training price '
    )
    omath(p, [_msub('p', 'T')])
    p.add_run(', inference prices ')
    omath(p, [_t('{'), _msubsup('p', 'I', 'f'), _t('('), _v('k'), _t(')}')])
    p.add_run(', and capacity allocations ')
    omath(p, [_t('{'), _msub('K', 'j'), _t('}')])
    p.add_run(
        ' such that: (i) each GPU-hour is allocated to its highest-margin use; '
        '(ii) training and inference markets clear; '
        '(iii) all allocations are feasible ('
    )
    omath(p, [_msub('K', 'j'), _t(' \u2264 '),
              _mbar_sub('K', 'j')])
    p.add_run(
        '). '
        'Existence follows from a fixed-point argument: the training supply curve is a '
        'step function with steps at '
    )
    omath(p, [_msub('c', '(i)')])
    p.add_run(
        ' and widths '
    )
    omath(p, [_mbar_sub('K', '(i)')])
    p.add_run('; intersection with the demand curve pins down ')
    omath(p, [_msub('p', 'T')])
    p.add_run('.')

    # B.6 Welfare
    cur = mkh(doc, body, cur, 'B.6 Welfare Cost of Sovereignty', level=2)
    p, cur = mkp(doc, body, cur)
    p.add_run('The welfare cost has two components. For each importing country ')
    omath(p, [_v('k')])
    p.add_run(', let ')
    omath(p, [_v('i'), _t(' = '), _msubsup('j', 'k', '*')])
    p.add_run(' denote its equilibrium supplier (the seller minimizing delivered cost). '
              'Import markup:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msub('DWL', 'import'), _t(' = '),
        _nary('\u2211', [_v('k'), _t(' \u2208 '), _msub('M', 'T')], [],
              [_msub('q', 'Tk'), _t(' \u00b7 '),
               _msub('\u03BB', 'ik'),
               _t(' \u00b7 '), _msub('p', 'T')]), _t('.'),
    ], eq_num='B.4')

    p, cur = mkp(doc, body, cur)
    p.add_run('Allocative inefficiency:')
    p.paragraph_format.space_after = Pt(2)

    _, cur = omath_display(doc, body, cur, [
        _msub('DWL', 'alloc'), _t(' = '),
        _nary('\u2211', [_v('k'), _t(' : '), _msub('p', 'T'), _t(' < '),
                         _msub('c', 'k'), _t(' \u2264 (1+'),
                         _limlow([_t('min')], [_v('j')]),
                         _t(' '), _msub('\u03BB', 'jk'), _t(')'),
                         _msub('p', 'T')], [],
              [_msub('q', 'Tk'), _t(' \u00b7 ('),
               _msub('c', 'k'), _t(' \u2212 '), _msub('p', 'T'), _t(').')]),
    ], eq_num='B.5')

    p, cur = mkp(doc, body, cur)
    p.add_run('Total: ')
    omath(p, [_t('DWL('), _msub('\u03BB', 'ij'), _t(') = '),
              _msub('DWL', 'import'), _t(' + '),
              _msub('DWL', 'alloc')])
    p.add_run(
        '. Under capacity constraints, both components are smaller because the higher '
    )
    omath(p, [_msub('p', 'T')])
    p.add_run(' narrows the gap between domestic and import costs.')

    return cur


def write_sensitivity_appendix(doc, body, last_el, demand_data):
    """Appendix C: Sensitivity Analysis with Table A3."""
    print("Inserting Appendix C (Sensitivity Analysis)...")

    pb = add_page_break(doc, body, last_el)
    cur = mkh(doc, body, pb, 'Appendix C: Sensitivity Analysis', level=1)

    sens = demand_data.get("sensitivity", [])
    if not sens:
        return cur

    p, cur = mkp(doc, body, cur)
    p._element.append(make_hyperlink('TableA3', 'Table A3'))
    p.add_run(
        ' reports equilibrium outcomes under seven robustness specifications. '
        'The baseline (Form B) applies the efficiency adjustment only to non-hardware costs, '
        'with equal weight on governance and grid reliability. The table varies the governance '
        'weight, institutional floor, hardware cost share, and functional form.'
    )

    headers = ['Scenario', 'Parameter change', 'Dev top 15',
               'Max markup', 'Spearman \u03c1 vs cr', 'Top 5']
    rows = []
    for s in sens:
        rows.append([
            s["label"],
            s.get("param_change", "\u2014"),
            str(s["dev_top15"]),
            f'{s["max_markup"]:.1f}%',
            f'{s["rank_corr"]:.2f}',
            s.get("top5_str", ""),
        ])

    col_widths = [2000, 2200, 700, 700, 900, 2300]
    tbl_el = add_table(
        doc, body, cur, headers, rows, col_widths=col_widths,
        title='Table A3. Sensitivity of efficiency-adjusted rankings to parameter variation',
        bookmark_id=142, bookmark_name='TableA3', backlink_name='TableA3txt',
    )

    # Notes paragraph
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(2)
    note.paragraph_format.space_after = Pt(0)
    note.paragraph_format.first_line_indent = Inches(0)
    note.paragraph_format.line_spacing = 1.0
    rn = note.add_run(
        'Notes: Each row applies Form B (c_adj = \u03C1 + (c_cr \u2212 \u03C1) / \u03BE_eff) '
        'unless stated otherwise. Dev top 15 = number of developing countries in top 15 '
        'under efficiency-adjusted ranking. Max markup = maximum percentage increase in '
        'adjusted cost over cost-recovery cost. Spearman \u03C1 vs cr = rank correlation '
        'between efficiency-adjusted and cost-recovery rankings.'
    )
    rn.font.size = Pt(10)
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_el = note._element
    body.remove(note_el)
    tbl_el.addnext(note_el)

    return note_el


def write_kyrgyzstan_appendix(doc, body, last_el):
    """Appendix D: Data Center Investment Model — Kyrgyzstan."""
    print("Inserting Appendix D (Kyrgyzstan DCF)...")

    pb = add_page_break(doc, body, last_el)
    cur = mkh(doc, body, pb, 'Appendix D: Data Center Investment Model \u2014 Kyrgyzstan', level=1)

    # ── Parameters ────────────────────────────────────────────────────────
    IT_MW = 40
    PUE_KGZ = 1.08
    TOTAL_MW = IT_MW * PUE_KGZ
    LIFE = 15
    GP = 25_000
    G_LIFE = 3
    G_UTIL = 0.70
    G_TDP_W = 700
    GPUS_MW = 1_000 / G_TDP_W * 1_000
    N_GPU = int(IT_MW * GPUS_MW)
    H = 365.25 * 24
    NET_COST = 2_000
    P_ELEC = 0.038
    P_CONSTR_W = 7.83
    CONSTR = IT_MW * 1e6 * P_CONSTR_W
    STAFF = 50 * 12_000
    MAINT_PCT = 0.02
    INS_PCT = 0.005
    BW_COST = 2_400_000
    REV_HR = 2.00
    RAMP = {0: 0.0, 1: 0.40, 2: 0.60, 3: 0.70}
    TAX_R = 0.10
    GPU_DECLINE = 0.10
    ELEC_ESC = 0.02

    # WACC
    RF = 0.05
    CRP = 0.04
    ERP = 0.06
    COE = RF + CRP + ERP  # 15%
    COD = 0.10
    DSHARE = 0.40
    ESHARE = 0.60
    WACC = ESHARE * COE + DSHARE * COD * (1 - TAX_R)

    # GPU refresh schedule
    gpu_refresh = [1, 4, 7, 10, 13]
    gpu_prices = [(yr, GP * (1 - GPU_DECLINE) ** i) for i, yr in enumerate(gpu_refresh)]
    net_refresh = [1, 6, 11]

    # ── DCF helpers ────────────────────────────────────────────────────────
    years = list(range(0, LIFE + 1))

    def _dcf_years(gpu_adj=0, elec_adj=0, price_adj=0, util_adj=0):
        """Compute year-by-year cash flows. Returns list of per-year dicts."""
        adj_prices = [(gy, gp * (1 + gpu_adj)) for gy, gp in gpu_prices]
        rows = []
        cum = 0
        for yr in years:
            cx = CONSTR if yr == 0 else 0
            for gy, gp in adj_prices:
                if yr == gy:
                    cx += N_GPU * gp
            if yr in net_refresh:
                cx += N_GPU * NET_COST

            if yr >= 1:
                util = RAMP.get(yr, G_UTIL)
                ep = (P_ELEC + elec_adj) * (1 + ELEC_ESC) ** (yr - 1)
                gpu_val = 0
                for gy, gp in reversed(adj_prices):
                    if gy <= yr:
                        gpu_val = N_GPU * gp * max(0, 1 - (yr - gy) / G_LIFE)
                        break
                ox = (TOTAL_MW * 1_000 * H * ep + STAFF * 1.03 ** (yr - 1)
                      + CONSTR * MAINT_PCT + (CONSTR + gpu_val) * INS_PCT + BW_COST)
                rev = N_GPU * H * min(max(util + util_adj, 0), 0.95) * (REV_HR + price_adj)
                depr_g = 0
                for gy, gp in adj_prices:
                    if gy <= yr < gy + G_LIFE:
                        depr_g = N_GPU * gp / G_LIFE
                        break
                depr = CONSTR / LIFE + depr_g
            else:
                ox = 0
                rev = 0
                depr = 0

            ebitda = rev - ox
            ebt = ebitda - depr
            tax = max(0, ebt * TAX_R)
            ni = ebt - tax
            fcf = ni + depr - cx
            cum += fcf
            rows.append(dict(year=yr, capex=cx, revenue=rev, opex=ox,
                             ebitda=ebitda, tax=tax, ni=ni, fcf=fcf, cum=cum))
        return rows

    def _npv_irr(rows, wacc):
        """Compute NPV at given WACC and IRR via bisection."""
        fcfs = [r['fcf'] for r in rows]
        npv_val = sum(f / (1 + wacc) ** y for f, y in zip(fcfs, years, strict=False))
        lo, hi = -0.50, 2.0
        for _ in range(200):
            mid = (lo + hi) / 2
            if sum(f / (1 + mid) ** y for f, y in zip(fcfs, years, strict=False)) > 0:
                lo = mid
            else:
                hi = mid
        return npv_val, mid

    # ── Compute year-by-year ──────────────────────────────────────────────
    results = _dcf_years()
    npv, irr = _npv_irr(results, WACC)
    payback = next((r['year'] for r in results if r['year'] >= 1 and r['cum'] > 0), None)

    tot_rev = sum(r['revenue'] for r in results)
    tot_cx = sum(r['capex'] for r in results)
    tot_ox = sum(r['opex'] for r in results)
    tot_elec = sum(TOTAL_MW * 1_000 * H * P_ELEC * (1 + ELEC_ESC) ** (y - 1)
                   for y in range(1, LIFE + 1))
    tot_gpu_cx = sum(N_GPU * gp for _, gp in gpu_prices)

    # ── Intro paragraph ───────────────────────────────────────────────────
    p, cur = mkp(doc, body, cur)
    p._element.append(make_bookmark(145, 'TableA4txt'))
    p._element.append(make_bookmark_end(145))
    p._element.append(make_bookmark(147, 'TableA5txt'))
    p._element.append(make_bookmark_end(147))
    p._element.append(make_bookmark(149, 'TableA6txt'))
    p._element.append(make_bookmark_end(149))
    p.add_run(
        'This appendix presents a 15-year discounted cash flow (DCF) analysis for a '
        'hypothetical 40\u2009MW data center in Kyrgyzstan, the lowest-cost seller in the '
        'cost-recovery-adjusted calibration. All parameters are drawn from the calibration '
        'or from industry benchmarks.'
    )

    # ── Table A4: Facility specification ──────────────────────────────────
    specs_rows = [
        ['IT capacity', f'{IT_MW} MW'],
        ['Total power (with cooling)', f'{TOTAL_MW:.1f} MW (PUE = {PUE_KGZ:.2f})'],
        ['GPU count', f'{N_GPU:,} (H100-class, {G_TDP_W}W each)'],
        ['GPU cost / lifetime', f'${GP:,} / {G_LIFE} yr (\u221210% per generation)'],
        ['Construction cost', f'${CONSTR / 1e6:.0f}M (${P_CONSTR_W:.2f}/W)'],
        ['Electricity price', f'${P_ELEC:.3f}/kWh (+2%/yr real)'],
        ['Revenue price', f'${REV_HR:.2f}/GPU-hr (wholesale)'],
        ['Utilization', f'{G_UTIL:.0%} steady-state (40% yr 1, 60% yr 2)'],
        ['WACC', f'{WACC:.1%}'],
    ]
    tbl_a4 = add_table(doc, body, cur, ['Parameter', 'Value'],
                       specs_rows, col_widths=[3500, 5300],
                       title='Table A4. Facility specification',
                       bookmark_id=144, bookmark_name='TableA4',
                       backlink_name='TableA4txt')

    # WACC note
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 1.0
    rn = p.add_run(
        f'Notes: WACC = {ESHARE:.0%} \u00d7 {COE:.0%} (cost of equity) '
        f'+ {DSHARE:.0%} \u00d7 {COD:.0%} \u00d7 (1 \u2212 {TAX_R:.0%}) (after-tax debt) '
        f'= {WACC:.1%}. Cost of equity includes a {CRP:.0%} country risk premium and '
        f'{ERP:.0%} emerging-market equity premium over the {RF:.0%} risk-free rate.'
    )
    rn.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    wacc_el = p._element
    body.remove(wacc_el)
    tbl_a4.addnext(wacc_el)
    cur = wacc_el

    # ── Table A5: Year-by-year cash flow ──────────────────────────────────
    cur = add_page_break(doc, body, cur)
    cf_headers = ['Year', 'CAPEX', 'Revenue', 'OPEX', 'EBITDA', 'FCF', 'Cum.\u2009CF']
    cf_rows = []
    for r in results:
        cf_rows.append([
            str(r['year']),
            f'{r["capex"] / 1e6:.1f}',
            f'{r["revenue"] / 1e6:.1f}',
            f'{r["opex"] / 1e6:.1f}',
            f'{r["ebitda"] / 1e6:.1f}',
            f'{r["fcf"] / 1e6:.1f}',
            f'{r["cum"] / 1e6:.1f}',
        ])
    # Totals row
    cf_rows.append([
        'Total',
        f'{tot_cx / 1e6:.1f}',
        f'{tot_rev / 1e6:.1f}',
        f'{tot_ox / 1e6:.1f}',
        f'{sum(r["ebitda"] for r in results) / 1e6:.1f}',
        f'{sum(r["fcf"] for r in results) / 1e6:.1f}',
        '',
    ])
    tbl_a5 = add_table(doc, body, cur, cf_headers, cf_rows,
                       col_widths=[700, 1400, 1400, 1400, 1400, 1400, 1100],
                       title='Table A5. Year-by-year cash flow ($ millions)',
                       bookmark_id=146, bookmark_name='TableA5',
                       backlink_name='TableA5txt')

    # ── Key metrics paragraph ─────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    p.add_run(
        f'The project yields an NPV of ${npv / 1e6:,.0f}M at a {WACC:.1%} WACC, '
        f'an IRR of {irr:.1%}, and a simple payback in year\u2009{payback}. '
        f'GPU hardware accounts for ${tot_gpu_cx / 1e6:.0f}M of the '
        f'${tot_cx / 1e6:.0f}M total CAPEX ({tot_gpu_cx / tot_cx:.0%}), '
        f'and electricity represents {tot_elec / tot_ox:.0%} of operating costs.'
    )
    met_el = p._element
    body.remove(met_el)
    tbl_a5.addnext(met_el)
    cur = met_el

    # ── Table A6: Sensitivity analysis ────────────────────────────────────
    cur = add_page_break(doc, body, cur)

    def _run_scen(label, wacc_adj=0, price_adj=0, elec_adj=0, gpu_adj=0, util_adj=0):
        rows = _dcf_years(gpu_adj=gpu_adj, elec_adj=elec_adj,
                          price_adj=price_adj, util_adj=util_adj)
        npv_s, irr_s = _npv_irr(rows, WACC + wacc_adj)
        return [label, f'${npv_s / 1e6:,.0f}', f'{irr_s:.1%}']

    sens_scenarios = [
        _run_scen('Base case'),
        _run_scen('GPU price \u221220%', gpu_adj=-0.20),
        _run_scen('GPU price +20%', gpu_adj=+0.20),
        _run_scen('Electricity +50%', elec_adj=+0.019),
        _run_scen('Electricity \u221225%', elec_adj=-0.0095),
        _run_scen('Revenue +5%', price_adj=+0.08),
        _run_scen('Revenue \u22125%', price_adj=-0.08),
        _run_scen('Utilization 80%', util_adj=+0.10),
        _run_scen('Utilization 60%', util_adj=-0.10),
        _run_scen('WACC 10%', wacc_adj=-0.026),
        _run_scen('WACC 16%', wacc_adj=+0.034),
    ]
    tbl_a6 = add_table(doc, body, cur, ['Scenario', 'NPV ($M)', 'IRR'],
                       sens_scenarios, col_widths=[3800, 2400, 2600],
                       title='Table A6. Sensitivity of investment returns to parameter variation',
                       center_cols=[1, 2],
                       bookmark_id=148, bookmark_name='TableA6',
                       backlink_name='TableA6txt')

    # ── Risks paragraph ───────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    r = p.add_run('Risks. ')
    r.bold = True
    p.add_run(
        'Kyrgyzstan depends on the Toktogul reservoir for over 80% of electricity; '
        'seasonal drawdowns and drought years create acute power shortages. '
        'GPU procurement faces US export-control uncertainty. '
        'The production-efficiency index assigns Kyrgyzstan a governance score of 0.50, '
        'reflecting underdeveloped contract enforcement and regulatory frameworks. '
        'Despite these risks, the engineering economics are clear: '
        'electricity at $0.038/kWh and a PUE of 1.08 yield production costs well below '
        'the global median, and the positive NPV survives eight of ten '
        'perturbations in Table\u2009A6.'
        ' The share of this surplus retained in Kyrgyzstan depends on the ownership and '
        'fiscal structure: if the facility is owned by a foreign hyperscaler, most operating '
        'surplus flows abroad as repatriated profits, and the host country retains only the '
        'electricity payment and construction-phase employment unless the government '
        'captures rent through taxation, equity participation, or resource royalties.'
    )
    risk_el = p._element
    body.remove(risk_el)
    tbl_a6.addnext(risk_el)

    return risk_el


def write_construction_regression_appendix(doc, body, last_el):
    """Appendix E: Construction Cost Regression."""
    import math as _math
    print("Inserting Appendix E (Construction Regression)...")

    pb = add_page_break(doc, body, last_el)
    cur = mkh(doc, body, pb, 'Appendix E: Construction Cost Regression', level=1)

    p, cur = mkp(doc, body, cur)
    p._element.append(make_bookmark(151, 'TableA7txt'))
    p._element.append(make_bookmark_end(151))
    p.add_run(
        'Data center construction costs per watt of IT capacity are observed for 37 countries '
        'from the Turner & Townsend Data Centre Construction Cost Index 2025 (52 markets). '
        'For the remaining countries, construction costs are predicted using the log-linear '
        'regression reported in '
    )
    p._element.append(make_hyperlink('TableA7', 'Table\u2009A7'))
    p.add_run(
        '. The dependent variable is ln($/W). '
        'Since construction accounts for only 3\u20136% of total per-GPU-hour costs, '
        'imputation error has limited impact on cost rankings.'
    )

    # ── Run regression inline ──────────────────────────────────────────────
    import numpy as _np
    _DATA = DATA

    MARKET_TO_ISO3 = {
        "Tokyo": "JPN", "Singapore": "SGP", "Zurich": "CHE", "Osaka": "JPN",
        "Silicon Valley": "USA", "New Jersey": "USA", "Oslo": "NOR",
        "Auckland": "NZL", "Stockholm": "SWE", "Helsinki": "FIN",
        "Copenhagen": "DNK", "London": "GBR", "Vienna": "AUT",
        "Cardiff": "GBR", "Frankfurt": "DEU", "Berlin": "DEU",
        "Kuala Lumpur": "MYS", "Kingdom of Saudi Arabia": "SAU",
        "Chicago": "USA", "Jakarta": "IDN", "North Virginia": "USA",
        "Portland": "USA", "Paris": "FRA", "Amsterdam": "NLD",
        "S\u00e3o Paulo": "BRA", "Sydney": "AUS", "Lagos": "NGA",
        "Melbourne": "AUS", "Quer\u00e9taro": "MEX", "Cape Town": "ZAF",
        "Lisbon": "PRT", "Seoul": "KOR", "Johannesburg": "ZAF",
        "Bordeaux": "FRA", "Dublin": "IRL", "Madrid": "ESP",
        "Atlanta": "USA", "Montevideo": "URY", "Phoenix": "USA",
        "Columbus": "USA", "Milan": "ITA", "Nairobi": "KEN",
        "Dallas": "USA", "Charlotte": "USA", "Toronto": "CAN",
        "UAE": "ARE", "Warsaw": "POL", "Santiago": "CHL",
        "Athens": "GRC", "Bogot\u00e1": "COL", "Mumbai": "IND",
        "Shanghai": "CHN",
    }

    dcci = {}
    with open(_DATA / "dcci_2025_construction_costs.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            iso3 = MARKET_TO_ISO3[row["market"]]
            cost = float(row["usd_per_watt"])
            if iso3 in dcci:
                dcci[iso3].append(cost)
            else:
                dcci[iso3] = [cost]
    for iso3 in dcci:
        dcci[iso3] = _np.mean(dcci[iso3])

    gdp_d = {}
    with open(_DATA / "wb_gdp_per_capita_ppp_2023.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            gdp_d[row["iso3"]] = float(row["gdp_pcap_ppp_2023"])
    reg_d = {}
    with open(_DATA / "wb_country_regions.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            reg_d[row["iso3"]] = row["region"]
    urban_d = {}
    with open(_DATA / "wb_urban_share_2023.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            urban_d[row["iso3"]] = float(row["urban_share_pct"]) / 100.0
    seismic_d = {}
    with open(_DATA / "seismic_zones.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            seismic_d[row["iso3"]] = int(row["seismic_high"])

    REF_REGION = "Europe & Central Asia"
    DUMMY_REGIONS = sorted(r for r in set(reg_d.values()) if r != REF_REGION)

    matched = []
    for iso3, avg_cost in dcci.items():
        if iso3 in gdp_d and iso3 in reg_d:
            matched.append({
                "iso3": iso3, "cost": avg_cost,
                "gdp_pcap": gdp_d[iso3], "region": reg_d[iso3],
                "urban_share": urban_d.get(iso3, 0.5),
                "seismic": seismic_d.get(iso3, 0),
            })

    n = len(matched)
    k = 5 + len(DUMMY_REGIONS)
    y = _np.array([_math.log(m["cost"]) for m in matched])
    X = _np.zeros((n, k))
    col_names = ["Intercept", "ln(GDP per capita)", "ln(Population)",
                 "Urban population share",
                 "Seismic zone indicator"] + [r.split(",")[0].strip() for r in DUMMY_REGIONS]
    pop_d = {}
    with open(_DATA / "wb_population_2023.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            pop_d[row["iso3"]] = int(row["population_2023"])
    for i, m in enumerate(matched):
        X[i, 0] = 1.0
        X[i, 1] = _math.log(m["gdp_pcap"])
        X[i, 2] = _math.log(pop_d.get(m["iso3"], 1_000_000))
        X[i, 3] = m["urban_share"]
        X[i, 4] = m["seismic"]
        for j2, reg in enumerate(DUMMY_REGIONS):
            X[i, 5 + j2] = 1.0 if m["region"] == reg else 0.0

    beta = _np.linalg.lstsq(X, y, rcond=None)[0]
    y_hat = X @ beta
    resid = y - y_hat
    ss_res = _np.sum(resid ** 2)
    ss_tot = _np.sum((y - _np.mean(y)) ** 2)
    r2 = 1 - ss_res / ss_tot
    adj_r2 = 1 - (1 - r2) * (n - 1) / (n - k)
    rmse = _math.sqrt(ss_res / (n - k))
    var_beta = ss_res / (n - k) * _np.diag(_np.linalg.inv(X.T @ X))
    se = _np.sqrt(_np.maximum(var_beta, 0))

    # Build table rows — coefficient with significance stars, SE in italic
    reg_rows = []
    for j2 in range(k):
        sig = ''
        t = beta[j2] / se[j2] if se[j2] > 0 else 0
        if abs(t) > 2.576:
            sig = '***'
        elif abs(t) > 1.96:
            sig = '**'
        elif abs(t) > 1.645:
            sig = '*'
        reg_rows.append([
            col_names[j2],
            f'{beta[j2]:.3f}{sig}',
            f'({se[j2]:.3f})',
        ])

    tbl = add_table(doc, body, cur,
                    ['Variable', 'Coefficient', 'Std. Error'],
                    reg_rows, col_widths=[3600, 2200, 2200],
                    title='Table A7. Construction cost regression: ln($/W)',
                    center_cols=[1, 2],
                    bookmark_id=150, bookmark_name='TableA7',
                    backlink_name='TableA7txt')

    # Post-process: make SE column numbers italic
    all_trs = tbl.findall(f'{{{W_NS}}}tr')
    for tr in all_trs[1:]:  # skip header
        tcs = tr.findall(f'{{{W_NS}}}tc')
        if len(tcs) > 2:
            tc_se = tcs[2]
            for r_el in tc_se.findall(f'.//{{{W_NS}}}r'):
                rPr = r_el.find(f'{{{W_NS}}}rPr')
                if rPr is None:
                    rPr = OxmlElement('w:rPr')
                    r_el.insert(0, rPr)
                i_el = OxmlElement('w:i')
                rPr.append(i_el)

    # Notes paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 1.0
    rn = p.add_run(
        f'Notes: OLS regression on {n} countries from the Turner & Townsend DCCI 2025. '
        f'Dependent variable: ln(construction cost in $/W). '
        f'R\u00b2 = {r2:.2f}, adjusted R\u00b2 = {adj_r2:.2f}, RMSE = {rmse:.3f}. '
        f'Reference region: Europe & Central Asia. '
        f'*** p < 0.01, ** p < 0.05, * p < 0.10.'
    )
    rn.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_el = p._element
    body.remove(note_el)
    tbl.addnext(note_el)

    return note_el


def write_workload_appendix(doc, body, last_el):
    """Appendix F: Workload Classification (Table A8)."""
    print("Inserting Appendix F (Workload Classification)...")

    pb = add_page_break(doc, body, last_el)
    cur = mkh(doc, body, pb, 'Appendix F: Workload Classification', level=1)

    # Introductory paragraph
    p, cur = mkp(doc, body, cur)
    p._element.append(make_bookmark(153, 'TableA8txt'))
    p._element.append(make_hyperlink('TableA8', 'Table A8'))
    p._element.append(make_bookmark_end(153))
    p.add_run(
        ' summarizes the latency sensitivity and offshorability of major AI workload '
        'types. The model collapses these into two categories \u2014 training '
        '(\u03C4 = 0) and inference (\u03C4 > 0) \u2014 but the intermediate workloads '
        'noted in footnote 7 occupy a middle ground that may narrow the effective '
        'offshorable share.'
    )

    # Table A8
    headers = ['Workload', 'Example', 'Latency tolerance',
               'Offshorable?', 'Model treatment']
    rows = [
        ['Large-scale training', 'Foundation model pre-training',
         'Days\u2013weeks', 'Fully', '\u03C4\u1D1B = 0'],
        ['Fine-tuning', 'Domain adaptation on proprietary data',
         'Hours', 'Mostly', 'Treated as training'],
        ['Agentic inference', 'Multi-step reasoning, tool use',
         '500\u20132,000 ms', 'Regionally', 'Intermediate (fn.\u00A07)'],
        ['Interactive inference', 'Chatbot, search, recommendation',
         '50\u2013200 ms', 'Limited', '\u03C4\u1D62 > 0, threshold l\u0304'],
        ['Real-time control', 'Autonomous vehicles, robotics',
         '< 20 ms', 'No', 'Domestic only'],
    ]
    tbl = add_table(doc, body, cur, headers, rows,
                    col_widths=[2000, 2800, 1500, 1200, 1860],
                    title='Table A8. Workload classification and offshorability',
                    bookmark_id=154, bookmark_name='TableA8',
                    backlink_name='TableA8txt')

    # Notes paragraph
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.line_spacing = 1.0
    rn = p.add_run(
        'Notes: Latency tolerance is approximate round-trip time. '
        '\u201COffshorable\u201D refers to whether the workload can be processed in a '
        'different country from the end user without significant quality degradation. '
        'The model treats fine-tuning and agentic inference as part of the training '
        'share \u03B1; footnote 7 notes this simplification.'
    )
    rn.font.size = Pt(10)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_el = p._element
    body.remove(note_el)
    tbl.addnext(note_el)

    return note_el


def write_figure1_calibration(doc, body, last_ref):
    """Insert Figure 1 (calibration strategy flowchart) after references."""
    print("Inserting Figure 1 (Calibration Strategy)...")

    fig_path = DOCS / "calibration_strategy_fig1.png"

    # Page break before figure
    pb_el = add_page_break(doc, body, last_ref)

    # Figure title with bookmark
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.paragraph_format.first_line_indent = Inches(0)
    title_p._element.append(make_bookmark(152, 'Figure1cals'))
    hl = OxmlElement('w:hyperlink')
    hl.set(qn('w:anchor'), 'Figure1calstxt')
    hl.set(qn('w:history'), '1')
    r_el = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    b_el = OxmlElement('w:b')
    rPr.append(b_el)
    sz_el = OxmlElement('w:sz')
    sz_el.set(qn('w:val'), '20')
    rPr.append(sz_el)
    clr_el = OxmlElement('w:color')
    clr_el.set(qn('w:val'), LINK_COLOR)
    uu_el = OxmlElement('w:u')
    uu_el.set(qn('w:val'), 'single')
    rPr.append(clr_el)
    rPr.append(uu_el)
    r_el.append(rPr)
    t_el = OxmlElement('w:t')
    t_el.text = 'Figure 1'
    r_el.append(t_el)
    hl.append(r_el)
    title_p._element.append(hl)
    title_p._element.append(make_bookmark_end(152))
    run_ft = title_p.add_run('. Calibration strategy')
    run_ft.bold = True
    run_ft.font.size = Pt(10)
    title_el = title_p._element
    body.remove(title_el)
    pb_el.addnext(title_el)

    # Embed image
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(4)
    pic_p.paragraph_format.space_after = Pt(4)
    run = pic_p.add_run()
    run.add_picture(str(fig_path), width=Inches(5.5))
    pic_el = pic_p._element
    body.remove(pic_el)
    title_el.addnext(pic_el)

    # Notes
    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(4)
    note_p.paragraph_format.space_after = Pt(6)
    note_p.paragraph_format.first_line_indent = Inches(0)
    note_p.paragraph_format.left_indent = Inches(0.5)
    note_p.paragraph_format.right_indent = Inches(0.5)
    rn1 = note_p.add_run('Notes: ')
    rn1.bold = True
    rn1.font.size = Pt(10)
    rn2 = note_p.add_run(
        'Step 1 adjusts production costs from observed electricity tariffs '
        'through cost-recovery pricing (removing subsidies) to efficiency-adjusted '
        'costs (penalizing weak governance and grid reliability). Step 2 applies '
        'trade frictions via bilateral sovereignty premiums and a uniform premium '
        'as a robustness check. '
        'Step 5 replaces the bilateral sovereignty premium with equation (2\u2032), '
        'in which the trust counterparty is the hyperscaler operator rather than '
        'the host country.'
    )
    rn2.font.size = Pt(10)
    note_p.paragraph_format.line_spacing = 1.0
    note_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_el = note_p._element
    body.remove(note_el)
    pic_el.addnext(note_el)

    return note_el


def write_figure4b(doc, body, last_ref, demand_data):
    """Generate and embed Figure 2 (efficiency rank scatter) after Figure 1."""
    # FLAG E1: Figure 2 auto-regenerated with Form B adjusted costs.
    # Max rank change is now +57 (Turkmenistan) vs previous +82 (Iran).
    print("Embedding Figure 2 (Efficiency Rank Scatter) [v27: Form B values]...")
    xi_adj = demand_data.get("xi_adjusted", {})
    baseline_rank = xi_adj.get("baseline_rank_map", {})
    xi_rank = xi_adj.get("xi_rank_map", {})
    iso_country = demand_data.get("iso_country", {})
    xi = demand_data.get("xi", {})
    if not baseline_rank or not xi_rank:
        return last_ref

    # Countries with active DC construction announcements (from Section 1)
    DC_ACTIVE = {
        'ARM', 'KEN', 'SAU', 'MAR', 'MYS', 'IDN', 'ARE',  # original
        'IND', 'BRA', 'MEX', 'CHN', 'THA', 'VNM', 'TUR',  # tier 1
        'PHL', 'ZAF', 'EGY', 'KAZ', 'UZB', 'NGA',         # tier 1-2
    }
    # Countries whose ξ value should be shown in the label
    XI_SHOW = {'IRN', 'PAK', 'CHN', 'FIN', 'KEN',
               'RUS', 'TJK', 'UKR', 'BIH', 'ARM', 'BGR'}

    common = [iso for iso in baseline_rank if iso in xi_rank]

    # Separate into active-construction and regular
    reg_isos = [iso for iso in common if iso not in DC_ACTIVE]
    act_isos = [iso for iso in common if iso in DC_ACTIVE]

    fig, ax = plt.subplots(figsize=(5.5, 5.5))
    # Regular countries: dots
    ax.scatter([baseline_rank[iso] + 1 for iso in reg_isos],
               [xi_rank[iso] + 1 for iso in reg_isos],
               s=20, c='#b2182b', alpha=0.7, marker='o',
               edgecolors='white', linewidth=0.3, label='Other countries', zorder=2)
    # Active construction: stars
    if act_isos:
        ax.scatter([baseline_rank[iso] + 1 for iso in act_isos],
                   [xi_rank[iso] + 1 for iso in act_isos],
                   s=60, c='#1a3a5c', alpha=0.85, marker='*',
                   edgecolors='white', linewidth=0.3,
                   label='Active DC construction', zorder=3)
    maxr = max(max(baseline_rank[iso] + 1 for iso in common),
               max(xi_rank[iso] + 1 for iso in common))
    ax.plot([1, maxr], [1, maxr], '--', color='gray', alpha=0.5, linewidth=0.8)

    # Build labels
    def _label(iso):
        name = iso_country.get(iso, iso)
        if iso in XI_SHOW:
            return f'{name} (\u03BE={rhup(xi.get(iso, 1.0)):.2f})'
        return name

    try:
        from adjustText import adjust_text
        texts = []
        for iso in common:
            shift = abs(baseline_rank[iso] - xi_rank[iso])
            if (shift > 15 or baseline_rank[iso] < 5 or xi_rank[iso] < 5
                    or iso in XI_SHOW or iso in DC_ACTIVE):
                texts.append(ax.text(baseline_rank[iso] + 1, xi_rank[iso] + 1,
                                     _label(iso), fontsize=5.5, alpha=0.85))
        adjust_text(texts, ax=ax,
                    arrowprops=dict(arrowstyle='-', color='gray', alpha=0.4, lw=0.4),
                    force_text=(0.4, 0.4), expand=(1.2, 1.4))
    except ImportError:
        for iso in common:
            shift = abs(baseline_rank[iso] - xi_rank[iso])
            if (shift > 15 or baseline_rank[iso] < 5 or xi_rank[iso] < 5
                    or iso in XI_SHOW or iso in DC_ACTIVE):
                ax.annotate(_label(iso),
                            (baseline_rank[iso] + 1, xi_rank[iso] + 1),
                            fontsize=5.5, alpha=0.85)

    ax.set_xlabel('Baseline cost rank', fontsize=9)
    ax.set_ylabel('Efficiency-adjusted rank', fontsize=9)
    # No legend – star/dot distinction explained in figure notes
    ax.grid(alpha=0.2)
    fig.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format='png', dpi=300, bbox_inches='tight')
    plt.close(fig)
    buf.seek(0)

    # Page break before figure
    pb_el = add_page_break(doc, body, last_ref)

    # Figure title with bookmark (outside the image)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(6)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.paragraph_format.first_line_indent = Inches(0)
    title_p._element.append(make_bookmark(120, 'Figure2'))
    hl_f1 = OxmlElement('w:hyperlink')
    hl_f1.set(qn('w:anchor'), 'Figure2txt')
    hl_f1.set(qn('w:history'), '1')
    r_f1 = OxmlElement('w:r')
    rPr_f1 = OxmlElement('w:rPr')
    b_f1 = OxmlElement('w:b')
    rPr_f1.append(b_f1)
    sz_f1 = OxmlElement('w:sz')
    sz_f1.set(qn('w:val'), '20')
    rPr_f1.append(sz_f1)
    clr_f1 = OxmlElement('w:color')
    clr_f1.set(qn('w:val'), LINK_COLOR)
    uu_f1 = OxmlElement('w:u')
    uu_f1.set(qn('w:val'), 'single')
    rPr_f1.append(clr_f1)
    rPr_f1.append(uu_f1)
    r_f1.append(rPr_f1)
    t_f1 = OxmlElement('w:t')
    t_f1.text = 'Figure 2'
    r_f1.append(t_f1)
    hl_f1.append(r_f1)
    title_p._element.append(hl_f1)
    title_p._element.append(make_bookmark_end(120))
    run_ft = title_p.add_run('. Rank change with efficiency adjustment')
    run_ft.bold = True
    run_ft.font.size = Pt(10)
    title_el = title_p._element
    body.remove(title_el)
    pb_el.addnext(title_el)

    # Embed image
    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(4)
    pic_p.paragraph_format.space_after = Pt(4)
    run = pic_p.add_run()
    run.add_picture(buf, width=Inches(4.5))
    pic_el = pic_p._element
    body.remove(pic_el)
    title_el.addnext(pic_el)

    # Notes (with 0.5" left and right indent)
    note_p = doc.add_paragraph()
    note_p.paragraph_format.space_before = Pt(4)
    note_p.paragraph_format.space_after = Pt(6)
    note_p.paragraph_format.first_line_indent = Inches(0)
    note_p.paragraph_format.left_indent = Inches(0.5)
    note_p.paragraph_format.right_indent = Inches(0.5)
    rn1 = note_p.add_run('Notes: ')
    rn1.bold = True
    rn1.font.size = Pt(10)
    rn2 = note_p.add_run(
        'Each point is one country. The dashed line marks unchanged rank. '
        'Countries above the line improve their position after efficiency '
        'adjustment; countries below it fall. Stars (\u2605) indicate countries '
        'with active data center construction announcements. '
        'Values in parentheses show the production-efficiency index \u03BE\u1d49\u1da0\u1da0. '
        'Even countries with \u03BE\u1d49\u1da0\u1da0 \u2248 1 shift off the diagonal because '
        'penalizing low-\u03BE\u1d49\u1da0\u1da0 competitors pushes them down, mechanically '
        'raising higher-\u03BE\u1d49\u1da0\u1da0 countries.'
    )
    rn2.font.size = Pt(10)
    note_p.paragraph_format.line_spacing = 1.0
    note_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_el = note_p._element
    body.remove(note_el)
    pic_el.addnext(note_el)

    return note_el


def write_table1(doc, body, after_el):
    """Table 1: Country regime taxonomy (5×5 grid), placed after Figure 2."""
    print("Inserting Table 1 (Country regime taxonomy)...")
    cur = after_el

    # Page break before Table 1
    cur = add_page_break(doc, body, cur)

    # Title
    tp_tax = doc.add_paragraph()
    tp_tax.paragraph_format.space_before = Pt(10)
    tp_tax.paragraph_format.space_after = Pt(4)
    tp_tax.paragraph_format.first_line_indent = Inches(0)
    tp_tax.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tp_tax._element.append(make_bookmark(131, 'Table1'))
    hl_tax = OxmlElement('w:hyperlink')
    hl_tax.set(qn('w:anchor'), 'Table1txt')
    hl_tax.set(qn('w:history'), '1')
    r_tax = OxmlElement('w:r')
    rPr_tax = OxmlElement('w:rPr')
    b_tax = OxmlElement('w:b')
    rPr_tax.append(b_tax)
    sz_tax = OxmlElement('w:sz')
    sz_tax.set(qn('w:val'), '20')
    rPr_tax.append(sz_tax)
    clr_tax = OxmlElement('w:color')
    clr_tax.set(qn('w:val'), LINK_COLOR)
    uu_tax = OxmlElement('w:u')
    uu_tax.set(qn('w:val'), 'single')
    rPr_tax.append(clr_tax)
    rPr_tax.append(uu_tax)
    r_tax.append(rPr_tax)
    t_tax = OxmlElement('w:t')
    t_tax.text = 'Table 1'
    r_tax.append(t_tax)
    hl_tax.append(r_tax)
    tp_tax._element.append(hl_tax)
    tp_tax._element.append(make_bookmark_end(131))
    run_tt = tp_tax.add_run('. Country regime taxonomy (Proposition 1)')
    run_tt.bold = True
    run_tt.font.size = Pt(10)
    tp_tax_el = tp_tax._element
    body.remove(tp_tax_el)
    cur.addnext(tp_tax_el)
    cur = tp_tax_el

    # Build the 5×5 taxonomy table
    tax_tbl = doc.add_table(rows=5, cols=5)
    tax_tbl.style = 'Table Grid'
    tax_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove all borders, then selectively add
    tblPr_t = tax_tbl._tbl.find(qn('w:tblPr'))
    if tblPr_t is None:
        tblPr_t = OxmlElement('w:tblPr')
        tax_tbl._tbl.insert(0, tblPr_t)
    old_bdr_t = tblPr_t.find(qn('w:tblBorders'))
    if old_bdr_t is not None:
        tblPr_t.remove(old_bdr_t)
    tblBorders_t = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        bb = OxmlElement(f'w:{side}')
        bb.set(qn('w:val'), 'single')
        bb.set(qn('w:sz'), '4')
        bb.set(qn('w:space'), '0')
        bb.set(qn('w:color'), 'auto')
        tblBorders_t.append(bb)
    tblPr_t.append(tblBorders_t)
    tblW_t = tblPr_t.find(qn('w:tblW'))
    if tblW_t is None:
        tblW_t = OxmlElement('w:tblW')
        tblPr_t.append(tblW_t)
    tblW_t.set(qn('w:w'), '4800')
    tblW_t.set(qn('w:type'), 'pct')

    def _tax_cell(row, col, text, bold=False, italic=False, gray=False,
                  font_size=8, align=WD_ALIGN_PARAGRAPH.CENTER):
        """Set content and formatting of a taxonomy table cell."""
        cell = tax_tbl.rows[row].cells[col]
        cell.text = ''
        p_c = cell.paragraphs[0]
        p_c.alignment = align
        if gray:
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'), 'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'), 'E0E0E0')
            cell._tc.get_or_add_tcPr().append(shading)
        for line_i, line in enumerate(text.split('\n')):
            if line_i > 0:
                p_c.add_run().add_break()
            run = p_c.add_run(line)
            run.font.size = Pt(font_size)
            run.bold = bold
            run.italic = italic

    # Row 0: header row — merge [0,2:4] for "Training" header
    _tax_cell(0, 0, '', font_size=7)
    _tax_cell(0, 1, '', font_size=7)
    _tax_cell(0, 2, 'Training', bold=True, font_size=8)
    _tax_cell(0, 3, 'domestic production cost \u2192', bold=True, font_size=7, italic=True)
    _tax_cell(0, 4, '', font_size=7)
    # Merge row 0 cells 2-4
    tax_tbl.rows[0].cells[2].merge(tax_tbl.rows[0].cells[4])
    # Re-set merged cell text
    merged_cell = tax_tbl.rows[0].cells[2]
    merged_cell.text = ''
    mp = merged_cell.paragraphs[0]
    mp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_train = mp.add_run('Training')
    r_train.bold = True
    r_train.font.size = Pt(9)
    mp.add_run().add_break()
    r_arrow = mp.add_run('domestic production cost \u2192')
    r_arrow.italic = True
    r_arrow.font.size = Pt(7)

    # Row 1: sub-headers
    _tax_cell(1, 0, '', font_size=7)
    _tax_cell(1, 1, '', font_size=7)
    _tax_cell(1, 2, 'Export', bold=True, font_size=8)
    _tax_cell(1, 3, 'Domestic', bold=True, font_size=8)
    _tax_cell(1, 4, 'Import', bold=True, font_size=8)
    # Merge row 0-1 col 0 and row 0-1 col 1 for the corner
    tax_tbl.rows[0].cells[0].merge(tax_tbl.rows[1].cells[0])
    tax_tbl.rows[0].cells[1].merge(tax_tbl.rows[1].cells[1])

    # Merge col 0 rows 2-4 for "Inference" header
    tax_tbl.rows[2].cells[0].merge(tax_tbl.rows[4].cells[0])
    inf_cell = tax_tbl.rows[2].cells[0]
    inf_cell.text = ''
    ip = inf_cell.paragraphs[0]
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_inf = ip.add_run('Inference')
    r_inf.bold = True
    r_inf.font.size = Pt(9)
    ip.add_run().add_break()
    r_iarrow = ip.add_run('\u2190 domestic production cost')
    r_iarrow.italic = True
    r_iarrow.font.size = Pt(7)
    # Vertical text direction for the Inference cell
    tcPr_inf = inf_cell._tc.get_or_add_tcPr()
    text_dir = OxmlElement('w:textDirection')
    text_dir.set(qn('w:val'), 'btLr')
    tcPr_inf.append(text_dir)
    vAlign_inf = OxmlElement('w:vAlign')
    vAlign_inf.set(qn('w:val'), 'center')
    tcPr_inf.append(vAlign_inf)

    # Row labels (col 1): Export, Domestic, Import
    _tax_cell(2, 1, 'Export', bold=True, font_size=8)
    _tax_cell(3, 1, 'Domestic', bold=True, font_size=8)
    _tax_cell(4, 1, 'Import', bold=True, font_size=8)

    # Content cells — 3×3 grid (rows 2-4, cols 2-4)
    _tax_cell(2, 2, '\u2713  (i) EE\nT+I exporter\nCheapest producers', font_size=7.5)
    _tax_cell(2, 3, '\u2717\nProp. 4', font_size=7.5, gray=True)
    _tax_cell(2, 4, '\u2713  (ii) IE\nInference hub\nRegional low-cost hubs', font_size=7.5)
    _tax_cell(3, 2, '\u2717\nProp. 4', font_size=7.5, gray=True)
    _tax_cell(3, 3, '\u2713  (iv) DD\nDomestic\nHigh sovereignty', font_size=7.5)
    _tax_cell(3, 4, '\u2713  (iii) ID\nHybrid\nIsolated / moderate cost', font_size=7.5)
    _tax_cell(4, 2, '\u2717\nProp. 4', font_size=7.5, gray=True)
    _tax_cell(4, 3, '\u2717\nSovereignty\ndominates', font_size=7.5, gray=True)
    _tax_cell(4, 4, '\u2713  (v) II\nFull importer\nHigh-cost countries', font_size=7.5)

    # Cell spacing
    for row in tax_tbl.rows:
        for cell in row.cells:
            for pp in cell.paragraphs:
                pPr = pp._element.get_or_add_pPr()
                sp = OxmlElement('w:spacing')
                sp.set(qn('w:before'), '20')
                sp.set(qn('w:after'), '20')
                pPr.append(sp)

    # Position table after title
    tax_tbl_el = tax_tbl._tbl
    body.remove(tax_tbl_el)
    tp_tax_el.addnext(tax_tbl_el)
    cur = tax_tbl_el

    # Table notes
    tn = doc.add_paragraph()
    tn.paragraph_format.space_before = Pt(2)
    tn.paragraph_format.space_after = Pt(8)
    tn.paragraph_format.first_line_indent = Inches(0)
    tn.paragraph_format.line_spacing = 1.0
    tn.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rn = tn.add_run(
        'Notes: \u2713 = feasible in equilibrium. \u2717 = ruled out. '
        'Grey cells cannot arise. Roman numerals correspond to regime labels '
        'in Proposition 1. '
        'Letter codes (EE, IE, ID, DD, II) denote training status '
        '(first letter: E\u2009=\u2009export, I\u2009=\u2009import, D\u2009=\u2009domestic) and inference status '
        '(second letter), used in '
    )
    rn.font.size = Pt(10)
    tn._element.append(make_hyperlink('Table3', 'Table 3a', rPr_orig=_rPr_pt(10)))
    rn = tn.add_run('.')
    rn.font.size = Pt(10)
    tn_el = tn._element
    body.remove(tn_el)
    tax_tbl_el.addnext(tn_el)

    return tn_el


def write_table2(doc, body, after_el, demand_data):
    """Table 2: Model parameters (formerly Table 1), placed in main body."""
    print("Inserting Table 2 (Model parameters)...")

    # Page break before Table 2
    pb_el = add_page_break(doc, body, after_el)

    # Table 2 title with bookmark
    tp1 = doc.add_paragraph()
    tp1.paragraph_format.space_before = Pt(6)
    tp1.paragraph_format.space_after = Pt(3)
    tp1.paragraph_format.first_line_indent = Inches(0)
    tp1._element.append(make_bookmark(110, 'Table2'))
    hl_a1 = OxmlElement('w:hyperlink')
    hl_a1.set(qn('w:anchor'), 'Table2txt')
    hl_a1.set(qn('w:history'), '1')
    r_a1 = OxmlElement('w:r')
    rPr_a1 = OxmlElement('w:rPr')
    b_a1 = OxmlElement('w:b')
    rPr_a1.append(b_a1)
    sz_a1 = OxmlElement('w:sz')
    sz_a1.set(qn('w:val'), '20')
    rPr_a1.append(sz_a1)
    clr_a1 = OxmlElement('w:color')
    clr_a1.set(qn('w:val'), LINK_COLOR)
    uu_a1 = OxmlElement('w:u')
    uu_a1.set(qn('w:val'), 'single')
    rPr_a1.append(clr_a1)
    rPr_a1.append(uu_a1)
    r_a1.append(rPr_a1)
    t_a1 = OxmlElement('w:t')
    t_a1.text = 'Table 2'
    r_a1.append(t_a1)
    hl_a1.append(r_a1)
    tp1._element.append(hl_a1)
    tp1._element.append(make_bookmark_end(110))
    run_tt1 = tp1.add_run('. Model parameters')
    run_tt1.bold = True
    run_tt1.font.size = Pt(10)
    tp1_el = tp1._element
    body.remove(tp1_el)
    pb_el.addnext(tp1_el)

    # Load parameters from CSV
    param_rows = []
    with open(DATA / "model_parameters.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            param_rows.append(row)

    _sym_map = {
        'gamma': '\u03B3', 'P_GPU': 'P_GPU', 'L': 'L',
        'beta': '\u03B2', 'H': 'H', 'rho': '\u03C1', 'eta': '\u03B7',
        'phi': '\u03C6', 'delta': '\u03B4', 'theta_bar': '\u03B8\u0304',  # rendered via OMML in cell
        'D': 'D', 'tau': '\u03C4', 'lambda': '\u03BB', 'alpha': '\u03B1',
        'Q': 'Q', 'xi_j': '\u03BE\u2C7C',
    }
    _source_to_bm = {
        'NVIDIA (2024)': 'NVIDIA2024',
        'Barroso et al. (2018)': 'Barroso2018',
        'Liu et al. (2023)': 'Liu2023',
        'Flucker et al. (2013)': 'Flucker2013',
        'Turner and Townsend (2025)': 'TurnerTownsend2025',
        'UNCTAD (2025)': 'UNCTAD2025',
        'Deloitte (2025)': 'Deloitte2025',
        'Bailey et al. (2017)': 'Bailey2017',
        'Epoch AI (2024)': 'EpochAI2024',
        'Google (2024)': 'Google2024',
        'WGI and Enterprise Surveys': 'WorldBank2024',
    }

    n_params = len(param_rows)
    param_tbl = doc.add_table(rows=n_params + 1, cols=4)
    param_tbl.style = 'Table Grid'
    param_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _tbl_clear_borders(param_tbl)

    _pcw = [Inches(2.6), Inches(0.6), Inches(1.5), Inches(1.8)]
    _pcw_labels = ['Parameter', 'Symbol', 'Value', 'Source']

    for j, lbl in enumerate(_pcw_labels):
        cell = param_tbl.rows[0].cells[j]
        cell.text = ''
        p_h = cell.paragraphs[0]
        p_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rh = p_h.add_run(lbl)
        rh.bold = True
        rh.font.size = Pt(10)
        rh.font.name = TIMES_NEW_ROMAN
        cell.width = _pcw[j]
        _tbl_border(cell._tc, ['top', 'bottom'])

    for i, pr in enumerate(param_rows):
        sym_display = _sym_map.get(pr['symbol'], pr['symbol'])
        val_str = pr['value']
        if pr['unit']:
            val_str = f"{val_str} {pr['unit']}"
        if pr['symbol'] == 'P_GPU':
            val_str = f"${int(float(pr['value'])):,}"
        elif pr['symbol'] == 'rho':
            val_str = f"${RHO:.2f}/hr"
        elif pr['symbol'] == 'eta':
            val_str = f"${ETA:.2f}/hr"
        elif pr['symbol'] == 'Q':
            val_str = "6\u00d710\u00b9\u2070 GPU-hr/yr"
        src_text = pr['source']
        src_bm = _source_to_bm.get(src_text)
        row_data = [pr['description'], sym_display, val_str]
        for j, txt in enumerate(row_data):
            cell = param_tbl.rows[i + 1].cells[j]
            cell.text = ''
            p_c = cell.paragraphs[0]
            if j == 0:
                p_c.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Special symbol rendering
            if j == 1 and pr['symbol'] == 'P_GPU':
                rc = p_c.add_run('P')
                rc.font.size = Pt(10)
                rc.font.name = TIMES_NEW_ROMAN
                rc.italic = True
                rc_sub = p_c.add_run('GPU')
                rc_sub.font.size = Pt(10)
                rc_sub.font.name = TIMES_NEW_ROMAN
                rPr_sub = rc_sub._element.get_or_add_rPr()
                vertAlign = OxmlElement('w:vertAlign')
                vertAlign.set(qn('w:val'), 'subscript')
                rPr_sub.append(vertAlign)
            elif j == 1 and pr['symbol'] == 'theta_bar':
                omath(p_c, [_mbar('\u03B8')])
            elif j == 1 and pr['symbol'] == 'lambda':
                omath(p_c, [_msub('\u03BB', 'ij')])
            elif j == 1 and pr['symbol'] == 'xi_j':
                omath(p_c, [_msubsup('\u03BE', 'j', 'eff')])
            else:
                rc = p_c.add_run(txt)
                rc.font.size = Pt(10)
                rc.font.name = TIMES_NEW_ROMAN
            cell.width = _pcw[j]
        src_cell = param_tbl.rows[i + 1].cells[3]
        src_cell.text = ''
        src_p = src_cell.paragraphs[0]
        src_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        src_cell.width = _pcw[3]
        if src_bm and src_text:
            rPr_src = OxmlElement('w:rPr')
            sz_src = OxmlElement('w:sz')
            sz_src.set(qn('w:val'), '20')
            rPr_src.append(sz_src)
            rFnt_src = OxmlElement('w:rFonts')
            rFnt_src.set(qn('w:ascii'), TIMES_NEW_ROMAN)
            rFnt_src.set(qn('w:hAnsi'), TIMES_NEW_ROMAN)
            rPr_src.append(rFnt_src)
            hl_src = make_hyperlink(src_bm, src_text, rPr_orig=rPr_src)
            src_p._element.append(hl_src)
        elif src_text:
            rc_src = src_p.add_run(src_text)
            rc_src.font.size = Pt(10)
            rc_src.font.name = TIMES_NEW_ROMAN
        if i == n_params - 1:
            for j in range(4):
                _tbl_border(param_tbl.rows[i + 1].cells[j]._tc, ['bottom'], style='double')

    for row in param_tbl.rows:
        for cell in row.cells:
            for pp in cell.paragraphs:
                pPr = pp._element.get_or_add_pPr()
                sp = OxmlElement('w:spacing')
                sp.set(qn('w:before'), '10')
                sp.set(qn('w:after'), '10')
                pPr.append(sp)

    param_tbl_el = param_tbl._tbl
    body.remove(param_tbl_el)
    tp1_el.addnext(param_tbl_el)

    # Table 2 notes
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(6)
    note.paragraph_format.first_line_indent = Inches(0)
    note.paragraph_format.line_spacing = 1.0
    rn1 = note.add_run('Notes: ')
    rn1.bold = True
    rn1.font.size = Pt(10)
    rn1 = note.add_run(
        'Hardware cost \u03C1 = P(GPU) / (L \u00b7 H \u00b7 \u03B2). '
        'PUE(\u03B8) = \u03C6 + \u03B4 \u00b7 max(0, \u03B8 \u2212 \u03B8\u0304). '
        'RTT = round-trip time, the network delay for a data packet to travel from '
        'client to server and back, measured in milliseconds. '
        'The production-efficiency index \u03BE\u2C7C\u1d49\u1da0\u1da0 is a weighted geometric mean of '
        'governance quality and grid reliability (\u03C9 = 0.50, \u03BE_floor = 0.30; equation 3). '
        'The efficiency adjustment applies to non-hardware costs: '
        'c_adj = \u03C1 + (c_cr \u2212 \u03C1) / \u03BE. '
        'Sanctions exposure is captured by the bilateral '
        'sovereignty premium \u03BB\u1d62\u2c7c (equation 2).'
    )
    rn1.font.size = Pt(10)
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_el = note._element
    body.remove(note_el)
    param_tbl_el.addnext(note_el)

    return note_el


def write_table3(doc, body, after_el, demand_data):
    """Table 3a: Cost Specifications + Table 3b: Sovereignty Specifications (landscape)."""
    print("Inserting Table 3a/3b (Country rankings, landscape)...")

    table3_data = demand_data["table3"]

    # ─── Attach portrait sectPr to previous paragraph (no empty page) ───
    sect_port = OxmlElement('w:sectPr')
    pg_sz_p = OxmlElement('w:pgSz')
    pg_sz_p.set(qn('w:w'), '12240')
    pg_sz_p.set(qn('w:h'), '15840')
    sect_port.append(pg_sz_p)
    pg_mar_p = OxmlElement('w:pgMar')
    pg_mar_p.set(qn('w:top'), '1440')
    pg_mar_p.set(qn('w:right'), '1440')
    pg_mar_p.set(qn('w:bottom'), '1440')
    pg_mar_p.set(qn('w:left'), '1440')
    pg_mar_p.set(qn('w:header'), '720')
    pg_mar_p.set(qn('w:footer'), '720')
    sect_port.append(pg_mar_p)
    prev_pPr = after_el.find(f'{{{W_NS}}}pPr')
    if prev_pPr is None:
        prev_pPr = etree.SubElement(after_el, f'{{{W_NS}}}pPr')
    prev_pPr.append(sect_port)

    # ─── Table 3 title with bookmark ───
    tp3 = doc.add_paragraph()
    tp3.paragraph_format.space_before = Pt(6)
    tp3.paragraph_format.space_after = Pt(3)
    tp3.paragraph_format.first_line_indent = Inches(0)
    tp3._element.append(make_bookmark(111, 'Table3'))
    hl_t3 = OxmlElement('w:hyperlink')
    hl_t3.set(qn('w:anchor'), 'Table3txt')
    hl_t3.set(qn('w:history'), '1')
    r_t3 = OxmlElement('w:r')
    rPr_t3 = OxmlElement('w:rPr')
    b_t3 = OxmlElement('w:b')
    rPr_t3.append(b_t3)
    sz_t3 = OxmlElement('w:sz')
    sz_t3.set(qn('w:val'), '20')
    rPr_t3.append(sz_t3)
    clr_t3 = OxmlElement('w:color')
    clr_t3.set(qn('w:val'), LINK_COLOR)
    uu_t3 = OxmlElement('w:u')
    uu_t3.set(qn('w:val'), 'single')
    rPr_t3.append(clr_t3)
    rPr_t3.append(uu_t3)
    r_t3.append(rPr_t3)
    t_t3 = OxmlElement('w:t')
    t_t3.text = 'Table 3a'
    r_t3.append(t_t3)
    hl_t3.append(r_t3)
    tp3._element.append(hl_t3)
    tp3._element.append(make_bookmark_end(111))
    run_tt3 = tp3.add_run('. Country rankings under alternative cost specifications')
    run_tt3.bold = True
    run_tt3.font.size = Pt(10)
    tp3_el = tp3._element
    body.remove(tp3_el)
    after_el.addnext(tp3_el)

    # ─── Short name lookup ───
    _short = {
        "United Arab Emirates": "UAE",
        "United Kingdom": "UK",
        "United States": "USA",
        "United States of America": "USA",
        "Bosnia and Herzegovina": "Bosnia & Herz.",
        "North Macedonia": "N. Macedonia",
        "Czech Republic": "Czechia",
    }

    def _sname(full):
        return _short.get(full, full[:18] + '.' if len(full) > 19 else full)

    # ─── Row selection: 25 curated countries sorted by rank_eff ───
    _show = [
        'CAN', 'FIN', 'NOR', 'CHN', 'KGZ',   # top 5 preferred spec (Form B)
        'SWE', 'XKX', 'MNE', 'USA', 'ETH',   # ranks 6-10
        'ISL', 'NZL', 'AUS', 'VNM', 'GBR',   # ranks 11-15
        'ARG', 'IND', 'FRA', 'COL', 'PRT',    # ranks 16-20
        'ARE', 'LVA', 'BEL', 'MLT', 'GEO',   # ranks 21-25
    ]
    _show_set = set(_show)
    top_eff = sorted(table3_data, key=lambda x: x["rank_eff"])
    top_rows = [d for d in top_eff if d["iso"] in _show_set]

    # ─── Build table ───
    n_data = len(top_rows)
    n_rows = 2 + n_data  # 2 header rows + data
    n_cols = 15  # added ξ sub-column under (3)
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = 'Table Grid'
    _tbl_clear_borders(tbl)

    _s = partial(_tbl_set, tbl, font_size=8)  # default fs=8 for Table 3a

    # ─── Row 0: Group headers ───
    _s(0, 0, '')
    _tbl_merge(tbl, 0, 1, 3)
    _s(0, 1, '(1) Raw Electricity', bold=True)
    _tbl_merge(tbl, 0, 4, 6)
    _s(0, 4, '(2) Cost-Recovery', bold=True)
    _tbl_merge(tbl, 0, 7, 10)
    _s(0, 7, '(3) Efficiency-Adjusted', bold=True)
    _tbl_merge(tbl, 0, 11, 13)
    _s(0, 11, '(4) Bilateral \u03bb\u1d62\u2c7c', bold=True)
    _s(0, 14, '')

    # Top + bottom border on row 0
    for j in range(n_cols):
        _tbl_border(tbl.cell(0, j)._tc, ['top', 'bottom'])

    # ─── Row 1: Sub-headers ───
    sub_headers = ['Country',
                   'c\u2c7c', 'Rank', 'Type',
                   'c\u2c7c', 'Rank', 'Type',
                   'c\u2c7c\u1d43\u1d48\u02b2', '\u03be\u2c7c\u1d49\u1da0\u1da0', 'Rank', 'Type',
                   'c\u2c7c\u1d43\u1d48\u02b2', 'Rank', 'Type',
                   '\u0394']
    for j, hdr in enumerate(sub_headers):
        _s(1, j, hdr, bold=True, align='left' if j == 0 else 'center')
        _tbl_border(tbl.cell(1, j)._tc, ['top', 'bottom'])

    # ─── Data rows ───
    row_idx = 2
    all_data_rows = list(top_rows)

    for d in all_data_rows:
        _s(row_idx, 0, _sname(d["country"]), align='left')
        _s(row_idx, 1, f'${d["cj_raw"]:.2f}')
        _s(row_idx, 2, str(d["rank_raw"]))
        _s(row_idx, 3, d["type_raw"])
        _s(row_idx, 4, f'${d["cj_cr"]:.2f}')
        _s(row_idx, 5, str(d["rank_cr"]))
        _s(row_idx, 6, d["type_cr"])
        _s(row_idx, 7, f'${d["cj_eff"]:.2f}')
        _s(row_idx, 8, f'{rhup(d["xi"]):.2f}')
        _s(row_idx, 9, str(d["rank_eff"]))
        _s(row_idx, 10, d["type_eff"])
        _s(row_idx, 11, f'${d["cj_eff"]:.2f}')
        _s(row_idx, 12, str(d["rank_eff"]), bold=True)
        _s(row_idx, 13, d.get("type_bilat", d.get("type_sov", "II")))
        delta_val = d["delta"]
        delta_str = f'+{delta_val}' if delta_val > 0 else str(delta_val)
        _s(row_idx, 14, delta_str)
        row_idx += 1

    # Double bottom border on last data row
    last_data_row = row_idx - 1
    for j in range(n_cols):
        _tbl_border(tbl.cell(last_data_row, j)._tc, ['bottom'], style='double')

    # Column widths and spacing
    _tbl_col_widths(tbl, [
        1800,                # Country
        900, 540, 540,       # (1)
        900, 540, 540,       # (2)
        900, 450, 540, 540,  # (3): cⱼ/ξ, ξ, Rank, Type
        900, 540, 540,       # (4)
        540,                 # Delta
    ])
    _tbl_cell_spacing(tbl)

    # Position table after title
    tbl_el = tbl._tbl
    body.remove(tbl_el)
    tp3_el.addnext(tbl_el)

    # ─── Table notes ───
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(2)
    note.paragraph_format.space_after = Pt(6)
    note.paragraph_format.first_line_indent = Inches(0)
    note.paragraph_format.line_spacing = 1.0
    rn3 = note.add_run(
        'Notes: '
    )
    rn3.bold = True
    rn3.font.size = Pt(10)
    rn3.font.name = 'Times New Roman'
    rn3 = note.add_run(
        'Columns show unit cost ($/GPU-hr), rank among 85 countries, and regime type '
        '(EE\u2009=\u2009training + inference exporter; IE\u2009=\u2009inference exporter; '
        'DD\u2009=\u2009domestic producer; II\u2009=\u2009full importer). '
        '(1)\u2009Raw: observed electricity tariffs. '
        '(2)\u2009Cost-recovery: subsidized tariffs replaced with LRMC. '
        '(3)\u2009Efficiency-adjusted: \u03C1 + (c_cr \u2212 \u03C1) / \u03be\u2c7c\u1d49\u1da0\u1da0, '
        'where \u03C1 = $1.36/hr is globally-priced hardware; the \u03be column reports each '
        'country\u2019s efficiency index (preferred specification). '
        '(4)\u2009Bilateral: efficiency-adjusted cost with bilateral sovereignty premium '
        '\u03bb\u1d62\u2c7c from equation (2). '
        'Ranks in columns (3) and (4) are the same, because the sovereignty premium '
        'does not affect production costs or cost rankings; it reshapes regime assignments. '
        '\u0394\u2009=\u2009rank change from (1) to (3); positive values indicate improvement. '
        '25 selected countries; see '
    )
    rn3.font.size = Pt(10)
    rn3.font.name = 'Times New Roman'
    note._element.append(make_hyperlink('TableA2', 'Table A2', rPr_orig=_rPr_pt(10)))
    rn3 = note.add_run(' for all 85 countries.')
    rn3.font.size = Pt(10)
    rn3.font.name = 'Times New Roman'
    note.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note_el = note._element
    body.remove(note_el)
    tbl_el.addnext(note_el)

    # ═══════════════════════════════════════════════════════════════════════
    # TABLE 3b: SOVEREIGNTY SPECIFICATIONS
    # ═══════════════════════════════════════════════════════════════════════
    print("  Inserting Table 3b (Sovereignty specifications)...")

    # Table 3b title (new page)
    tp3b = doc.add_paragraph()
    tp3b.paragraph_format.page_break_before = True
    tp3b.paragraph_format.space_before = Pt(12)
    tp3b.paragraph_format.space_after = Pt(3)
    tp3b.paragraph_format.first_line_indent = Inches(0)
    tp3b._element.append(make_bookmark(114, 'Table3b'))
    hl_t3b = OxmlElement('w:hyperlink')
    hl_t3b.set(qn('w:anchor'), 'Table3btxt')
    hl_t3b.set(qn('w:history'), '1')
    r_t3b = OxmlElement('w:r')
    rPr_t3b = OxmlElement('w:rPr')
    b_t3b = OxmlElement('w:b')
    rPr_t3b.append(b_t3b)
    sz_t3b = OxmlElement('w:sz')
    sz_t3b.set(qn('w:val'), '20')
    rPr_t3b.append(sz_t3b)
    clr_t3b = OxmlElement('w:color')
    clr_t3b.set(qn('w:val'), LINK_COLOR)
    uu_t3b = OxmlElement('w:u')
    uu_t3b.set(qn('w:val'), 'single')
    rPr_t3b.append(clr_t3b)
    rPr_t3b.append(uu_t3b)
    r_t3b.append(rPr_t3b)
    t_t3b = OxmlElement('w:t')
    t_t3b.text = 'Table 3b'
    r_t3b.append(t_t3b)
    hl_t3b.append(r_t3b)
    tp3b._element.append(hl_t3b)
    r_3b_sub = tp3b.add_run('. Country rankings under alternative sovereignty specifications')
    r_3b_sub.bold = True
    r_3b_sub.font.size = Pt(10)
    tp3b._element.append(make_bookmark_end(114))
    tp3b_el = tp3b._element
    body.remove(tp3b_el)
    note_el.addnext(tp3b_el)

    # Table 3b: 7 columns — Country, c/ξ^eff, (4)/(5) Bilateral, (6) Uniform, Rank_6, (7) FDI, λ_k^*
    n_rows_3b = 2 + n_data
    n_cols_3b = 7
    tbl3b = doc.add_table(rows=n_rows_3b, cols=n_cols_3b)
    tbl3b.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3b.style = 'Table Grid'
    _tbl_clear_borders(tbl3b)

    _s3b = partial(_tbl_set, tbl3b, font_size=8)  # default fs=8 for Table 3b

    # Row 0: Group headers (merged)
    _s3b(0, 0, '')
    _tbl_merge(tbl3b, 0, 1, 2)
    _s3b(0, 1, '(4)/(5) Bilateral', bold=True)
    _tbl_merge(tbl3b, 0, 3, 4)
    _s3b(0, 3, '(6) Uniform (\u03bb = 10%)', bold=True)
    _s3b(0, 5, '(7) FDI', bold=True)
    _s3b(0, 6, '')
    for j in range(n_cols_3b):
        _tbl_border(tbl3b.cell(0, j)._tc, ['top', 'bottom'])

    # Row 1: Sub-headers
    sub_hdr_3b = ['Country', 'c\u2c7c\u1d43\u1d48\u02b2',
                  'Type', 'Type', 'Rank\u2086', 'Type', '\u03bb\u2096*']
    for j, hdr in enumerate(sub_hdr_3b):
        _s3b(1, j, hdr, bold=True, align='left' if j == 0 else 'center')
        _tbl_border(tbl3b.cell(1, j)._tc, ['top', 'bottom'])

    # Data rows
    row_3b = 2
    for d in all_data_rows:
        _s3b(row_3b, 0, _sname(d["country"]), align='left')
        _s3b(row_3b, 1, f'${d["cj_eff"]:.2f}')
        _s3b(row_3b, 2, d.get("type_bilat", "II"))
        _s3b(row_3b, 3, d.get("type_uniform", "II"))
        _s3b(row_3b, 4, str(d.get("rank_sov", d["rank_eff"])))
        _s3b(row_3b, 5, d.get("type_fdi", "II"))
        lks = d.get("lam_k_star", 0)
        lks_str = f'{lks * 100:.1f}%' if lks >= 0 else f'\u2212{abs(lks) * 100:.1f}%'
        _s3b(row_3b, 6, lks_str)
        row_3b += 1

    # Bottom border
    for j in range(n_cols_3b):
        _tbl_border(tbl3b.cell(row_3b - 1, j)._tc, ['bottom'], style='double')

    # Column widths and spacing
    _tbl_col_widths(tbl3b, [2000, 1100, 900, 900, 800, 900, 1100])
    _tbl_cell_spacing(tbl3b)

    tbl3b_el = tbl3b._tbl
    body.remove(tbl3b_el)
    tp3b_el.addnext(tbl3b_el)

    # Table 3b notes
    note3b = doc.add_paragraph()
    note3b.paragraph_format.space_before = Pt(2)
    note3b.paragraph_format.space_after = Pt(6)
    note3b.paragraph_format.first_line_indent = Inches(0)
    note3b.paragraph_format.line_spacing = 1.0
    rn3b = note3b.add_run('Notes: ')
    rn3b.bold = True
    rn3b.font.size = Pt(10)
    rn3b.font.name = 'Times New Roman'
    rn3b = note3b.add_run(
        'Type codes: EE\u2009=\u2009training + inference exporter; '
        'IE\u2009=\u2009inference exporter; DD\u2009=\u2009domestic producer; '
        'II\u2009=\u2009full importer. '
        '(4)/(5)\u2009Bilateral: bilateral sovereignty premium \u03bb\u1d62\u2c7c '
        'from equation (2), with geopolitical, regulatory, and sanctions components. '
        'Under the calibrated parameters, demand tiering (sovereign 10%, regulated 20%, '
        'commercial 70%) leaves regime type assignments unchanged for all countries; '
        'tiering affects within-country demand allocation across tiers but not the '
        'equilibrium set of exporters. '
        '(6)\u2009Uniform: uniform \u03bb\u2009=\u200910% premium (robustness check). '
        '(7)\u2009Hyperscaler FDI: bilateral premium replaced by \u03bb\u1da0\u1d48\u1d49 '
        'from equation (2\u2032), where h is the hyperscaler\u2019s home country (assumed US). '
        'Trust premium reflects buyer\u2013operator relationship. Sanctions on host still apply. '
        '* = sanctioned/GPU-blocked. \u2020 = developing-country exporter. '
        '\u03bb\u2096*\u2009=\u2009c\u2096/p\u209c\u2009\u2212\u20091 is the minimum '
        'bilateral sovereignty premium at which country k switches from importing to '
        'domestic training production (Proposition 3); negative values indicate exporters.'
    )
    rn3b.font.size = Pt(10)
    rn3b.font.name = 'Times New Roman'
    note3b.alignment = WD_ALIGN_PARAGRAPH.LEFT
    note3b_el = note3b._element
    body.remove(note3b_el)
    tbl3b_el.addnext(note3b_el)

    # ─── Attach landscape sectPr to Table 3b notes paragraph ───
    sect_pr = OxmlElement('w:sectPr')
    pg_sz = OxmlElement('w:pgSz')
    pg_sz.set(qn('w:w'), '15840')
    pg_sz.set(qn('w:h'), '12240')
    pg_sz.set(qn('w:orient'), 'landscape')
    sect_pr.append(pg_sz)
    pg_mar = OxmlElement('w:pgMar')
    pg_mar.set(qn('w:top'), '1080')
    pg_mar.set(qn('w:right'), '1080')
    pg_mar.set(qn('w:bottom'), '1080')
    pg_mar.set(qn('w:left'), '1080')
    pg_mar.set(qn('w:header'), '720')
    pg_mar.set(qn('w:footer'), '720')
    sect_pr.append(pg_mar)
    note3b_pPr = note3b_el.find(f'{{{W_NS}}}pPr')
    if note3b_pPr is None:
        note3b_pPr = etree.SubElement(note3b_el, f'{{{W_NS}}}pPr')
    note3b_pPr.append(sect_pr)

    return note3b_el


def write_references(doc, body, refs):
    print("Updating references...")

    # Page break before References heading
    add_page_break(doc, body, refs.getprevious())

    all_now = list(body)
    ri = all_now.index(refs)
    ref_els = []
    ref_txts = []
    for i in range(ri + 1, len(all_now)):
        el = all_now[i]
        if el.tag == qn('w:p'):
            t = "".join(r.text or "" for r in el.findall(f'.//{qn("w:t")}'))
            if t.strip():
                ref_txts.append(t.strip())
                ref_els.append(el)
        elif el.tag == qn('w:sectPr'):
            break

    new_refs = [
        'Asian Development Bank. (2020). \u201CToktogul Rehabilitation Project Phase 3: '
        'Sector Assessment.\u201D Asian Development Bank. Report No. 49013-002.',

        'Barroso, L., U. H\u00F6lzle, and P. Ranganathan. (2018). '
        'The Datacenter as a Computer: Designing Warehouse-Scale Machines, '
        '3rd ed. San Rafael, CA: Morgan & Claypool.',

        'Brainard, S. (1997). \u201CAn Empirical Assessment of the Proximity-Concentration '
        'Trade-off.\u201D American Economic Review, 87(4): 520\u2013544.',

        'Calcaterra, M., L. Reis, P. Fragkos, T. Briera, H. Boer, F. Egli, '
        'J. Emmerling, G. Iyer, S. Mittal, F. Polzin, M. Sanders, T. Schmidt, '
        'A. Serebriakova, and B. Steffen. (2024). \u201CReducing the Cost of Capital '
        'to Finance the Energy Transition in Developing Countries.\u201D '
        'Nature Energy, 9(10): 1241\u20131251.',

        'Cloudscene. (2025). Global Data Center Directory. cloudscene.com.',

        'Deloitte. (2025). \u201CTechnology, Media, and Telecommunications Predictions 2026.\u201D '
        'Deloitte Insights.',

        'Deloitte and Google. (2020). \u201CMilliseconds Make Millions.\u201D '
        'Deloitte Digital and Google.',

        'EIA. (2025). Electric Power Monthly. U.S. Energy Information Administration.',

        'Epoch AI. (2024). \u201CThe Training Compute of Notable AI Models.\u201D epochai.org.',

        'Eurostat. (2025). Electricity Prices for Non-Household Consumers '
        '(nrg_pc_205). Luxembourg: Eurostat.',

        'Firebird. (2026). \u201CPhase 2 of Armenia AI Megaproject, Scaling to $4 Billion '
        'and 50,000 GPUs.\u201D Press release, January 2026.',

        'Flucker, S., R. Tozer, and R. Whitehead. (2013). \u201CData Centre Energy '
        'Efficiency Analysis to Minimize Total Cost of Ownership.\u201D Building Services '
        'Engineering Research and Technology, 34(1): 103\u2013117.',

        'GlobalPetrolPrices. (2025). Electricity Prices Around the World. '
        'globalpetrolprices.com.',

        'Goldfarb, A., and D. Trefler. (2018). \u201CAI and International Trade.\u201D '
        'In The Economics of Artificial Intelligence. Chicago: Univ. of Chicago Press, '
        'pp. 463\u2013492.',

        'Google. (2024). 2024 Environmental Report. sustainability.google.',

        'Hausmann, R., J. Hwang, and D. Rodrik. (2007). \u201CWhat You Export Matters.\u201D '
        'Journal of Economic Growth, 12(1): 1\u201325.',

        'Helpman, E., M. Melitz, and S. Yeaple. (2004). \u201CExport Versus FDI with '
        'Heterogeneous Firms.\u201D American Economic Review, 94(1): 300\u2013316.',

        'Hersbach, H., et al. (2020). \u201CThe ERA5 Global Reanalysis.\u201D '
        'Quarterly Journal of the Royal Meteorological Society, 146(730): 1999\u20132049.',


        'IEA. (2025). \u201CEnergy Demand from AI.\u201D Published online at iea.org.',

        'IMF. (2025). \u201CFossil Fuel Subsidies Data: 2025 Update.\u201D '
        'IMF Working Paper WP/25/270.',

        'Korinek, A., and J. Stiglitz. (2021). \u201CAI, Globalization, and Strategies for '
        'Economic Development.\u201D NBER Working Paper No. 28453.',

        'Lazard. (2025). Lazard\u2019s Levelized Cost of Energy Analysis, Version 17.0. '
        'lazard.com.',

        'Krugman, P. (1991). \u201CIncreasing Returns and Economic Geography.\u201D '
        'Journal of Political Economy, 99(3): 483\u2013499.',

        'Lim\u00E3o, N., and A. Venables. (2001). \u201CInfrastructure, Geographical '
        'Disadvantage, Transport Costs, and Trade.\u201D '
        'World Bank Economic Review, 15(3): 451\u2013479.',

        'Liu, Z., A. Wierman, Y. Chen, B. Raber, and J. Moriarty. (2023). '
        '\u201CSustainability of Data Center Digital Twins.\u201D '
        'Proceedings of ACM e-Energy, pp. 178\u2013189.',

        'NVIDIA. (2024). NVIDIA H100 Tensor Core GPU Datasheet. nvidia.com.',

        'Eaton, J., and S. Kortum. (2002). \u201CTechnology, Geography, and Trade.\u201D '
        'Econometrica, 70(5): 1741\u20131779.',

        'Turner & Townsend. (2025). Data Centre Construction Cost Index 2025. '
        'turnerandtownsend.com.',

        'Turner Lee, N., and D. West. (2025). \u201CThe Future of Data Centers.\u201D '
        'Brookings Institution, November 2025.',

        'UNCTAD. (2025). Technology and Innovation Report 2025. Geneva: United Nations.',


        'Uptime Institute. (2024). Global Data Center Survey Results 2024. uptimeinstitute.com.',

        'WonderNetwork. (2024). Global Ping Statistics. wondernetwork.com.',

        'World Bank. (2024). World Development Indicators. Washington, DC.',

        'Lehdonvirta, V., B. Wu, and Z. Hawkins. (2024). \u201CCompute North vs. Compute South: '
        'The Uneven Possibilities of Compute-Based AI Governance Around the Globe.\u201D '
        'Proceedings of the AAAI/ACM Conference on AI, Ethics, and Society, 7(1): 828\u2013838.',

        'Pilz, K., Y. Mahmood, and L. Heim. (2025). AI\u2019s Power Requirements Under '
        'Exponential Growth. Santa Monica, CA: RAND Corporation, RR-A3572-1.',

        'Sastry, G., L. Heim, et al. (2024). \u201CComputing Power and the Governance of '
        'Artificial Intelligence.\u201D arXiv:2402.08797.',


        'Arkolakis, C., A. Costinot, and A. Rodr\u00EDguez-Clare. (2012). \u201CNew Trade '
        'Models, Same Old Gains?\u201D American Economic Review, 102(1): 94\u2013130.',

        'Bailey, M., A. Strezhnev, and E. Voeten. (2017). \u201CEstimating Dynamic '
        'State Preferences from United Nations Voting Data.\u201D '
        'Journal of Conflict Resolution, 61(2): 430\u2013456.',

        'van der Ploeg, F. (2011). \u201CNatural Resources: Curse or Blessing?\u201D '
        'Journal of Economic Literature, 49(2): 366\u2013420.',

        'Ohlin, B. (1933). Interregional and International Trade. '
        'Cambridge, MA: Harvard University Press.',

        'Biglaiser, G., J. Cr\u00E9mer, and A. Mantovani. (2024). \u201CThe Economics of the Cloud.\u201D '
        'Toulouse School of Economics Working Paper No. 24-1520.',

        'Blinder, A. (2006). \u201COffshoring: The Next Industrial Revolution?\u201D '
        'Foreign Affairs, 85(2): 113\u2013128.',

        'Stojkoski, V., P. Koch, E. Coll, and C. A. Hidalgo. (2024). '
        '\u201CEstimating Digital Product Trade through Corporate Revenue Data.\u201D '
        'Nature Communications, 15: 5262.',

        'World Bank. (2025). Digital Progress and Trends Report 2025: '
        'Strengthening AI Foundations. Washington, DC: World Bank.',

        'World Bank. (2025). Enterprise Surveys. '
        'Washington, DC: World Bank. enterprisesurveys.org.',

        'Farole, T. (2011). Special Economic Zones in Africa: '
        'Comparing Performance and Learning from Global Experiences. '
        'Washington, DC: World Bank.',

        'Frick, S., A. Rodr\u00EDguez-Pose, and M. Wong. (2019). \u201CToward Economically '
        'Dynamic Special Economic Zones in Emerging Countries.\u201D '
        'Economic Geography, 95(1): 30\u201364.',

        'World Bank. (2017). Special Economic Zones: An Operational Review '
        'of Their Impacts. Washington, DC: World Bank.',
    ]

    ref_txts = sorted(new_refs, key=lambda x: x.lower())
    for el in ref_els:
        body.remove(el)

    # Build reverse map: reference text prefix -> key
    def find_ref_key(ref_text):
        """Find the citation key for a reference text."""
        for key, prefix in REF_KEY_MAP.items():
            if ref_text.startswith(prefix):
                return key
        return None

    bm_id_refs = [500]  # bookmark IDs for references
    cur = refs
    for rt in ref_txts:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = Pt(12)
        italic_portion = find_italic_portion(rt)
        key = find_ref_key(rt)
        if key:
            # Add bookmark target for in-text citation links
            p._element.append(make_bookmark(bm_id_refs[0], key))
            # Author portion as hyperlink back to in-text citation
            split_pos = rt.find('\u201C')
            if split_pos < 0:
                split_pos = rt.find('(', 3)
            if split_pos < 0:
                split_pos = len(rt)
            author_part = rt[:split_pos]
            remaining = rt[split_pos:]
            p._element.append(make_hyperlink(f'{key}txt', author_part))
            if remaining:
                _write_ref_segments(p, remaining, italic_portion)
            p._element.append(make_bookmark_end(bm_id_refs[0]))
            bm_id_refs[0] += 1
        else:
            _write_ref_segments(p, rt, italic_portion)
        el = p._element
        body.remove(el)
        cur.addnext(el)
        cur = el
    print(f"  {len(ref_txts)} references")
    return cur  # last reference element


def link_citations(body):
    print("Linking citations...")
    bm_id_cite = [200]
    passes = 0
    while True:
        n = link_citations_pass(body, CITE_MAP, bm_id_cite)
        passes += 1
        if n == 0 or passes > 10:
            break
    print(f"  {bm_id_cite[0] - 200} citation links created in {passes} passes")


def link_equations(body):
    """Link 'equation (N)' mentions in text to their display equation bookmarks."""
    print("Linking equation references...")
    import re
    count = 0
    bm_id_eq = [900]
    eq_pattern = re.compile(r'(?:equation|eq\.) \((\d+)\)')
    for p_el in list(body.findall(qn('w:p'))):
        for child in list(p_el):
            if child.tag != qn('w:r'):
                continue
            t_el = child.find(qn('w:t'))
            if t_el is None or not t_el.text:
                continue
            text = t_el.text
            m = eq_pattern.search(text)
            if not m:
                continue
            eq_num = m.group(1)
            anchor = f'Eq{eq_num}'
            match_text = m.group(0)
            idx = m.start()
            before = text[:idx]
            after = text[idx + len(match_text):]
            rPr_orig = child.find(qn('w:rPr'))
            t_el.text = before
            t_el.set(XML_SPACE, SPACE_PRESERVE)
            ins = child
            bm_start = make_bookmark(bm_id_eq[0], f'{anchor}txt{bm_id_eq[0]}')
            ins.addnext(bm_start)
            ins = bm_start
            hyperlink = make_hyperlink(anchor, match_text, rPr_orig)
            ins.addnext(hyperlink)
            ins = hyperlink
            bm_end = make_bookmark_end(bm_id_eq[0])
            ins.addnext(bm_end)
            ins = bm_end
            bm_id_eq[0] += 1
            if after:
                ra = OxmlElement('w:r')
                if rPr_orig is not None:
                    ra.append(copy.deepcopy(rPr_orig))
                ta = OxmlElement('w:t')
                ta.set(XML_SPACE, SPACE_PRESERVE)
                ta.text = after
                ra.append(ta)
                ins.addnext(ra)
            count += 1
    print(f"  {count} equation links created")


def fix_orphan_backlinks(body, refs):
    """Remove hyperlink wrappers in references whose back-link targets don't exist in the body."""
    # Collect all bookmark names in the document
    all_bookmarks = set()
    for el in body:
        for bm in el.findall(f'.//{{{W_NS}}}bookmarkStart'):
            name = bm.get(f'{{{W_NS}}}name', '')
            if name:
                all_bookmarks.add(name)

    # Scan reference paragraphs for hyperlinks with missing targets
    refs_idx = list(body).index(refs)
    fixed = 0
    for el in list(body)[refs_idx + 1:]:
        if el.tag == f'{{{W_NS}}}sectPr':
            break
        # Stop at headings (e.g. Appendix) that follow references
        if el.tag == qn('w:p'):
            pPr = el.find(qn('w:pPr'))
            if pPr is not None:
                pS = pPr.find(qn('w:pStyle'))
                if pS is not None and 'Heading' in pS.get(qn('w:val'), ''):
                    break
                if pPr.find(f'{{{W_NS}}}sectPr') is not None:
                    break
        for hl in el.findall(f'.//{{{W_NS}}}hyperlink'):
            anchor = hl.get(f'{{{W_NS}}}anchor', '')
            if anchor and anchor not in all_bookmarks:
                # Replace hyperlink element with its child runs (keep text, drop link)
                parent = hl.getparent()
                idx = list(parent).index(hl)
                children = list(hl)
                for child in children:
                    hl.remove(child)
                    parent.insert(idx, child)
                    idx += 1
                parent.remove(hl)
                fixed += 1
    if fixed:
        print(f"  Fixed {fixed} orphan back-link(s) in references")


def apply_formatting(doc, body, refs, title_el, author_el, ver_el, abs_text_el):
    print("Applying formatting...")
    # Set Normal style defaults
    normal = doc.styles['Normal']
    normal.font.name = TIMES_NEW_ROMAN
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.5

    # Fix Heading 1 style: remove theme font, set blue color at style level
    h1 = doc.styles['Heading 1']
    h1.font.name = TIMES_NEW_ROMAN
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = HEADING_BLUE
    # Remove theme font attributes that override explicit font names
    h1_rPr = h1._element.find(qn('w:rPr'))
    if h1_rPr is not None:
        h1_fonts = h1_rPr.find(qn('w:rFonts'))
        if h1_fonts is not None:
            for attr in ['asciiTheme', 'hAnsiTheme', 'eastAsiaTheme', 'cstheme']:
                h1_fonts.attrib.pop(qn(f'w:{attr}'), None)

    # Identify reference paragraphs to protect their spacing
    refs_idx = list(body).index(refs)
    ref_elements = set()
    for el in list(body)[refs_idx + 1:]:
        if el.tag == qn('w:sectPr'):
            break
        if el.tag == qn('w:p'):
            pPr = el.find(qn('w:pPr'))
            if pPr is not None:
                pS = pPr.find(qn('w:pStyle'))
                if pS is not None and 'Heading' in pS.get(qn('w:val'), ''):
                    break
                if pPr.find(f'{{{W_NS}}}sectPr') is not None:
                    break
            ref_elements.add(el)

    # Paragraphs to protect from global formatting (centered title page elements)
    _protected = {title_el, author_el, ver_el, abs_text_el}

    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        # Heading 1: Times New Roman, blue, 14pt, bold
        if style == 'Heading 1':
            for run in p.runs:
                run.font.color.rgb = HEADING_BLUE
                run.font.name = TIMES_NEW_ROMAN
                run.font.size = Pt(14)
                run.bold = True
            continue
        # Heading 2: Times New Roman, blue, 12pt, italic, no bold
        if style == 'Heading 2':
            for run in p.runs:
                run.font.color.rgb = HEADING_BLUE
                run.font.name = TIMES_NEW_ROMAN
                run.font.size = Pt(12)
                run.italic = True
                run.bold = False
            continue
        if 'Heading' not in style and p.text.strip():
            # Skip title page elements (centered)
            if p._element not in _protected and p.alignment != WD_ALIGN_PARAGRAPH.LEFT:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                if p.paragraph_format.first_line_indent is None or p.paragraph_format.first_line_indent > 0:
                    p.paragraph_format.first_line_indent = Inches(0)
            # Subtitle runs: italic (non-bold) first run ending with "." → font 12, TNR, not bold
            runs = [r for r in p.runs if r.text.strip()]
            if runs and runs[0].italic and not runs[0].bold and runs[0].text.rstrip().endswith('.'):
                runs[0].font.size = Pt(12)
                runs[0].font.name = TIMES_NEW_ROMAN
                runs[0].bold = False
            # Preserve reference formatting (hanging indent + 4pt spacing)
            if p._element in ref_elements:
                continue
            # Preserve title page spacing
            if p._element in _protected:
                continue
            p.paragraph_format.space_before = Pt(0)
            # Preserve Pt(2) spacing on paragraphs immediately before equations
            if p.paragraph_format.space_after is None or p.paragraph_format.space_after >= Pt(8):
                p.paragraph_format.space_after = Pt(8)


def add_page_numbers_and_break(doc, body, kw_el):
    print("Adding page numbers...")
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.different_first_page_header_footer = True

    # Default footer: right-aligned page number
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.clear()
    # Insert PAGE field: w:fldSimple or fldChar sequence
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), ' PAGE ')
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = '1'
    r.append(t)
    fld.append(r)
    fp._element.append(fld)

    # First page footer: empty (no page number on title page)
    first_footer = section.first_page_footer
    first_footer.is_linked_to_previous = False
    if first_footer.paragraphs:
        first_footer.paragraphs[0].clear()

    # ═══════════════════════════════════════════════════════════════════════
    # STEP 17: PAGE BREAK AFTER ABSTRACT (title on separate page)
    # ═══════════════════════════════════════════════════════════════════════
    print("Adding page break after keywords...")
    add_page_break(doc, body, kw_el)


def _add_word_comments(docx_path, comments):
    """Add Word comments to a saved docx file (post-processing).

    comments: list of (id, author, text, anchor_text) tuples.
    anchor_text is a substring to search for in paragraph text; the comment
    is anchored to the first paragraph containing that substring.
    """
    import zipfile
    from io import BytesIO

    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    R_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
    CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
    COMMENT_REL = ('http://schemas.openxmlformats.org/officeDocument/2006/'
                   'relationships/comments')
    NSMAP = {'w': W, 'r': R_NS}

    # Read all files from the docx
    files = {}
    with zipfile.ZipFile(docx_path, 'r') as zin:
        for name in zin.namelist():
            files[name] = zin.read(name)

    # Parse document.xml
    doc_xml = etree.fromstring(files['word/document.xml'])
    doc_body = doc_xml.find(f'{{{W}}}body')

    # Build comments.xml
    comments_root = etree.Element(f'{{{W}}}comments', nsmap=NSMAP)
    now_str = datetime.now().strftime('%Y-%m-%dT%H:%M:%S') + 'Z'

    placed = 0
    for cid, author, text, anchor_text in comments:
        # Create comment element
        c_el = etree.SubElement(comments_root, f'{{{W}}}comment')
        c_el.set(f'{{{W}}}id', str(cid))
        c_el.set(f'{{{W}}}author', author)
        c_el.set(f'{{{W}}}date', now_str)
        c_el.set(f'{{{W}}}initials', author[0])
        cp = etree.SubElement(c_el, f'{{{W}}}p')
        cr = etree.SubElement(cp, f'{{{W}}}r')
        ct = etree.SubElement(cr, f'{{{W}}}t')
        ct.text = text

        # Find anchor paragraph in document body
        for p_el in doc_body.iter(f'{{{W}}}p'):
            p_text = ''.join(t.text or '' for t in p_el.iter(f'{{{W}}}t'))
            if anchor_text in p_text:
                # Insert commentRangeStart before first run
                first_r = p_el.find(f'{{{W}}}r')
                if first_r is None:
                    continue
                cs = etree.Element(f'{{{W}}}commentRangeStart')
                cs.set(f'{{{W}}}id', str(cid))
                first_r.addprevious(cs)
                # Insert commentRangeEnd after last run
                runs = list(p_el.findall(f'{{{W}}}r'))
                last_r = runs[-1] if runs else first_r
                ce = etree.Element(f'{{{W}}}commentRangeEnd')
                ce.set(f'{{{W}}}id', str(cid))
                last_r.addnext(ce)
                # Insert commentReference run after rangeEnd
                ref_r = etree.Element(f'{{{W}}}r')
                ref_rPr = etree.SubElement(ref_r, f'{{{W}}}rPr')
                ref_style = etree.SubElement(ref_rPr, f'{{{W}}}rStyle')
                ref_style.set(f'{{{W}}}val', 'CommentReference')
                ref_cr = etree.SubElement(ref_r, f'{{{W}}}commentReference')
                ref_cr.set(f'{{{W}}}id', str(cid))
                ce.addnext(ref_r)
                placed += 1
                break

    if placed == 0:
        print("  Warning: no comment anchors found, skipping comments")
        return

    # Serialize comments.xml
    files['word/comments.xml'] = etree.tostring(
        comments_root, xml_declaration=True, encoding='UTF-8', standalone=True)

    # Update [Content_Types].xml
    ct_xml = etree.fromstring(files['[Content_Types].xml'])
    ct_override = etree.SubElement(ct_xml, 'Override')
    ct_override.set('PartName', '/word/comments.xml')
    ct_override.set('ContentType',
                    'application/vnd.openxmlformats-officedocument.'
                    'wordprocessingml.comments+xml')
    files['[Content_Types].xml'] = etree.tostring(
        ct_xml, xml_declaration=True, encoding='UTF-8', standalone=True)

    # Update word/_rels/document.xml.rels
    rels_key = 'word/_rels/document.xml.rels'
    rels_xml = etree.fromstring(files[rels_key])
    rel_el = etree.SubElement(rels_xml, 'Relationship')
    rel_el.set('Id', 'rIdComments')
    rel_el.set('Type', COMMENT_REL)
    rel_el.set('Target', 'comments.xml')
    files[rels_key] = etree.tostring(
        rels_xml, xml_declaration=True, encoding='UTF-8', standalone=True)

    # Serialize updated document.xml
    files['word/document.xml'] = etree.tostring(
        doc_xml, xml_declaration=True, encoding='UTF-8', standalone=True)

    # Write back
    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
        for name, data in files.items():
            zout.writestr(name, data)

    print(f"  Added {placed} Word comment(s)")


def main():
    # ═══════════════════════════════════════════════════════════════════════
    # LOAD DATA (v3)
    # ═══════════════════════════════════════════════════════════════════════

    print("Loading data...")
    with open(DATA / "calibration_results_v3.csv", encoding="utf-8") as f:
        cal = list(csv.DictReader(f))
    reg = {}
    with open(DATA / "calibration_regimes_v3.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            reg[row["iso3"]] = row
    # World Bank operational ECA region (developing Europe & Central Asia)
    eca = {
        'ALB', 'ARM', 'AZE', 'BLR', 'BIH', 'BGR', 'HRV', 'CZE', 'EST',
        'GEO', 'HUN', 'KAZ', 'XKX', 'KGZ', 'LVA', 'LTU', 'MDA', 'MNE',
        'MKD', 'POL', 'ROU', 'RUS', 'SRB', 'SVK', 'SVN', 'TJK', 'TUR',
        'TKM', 'UKR', 'UZB',
    }

    eca_cal = [row for row in cal if row["iso3"] in eca]
    non_eca_cal = [row for row in cal if row["iso3"] not in eca]
    n_eca = len(eca_cal)
    n_total = len(cal)

    # v27: Production-efficiency index ξ_j^{eff} ∈ (0, 1]
    # ξ_raw = governance^ω × grid^(1−ω), ω = OMEGA_XI = 0.50
    # ξ_eff = XI_FLOOR + (1 − XI_FLOOR) × ξ_raw  (institutional enclave floor)
    # Data source: raw WGI Rule of Law percentile (G) and grid reliability (R)
    # from xi_scenarios.xlsx (replacing reliability_index.csv)
    xi = {}          # ξ_j^{eff} (with floor)
    xi_raw = {}      # ξ_j^{raw} (before floor, for robustness)
    xi_old = {}      # old ξ (for comparison, from reliability_index.csv)
    xi_components = {}  # raw components for Table A1
    import openpyxl
    _xi_wb = openpyxl.load_workbook(DATA / "xi_scenarios.xlsx", read_only=True)
    _xi_ws = _xi_wb['Data']
    _xi_hdr = [c.value for c in next(_xi_ws.iter_rows(max_row=1))]
    for _xi_row in _xi_ws.iter_rows(min_row=2, values_only=True):
        _xi_d = dict(zip(_xi_hdr, _xi_row))
        iso = _xi_d["ISO3"]
        gov = float(_xi_d["G_RoL"])       # WGI Rule of Law percentile [0,1]
        grid = float(_xi_d["R_grid"])      # grid reliability [0,1]
        xi_old[iso] = float(_xi_d.get("xi_eff_v26", 0.5))
        if gov > 0 and grid > 0:
            xi_raw[iso] = (gov ** OMEGA_XI) * (grid ** (1 - OMEGA_XI))
            xi[iso] = compute_xi_eff(gov, grid)
        else:
            xi_raw[iso] = 0.01
            xi[iso] = XI_FLOOR + (1 - XI_FLOOR) * 0.01
        xi_components[iso] = {"governance": gov, "grid": grid}
    _xi_wb.close()

    # Regime counts computed later from equilibrium results (see regime_5)

    print(f"  Total: {n_total}, ECA: {n_eca}")

    # ═══════════════════════════════════════════════════════════════════════
    # DEMAND CALIBRATION (MW-capacity-based shares)
    # ═══════════════════════════════════════════════════════════════════════
    print("Loading data center capacity estimates...")
    dc_counts = {}
    dc_capacity = {}
    dc_sources = {}
    with open(DATA / "dc_capacity_estimates.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            dc_counts[row["iso3"]] = int(row["n_datacenters"])
            dc_capacity[row["iso3"]] = float(row["capacity_mw"])
            dc_sources[row["iso3"]] = row["source"]

    # Capacity for each calibration country
    dc_k = {}
    for row in cal:
        iso = row["iso3"]
        dc_k[iso] = dc_capacity.get(iso, 5.0)  # minimum 5 MW if no data
    total_dc = sum(dc_k.values())
    omega = {iso: d / total_dc for iso, d in dc_k.items()}

    # Top demand centers
    sorted_omega = sorted(omega.items(), key=lambda x: -x[1])
    top5_labels = []
    for iso, w in sorted_omega[:5]:
        co = next(r["country"] for r in cal if r["iso3"] == iso)
        top5_labels.append((iso, co, w))
    top5_share = sum(w for _, _, w in top5_labels)

    # Training export revenue shares
    train_revenue = {}
    for iso in dc_k:
        if iso in reg:
            src = reg[iso]["best_train_source"]
            train_revenue[src] = train_revenue.get(src, 0) + omega[iso]

    # Inference export revenue shares
    inf_revenue = {}
    for iso in dc_k:
        if iso in reg:
            src = reg[iso]["best_inf_source"]
            inf_revenue[src] = inf_revenue.get(src, 0) + omega[iso]

    # HHI
    hhi_t = sum(s**2 for s in train_revenue.values())
    hhi_i = sum(s**2 for s in inf_revenue.values())

    # Lambda* for each country (cost-based threshold)
    costs_dict = {row["iso3"]: float(row["c_j_total"]) + ETA for row in cal}
    lambda_star = {}
    for iso, c_k in costs_dict.items():
        min_foreign = min(c for i, c in costs_dict.items() if i != iso)
        lambda_star[iso] = c_k / min_foreign - 1

    # v24: Pre-compute bilateral lambdas for key pairs
    # For each buyer k, find min bilateral lambda across all non-sanctioned suppliers
    lambda_min_bilateral = {}  # iso_k → min λ_{k,j} over eligible j
    for iso_k in dc_k:
        min_lam = float('inf')
        for iso_j in costs_dict:
            if iso_j == iso_k:
                continue
            lam_kj = compute_bilateral_lambda(iso_k, iso_j)
            if lam_kj < min_lam:
                min_lam = lam_kj
        lambda_min_bilateral[iso_k] = min_lam

    # v24: Compute tier-specific bilateral lambdas
    # For each buyer, find best supplier per tier
    def _tier_lambda(iso_k, tier):
        """Min bilateral λ for buyer k in tier t, considering tier-specific weights."""
        if tier == 1:
            return float('inf')  # all domestic
        min_lam = float('inf')
        for iso_j in costs_dict:
            if iso_j == iso_k:
                continue
            if iso_k in SANCTIONED or iso_j in SANCTIONED:
                if not (iso_k in SANCTIONED and iso_j in SANCTIONED
                        and _get_bloc(iso_k) == _get_bloc(iso_j)):
                    lam = float('inf')
                else:
                    G = compute_geo_distance(iso_k, iso_j)
                    R = compute_reg_compat(iso_k, iso_j)
                    lam = ALPHA_GEO * G + ALPHA_REG * (1 - R)
            else:
                G = compute_geo_distance(iso_k, iso_j)
                R = compute_reg_compat(iso_k, iso_j)
                if tier == 2:
                    # Regulated: higher weight on regulatory compatibility
                    lam = 0.04 * G + 0.20 * (1 - R)
                else:  # tier 3
                    # Commercial: only geopolitical alignment
                    lam = ALPHA_GEO * G
            if lam < min_lam:
                min_lam = lam
        return min_lam

    # v24: Welfare computation deferred to after equilibrium (needs p_T)
    # (old uniform welfare moved to bilateral welfare below)

    # Counterfactual: doubling sovereignty to 20% (uniform, for comparison)
    min_cost = min(costs_dict.values())
    count_dom_10 = sum(
        1 for iso in dc_k
        if iso in costs_dict and costs_dict[iso] <= 1.10 * min_cost)
    count_dom_20 = sum(
        1 for iso in dc_k
        if iso in costs_dict and costs_dict[iso] <= 1.20 * min_cost)
    extra_dom = count_dom_20 - count_dom_10
    export_share_10 = sum(
        omega[iso] for iso in dc_k
        if iso in costs_dict and costs_dict[iso] > 1.10 * min_cost)
    export_share_20 = sum(
        omega[iso] for iso in dc_k
        if iso in costs_dict and costs_dict[iso] > 1.20 * min_cost)

    # v24: Use SANCTIONED constant instead of hard-coded set
    sanctioned = SANCTIONED

    # Kyrgyzstan inference clients
    kgz_inf_clients = []
    for iso in dc_k:
        if iso in reg and reg[iso]["best_inf_source"] == "KGZ":
            co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
            kgz_inf_clients.append((iso, co, omega[iso] * 100))

    # Build demand_data dict for passing to write functions
    demand_data = {
        "omega": omega, "sorted_omega": sorted_omega,
        "top5_labels": top5_labels, "top5_share": top5_share,
        "train_revenue": train_revenue, "inf_revenue": inf_revenue,
        "hhi_t": hhi_t, "hhi_i": hhi_i,
        "lambda_star": lambda_star, "costs_dict": costs_dict,
        # welfare_total, welfare_pct computed after bilateral equilibrium (below)
        "welfare_train": 0, "welfare_inf": 0,
        "weighted_avg_cost": 0,  # computed after equilibrium
        "count_dom_10": count_dom_10, "count_dom_20": count_dom_20,
        "extra_dom": extra_dom,
        "export_share_10": export_share_10, "export_share_20": export_share_20,
        "kgz_inf_clients": kgz_inf_clients,
        "dc_k": dc_k, "dc_counts": dc_counts, "dc_sources": dc_sources,
    }
    print(f"  DC data for {len(dc_k)} countries, HHI_T={hhi_t:.4f}, HHI_I={hhi_i:.4f}")

    # ═══════════════════════════════════════════════════════════════════════
    # CAPACITY-CONSTRAINED EQUILIBRIUM
    # ═══════════════════════════════════════════════════════════════════════
    print("Computing capacity-constrained equilibrium...")

    # Load grid capacity data (apply scale correction)
    k_bar = {}
    with open(DATA / "grid_capacity_estimates.csv", encoding="utf-8") as f:
        for row in csv.DictReader(f):
            k_bar[row["iso3"]] = float(row["K_bar_gpu_hours"]) * K_BAR_SCALE

    # Training supply stack: rank countries by c_j, compute cumulative capacity
    supply_stack = sorted(
        [(iso, costs_dict[iso], k_bar.get(iso, 1e12))
         for iso in costs_dict if iso in k_bar],
        key=lambda x: x[1]
    )

    def solve_capacity_equilibrium(lam, label, bilateral=False, tiered=False):
        """Solve for capacity-constrained training equilibrium.

        If bilateral=True, uses bilateral λ_{ij} (buyers face pair-specific premia).
        If tiered=True, uses demand tiers (Tier 1 domestic, Tier 2/3 bilateral).
        If bilateral=False, uses uniform scalar lam (old specification).
        """
        p_T = supply_stack[0][1]  # start with cheapest non-sanctioned
        for iso_j, c_j, k_j in supply_stack:
            if iso_j not in sanctioned:
                p_T = c_j
                break
        m_T = 0
        Q_TX = 0
        for _ in range(30):
            Q_TX = 0
            for iso_k in dc_k:
                if iso_k not in costs_dict:
                    continue
                c_k = costs_dict[iso_k]
                w_k = omega.get(iso_k, 0)
                if bilateral or tiered:
                    # v24: tier-specific import decision
                    for tier, w_t in [(1, W_TIER1), (2, W_TIER2), (3, W_TIER3)]:
                        if not tiered:
                            w_t = 1.0  # bilateral but no tiering
                            tier = 3   # use tier-3 lambdas (geopolitical only)
                        if tier == 1:
                            pass  # Tier 1: all domestic, no exports
                        else:
                            lam_k = _tier_lambda(iso_k, tier)
                            if lam_k < float('inf') and c_k > (1 + lam_k) * p_T:
                                Q_TX += w_t * ALPHA * w_k * Q_TOTAL
                        if not tiered:
                            break  # only one pass when not tiered
                else:
                    # Old uniform specification
                    if c_k > (1 + lam) * p_T:
                        Q_TX += ALPHA * w_k * Q_TOTAL
            cum_cap = 0
            found = False
            p_T_new = p_T
            for idx, (iso_j, c_j, k_j) in enumerate(supply_stack):
                if iso_j in sanctioned:
                    continue
                cap_available = k_j * ALPHA
                cum_cap += cap_available
                if cum_cap >= Q_TX and Q_TX > 0:
                    p_T_new = c_j
                    m_T = idx
                    found = True
                    break
            if found and abs(p_T_new - p_T) < 0.0001:
                p_T = p_T_new
                break
            if found:
                p_T = p_T_new

        # Compute shares
        shares = {}
        remaining = Q_TX
        for iso_j, c_j, k_j in supply_stack:
            if iso_j in sanctioned:
                continue
            if c_j > p_T:
                break
            ca = min(k_j * ALPHA, remaining)
            if ca > 0:
                shares[iso_j] = ca
                remaining -= ca
            if remaining <= 0:
                break
        total_exp = sum(shares.values())
        hhi = sum((s / total_exp) ** 2 for s in shares.values()) if total_exp > 0 else 1.0
        # Shadow values
        mu = {}
        for iso_j, c_j, k_j in supply_stack:
            if iso_j in sanctioned:
                continue
            if c_j < p_T:
                allocated = shares.get(iso_j, 0)
                if allocated >= k_j * ALPHA * 0.99:
                    mu[iso_j] = p_T - c_j
        # Lambda_star under capacity constraints
        ls_cap = {iso: c_k / p_T - 1 for iso, c_k in costs_dict.items()}

        print(f"  [{label}] p_T = ${p_T:.3f}/hr, {len(shares)} exporters, "
              f"HHI_T = {hhi:.4f}, {len(mu)} constrained")
        for iso_m, mu_v in sorted(mu.items(), key=lambda x: -x[1])[:5]:
            co = next((r["country"] for r in cal if r["iso3"] == iso_m), iso_m)
            print(f"    {co}: \u03bc = ${mu_v:.3f}/hr")
        return p_T, m_T, shares, hhi, mu, ls_cap, len(shares)

    # Pass 1: pure cost minimization (lambda=0) — main capacity result
    (p_T_0, _, shares_0_orig, cap_hhi_0, mu_0, ls_0, n_exp_0
     ) = solve_capacity_equilibrium(0.0, "\u03bb=0")
    p_T_0_orig = p_T_0  # save pre-cost-recovery training price

    # Pass 2: uniform sovereignty (lambda=LAMBDA) — robustness comparison
    (p_T_sov, _, _, cap_hhi_sov, _, _, n_exp_sov
     ) = solve_capacity_equilibrium(LAMBDA_UNIFORM, f"\u03bb={LAMBDA_UNIFORM}")

    # v24 Pass 3: bilateral λ_{ij} without demand tiering
    (p_T_bilat, _, shares_bilat, cap_hhi_bilat, mu_bilat, ls_bilat, n_exp_bilat
     ) = solve_capacity_equilibrium(0, "bilateral \u03bb_{ij}", bilateral=True)

    # v24 Pass 4: bilateral λ_{ij} with demand tiering
    (p_T_tiered, _, shares_tiered, cap_hhi_tiered, mu_tiered, ls_tiered, n_exp_tiered
     ) = solve_capacity_equilibrium(0, "bilateral tiered", bilateral=True, tiered=True)

    # Store all equilibrium results
    demand_data["p_T"] = p_T_0                       # pure-cost training price
    demand_data["p_T_sov"] = p_T_sov                 # uniform sovereignty training price
    demand_data["cap_hhi_t"] = cap_hhi_0             # pure-cost HHI
    demand_data["cap_hhi_t_sov"] = cap_hhi_sov       # uniform sovereignty HHI
    demand_data["n_train_exporters"] = n_exp_0
    demand_data["n_train_exporters_sov"] = n_exp_sov
    demand_data["mu_j"] = mu_0
    # v24: bilateral equilibrium results
    demand_data["p_T_bilat"] = p_T_bilat
    demand_data["cap_hhi_bilat"] = cap_hhi_bilat
    demand_data["n_train_exporters_bilat"] = n_exp_bilat
    demand_data["shares_bilat"] = shares_bilat
    # v24: tiered bilateral equilibrium results
    demand_data["p_T_tiered"] = p_T_tiered
    demand_data["cap_hhi_tiered"] = cap_hhi_tiered
    demand_data["n_train_exporters_tiered"] = n_exp_tiered
    demand_data["shares_tiered"] = shares_tiered

    # ═══════════════════════════════════════════════════════════════════════
    # COST-RECOVERY ADJUSTMENT (PREFERRED BASELINE)
    # ═══════════════════════════════════════════════════════════════════════
    print("Computing cost-recovery adjustment...")

    # Load latency data for inference recomputation
    latency_data = {}
    with open(DATA / "country_pair_latency.csv", encoding="utf-8") as f:
        for lrow in csv.DictReader(f):
            latency_data[(lrow["iso3_from"], lrow["iso3_to"])] = float(lrow["avg_ms"])
    DOMESTIC_LATENCY_DEFAULT = 5.0

    def _get_latency(j, k):
        if j == k:
            return latency_data.get((j, k), DOMESTIC_LATENCY_DEFAULT)
        if (j, k) in latency_data:
            return latency_data[(j, k)]
        if (k, j) in latency_data:
            return latency_data[(k, j)]
        return None

    adj_costs = dict(costs_dict)  # copy baseline
    adj_changes = {}
    for iso, p_E_adj in SUBSIDY_ADJ.items():
        if iso not in adj_costs:
            continue
        row = next(r for r in cal if r["iso3"] == iso)
        p_E_orig = float(row["p_E_usd_kwh"])
        pue = float(row["pue"])
        delta_elec = pue * GAMMA * (p_E_adj - p_E_orig)
        adj_costs[iso] = costs_dict[iso] + delta_elec
        subsidy_gap = p_E_adj - p_E_orig  # $/kWh gap
        # Fiscal transfer for a hypothetical 100 MW IT-load data center ($/year)
        # Total facility power = IT load × PUE (includes cooling overhead)
        fiscal_transfer_100mw = subsidy_gap * 1000 * 100 * pue * H_YR  # kWh/yr * $/kWh
        adj_changes[iso] = {
            "country": row["country"],
            "p_E_orig": p_E_orig, "p_E_adj": p_E_adj,
            "c_j_orig": costs_dict[iso], "c_j_adj": adj_costs[iso],
            "subsidy_gap": subsidy_gap,
            "fiscal_transfer_100mw": fiscal_transfer_100mw,
        }

    adj_ranked = sorted(adj_costs.items(), key=lambda x: x[1])
    adj_rank_map = {iso: rank for rank, (iso, _) in enumerate(adj_ranked, 1)}

    # Regime changes computed later using 5-type classification (see regime_5)

    adj_cheapest = adj_ranked[0][0]
    # Top 5 adjusted ranking
    adj_top5 = []
    for iso, c in adj_ranked[:5]:
        co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
        adj_top5.append((iso, co, c))

    # Subsidy gap statistics
    gaps = [v["subsidy_gap"] * 1000 for v in adj_changes.values()]  # $/MWh
    max_gap_iso = max(adj_changes, key=lambda x: adj_changes[x]["subsidy_gap"])
    max_gap_entry = adj_changes[max_gap_iso]

    demand_data["adj_top5"] = adj_top5
    demand_data["adj_cheapest_name"] = next(
        (r["country"] for r in cal if r["iso3"] == adj_cheapest), adj_cheapest)
    demand_data["adj_rank_map"] = adj_rank_map
    demand_data["adj_costs"] = adj_costs
    demand_data["n_adjusted"] = len(adj_changes)
    demand_data["max_gap_country"] = max_gap_entry["country"]
    demand_data["max_fiscal_transfer"] = max_gap_entry["fiscal_transfer_100mw"]
    demand_data["min_gap_mwh"] = min(gaps)
    demand_data["max_gap_mwh_val"] = max(gaps)

    print(f"  Adjusted {len(adj_changes)} countries; new cheapest: "
          f"{demand_data['adj_cheapest_name']} (${adj_costs[adj_cheapest]:.3f}/hr)")
    print(f"  Subsidy gap range: ${min(gaps):.0f}\u2013${max(gaps):.0f}/MWh")
    print(f"  Max fiscal transfer (100 MW): {max_gap_entry['country']} "
          f"${max_gap_entry['fiscal_transfer_100mw'] / 1e6:.0f}M/yr")
    for iso, co, c in adj_top5:
        flag = " *" if iso in adj_changes else ""
        print(f"    {adj_rank_map[iso]:>2}. {co:<24} ${c:.3f}/hr{flag}")

    # ═══════════════════════════════════════════════════════════════════════
    # RE-COMPUTE EQUILIBRIUM ON COST-RECOVERY BASELINE
    # ═══════════════════════════════════════════════════════════════════════
    print("Re-computing equilibrium on cost-recovery baseline...")

    # Build efficiency-adjusted cost-recovery costs for training equilibrium
    # Sensitivity analysis (Table A3) uses c/ξ; main equilibrium must match
    adj_costs_eff = {}
    for iso in adj_costs:
        xi_j = xi.get(iso, 1.0)
        adj_costs_eff[iso] = RHO + (adj_costs[iso] - RHO) / xi_j if xi_j > 0 else 999
    adj_supply_stack = sorted(
        [(iso, adj_costs_eff[iso], k_bar.get(iso, 1e12))
         for iso in adj_costs_eff if iso in k_bar],
        key=lambda x: x[1]
    )
    supply_stack = adj_supply_stack  # noqa: F841
    costs_dict = adj_costs_eff

    # Re-run capacity equilibrium on cost-recovery costs
    (p_T_0, _, shares_0, cap_hhi_0, mu_0, ls_0, n_exp_0
     ) = solve_capacity_equilibrium(0.0, "\u03bb=0 cost-recovery")
    (p_T_sov, _, shares_sov, cap_hhi_sov, _, _, n_exp_sov
     ) = solve_capacity_equilibrium(LAMBDA_UNIFORM, f"\u03bb={LAMBDA_UNIFORM} cost-recovery")
    # v24: bilateral on cost-recovery
    (p_T_bilat, _, shares_bilat, cap_hhi_bilat, mu_bilat, ls_bilat, n_exp_bilat
     ) = solve_capacity_equilibrium(0, "bilateral CR", bilateral=True)
    (p_T_tiered, _, shares_tiered, cap_hhi_tiered, _, ls_tiered, n_exp_tiered
     ) = solve_capacity_equilibrium(0, "bilateral tiered CR", bilateral=True, tiered=True)

    demand_data["p_T"] = p_T_0
    demand_data["p_T_sov"] = p_T_sov
    demand_data["p_T_bilat"] = p_T_bilat
    demand_data["p_T_tiered"] = p_T_tiered
    demand_data["cap_hhi_t"] = cap_hhi_0
    demand_data["cap_hhi_t_sov"] = cap_hhi_sov
    demand_data["cap_hhi_bilat"] = cap_hhi_bilat
    demand_data["cap_hhi_tiered"] = cap_hhi_tiered
    demand_data["n_train_exporters"] = n_exp_0
    demand_data["n_train_exporters_sov"] = n_exp_sov
    demand_data["n_train_exporters_bilat"] = n_exp_bilat
    demand_data["n_train_exporters_tiered"] = n_exp_tiered
    demand_data["shares_bilat"] = shares_bilat
    demand_data["shares_tiered"] = shares_tiered
    demand_data["mu_j"] = mu_0
    demand_data["lambda_star"] = ls_0

    # Recompute inference sourcing under cost-recovery costs
    # v24: includes bilateral λ_{ij} and ξ_j^{eff} in delivered cost
    adj_reg = {}
    adj_reg_bilat = {}  # bilateral inference sourcing per tier
    for iso_k in dc_k:
        c_k = adj_costs.get(iso_k)
        if c_k is None:
            continue
        xi_k = xi.get(iso_k, 1.0)
        l_kk = _get_latency(iso_k, iso_k)
        P_I_dom = (1 + TAU * (l_kk or 0)) * (RHO + (c_k - RHO) / xi_k)
        # Free-trade inference (no sovereignty, for comparison)
        best_inf_cost = P_I_dom
        best_inf_src = iso_k
        for iso_j, c_j in adj_costs.items():
            if iso_j == iso_k:
                continue
            l_jk = _get_latency(iso_j, iso_k)
            if l_jk is None:
                continue
            xi_j = xi.get(iso_j, 1.0)
            cost_del = (1 + TAU * l_jk) * (RHO + (c_j - RHO) / xi_j)
            if cost_del < best_inf_cost:
                best_inf_cost = cost_del
                best_inf_src = iso_j
        adj_reg[iso_k] = {
            'best_inf_source': best_inf_src,
            'best_inf_cost': f'{best_inf_cost:.4f}',
            'P_I_domestic': f'{P_I_dom:.4f}',
        }
        # v24: bilateral inference sourcing per tier
        tier_inf = {}
        for tier in (1, 2, 3):
            if tier == 1:
                tier_inf[tier] = {'source': iso_k, 'cost': P_I_dom}
                continue
            best_cost_t = P_I_dom  # domestic always available
            best_src_t = iso_k
            for iso_j, c_j in adj_costs.items():
                if iso_j == iso_k:
                    continue
                lam_kj = compute_bilateral_lambda(iso_k, iso_j)
                if lam_kj >= float('inf'):
                    continue
                # Tier-specific lambda adjustment
                if tier == 2:
                    G = compute_geo_distance(iso_k, iso_j)
                    R = compute_reg_compat(iso_k, iso_j)
                    lam_eff = 0.04 * G + 0.20 * (1 - R)
                else:  # tier 3
                    G = compute_geo_distance(iso_k, iso_j)
                    lam_eff = ALPHA_GEO * G
                l_jk = _get_latency(iso_j, iso_k)
                if l_jk is None:
                    continue
                xi_j = xi.get(iso_j, 1.0)
                cost_del = (1 + lam_eff) * (1 + TAU * l_jk) * (RHO + (c_j - RHO) / xi_j)
                if cost_del < best_cost_t:
                    best_cost_t = cost_del
                    best_src_t = iso_j
            tier_inf[tier] = {'source': best_src_t, 'cost': best_cost_t}
        adj_reg_bilat[iso_k] = tier_inf

    # Recompute inference revenue shares
    adj_inf_revenue = {}
    for iso in dc_k:
        if iso in adj_reg:
            src = adj_reg[iso]['best_inf_source']
            adj_inf_revenue[src] = adj_inf_revenue.get(src, 0) + omega.get(iso, 0)
    adj_hhi_i = sum(s**2 for s in adj_inf_revenue.values())
    demand_data["inf_revenue"] = adj_inf_revenue
    demand_data["hhi_i"] = adj_hhi_i

    # Recompute welfare — uniform specification (for comparison)
    adj_welfare_train = 0
    adj_welfare_inf = 0
    for iso in dc_k:
        if iso in adj_reg and iso in adj_costs:
            c_k = adj_costs[iso]
            min_foreign = min(
                c for i, c in adj_costs.items()
                if i != iso and i not in sanctioned)
            adj_welfare_train += omega.get(iso, 0) * max(0, c_k - min_foreign)
            best_inf = float(adj_reg[iso]["best_inf_cost"])
            P_I_dom = float(adj_reg[iso]["P_I_domestic"])
            adj_welfare_inf += omega.get(iso, 0) * max(0, P_I_dom - best_inf)
    adj_welfare_total = adj_welfare_train + adj_welfare_inf
    adj_weighted_avg = sum(
        omega.get(iso, 0) * adj_costs[iso]
        for iso in dc_k if iso in adj_costs)
    adj_welfare_pct = (adj_welfare_total / adj_weighted_avg * 100
                       if adj_weighted_avg > 0 else 0)
    demand_data["welfare_total_uniform"] = adj_welfare_total
    demand_data["welfare_pct_uniform"] = adj_welfare_pct
    demand_data["weighted_avg_cost"] = adj_weighted_avg

    # v24: Bilateral tiered welfare computation
    bilat_welfare_train = 0
    bilat_welfare_inf = 0
    for iso_k in dc_k:
        if iso_k not in adj_costs:
            continue
        c_k = adj_costs[iso_k]
        w_k = omega.get(iso_k, 0)
        xi_k = xi.get(iso_k, 1.0)
        # Training welfare: tier-weighted
        for tier, w_t in [(1, W_TIER1), (2, W_TIER2), (3, W_TIER3)]:
            if tier == 1:
                # Tier 1: domestic, welfare cost = c_k - min accessible foreign
                min_foreign_t = min(
                    (adj_costs[j] for j in adj_costs if j != iso_k
                     and j not in sanctioned), default=c_k)
                bilat_welfare_train += w_t * w_k * max(0, c_k - min_foreign_t)
            else:
                # Find cheapest accessible foreign supplier for this tier
                min_foreign_t = c_k  # default: domestic
                for iso_j, c_j in adj_costs.items():
                    if iso_j == iso_k:
                        continue
                    lam_kj = compute_bilateral_lambda(iso_k, iso_j)
                    if lam_kj >= float('inf'):
                        continue
                    delivered = (1 + lam_kj) * c_j
                    if delivered < min_foreign_t:
                        min_foreign_t = delivered
                if min_foreign_t < c_k:
                    # Country would import if λ allows; welfare cost = 0
                    pass
                else:
                    # Forced domestic: welfare cost = c_k - best foreign under λ=0
                    best_free = min(
                        (c_j for j, c_j in adj_costs.items()
                         if j != iso_k and j not in sanctioned), default=c_k)
                    bilat_welfare_train += w_t * w_k * max(0, c_k - best_free)
        # Inference welfare: tier-weighted
        for tier, w_t in [(1, W_TIER1), (2, W_TIER2), (3, W_TIER3)]:
            if iso_k in adj_reg_bilat:
                tier_info = adj_reg_bilat[iso_k].get(tier, {})
                P_I_dom = float(adj_reg[iso_k]["P_I_domestic"])
                if tier_info.get('source', iso_k) == iso_k:
                    # Domestic: welfare cost = domestic - best free-trade foreign
                    best_free_inf = float(adj_reg[iso_k]["best_inf_cost"])
                    bilat_welfare_inf += w_t * w_k * max(0, P_I_dom - best_free_inf)
    bilat_welfare_total = bilat_welfare_train + bilat_welfare_inf
    bilat_welfare_pct = (bilat_welfare_total / adj_weighted_avg * 100
                         if adj_weighted_avg > 0 else 0)
    demand_data["welfare_total"] = bilat_welfare_total
    demand_data["welfare_pct"] = bilat_welfare_pct
    demand_data["welfare_train"] = bilat_welfare_train
    demand_data["welfare_inf"] = bilat_welfare_inf

    print(f"  Bilateral tiered welfare: {bilat_welfare_pct:.1f}% "
          f"(vs uniform {adj_welfare_pct:.1f}%)")
    demand_data["adj_reg_bilat"] = adj_reg_bilat

    # Recompute counterfactual
    adj_min_cost = min(adj_costs.values())
    adj_count_dom_10 = sum(
        1 for iso in dc_k
        if iso in adj_costs and adj_costs[iso] <= 1.10 * adj_min_cost)
    adj_count_dom_20 = sum(
        1 for iso in dc_k
        if iso in adj_costs and adj_costs[iso] <= 1.20 * adj_min_cost)
    demand_data["extra_dom"] = adj_count_dom_20 - adj_count_dom_10
    demand_data["export_share_10"] = sum(
        omega.get(iso, 0) for iso in dc_k
        if iso in adj_costs and adj_costs[iso] > 1.10 * adj_min_cost)
    demand_data["export_share_20"] = sum(
        omega.get(iso, 0) for iso in dc_k
        if iso in adj_costs and adj_costs[iso] > 1.20 * adj_min_cost)

    # Recompute KGZ inference clients
    adj_kgz_clients = []
    for iso in dc_k:
        if iso in adj_reg and adj_reg[iso]["best_inf_source"] == "KGZ":
            co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
            adj_kgz_clients.append((iso, co, omega.get(iso, 0) * 100))
    demand_data["kgz_inf_clients"] = adj_kgz_clients

    # Store adj_reg and adj_costs for write functions
    demand_data["adj_reg"] = adj_reg
    demand_data["costs_dict"] = adj_costs

    # ── 5-type regime classification (Proposition 1) ──────────────────────
    def classify_regime_5type(iso, train_exporter_isos, inf_exporter_isos,
                              is_dom_train, is_dom_inf):
        """Classify country into Prop 1's 5-type taxonomy."""
        exports_train = iso in train_exporter_isos
        exports_inf = iso in inf_exporter_isos
        if exports_train:
            # Type (i): training exporter → must also export inference (Prop 4)
            return "T+I exporter"
        if exports_inf and not is_dom_train:
            # Type (ii): inference hub — imports training, exports inference
            return "inference hub"
        if not is_dom_train and is_dom_inf:
            # Type (iii): imports training, produces inference domestically
            return "hybrid"
        if is_dom_train and is_dom_inf:
            # Type (iv): domestic both
            return "domestic"
        # Type (v): imports both
        return "full importer"

    # Inference exporters: countries that serve at least one other country
    inf_exporter_isos = set()
    for iso_k, info in adj_reg.items():
        src = info['best_inf_source']
        if src != iso_k:
            inf_exporter_isos.add(src)

    # v24: regime classification under bilateral λ (preferred specification)
    _reg_5type = {"T+I exporter": 0, "inference hub": 0, "hybrid": 0,
                  "domestic": 0, "full importer": 0}
    regime_5 = {}  # iso → regime label
    # Under bilateral: use tiered bilateral equilibrium
    bilat_train_exporters = set(shares_tiered.keys())
    bilat_inf_exporters = set()
    for iso_k in dc_k:
        if iso_k in adj_reg_bilat:
            # Use tier 3 (majority of demand) for inference export determination
            src = adj_reg_bilat[iso_k].get(3, {}).get('source', iso_k)
            if src != iso_k:
                bilat_inf_exporters.add(src)
    for iso_k in dc_k:
        c_k = adj_costs.get(iso_k)
        if c_k is None:
            continue
        # Under bilateral: country k produces training domestically if
        # for all suppliers j, λ_{kj} >= λ_k^* = c_k/p_T - 1
        lam_k_min = lambda_min_bilateral.get(iso_k, float('inf'))
        lam_star_k = c_k / p_T_tiered - 1 if p_T_tiered > 0 else 0
        is_dom_train = (lam_k_min >= lam_star_k) or (c_k <= p_T_tiered)
        # Inference: check tier 3 (dominant tier)
        is_dom_inf = (adj_reg_bilat.get(iso_k, {}).get(3, {}).get('source', iso_k) == iso_k)
        r = classify_regime_5type(iso_k, bilat_train_exporters,
                                  bilat_inf_exporters, is_dom_train, is_dom_inf)
        regime_5[iso_k] = r
        _reg_5type[r] += 1

    demand_data["regime_5"] = regime_5
    demand_data["reg_5type_counts"] = _reg_5type
    print(f"  5-type regimes (bilateral): {dict((k, v) for k, v in _reg_5type.items() if v)}")

    # Also compute uniform-λ regimes for comparison
    _reg_5type_uniform = {"T+I exporter": 0, "inference hub": 0, "hybrid": 0,
                          "domestic": 0, "full importer": 0}
    regime_5_uniform = {}
    uniform_train_exporters = set(shares_sov.keys()) if shares_sov else set()
    uniform_inf_exporters = set()
    for iso_k in dc_k:
        if iso_k in adj_reg:
            src = adj_reg[iso_k].get('best_inf_source', iso_k)
            if src != iso_k:
                uniform_inf_exporters.add(src)
    for iso_k in dc_k:
        c_k = adj_costs.get(iso_k)
        if c_k is None:
            continue
        is_dom_train = (c_k <= (1 + LAMBDA_UNIFORM) * p_T_0)
        is_dom_inf = (adj_reg.get(iso_k, {}).get('best_inf_source') == iso_k)
        r = classify_regime_5type(iso_k, uniform_train_exporters,
                                  uniform_inf_exporters, is_dom_train, is_dom_inf)
        regime_5_uniform[iso_k] = r
        _reg_5type_uniform[r] += 1
    demand_data["regime_5_uniform"] = regime_5_uniform
    demand_data["reg_5type_counts_uniform"] = _reg_5type_uniform
    print(f"  5-type regimes (uniform): {dict((k, v) for k, v in _reg_5type_uniform.items() if v)}")

    # Count regime changes: pre-cost-recovery vs cost-recovery 5-type regimes
    orig_train_exporters = set(shares_0_orig.keys())
    orig_inf_exporters = set()
    for iso in dc_k:
        if iso in reg:
            src = reg[iso]["best_inf_source"]
            if src != iso:
                orig_inf_exporters.add(src)
    orig_costs = {row["iso3"]: float(row["c_j_total"]) + ETA for row in cal}
    regime_changes = 0
    for iso_k in dc_k:
        if iso_k not in regime_5:
            continue
        new_r = regime_5[iso_k]
        c_k_orig = orig_costs.get(iso_k)
        if c_k_orig is None:
            continue
        is_dom_train_orig = (c_k_orig <= (1 + LAMBDA) * p_T_0_orig)
        is_dom_inf_orig = (reg.get(iso_k, {}).get("best_inf_source") == iso_k)
        orig_r = classify_regime_5type(iso_k, orig_train_exporters,
                                       orig_inf_exporters,
                                       is_dom_train_orig, is_dom_inf_orig)
        if orig_r != new_r:
            regime_changes += 1
    demand_data["regime_changes"] = regime_changes
    print(f"  Regime changes (pre-CR → CR): {regime_changes}")

    # Print summary
    print(f"  Cost-recovery inference HHI_I = {adj_hhi_i:.4f}")
    adj_inf_top5 = sorted(adj_inf_revenue.items(), key=lambda x: -x[1])[:5]
    for iso, share in adj_inf_top5:
        co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
        print(f"    {co}: {share * 100:.1f}%")

    # ═══════════════════════════════════════════════════════════════════════
    # v28: HYPERSCALER FDI REGIME CLASSIFICATION
    # λ^FDI replaces λ_{ij} — trust attaches to operator, not host country.
    # Uses pairwise cost comparison to identify the set of potential exporters
    # and importers under hyperscaler intermediation (equation 2').
    # ═══════════════════════════════════════════════════════════════════════
    print("\nComputing hyperscaler FDI equilibrium...")

    # FDI training: capacity-constrained equilibrium with FDI lambda.
    # Under FDI, the hyperscaler operates the facility and partially mitigates
    # the host's governance penalty. We use cost-recovery costs for the supply
    # stack (hosts) and ξ-adjusted costs for the demand side (buyers).
    # Build FDI supply stack from cost-recovery costs (non-sanctioned)
    fdi_supply_stack = sorted(
        [(iso, adj_costs[iso], k_bar.get(iso, 1e12))
         for iso in adj_costs if iso in k_bar and iso not in SANCTIONED],
        key=lambda x: x[1]
    )

    def _solve_fdi_equilibrium():
        """Capacity-constrained FDI training equilibrium."""
        p_T = fdi_supply_stack[0][1]
        for iso_j, c_j, k_j in fdi_supply_stack:
            p_T = c_j
            break
        Q_TX = 0
        for _ in range(30):
            Q_TX = 0
            for iso_k in dc_k:
                c_k = costs_dict.get(iso_k)  # buyer's ξ-adjusted domestic cost
                if c_k is None:
                    continue
                w_k = omega.get(iso_k, 0)
                # FDI lambda: min across non-sanctioned hosts
                lam_fdi_min = float('inf')
                for iso_j in adj_costs:
                    if iso_j == iso_k or iso_j in SANCTIONED:
                        continue
                    lam = compute_fdi_lambda(iso_j, iso_k, hyperscaler_h='USA')
                    if lam < lam_fdi_min:
                        lam_fdi_min = lam
                if lam_fdi_min < float('inf') and c_k > (1 + lam_fdi_min) * p_T:
                    Q_TX += ALPHA * w_k * Q_TOTAL
            # Walk up supply stack
            cum_cap = 0
            p_T_new = p_T
            for iso_j, c_j, k_j in fdi_supply_stack:
                cum_cap += k_j * ALPHA
                if cum_cap >= Q_TX and Q_TX > 0:
                    p_T_new = c_j
                    break
            if abs(p_T_new - p_T) < 0.0001:
                p_T = p_T_new
                break
            p_T = p_T_new
        # Compute shares
        shares = {}
        remaining = Q_TX
        for iso_j, c_j, k_j in fdi_supply_stack:
            if c_j > p_T:
                break
            ca = min(k_j * ALPHA, remaining)
            if ca > 0:
                shares[iso_j] = ca
                remaining -= ca
            if remaining <= 0:
                break
        return p_T, shares

    p_T_fdi, shares_fdi = _solve_fdi_equilibrium()
    fdi_can_export_train = set(shares_fdi.keys())

    # Also identify countries that would import training under FDI
    fdi_would_import_train = {}
    for iso_k in dc_k:
        c_k = costs_dict.get(iso_k)
        if c_k is None:
            continue
        best_supplier = None
        best_delivered = c_k
        for iso_j in adj_costs:
            if iso_j == iso_k or iso_j in SANCTIONED:
                continue
            lam_fdi = compute_fdi_lambda(iso_j, iso_k, hyperscaler_h='USA')
            if lam_fdi >= float('inf'):
                continue
            delivered = (1 + lam_fdi) * adj_costs[iso_j]
            if delivered < best_delivered:
                best_delivered = delivered
                best_supplier = iso_j
        if best_supplier is not None:
            fdi_would_import_train[iso_k] = best_supplier
    demand_data["p_T_fdi"] = p_T_fdi
    demand_data["shares_fdi"] = {j: 1 for j in fdi_can_export_train}
    demand_data["n_train_exporters_fdi"] = len(fdi_can_export_train)

    # FDI inference sourcing: use hyperscaler trust channel
    # Host cost = cost-recovery (hyperscaler mitigates ξ); buyer cost = ξ-adjusted
    fdi_inf_src = {}
    for iso_k in dc_k:
        c_k_eff = costs_dict.get(iso_k)  # buyer's ξ-adjusted domestic cost
        if c_k_eff is None:
            continue
        l_kk = _get_latency(iso_k, iso_k)
        P_I_dom = (1 + TAU * (l_kk or 0)) * c_k_eff
        best_cost = P_I_dom
        best_src = iso_k
        for iso_j in adj_costs:
            if iso_j == iso_k:
                continue
            if iso_j not in dc_k:
                continue
            lam_fdi = compute_fdi_lambda(iso_j, iso_k, hyperscaler_h='USA')
            if lam_fdi >= float('inf'):
                continue
            l_jk = _get_latency(iso_j, iso_k)
            if l_jk is None:
                continue
            c_j_cr = adj_costs[iso_j]  # host cost-recovery (hyperscaler mitigates ξ)
            cost_del = (1 + lam_fdi) * (1 + TAU * l_jk) * c_j_cr
            if cost_del < best_cost:
                best_cost = cost_del
                best_src = iso_j
        fdi_inf_src[iso_k] = best_src
    demand_data["fdi_inf_src"] = fdi_inf_src

    # FDI regime classification (5-type) using pairwise trade patterns
    fdi_train_exporters = fdi_can_export_train
    fdi_inf_exporters = set()
    for iso_k, src in fdi_inf_src.items():
        if src != iso_k:
            fdi_inf_exporters.add(src)

    regime_5_fdi = {}
    _reg_5type_fdi = {"T+I exporter": 0, "inference hub": 0, "hybrid": 0,
                      "domestic": 0, "full importer": 0}
    for iso_k in dc_k:
        c_k = costs_dict.get(iso_k)
        if c_k is None:
            continue
        exports_train = iso_k in fdi_train_exporters
        exports_inf = iso_k in fdi_inf_exporters
        imports_train = iso_k in fdi_would_import_train
        imports_inf = (fdi_inf_src.get(iso_k, iso_k) != iso_k)
        # Classify:
        if exports_train and (exports_inf or not imports_inf):
            r = "T+I exporter"
        elif exports_inf and imports_train:
            r = "inference hub"
        elif imports_train and not imports_inf:
            r = "hybrid"
        elif not imports_train and not imports_inf:
            r = "domestic"
        else:
            r = "full importer"
        regime_5_fdi[iso_k] = r
        _reg_5type_fdi[r] += 1

    demand_data["regime_5_fdi"] = regime_5_fdi
    demand_data["reg_5type_counts_fdi"] = _reg_5type_fdi

    # Count developing-country exporters under FDI
    DEVELOPING = {
        'CHN', 'KGZ', 'XKX', 'MNE', 'ETH', 'VNM', 'IND', 'KEN', 'ARE',
        'EGY', 'DZA', 'UZB', 'TJK', 'TKM', 'ALB', 'MKD', 'GEO', 'ARM',
        'MDA', 'UKR', 'BIH', 'SRB', 'IDN', 'MYS', 'PHL', 'THA', 'COL',
        'MEX', 'BRA', 'ARG', 'CHL', 'PER', 'NGA', 'ZAF', 'MAR', 'TUN',
        'SEN', 'BGD', 'PAK', 'LKA', 'MMR', 'LAO', 'KHM',
    }
    n_dev_fdi_exporters = sum(
        1 for iso, r in regime_5_fdi.items()
        if iso in DEVELOPING and r in ("T+I exporter", "inference hub"))
    demand_data["n_dev_fdi_exporters"] = n_dev_fdi_exporters
    demand_data["DEVELOPING"] = DEVELOPING

    print(f"  FDI p_T = ${p_T_fdi:.3f}/hr, {len(fdi_can_export_train)} training exporters")
    print(f"  FDI 5-type regimes: {dict((k, v) for k, v in _reg_5type_fdi.items() if v)}")
    print(f"  Developing-country FDI exporters: {n_dev_fdi_exporters}")
    fdi_exp_names = []
    for iso in sorted(fdi_train_exporters | fdi_inf_exporters):
        co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
        dev_mark = "\u2020" if iso in DEVELOPING else ""
        fdi_exp_names.append(f"{co}{dev_mark}")
    print(f"  FDI exporters: {', '.join(fdi_exp_names)}")

    # ═══════════════════════════════════════════════════════════════════════
    # SENSITIVITY ANALYSIS
    # ═══════════════════════════════════════════════════════════════════════
    print("\nRunning sensitivity analysis...")
    sens_results = run_sensitivity(cal, omega, dc_k, k_bar, SANCTIONED, xi, xi_raw)
    demand_data["sensitivity"] = sens_results

    # ═══════════════════════════════════════════════════════════════════════
    # EFFICIENCY-ADJUSTED COST RANKINGS (v24: ξ_j^{eff})
    # ═══════════════════════════════════════════════════════════════════════
    print("Computing efficiency-adjusted rankings...")
    # Use cost-recovery adjusted costs (preferred baseline)
    xi_costs = {}
    for row in cal:
        iso = row["iso3"]
        c_j = float(row["c_j_total"]) + ETA
        xi_j = xi.get(iso, 1.0)
        # Apply cost-recovery adjustment if applicable
        if iso in SUBSIDY_ADJ:
            p_E_cr = SUBSIDY_ADJ[iso]
            p_E_raw = float(row["p_E_usd_kwh"])
            c_j = c_j + (p_E_cr - p_E_raw) * float(row["pue"]) * GAMMA
        xi_costs[iso] = RHO + (c_j - RHO) / xi_j

    # Baseline (no xi adjustment) costs for comparison
    baseline_costs = {}
    for row in cal:
        iso = row["iso3"]
        c_j = float(row["c_j_total"]) + ETA
        if iso in SUBSIDY_ADJ:
            p_E_cr = SUBSIDY_ADJ[iso]
            p_E_raw = float(row["p_E_usd_kwh"])
            c_j = c_j + (p_E_cr - p_E_raw) * float(row["pue"]) * GAMMA
        baseline_costs[iso] = c_j

    # Rank both
    baseline_rank = sorted(baseline_costs.items(), key=lambda x: x[1])
    xi_rank = sorted(xi_costs.items(), key=lambda x: x[1])
    baseline_order = [iso for iso, _ in baseline_rank]
    xi_order = [iso for iso, _ in xi_rank]

    # Top 5 with names
    xi_top5 = []
    for iso, cost in xi_rank[:5]:
        co = next((r["country"] for r in cal if r["iso3"] == iso), iso)
        xi_top5.append((co, cost))

    # Spearman rank correlation
    n_r = len(baseline_order)
    rank_base = {iso: i for i, iso in enumerate(baseline_order)}
    rank_xi = {iso: i for i, iso in enumerate(xi_order)}
    d_sq = sum((rank_base[iso] - rank_xi[iso]) ** 2
               for iso in baseline_order if iso in rank_xi)
    spearman = 1 - 6 * d_sq / (n_r * (n_r ** 2 - 1))

    # Count how many top-10 baseline producers fall out of top-10
    base_top10 = set(baseline_order[:10])
    xi_top10 = set(xi_order[:10])
    n_changed_top10 = len(base_top10 - xi_top10)

    demand_data["xi_adjusted"] = {
        "top5": xi_top5,
        "rank_corr": spearman,
        "n_changed_top10": n_changed_top10,
        "xi_order": xi_order[:10],
        "baseline_rank_map": {iso: i for i, iso in enumerate(baseline_order)},
        "xi_rank_map": {iso: i for i, iso in enumerate(xi_order)},
    }
    demand_data["xi"] = xi
    demand_data["xi_raw"] = xi_raw
    demand_data["xi_old"] = xi_old
    demand_data["xi_components"] = xi_components
    demand_data["lambda_min_bilateral"] = lambda_min_bilateral
    # Country name map for figure labels
    demand_data["iso_country"] = {r["iso3"]: r["country"] for r in cal}
    print(f"  Efficiency-adjusted top 5: {[f'{co} (${c:.2f})' for co, c in xi_top5]}")
    print(f"  Spearman rank corr: {spearman:.4f}, top-10 changes: {n_changed_top10}")

    # ═══════════════════════════════════════════════════════════════════════
    # TABLE 3 DATA: Country Rankings Under Alternative Pricing Assumptions
    # v24: Table 3a (cost specs) + Table 3b (sovereignty specs)
    # ═══════════════════════════════════════════════════════════════════════
    print("\nComputing Table 3 data...")
    rho_hw = GPU_PRICE / (GPU_LIFE * H_YR * GPU_UTIL)

    table3_data = []
    for r_row in cal:
        iso = r_row["iso3"]
        p_E_raw = float(r_row["p_E_usd_kwh"])
        pue = float(r_row["pue"])
        constr = float(r_row["p_L_usd_per_W"])
        xi_j = xi.get(iso, 1.0)  # v24: ξ_j^{eff} (no sanctions)

        elec_raw = GAMMA * p_E_raw * pue
        cr_price = SUBSIDY_ADJ.get(iso, p_E_raw)
        elec_cr = GAMMA * cr_price * pue
        constr_cost = (constr * GAMMA * 1000) / (DC_LIFE * H_YR * GPU_UTIL)

        # Back out networking residual from reported c_j
        cj_reported = float(r_row["c_j_total"])
        residual = cj_reported - (elec_raw + rho_hw + constr_cost)

        table3_data.append({
            "iso": iso, "country": r_row["country"],
            "p_E_raw": p_E_raw, "cr_price": cr_price,
            "pue": pue, "constr": constr, "xi": xi_j,
            "elec_raw": elec_raw, "elec_cr": elec_cr,
            "constr_cost": constr_cost,
            "cj_reported": cj_reported, "residual": residual,
        })

    # Mean networking cost (ρ_net)
    rho_net = sum(d["residual"] for d in table3_data) / len(table3_data)

    # ── Table 3a: Cost Specifications (1)-(3) ──
    for d in table3_data:
        d["cj_raw"] = d["elec_raw"] + rho_hw + d["constr_cost"] + rho_net    # (1) Raw
        d["cj_cr"] = d["elec_cr"] + rho_hw + d["constr_cost"] + rho_net      # (2) Cost-recovery
        d["cj_eff"] = RHO + (d["cj_cr"] - RHO) / d["xi"] if d["xi"] > 0 else 999  # (3) Form B

    # Validate: cj_raw should approximately match reported
    for d in table3_data:
        assert abs(d["cj_raw"] - d["cj_reported"]) < 0.015, \
            f'{d["iso"]}: raw={d["cj_raw"]:.4f} vs reported={d["cj_reported"]:.4f}'

    # v27: Override spec (3) with exact values from form_b_simulations.xlsx (C2 scenario)
    # This ensures rankings match the protocol lookup table exactly
    _sim_wb = openpyxl.load_workbook(DATA / "form_b_simulations.xlsx", read_only=True)
    _sim_ws = _sim_wb['Rankings']
    _sim_hdr = [c.value for c in next(_sim_ws.iter_rows(max_row=1))]
    _c2_cadj_i = _sim_hdr.index('c_adj\nC2')
    _c2_rank_i = _sim_hdr.index('rank\nC2')
    _c2_xieff_i = _sim_hdr.index('xi_eff\nC2')
    _sim_data = {}
    # Need country-to-ISO mapping
    _sim_data_ws = _sim_wb['Data']
    _sim_data_hdr = [c.value for c in next(_sim_data_ws.iter_rows(max_row=1))]
    _ctry_to_iso = {}
    for _r in _sim_data_ws.iter_rows(min_row=2, values_only=True):
        _dd = dict(zip(_sim_data_hdr, _r))
        # form_b_simulations Data sheet doesn't have ISO, but xi_scenarios does
    _sim_data_ws = None
    # Use country name matching from Rankings sheet
    for _r in _sim_ws.iter_rows(min_row=2, values_only=True):
        _country = _r[0]  # Column 0 is Country
        _sim_data[_country] = {
            'c_adj': float(_r[_c2_cadj_i]),
            'rank': int(_r[_c2_rank_i]),
            'xi_eff': float(_r[_c2_xieff_i]),
        }
    _sim_wb.close()
    # Map country names to table3_data countries
    _t3_country_map = {d["country"]: d for d in table3_data}
    _matched = 0
    for _sname, _svals in _sim_data.items():
        # Try exact match or common abbreviation
        d = _t3_country_map.get(_sname)
        if d is None:
            # Try partial match
            for _tname, _td in _t3_country_map.items():
                if _sname.startswith(_tname[:10]) or _tname.startswith(_sname[:10]):
                    d = _td
                    break
        if d is not None:
            d["cj_eff"] = _svals["c_adj"]
            d["xi"] = _svals["xi_eff"]
            _matched += 1
    print(f"  Table 3: matched {_matched}/{len(_sim_data)} countries from C2 scenario")

    # Rank under specs (1)-(3)
    for key, spec in [("rank_raw", "cj_raw"), ("rank_cr", "cj_cr"), ("rank_eff", "cj_eff")]:
        sorted_by = sorted(table3_data, key=lambda x: x[spec])
        for rank, d in enumerate(sorted_by, 1):
            d[key] = rank

    # Override adj_rank_map with table-consistent CR ranks (table3_data uses
    # recomputed costs that may differ slightly from the costs_dict path)
    demand_data["adj_rank_map"] = {d["iso"]: d["rank_cr"] for d in table3_data}

    # Type assignment for specs (1)-(3): free trade
    for d in table3_data:
        for rank_key, type_key in [("rank_raw", "type_raw"),
                                   ("rank_cr", "type_cr"),
                                   ("rank_eff", "type_eff")]:
            if d[rank_key] <= 5:
                d[type_key] = "EE"
            elif d[rank_key] <= 12 and dc_k.get(d["iso"], 9999) < 1000:
                d[type_key] = "IE"
            else:
                d[type_key] = "II"
        # Delta: rank improvement from raw to efficiency-adjusted
        d["delta"] = d["rank_raw"] - d["rank_eff"]

    # ── Table 3b: Sovereignty Specifications ──
    # Spec (4): Bilateral λ_{ij}, full demand (no tiering)
    for d in table3_data:
        iso = d["iso"]
        d["regime_bilat"] = regime_5.get(iso, "full importer")
        # Use bilateral equilibrium regime assignment
        if iso in shares_bilat:
            d["type_bilat"] = "EE"
        elif regime_5.get(iso) == "domestic":
            d["type_bilat"] = "DD"
        elif regime_5.get(iso) in ("inference hub",):
            d["type_bilat"] = "IE"
        else:
            d["type_bilat"] = "II"
        # Switching threshold: λ_k^* = c_k^{eff} / p_T^{bilat} - 1
        d["lam_k_star"] = d["cj_eff"] / p_T_bilat - 1 if p_T_bilat > 0 else 0

    # Spec (5): Bilateral λ_{ij}, tiered demand
    for d in table3_data:
        iso = d["iso"]
        if iso in shares_tiered:
            d["type_tiered"] = "EE"
        elif regime_5.get(iso) == "domestic":
            d["type_tiered"] = "DD"
        elif regime_5.get(iso) in ("inference hub",):
            d["type_tiered"] = "IE"
        else:
            d["type_tiered"] = "II"
        # Compute fraction of demand served domestically per country
        dom_share = W_TIER1  # Tier 1 always domestic
        for tier, w_t in [(2, W_TIER2), (3, W_TIER3)]:
            lam_k = _tier_lambda(iso, tier)
            lam_star_k = d["cj_eff"] / p_T_tiered - 1 if p_T_tiered > 0 else 0
            if lam_k >= lam_star_k or d["cj_eff"] <= p_T_tiered:
                dom_share += w_t  # domestic for this tier
        d["dom_share_tiered"] = dom_share

    # Spec (6): Uniform λ = 10% (comparison with old specification)
    p_star = sorted(table3_data, key=lambda x: x["cj_eff"])[4]["cj_eff"]
    for d in table3_data:
        d["sigma_bar"] = max(0, (d["cj_eff"] - p_star) / p_star)
        if d["sigma_bar"] <= 0.10:
            d["pj_sov"] = d["cj_eff"]   # domestic
            d["type_uniform"] = "DD" if d["rank_eff"] > 5 else "EE"
        else:
            d["pj_sov"] = p_star * 1.10  # importer
            d["type_uniform"] = "II"
        # Override: top exporters
        if d["rank_eff"] <= 5:
            d["type_uniform"] = "EE"

    # Rank spec (6) — tiebreak importers by cj_eff
    sorted_sov = sorted(table3_data, key=lambda x: (x["pj_sov"], x["cj_eff"]))
    for rank, d in enumerate(sorted_sov, 1):
        d["rank_sov"] = rank

    # v28: Spec (7): Hyperscaler FDI — λ^FDI from equation (2')
    shares_fdi = demand_data["shares_fdi"]
    for d in table3_data:
        iso = d["iso"]
        fdi_regime = regime_5_fdi.get(iso, "full importer")
        if iso in shares_fdi:
            d["type_fdi"] = "EE"
        elif fdi_regime == "domestic":
            d["type_fdi"] = "DD"
        elif fdi_regime == "inference hub":
            d["type_fdi"] = "IE"
        else:
            d["type_fdi"] = "II"
        # Mark sanctioned/GPU-blocked
        if iso in SANCTIONED:
            d["type_fdi"] += "*"
        elif iso in GPU_EXPORT_CONTROLLED:
            d["type_fdi"] += "*"
        # Mark developing-country exporter
        if iso in DEVELOPING and d["type_fdi"].rstrip("*") in ("EE", "IE"):
            d["type_fdi"] += "\u2020"  # dagger

    n_fdi_exporters = sum(1 for d in table3_data
                          if d.get("type_fdi", "").rstrip("*\u2020") in ("EE", "IE"))
    n_fdi_dev_exp = sum(1 for d in table3_data
                        if d["iso"] in DEVELOPING
                        and d.get("type_fdi", "").rstrip("*\u2020") in ("EE", "IE"))
    print(f"  Spec (7) FDI: {n_fdi_exporters} exporters, {n_fdi_dev_exp} developing")

    demand_data["table3"] = table3_data
    demand_data["p_star"] = p_star

    # Override xi_top5 to match Table 3 ranking (same cj_eff formula)
    t3_sorted_eff = sorted(table3_data, key=lambda x: x["cj_eff"])[:5]
    demand_data["xi_adjusted"]["top5"] = [(d["country"], d["cj_eff"]) for d in t3_sorted_eff]

    print(f"  Table 3: {len(table3_data)} countries, p* = ${p_star:.4f}/hr")
    print(f"  Spec (1) top 5: {[d['country'] for d in sorted(table3_data, key=lambda x: x['cj_raw'])[:5]]}")
    print(f"  Spec (3) top 5: {[d['country'] for d in sorted(table3_data, key=lambda x: x['cj_eff'])[:5]]}")
    n_dom_bilat = sum(1 for d in table3_data if d.get("type_bilat") == "DD")
    n_dom_tiered = sum(1 for d in table3_data if d.get("type_tiered") == "DD")
    n_dom_uniform = sum(1 for d in table3_data if d.get("type_uniform") == "DD")
    print(f"  Domestic: bilateral={n_dom_bilat}, tiered={n_dom_tiered}, uniform={n_dom_uniform}")

    # ═══════════════════════════════════════════════════════════════════════
    # LOAD v8 AND INDEX HEADINGS
    # ═══════════════════════════════════════════════════════════════════════

    print("\nLoading v8...")
    doc = Document(str(DOCS / "flop_trade_model_v8.docx"))
    body = doc.element.body
    all_el = list(body)
    init_footnotes(doc)

    hmap = {}
    for el in all_el:
        if el.tag == qn('w:p'):
            pPr = el.find(qn('w:pPr'))
            if pPr is not None:
                pS = pPr.find(qn('w:pStyle'))
                if pS is not None and 'Heading' in pS.get(qn('w:val'), ''):
                    ft = "".join(r.text or "" for r in el.findall(f'.//{qn("w:t")}'))
                    if '1.2' in ft:
                        hmap['1.2'] = el
                    elif '1.1' in ft:
                        hmap['1.1'] = el
                    elif '1.' in ft and 'Model' in ft:
                        hmap['1'] = el
                    elif '2.' in ft and 'Comp' in ft:
                        hmap['2'] = el
                    elif '3.' in ft and 'Make' in ft:
                        hmap['3'] = el
                    elif '4.' in ft and 'Calib' in ft:
                        hmap['4'] = el
                    elif '5.' in ft and 'Conc' in ft:
                        hmap['5'] = el
                    elif ft.strip() == 'References':
                        hmap['refs'] = el
                    elif ft.strip() == 'Abstract':
                        hmap['abs'] = el

    # ═══════════════════════════════════════════════════════════════════════
    # STEPS
    # ═══════════════════════════════════════════════════════════════════════

    title_el, author_el, ver_el, abs_text_el, kw_el = write_title_and_abstract(doc, body, all_el, hmap, demand_data)
    write_introduction(doc, body, hmap)
    write_literature(doc, body, hmap)
    write_production_technology(doc, body, hmap)
    write_trade_costs(doc, body, hmap)
    renumber_sections(hmap)
    # New Section 3 subsections (3.3, 3.4) — inserted before the renumbered Section 4 heading
    write_demand(doc, body, hmap, demand_data)
    write_sourcing_and_equilibrium(doc, body, hmap, demand_data)
    # Section 4: Equilibrium Properties (replaces old Comparative Advantage + Make-or-Buy)
    write_equilibrium_properties(doc, body, hmap, demand_data)
    write_data_section(doc, body, hmap, demand_data)
    write_calibration(doc, body, hmap, cal, reg, n_eca, n_total, demand_data)
    write_conclusion(doc, body, hmap, demand_data)

    refs = hmap['refs']
    last_ref = write_references(doc, body, refs)
    last_fig1 = write_figure1_calibration(doc, body, last_ref)
    last_fig2 = write_figure4b(doc, body, last_fig1, demand_data)
    last_table1_tax = write_table1(doc, body, last_fig2)
    last_table1 = write_table2(doc, body, last_table1_tax, demand_data)
    last_table3 = write_table3(doc, body, last_table1, demand_data)
    last_app_note = write_appendix(doc, body, last_table3, eca_cal, non_eca_cal, reg, demand_data)
    last_table_a2 = write_table_a2(doc, body, last_app_note, demand_data)
    last_model_app = write_model_appendix(doc, body, last_table_a2)
    last_sens_app = write_sensitivity_appendix(doc, body, last_model_app, demand_data)
    last_dcf_app = write_kyrgyzstan_appendix(doc, body, last_sens_app)
    last_reg_app = write_construction_regression_appendix(doc, body, last_dcf_app)
    write_workload_appendix(doc, body, last_reg_app)
    link_citations(body)
    link_equations(body)
    fix_orphan_backlinks(body, refs)
    apply_formatting(doc, body, refs, title_el, author_el, ver_el, abs_text_el)
    add_page_numbers_and_break(doc, body, kw_el)

    # ═══════════════════════════════════════════════════════════════════════
    # SAVE
    # ═══════════════════════════════════════════════════════════════════════
    flush_footnotes()
    doc.core_properties.author = 'Michael Lokshin'
    out = DOCS / "flop_trade_model_v28.docx"
    for _attempt in range(30):
        try:
            doc.save(str(out))
            break
        except PermissionError:
            if _attempt == 0:
                print(f"\nFile locked — waiting for Word to release {out.name}...")
            import time
            time.sleep(2)
    else:
        raise PermissionError(f"Could not save {out} after 60 seconds. Close Word and retry.")
    print(f"\nSaved {out}")

    # ═══════════════════════════════════════════════════════════════════════
    # POST-PROCESSING: Word comments removed per author request

    # ═══════════════════════════════════════════════════════════════════════
    # AUTO-COMMIT to git (preserves every successful generation)
    # ═══════════════════════════════════════════════════════════════════════
    try:
        import datetime as _dt
        import subprocess
        _repo = str(DOCS.parent)  # F:\onedrive\...\FLOPsExport
        _script = str(pathlib.Path(__file__).resolve())
        _ts = _dt.datetime.now().strftime('%Y-%m-%d %H:%M')
        subprocess.run(['git', 'add', _script, str(out)],
                       cwd=_repo, capture_output=True, timeout=10)
        result = subprocess.run(
            ['git', 'commit', '-m', f'Auto-save v28: {_ts}'],
            cwd=_repo, capture_output=True, text=True, timeout=10)
        if result.returncode == 0:
            _hash = result.stdout.strip().split()[1].rstrip(']')
            print(f"Git auto-commit: {_hash}")
        elif 'nothing to commit' in result.stdout:
            print("Git: no changes to commit")
        else:
            print(f"Git: {result.stdout.strip()}")
    except Exception as e:
        print(f"Git auto-commit skipped: {e}")


if __name__ == '__main__':
    main()
