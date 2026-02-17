"""
Create a language-cleaned side version of the paper.

Removes typical AI-generated words and expressions, making the prose
sound more natural and human-written. Saves as v20_language_edit.docx.

The user can then use Word's Review > Compare Documents to see all
changes as tracked changes against the original v20.
"""

import pathlib
import shutil
import tempfile

from docx import Document
from docx.oxml.ns import qn

DOCS = pathlib.Path(r"F:\onedrive\__documents\papers\FLOPsExport\Documents")
SRC = DOCS / "flop_trade_model_v20.docx"
DST = DOCS / "flop_trade_model_v20_language_edit.docx"

# Each entry: (old_text, new_text, description)
REPLACEMENTS = [
    (
        "The economic stakes are substantial. A 40 MW",
        "A 40 MW",
        "Intro [12]: remove generic AI intensifier",
    ),
    (
        "Several recent papers address compute governance.",
        "Several papers address compute governance directly.",
        "Lit review [18]: drop 'recent', add 'directly'",
    ),
    (
        "A growing literature examines the determinants of data center location",
        "Several studies examine the determinants of data center location",
        "Lit review [17]: 'growing literature' is AI-speak",
    ),
    (
        # Word uses smart right-single-quote (U+2019)
        "the model\u2019s core prediction that",
        "the prediction that",
        "Discussion [68]: drop AI-filler 'core'",
    ),
]


def replace_in_paragraph(para_element, old_text, new_text):
    """Replace *old_text* with *new_text* inside a paragraph's XML element.

    Handles both the simple case (target sits inside one <w:t>) and the
    complex case (target spans multiple <w:t> elements across runs).
    """
    t_elements = [
        t
        for r in para_element.findall(f'.//{qn("w:r")}')
        for t in r.findall(qn("w:t"))
    ]
    full_text = "".join(t.text or "" for t in t_elements)

    if old_text not in full_text:
        return False

    # Simple case: entirely within one <w:t>
    for t in t_elements:
        if t.text and old_text in t.text:
            t.text = t.text.replace(old_text, new_text, 1)
            return True

    # Complex case: spans multiple <w:t> elements
    pos = 0
    char_map = []
    for t in t_elements:
        length = len(t.text or "")
        char_map.append((t, pos, pos + length))
        pos += length

    idx = full_text.find(old_text)
    if idx == -1:
        return False
    end_idx = idx + len(old_text)

    involved = [(t, s, e) for t, s, e in char_map if e > idx and s < end_idx]
    if not involved:
        return False

    first_t, first_start, _ = involved[0]
    first_t.text = (first_t.text or "")[:idx - first_start] + new_text

    last_t, last_start, _ = involved[-1]
    if last_t is not first_t:
        last_t.text = (last_t.text or "")[end_idx - last_start:]

    for t, _, _ in involved[1:-1]:
        if t is not first_t and t is not last_t:
            t.text = ""

    return True


def search_paragraphs(paragraphs, old_text, new_text):
    """Try to apply a replacement in a list of paragraphs. Return True on success."""
    for para in paragraphs:
        if replace_in_paragraph(para._element, old_text, new_text):
            return True
    return False


def main():
    print("Creating language-cleaned side version...")
    print(f"  Source: {SRC}")
    print(f"  Destination: {DST}")

    # Copy to temp first (avoids OneDrive / Word lock issues)
    tmp = pathlib.Path(tempfile.gettempdir()) / "v20_language_edit_tmp.docx"
    shutil.copy2(SRC, tmp)
    doc = Document(str(tmp))

    n_applied = 0
    for old_text, new_text, desc in REPLACEMENTS:
        # Body paragraphs
        if search_paragraphs(doc.paragraphs, old_text, new_text):
            n_applied += 1
            print(f"  [OK] {desc}")
            continue

        # Table cells
        found = False
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if search_paragraphs(cell.paragraphs, old_text, new_text):
                        n_applied += 1
                        print(f"  [OK] {desc} (in table)")
                        found = True
                        break
                    if found:
                        break
                if found:
                    break
            if found:
                break

        if not found:
            print(f"  [!!] NOT FOUND: {desc}")
            print(f"        Searched for: {old_text[:80]}...")

    doc.save(str(tmp))
    shutil.copy2(tmp, DST)
    tmp.unlink()

    print(f"\nDone. Applied {n_applied}/{len(REPLACEMENTS)} replacements.")
    print(f"Saved: {DST}")
    print()
    print("To review changes with track changes:")
    print("  1. Open flop_trade_model_v20.docx in Word")
    print("  2. Review > Compare > Compare Documents")
    print("  3. Select v20_language_edit.docx as the revised document")


if __name__ == "__main__":
    main()
