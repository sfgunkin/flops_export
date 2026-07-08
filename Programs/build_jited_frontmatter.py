"""
Build the JITED submission front-matter artifacts (separate from the
double-blind manuscript flop_trade_model_v34_anon.docx):

  Documents/JITED_submission/JITED_title_page.docx   (non-anonymous)
  Documents/JITED_submission/JITED_cover_letter.docx

Author decisions baked in (2026-07-08):
  - Generative-AI declaration: AI tools assisted TEXT EDITING only.
  - Funding: none (World Bank staff research).
  - Data availability: World Bank Reproducible Research Repository (RRR).
  - Keywords trimmed 7 -> 6 (dropped "electricity costs") to meet JITED's max 6.

[TODO] markers flag the few items the author must confirm/fill.
"""
import pathlib
import zipfile
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches

ROOT = pathlib.Path(r"F:\onedrive\__documents\papers\_Submitted\FLOPsExport")
OUT = ROOT / "Documents" / "JITED_submission"
OUT.mkdir(parents=True, exist_ok=True)
ANON = ROOT / "Documents" / "flop_trade_model_v34_anon.docx"

TNR = "Times New Roman"
TITLE = "Cheap Energy Might Not Be Enough: A Trade Model of AI Compute Services"
ABSTRACT = (
    "Can energy-rich developing countries turn cheap electricity into AI "
    "compute exports? We develop a capacity-constrained trade model of AI "
    "compute services with bilateral frictions in delivery, regulation, and "
    "trust. Calibrating the model across 85 countries shows that several "
    "developing economies can produce compute at low cost, but this advantage "
    "rarely becomes export competitiveness. Hardware dominates unit costs and "
    "is globally priced, leaving only a 12–20 percent cross-country "
    "production-cost spread. Modest regulatory, financing, and trust frictions "
    "can therefore erase the gains from cheap power. The binding constraint is "
    "institutional credibility rather than electricity prices: enforceable "
    "data governance, stable regulation, credible power contracts, access to "
    "finance, and geopolitical alignment with buyers."
)
KEYWORDS = ("compute trade; FLOPs; artificial intelligence; data centers; "
            "comparative advantage; developing countries")
JEL = "F14, F18, L86, O14, O33, Q40"


def word_count(docx_path):
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    z = zipfile.ZipFile(docx_path)
    root = ET.fromstring(z.read("word/document.xml"))
    text = " ".join(t.text or "" for t in root.iter(ns + "t"))
    return len(text.split())


def base_style(doc):
    st = doc.styles["Normal"]
    st.font.name = TNR
    st.font.size = Pt(12)
    doc.sections[0].left_margin = Inches(1)
    doc.sections[0].right_margin = Inches(1)


def para(doc, text="", *, bold=False, align=None, before=0, after=6,
         size=12, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        r.bold = bold
        r.italic = italic
        r.font.name = TNR
        r.font.size = Pt(size)
    return p


def labelled(doc, label, body, *, before=6):
    """A statement paragraph: bold run-in label, then body text."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(label)
    r.bold = True
    r.font.name = TNR
    r.font.size = Pt(12)
    r2 = p.add_run(body)
    r2.font.name = TNR
    r2.font.size = Pt(12)
    return p


# ─────────────────────────────────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────────────────────────────────
def build_title_page(wc):
    doc = Document()
    base_style(doc)

    para(doc, TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=15,
         after=14)

    para(doc, "Michael Lokshin", align=WD_ALIGN_PARAGRAPH.CENTER, after=2,
         size=12)
    para(doc, "World Bank, Washington, DC, USA",
         align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, after=2, size=11)
    para(doc, "Corresponding author: mlokshin@worldbank.org",
         align=WD_ALIGN_PARAGRAPH.CENTER, after=2, size=11)
    para(doc, "ORCiD: [TODO – insert ORCiD iD]",
         align=WD_ALIGN_PARAGRAPH.CENTER, after=14, size=11)

    ap = doc.add_paragraph()
    ap.paragraph_format.space_after = Pt(6)
    r = ap.add_run("Abstract. ")
    r.bold = True
    r.font.name = TNR
    r.font.size = Pt(12)
    r2 = ap.add_run(ABSTRACT)
    r2.font.name = TNR
    r2.font.size = Pt(12)
    ap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    labelled(doc, "Keywords: ", KEYWORDS)
    labelled(doc, "JEL classification: ", JEL)
    labelled(doc, "Article type: ", "Original Article")
    labelled(doc, "Word count: ",
             f"approximately {wc:,} words (including tables, notes, and "
             f"appendices).")

    # Required statements
    para(doc, "Statements and Declarations", bold=True, before=14, after=6,
         size=12)
    labelled(doc, "Acknowledgements. ",
             "The author thanks Ivan Torre, Christine Ann McDaniel, and Erhan "
             "Artuc for constructive comments. The findings, interpretations, "
             "and conclusions are entirely those of the author and do not "
             "necessarily represent the views of the World Bank, its Executive "
             "Directors, or the countries they represent.")
    labelled(doc, "Funding. ",
             "This work received no external funding. It was prepared as part "
             "of the author’s research at the World Bank.")
    labelled(doc, "Disclosure of interest. ",
             "The author reports there are no competing interests to declare.")
    labelled(doc, "Declaration of generative AI use. ",
             "The author used generative AI tools to assist with text editing. "
             "The model, the analysis, and all conclusions are the author’s "
             "own.")
    labelled(doc, "Data availability. ",
             "The data and code that support the findings of this study are "
             "available in the World Bank Reproducible Research Repository "
             "(reproducibility report RR_EUR_2026_669). "
             "[TODO – insert the RRR DOI/URL once assigned.]")
    labelled(doc, "CRediT author statement. ",
             "Michael Lokshin: Conceptualization; Methodology; Formal "
             "analysis; Software; Data curation; Writing – original draft; "
             "Writing – review & editing.")

    doc.core_properties.author = "Michael Lokshin"
    path = OUT / "JITED_title_page.docx"
    doc.save(str(path))
    return path


# ─────────────────────────────────────────────────────────────────────────
# COVER LETTER
# ─────────────────────────────────────────────────────────────────────────
def build_cover_letter():
    doc = Document()
    base_style(doc)

    para(doc, "8 July 2026", after=12)
    para(doc, "The Editors", after=0)
    para(doc, "The Journal of International Trade & Economic Development",
         after=12)

    rp = doc.add_paragraph()
    rp.paragraph_format.space_after = Pt(12)
    r = rp.add_run("Re: ")
    r.bold = True
    r.font.name = TNR
    r.font.size = Pt(12)
    r2 = rp.add_run(
        "Submission of “Cheap Energy Might Not Be Enough: A Trade Model "
        "of AI Compute Services” (Original Article)")
    r2.font.name = TNR
    r2.font.size = Pt(12)

    para(doc, "Dear Editors,", after=8)

    para(doc,
         "I am pleased to submit the manuscript “Cheap Energy Might Not Be "
         "Enough: A Trade Model of AI Compute Services” for consideration "
         "as an Original Article in The Journal of International Trade & "
         "Economic Development.", align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8)

    para(doc,
         "The paper develops the first formal trade model of compute services "
         "(“FLOP exporting”) and asks whether energy-rich developing "
         "countries can convert cheap electricity into exports of AI compute. "
         "It builds a capacity-constrained model with bilateral frictions in "
         "delivery, regulation, and trust, and calibrates it across 85 "
         "countries. The central finding is that cheap power is not enough: "
         "because globally priced hardware dominates unit cost, the "
         "cross-country production-cost spread is only 12–20 percent, so "
         "modest institutional and financing frictions erase the advantage of "
         "low-cost electricity. The binding constraints are institutional "
         "credibility and the cost of capital rather than the power tariff, "
         "and the paper draws out the financing and regulatory levers that "
         "would let developing countries compete. Combining a formal model "
         "with a multi-country calibration at the trade–development "
         "interface, the paper fits squarely within the journal’s aims and "
         "scope.", align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8)

    para(doc,
         "This is original work. The manuscript is not published, is not under "
         "consideration for publication elsewhere, and has a single author. "
         "The data and code that reproduce every exhibit are available in the "
         "World Bank Reproducible Research Repository, and a data-availability "
         "statement is included on the title page.",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8)

    para(doc,
         "In accordance with the journal’s double-anonymous review policy, "
         "I have uploaded a fully anonymized manuscript together with a "
         "separate title page carrying author details and the required "
         "statements. I confirm the submission fee and understand the "
         "manuscript will be screened for originality. I would be glad to "
         "suggest potential reviewers if that would be helpful "
         "[TODO – optional: add 2–3 names/affiliations].",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8)

    para(doc, "Thank you for considering the manuscript.",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=12)

    para(doc, "Sincerely,", after=2)
    para(doc, "Michael Lokshin", after=0)
    para(doc, "World Bank, Washington, DC, USA", after=0, italic=True, size=11)
    para(doc, "mlokshin@worldbank.org", after=0, size=11)

    doc.core_properties.author = "Michael Lokshin"
    path = OUT / "JITED_cover_letter.docx"
    doc.save(str(path))
    return path


if __name__ == "__main__":
    wc = word_count(ANON)
    tp = build_title_page(wc)
    cl = build_cover_letter()
    print(f"Manuscript word count (anon): {wc:,}")
    print(f"Saved: {tp}")
    print(f"Saved: {cl}")
